from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, size=10.5):
    run.font.name = "等线"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(size)


def append_record(stem, heading, paragraphs):
    docx_path = DOCS / f"{stem}.docx"
    document = Document(docx_path)
    if not any(paragraph.text.strip() == heading for paragraph in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, 16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(docx_path)
    with ZipFile(docx_path) as archive:
        assert archive.testzip() is None

    txt_path = DOCS / f"{stem}.txt"
    text = txt_path.read_text(encoding="utf-8")
    if heading not in text:
        txt_path.write_text(
            text + "\n\n" + heading + "\n" + "\n".join(paragraphs) + "\n",
            encoding="utf-8",
        )


records = {
    "AI开发项目_Bug记录模板": (
        "v968 共享反应、人声、角色日记与日常变化修复记录（2026-08-18）",
        [
            "屏幕共享开始和结束不再只给模型一个状态词。v968 会把本通电话近期真实对话、当前共享状态和可用画面信息一起交给角色，让角色按人设自然接住上下文；通知式复述、无自然反应或外语台词缺少逐句中文翻译时，只用完全相同的上下文重试一次，仍不合格才放弃，避免无限重试和机械台词。",
            "共享外部视频时，网页的角色音量设置按 0–300% 完整传给私人 App。原生端仅对角色 TTS 使用动态压缩和增益，不重新配置正在承载抖音、B站等媒体的共享音频会话，目标是在保留后台视频声音的同时提高角色可听度。Windows 已完成静态与回归验证；实际扬声器、蓝牙和不同媒体音量比例仍须真机验收。",
            "角色手机日记改为独立四位密码，默认取相遇日期 MMDD，角色可用独立控制指令修改，不影响手机总密码。第一次输错即记录并让角色在后续消息中按人设自然发现。正确解锁日记应用时才按需准备当天一篇日记，只依据当天真实微信、语音／视频通话和共同生活记录，长度可长可短、时间随机；界面只显示“正在打开”，不暴露生成、模型、系统或字幕提示。",
            "日记首页保持紧凑列表，点击当天条目后进入全屏横线纸阅读；未收藏日记自生成起 24 小时后消失，收藏后永久保留，取消收藏则重新开始 24 小时期限。旧数据中没有 expiresAt 的日记不会被批量删除。通话记录内容默认折叠，点击后展开。角色微信允许修改“我”的备注，但角色只在下一次真实聊天时发现并自然反应。",
            "角色推特结合角色人设、X／圈子设置和当天真实生活，每次优先换具体细节和角度。去重采用宽松门槛：只有完全相同或几乎逐字重复才硬拦截；相似主题只触发最多一次温和改写，第二次只要不是硬重复就允许发布，避免规则过严导致模型写不出来。",
            "网页版本为 v968，私人 iOS 为 1.0.90 (90)，原生桥保持 23。公开网页仍只隔离真实 iPhone 伴生数据和原生控制上下文，不删除原有查岗、权限、锁定、限额或其他应用功能。PhoneWeb.bundle 已从共享源码重建，Windows 全量 Node 回归通过；当前环境没有 Mac/Xcode 和真实 iPhone，未宣称编译、签名或真机声音比例已经验收。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "发布补充规范｜上下文事件、按需日记与宽松内容去重（1.0.90 起）",
        [
            "向角色注入屏幕共享开始、结束等状态事件时，必须同时提供本轮真实上下文，并要求角色表现自然反应；不能把“开始了／结束了”当作可直接展示的角色回复。外语通话仍须保持原文与逐句中文翻译。质量修复最多一次，且重试必须复用原上下文，不能把上下文替换成孤立状态词。",
            "角色私密应用的密码域必须彼此独立。日记默认密码可由相遇日期 MMDD 派生，角色修改日记密码不得连带修改手机密码；第一次错误输入即形成可被角色发现的真实事件。日记只在用户打开应用时按需准备当天内容，不得在后台累计空白日期或自动补写多日；清理只删除明确带 expiresAt 且到期的未收藏条目，旧无期限数据必须保留。",
            "生成日常推特时，必须结合角色设置和近期真实生活，但内容去重不能阻断正常表达。仅完全相同或近乎逐字重复可硬拦截；宽泛主题相似只允许触发一次改写提示，改写后只要没有达到硬重复就应放行。不得因为角色连续谈到同一生活主题而反复要求生成，造成无回复。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "当前交付基线｜网页 v968／私人 iOS 1.0.90 (90)／原生桥 23（2026-08-18）",
        [
            "v968 让屏幕共享开始／结束事件携带本通电话真实上下文，并对机械状态复述、缺少自然反应或外语无逐句翻译的结果按同一上下文重试一次。共享外部媒体期间，角色 TTS 使用私人 App 原生压缩增益链提升人声，不重置正在播放视频的共享音频会话。",
            "角色手机新增独立日记密码和按需当天日记：默认密码取相遇日期 MMDD，第一次输错会被角色发现，角色可独立改密。日记只在打开时准备当天一篇，采用紧凑列表和全屏横线纸阅读；未收藏 24 小时到期，收藏永久保留，旧无期限条目不误删。通话记录默认折叠。",
            "角色微信支持修改用户自己的备注，角色到下一次真实聊天才发现；角色推特结合人设、圈子设置和真实日常，并使用最多一次温和改写的宽松去重，只有完全相同或几乎逐字重复才硬拦截。公开网页的数据隔离不删除任何原有业务功能。",
            "当前版本为网页 v968、私人 iOS 1.0.90 (90)、原生桥 23。PhoneWeb.bundle 已从共享源码重建并通过 Windows 自动测试；仍须在 Mac 编译签名，并在真实 iPhone 验证屏幕共享反应、逐句翻译、后台视频与角色人声比例、日记密码和跨 24 小时保留行为。",
        ],
    ),
}


for stem, (heading, paragraphs) in records.items():
    append_record(stem, heading, paragraphs)
