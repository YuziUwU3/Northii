from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, name="等线", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def append(name, heading, paragraphs):
    path = DOCS / name
    document = Document(path)
    if not any(p.text.strip() == heading for p in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, size=16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(path)
    with ZipFile(path) as archive:
        assert archive.testzip() is None


append("AI开发项目_Bug记录模板.docx", "v920／私人版 1.0.44 Bug 记录｜角色声音稳定回退（2026-08-14）", [
    "真机结果：v919 未解决。用户在未开启屏幕共享的普通语音电话和普通视频通话中，只能看到自己的字幕，角色没有声音，也没有角色字幕。因此 v919 不得记为真机完成。",
    "版本证据：逐项对照 v800（cc4e7b4）、v850（caac487）和 v907（bfcb17b）。v800、v850 在角色忙碌、角色生成或播音期间直接忽略识别结果，不把回声和杂音排队成下一条用户消息；角色回复后分别保留约 1.2 秒和 1.5 秒的尾音隔离。v907 还会在角色播音前暂停原生识别，播音后等待音频路由稳定并重建识别。",
    "修复：普通前台通话恢复 v907 已使用的网页 audio 播放器；只有应用进入后台或画中画时才保留 v910 以后需要的原生 call.audio.play。免提识别恢复 v800／v850 的忙碌期丢弃和尾音隔离，并恢复 v907 的暂停、等待、重建流程；所有等待都增加上限和失败回退，避免识别桥异常时卡住角色回复。v918 新增的 RMS、置信度和 voiceActivity 门槛已撤回。",
    "保护范围：本轮不修改用户字幕、角色字幕、共享小框字幕的 DOM、CSS、关键帧和动画参数；不修改角色回复文本、TTS 文案、屏幕共享逻辑、宠物逻辑和微信输入框。前台旧播放器与后台原生播放器被明确分流，避免为后台能力再次替换普通通话的稳定链路。",
    "验证：Windows 全量自动测试 566 项通过、0 项失败；版本号、缓存号和 Xcode 工程静态检查通过后方可打包。Windows 结果不能代替 Mac 编译和真实 iPhone；新包必须先测不共享屏幕的普通电话与普通视频，确认角色首句、连续回复、声音和角色字幕，再测静音环境误识别。",
])

append("AI开发项目_Bug修改规范.docx", "新增强制规范｜旧版三点证据与前后台音频分流（v920 起）", [
    "用户指出历史正常版本时，必须至少核对：最后正常版本的输入条件、核心执行链和收尾恢复三个位置。不能只搜索同名函数，也不能把用户的猜测直接当作根因；结论必须能够指出旧版与当前版的实际代码差异。",
    "涉及通话声音时，普通前台通话、后台、画中画、屏幕共享必须分别列为不同场景。为后台续播加入的原生播放器不得自动取代普通前台已经验证的网页播放器；需要共存时必须用明确的可测试条件分流。",
    "涉及语音识别和降噪时，先恢复用户确认正常版本的忙碌期隔离、尾音窗口、暂停和重建顺序，再做独立实验。不得在没有真机证据时叠加 RMS、置信度、文字相似度等猜测阈值，也不得把角色播音期间的识别结果排队后自动发送。",
    "声音修复不得顺带修改字幕动画。必须用回归测试锁定字幕关键帧、完整短句显现方式、PiP 字号和翻译布局；只有用户重新明确要求时才能单独修改。自动测试通过必须表述为代码检查通过，不得写成真机已修复。",
])

append("AI开发项目_项目说明文档.docx", "v920／私人版 1.0.44｜通话角色声音稳定回退（2026-08-14）", [
    "当前版本：网页 v920；私人 iOS 1.0.44 (44)；原生桥契约 18。PhoneWeb.bundle 由同一份 v920 网页核心重新生成。v919 在真实 iPhone 上仍然没有角色声音和角色字幕，故本版是基于历史正常代码的保守回退，不宣称真机完成。",
    "前台角色声音链恢复 v907 结构：TTS 音频由网页 audio 播放；原生语音识别在播音前暂停，播音结束后等待约 760 毫秒再重建。暂停和重建均有超时及兜底，不能无限阻塞 callAI 的 finally 收尾。后台或画中画仍使用原生 call.audio.play，以保留离开应用后的续播能力。",
    "免提识别恢复 v800／v850 规则：角色回复忙碌期间不消费结果、不显示临时用户字幕、不排队自动发送；角色回复结束后保留尾音隔离窗口。相同最终文本只有在没有新 interim、且像系统重复上报时才丢弃，用户真实重复说同一句仍可发送。",
    "本轮没有调整字幕动画。v918 引入但未获真机证明的 RMS、识别置信度和 voiceActivity 过滤已经移除，原生 voice processing 和 playAndRecord／voiceChat／mixWithOthers 音频会话保留。",
])

append("AI开发项目_新聊天启动说明.docx", "新聊天接手状态｜v920／私人版 1.0.44（2026-08-14）", [
    "当前代码基线：网页 v920、私人 iOS 1.0.44 (44)、原生桥契约 18、PhoneWeb.bundle v920。v919 真机失败，表现为普通电话和普通视频只有用户字幕，角色没有声音和角色字幕。不要把 v919 写成成功版本。",
    "v920 已对照 v800、v850、v907：恢复忙碌期丢弃识别、1.2／1.5 秒尾音隔离、角色播音前暂停识别、播音后等待并重建；普通前台恢复网页 audio，后台／PiP 才走原生 call.audio.play。字幕动画完全不在本轮修改范围。",
    "Windows 全量自动测试为 566 项通过、0 项失败；仍必须在 Mac 编译并在真实 iPhone 验收。第一顺序：不共享屏幕的普通电话；第二顺序：不共享屏幕的普通视频；每种都验证角色首句、连续两句、声音、角色字幕和用户字幕。第三顺序才测静音环境是否误识别。",
    "若真机仍无角色声音，必须收集 callAI lastError、TTS 请求结果、audio onplaying／onerror 和原生 speech.pause／rebuild 的日志，再定位具体停点；不得再次猜测阈值、不得修改字幕动画，也不得把前台播放无条件切回原生播放器。",
])
