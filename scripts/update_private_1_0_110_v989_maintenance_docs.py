from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, size=10.5):
    run.font.name = "等线"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(size)


def append_record(stem, heading, paragraphs):
    docx_path = DOCS / f"{stem}.docx"
    document = Document(docx_path)
    if not any(paragraph.text.strip() == heading for paragraph in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, 16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(docx_path)
    with ZipFile(docx_path) as archive:
        assert archive.testzip() is None

    txt_path = DOCS / f"{stem}.txt"
    text = txt_path.read_text(encoding="utf-8")
    if heading not in text:
        txt_path.write_text(
            text + "\n\n" + heading + "\n" + "\n".join(paragraphs) + "\n",
            encoding="utf-8",
        )


records = {
    "AI开发项目_Bug记录模板": (
        "v989 婚礼单章重做与第一幕身份锚点修复记录（2026-08-19）",
        [
            "现象：整场婚礼的旧生成逻辑会从任意一张已经成功的画面中寻找人物参考，导致后续章节或失败重试可能参考相邻章节，而不是统一参考第一幕；用户对单张画面不满意时只能重做整场，已经满意的五张也会被覆盖。",
            "根因：weddingIdentityReference 会遍历六幕并返回第一张可用的非当前画面；整场生成先并行完成所有首次尝试，再统一处理失败幕，因此第一幕尚未确立时，第 2～6 幕可能没有稳定图像锚点。准备数据也没有单幕版本键和只更新一个 sceneKeys 项的入口。",
            "修复：第一幕改为整场唯一身份锚点。整场按章节顺序生成，每一幕首次失败后立即完成两次自动重试，再开始下一幕；第 2～6 幕的首次生成与重试只引用第一幕。参考图先转换为可传输 data URL；若接口不支持参考图，不再静默退回无参考的纯文字生图，而是进入该幕重试与预设图回退。",
            "新增入口：线下约会的已完成现代／中式婚礼卡片增加“单独重做某一章”。六章可独立选择；重做第 2～6 章只参考当前第一幕，重做第一章则参考当前已保存的第一幕。每章最多三次生成，旧图在新图通过动作和人物一致性复核、成功写入独立缓存键之前始终保留；三次失败不覆盖原图，其他五章和婚礼记忆均不改变。",
            "版本与验证边界：网页 v989，私人 iOS 1.0.110 (110)，原生桥 23。PhoneWeb.bundle 必须由 v989 共享清单重建；Windows 全量回归通过后仍须在 Mac 编译签名，并在真实 iPhone 使用实际生图接口检查 data URL 参考图、长耗时重试和单幕覆盖。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "婚礼人物一致性与单幕覆盖规范（v989／1.0.110 起）",
        [
            "同一场婚礼必须只有一个人物图像身份锚点，固定为第一幕。第 2～6 幕的首次生成、动作复核失败重试和用户主动单幕重做都只能参考第一幕，不得改为最近成功幕、相邻幕或随机成功幕。提示词身份摘要不能替代实际第一幕图片输入。",
            "整场准备必须先让第一幕进入成功图或预设回退终态，再生成第二幕。每一幕的首次失败应当在进入下一幕前完成最多两次重试，避免后续幕在首幕身份尚未确定时独立随机人物。接口不能接收参考图时，不得静默降级为无参考生图并假装保持一致。",
            "单幕重做属于安全覆盖事务：开始前读取并保留原 src、verified 与准备记录；新图最多尝试三次，只有通过当前幕动作校验与第一幕人物一致性校验、写入新缓存键后，才更新 prepared.sceneKeys[scene]。任一步失败都恢复原图，不改其他 sceneKeys、婚书、关系或记忆。",
            "第一幕自身单独重做时必须把当前旧第一幕作为输入参考，以保持原人物，只修正动作与构图。若用户希望换成全新人物，必须走整场重新生成，不能借单幕按钮暗中换掉身份锚点。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "当前交付基线｜网页 v989／私人 iOS 1.0.110 (110)／原生桥 23（2026-08-19）",
        [
            "现代与中式婚礼的第一幕现在是整场唯一人物身份基准。第 2～6 幕首次生成、每幕两次自动重试和后续单幕重做均使用第一幕真实图片作为输入参考，锁定同一张脸、五官、发型、发色、发冠与礼服；不再从任意成功章节选参考。",
            "线下约会中，已经准备完成的现代或中式婚礼新增“单独重做某一章”。用户可选择六章中的任意一章。每次重做最多三次，旧图直到新图通过动作与人物一致性复核并写入独立缓存后才覆盖；连续失败时原图与其他五章保持不变。第一章重做使用当前旧第一章自身作为参考。",
            "整场新生成调整为严格顺序：第一章完成首次生成及必要重试后才开始第二章，以此类推。若生图接口不支持参考图片，后续章节进入重试／预设图安全回退，不再无参考随机生成并冒充同一人物。现代与中式的婚书、关系、两场独立记忆、固定旁白和角色主模型台词逻辑保持不变。",
            "当前版本为网页 v989、私人 iOS 1.0.110 (110)、原生桥 23。共享网页与 PhoneWeb.bundle 从同一清单同步；Windows 自动测试、包结构与文档渲染通过后，仍须在 Mac 编译签名，并在真实 iPhone 验证实际生图接口的首幕参考传输、单幕三次尝试、失败保留原图和成功只覆盖所选章节。",
        ],
    ),
}


for stem, (heading, paragraphs) in records.items():
    append_record(stem, heading, paragraphs)
