from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, name="等线", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def append(name, heading, paragraphs):
    path = DOCS / name
    document = Document(path)
    if not any(p.text.strip() == heading for p in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, size=16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(path)
    with ZipFile(path) as archive:
        assert archive.testzip() is None


append("AI开发项目_Bug记录模板.docx", "v908／私人 1.0.32 Bug 记录｜后台唤醒未注册、显式任务静默与宠物回窝（2026-08-13）", [
    "现象：查看当前软件立即测试、1 分钟后台通知以及角色生成中退出 App 连续三版实机失败；任务虽已创建，iPhone 退出页面后没有继续处理或通知。宠物睡觉仍叠在窝顶和窗台，而不是粉色窝垫内。",
    "根因：私人 App 的后台静默推送处理器被注册在 CompanionSyncView.task 内，但小手机主界面并不常驻该页面，因此冷启动或普通使用时处理器可能从未注册。服务器的显式 one_minute_test、app_watch_test、reply_handoff、device_handoff 又复用了允许模型输出“保持安静”的主动消息函数，任务会在真正发 APNs 前被取消，形成任务入库但无通知。宠物 CSS 的 top 坐标方向和数量布局与实际背景窝位不一致。",
    "修复：后台处理器改为 App 启动时注册，私人版唤醒后继续执行真实设备同步；四类显式任务禁止静默候选并强制生成真实回复。可见 APNs 保存状态码和 apns-id，静默唤醒也把状态、apns-id、错误和请求时间写入命令 result；命令完成改为合并 result，避免覆盖唤醒证据。宠物按物种和睡眠数量重新计算尺寸与坐标，名字隐藏，全部放进左侧粉色窝垫。",
    "验证：Windows Node 完整自动回归 503/503 通过；宠物四只同时睡觉的无头浏览器截图已人工确认在窝垫内；迁移 202608130003 已单独应用并登记，phone-role-push 已部署。Windows 无法执行 Xcode 五 Target 编译和真实 iPhone APNs，仍需 Mac 与真机验收；用户手动强制划掉 App 后，iOS 不保证 silent push 唤醒。",
])

append("AI开发项目_Bug修改规范.docx", "新增强制规范｜后台入口、显式任务与 APNs 可观测性（v908 起）", [
    "依赖静默推送的后台处理器必须在 UIApplication 启动阶段注册，禁止只放在某个可选页面、弹窗或 SwiftUI View.task 中。处理器应在系统允许的后台时间内完成必要工作并准确回报成功或失败。",
    "用户点击测试、正在回复接力、正在读取设备接力均为显式任务，不得复用允许“保持安静”的自主消息决策。显式任务要么生成并持久化角色回复，要么保留可重试失败状态；不得以 canceled 冒充成功。",
    "每次 APNs 请求必须保留 HTTP 状态、apns-id 或真实错误；设备回执只能与原唤醒诊断合并，禁止整段覆盖。排查通知时先区分任务生成、APNs 接受、设备唤醒、设备执行和可见通知五个阶段。",
    "带背景定位的宠物、人物或装饰改动必须用实际页面截图验收。CSS top 增大表示向下，坐标不得只凭源码猜测；多对象布局必须覆盖最大数量。",
])

append("AI开发项目_项目说明文档.docx", "v908／私人 1.0.32｜后台根因修复与宠物精确回窝（2026-08-13）", [
    "网页版升级为 v908，私人 iOS 安装版升级为 1.0.32 (32)。本版不扩大功能范围，集中修复三项连续失败的后台核心链路：1 分钟测试、查看当前软件立即测试，以及角色回复或设备查看中途退出 App 后的服务器接力与通知。",
    "私人 App 在启动时安装后台推送处理器；后台任务不再依赖用户是否打开过伴生同步页面。服务器显式任务不允许角色选择静默，同时记录 APNs 接受状态和设备回执，便于区分上游生成、推送和本机执行问题。迁移 202608130003 与 phone-role-push 已在线部署。",
    "宠物睡眠改为按同时睡觉数量和物种缩放，四只宠物都落在左侧粉色窝垫内，不再停在窝顶或窗台；已经通过实际渲染截图检查。共享 PhoneWeb.bundle 已重新生成。Windows 自动回归 503/503 通过，Mac 编译与真实 iPhone 后台通知仍需用户验收。",
])

append("AI开发项目_新聊天启动说明.docx", "新聊天接手状态｜v908／私人 1.0.32 已部署候选（2026-08-13）", [
    "当前基线：网页版 v908、私人 iOS 1.0.32 (32)、原生桥契约 11。迁移 202608130003_background_delivery_diagnostics.sql 已单独在线应用并登记，phone-role-push 已部署；不要为了同步迁移历史而重放远端未登记的旧迁移。",
    "本版后台根因：推送处理器曾只在非主入口 CompanionSyncView.task 注册，显式任务又允许模型静默。现已在 App 启动注册处理器，one_minute_test、app_watch_test、reply_handoff、device_handoff 强制真实结果，APNs 与设备回执均可诊断。不要退回仅靠页面 timer 或页面生命周期的方案。",
    "真机优先验收：开始 1 分钟测试后回桌面；角色正在回复时回桌面；立即查看当前软件后切到另一 App 并等待角色消息；最后查看数据库 outbox 的 push_diagnostic 和 command result.wake/deviceAcknowledgedAt。普通回桌面和锁屏应可工作，用户手动强制划掉 App 后 iOS 不保证 silent push。",
    "宠物四只睡觉的最终坐标已通过截图确认都在粉色窝垫内。Windows 自动回归 503/503 已通过；Mac 五 Target 编译、签名和真机 APNs 是剩余证据边界。音乐在线搜索尚未加入本版，后续须独立设计，不能影响现有本地导入和播放。",
])
