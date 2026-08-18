from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, name="等线", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def append(name, heading, paragraphs):
    path = DOCS / name
    document = Document(path)
    if not any(p.text.strip() == heading for p in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, size=16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(path)
    with ZipFile(path) as archive:
        assert archive.testzip() is None


append("AI开发项目_Bug记录模板.docx", "v906／私人 1.0.30 Bug 记录｜后台测试无反馈与心情延迟（2026-08-13）", [
    "现象：角色回复生成期间退出页面仍会中断且没有系统通知；一分钟真实测试点击后没有可见进度。聊天页顶部心情必须重新进入或点击后才显示新内容。",
    "根因：回复接管任务初次登记延后十分钟，页面隐藏时的 keepalive 提前请求若被 iOS 冻结，就失去及时兜底；测试入口只用短提示返回，且没有把主动开关、服务器资料、通知令牌和分钟任务的缺项展示出来。心情节点只在整页 render 时生成，新增消息只刷新消息列表。",
    "修复：回复接管初次持久任务改为一分钟兜底，页面隐藏仍尝试提前到五秒；一分钟测试先同步角色资料并核验通知令牌、资料和分钟任务，失败弹出具体缺项，成功显示任务编号。聊天心情栏始终保留可更新节点，每次消息追加或列表刷新同步文本、显示状态和点击动作。",
    "新增真实 App 立即测试：独立 app_watch_test 任务立即请求一次新鲜授权软件用量，不计正式每日次数和九十分钟冷却；识别后通知角色，用户五分钟仍未回复时沿用单选询问或锁定。没有新鲜变化时明确测试失败，禁止旧快照或软件名猜测。",
])
append("AI开发项目_Bug修改规范.docx", "新增强制规范｜后台测试必须有可观察回执（v906起）", [
    "任何面向用户的后台、通知或设备测试按钮都必须在开始前校验所有硬前提，并在界面中显示成功任务编号或逐项失败原因。不得只发短 toast 后静默返回，也不得把设置值已保存当成任务已入库。",
    "iOS 页面隐藏事件只能作为加速信号，不能作为唯一持久化时机。需要后台续跑的任务必须在前台操作发生时先写入一个保守兜底到期时间，隐藏时再 best-effort 提前；前台正常完成必须取消，避免双回复。",
    "实时标题、心情和状态栏若要求随消息变化，局部消息追加路径必须同步更新对应节点；不能依赖整页 render、重新进入页面或用户点击触发刷新。",
])
append("AI开发项目_项目说明文档.docx", "v906／私人 1.0.30｜后台实测、实时心情与宠物健康（2026-08-13）", [
    "苹果兼容适配开启并在私人 App／苹果主屏版生效时，视频通话顶部角色备注和通话时间整体下移十四像素；关闭开关、安卓和普通浏览器保持原位置。聊天页心情在角色消息到达后即时刷新。",
    "电子宠物家庭从最多两只扩为四只，四只拥有独立房间、睡眠、追球和一起玩站位。按最后喂食时间计算：两天未喂体态缩瘦，超过三天进入生病状态并显示头顶气泡；普通喂食不代替治疗，商城恢复药在成长手册中使用后恢复健康。",
    "后台接管的前台兜底改为一分钟，一分钟通知测试新增完整链路预检和任务编号。App 感知增加立即真实测试，不消耗正式次数或九十分钟冷却；取得新鲜软件变化后角色通知，五分钟未回复继续询问或锁定。需要应用迁移 202608130001 并重新部署 phone-role-push。",
])
append("AI开发项目_新聊天启动说明.docx", "新聊天接手状态｜v906／私人 1.0.30 候选（2026-08-13）", [
    "当前候选基线为网页 v906、私人 iOS 1.0.30 (30)。本版包含苹果兼容通话标题下移、聊天心情实时更新、四只宠物与生病用药，以及后台测试可诊断和 App 感知立即测试。",
    "发布必须应用 202608130001_background_app_watch_test.sql 并重新部署 phone-role-push。真机依次验证：苹果开关开关前后标题坐标、心情不重进即更新、回复中退出通知、一分钟测试、立即切换授权 App 后识别通知、五分钟询问或锁定，以及两天／三天宠物时间模拟与用药。",
    "普通后台或锁屏可由静默推送唤醒；用户上划强退后 iOS 仍可能拒绝设备采集唤醒。立即测试只有在切到已授权软件产生新用量后才能识别，禁止用旧快照假报成功。",
])

append("AI开发项目_Bug记录模板.docx", "v906 发布与验证补充（2026-08-13）", [
    "远端迁移 202608130001 已实际应用并登记，phone-role-push 已重新部署。Windows 完整自动回归 493/493 通过；DOCX 结构校验通过。当前机器没有 Xcode 和 LibreOffice，因此五 Target 编译、真实 iPhone 前后台通知／授权 App 读取／锁定以及维护文档的分页渲染仍需在 Mac 或真机验收。",
])
append("AI开发项目_项目说明文档.docx", "v906 发布状态补充（2026-08-13）", [
    "迁移 202608130001_background_app_watch_test.sql 已应用到远端并登记，phone-role-push 已完成部署。Windows 完整自动回归 493/493 通过。发布后的剩余证据边界是 Mac 五 Target 编译与真实 iPhone 的通知、Screen Time 和 Managed Settings 链路。",
])
append("AI开发项目_新聊天启动说明.docx", "v906 已部署状态补充（2026-08-13）", [
    "迁移 202608130001 与 phone-role-push 已部署，不要重复把它们描述为待部署。Windows 自动测试 493/493 通过。下一轮先读取远端任务诊断和真机回执，再判断是否需要改代码；不得用网页版结果替代私人 App 的 Mac 编译与真机证据。",
])
