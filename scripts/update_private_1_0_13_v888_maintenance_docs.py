from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, name="等线", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def append_section(name, heading, paragraphs):
    path = DOCS / name
    document = Document(path)
    if any(p.text.strip() == heading for p in document.paragraphs):
        return
    document.add_page_break()
    title = document.add_heading(heading, level=1)
    for run in title.runs:
        set_font(run, size=16)
    for text in paragraphs:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.25
        set_font(paragraph.add_run(text))
    document.save(path)


SECTIONS = {
    "AI开发项目_Bug修改规范.docx": (
        "新增强制规范｜私人手机号账号不得强依赖境外短信（v888／私人 1.0.13 起）",
        [
            "中国大陆私人单用户 App 在没有已验证短信供应商、发送资质和真机到达率证据时，不得把 Supabase Phone OTP 当成可交付登录方案。手机号仍可作为用户可见账号，但底层认证方式应选择实际可用且可恢复的路径。",
            "本项目固定使用“手机号 + 密码”：原生桥把规范化的 +86 号码确定性映射为仅供 Supabase Auth 使用的内部邮箱，界面、公开 North 和网页版不显示该邮箱，也不持有私人登录态。Phone Provider、Twilio 和公共自助注册保持关闭。",
            "密码只能存在于当前登录请求内，不得写入网页状态、日志、UserDefaults、本地文件或 Keychain；Keychain 只保存访问令牌、刷新令牌、到期时间、用户 ID 和遮罩手机号。云数据继续按 auth.uid() 开启 RLS。",
            "账号回归必须断言：不存在 /auth/v1/otp、account.otp.send、account.otp.verify；密码 grant 使用确定性内部邮箱；旧采集时间不能覆盖新备份；首次登录在云端已有数据时必须由用户明确选择保留本机或恢复云端。",
        ],
    ),
    "AI开发项目_Bug记录模板.docx": (
        "v888／私人 1.0.13 Bug 记录｜大陆手机号 OTP 前置条件不成立（2026-08-11）",
        [
            "v887 已完成私人账号入口、Keychain 会话和按 auth.uid() 隔离的云备份，但登录契约采用 Supabase Phone OTP。真机配置阶段确认用户使用中国大陆手机号、没有 Twilio 或其他短信供应商凭据，Supabase Phone Provider 也保持关闭，因此原方案即使代码通过也无法发送真实验证码。",
            "高风险根因不是网络刷新，而是把“手机号登录”误等同于“必须短信验证码”。境外短信对中国大陆号码的到达率和资质要求没有经过本项目真机验证，继续要求用户配置 Twilio会把恢复入口建立在不可靠外部条件上。",
            "v888 修复为手机号密码登录。App 把 +86 号码映射为 smallphone.86手机号@example.com，通过 Supabase Email/Password grant 登录；密码不落盘，令牌只进 Keychain。公开 North、网页版、锁定状态机、共同生活、时间轴、话筒和键盘布局均不随本次账号替换改动。",
            "首次上线仍需执行 private_phone_accounts 迁移并手动创建一名已确认用户。云端为空才自动上传当前本机；云端已有备份必须让用户明确选择，禁止用空账号或旧快照覆盖已导入的约 4.5 MB 备份。",
        ],
    ),
    "AI开发项目_项目说明文档.docx": (
        "共享网页 v888／私人小手机 1.0.13｜手机号密码与云恢复（2026-08-11）",
        [
            "当前架构仍是公开 North、私人小手机原生 App、安卓/浏览器网页共享业务核心并彼此隔离；公开审核中的 North 工程保持不动，本人设备最终只由私人小手机控制。",
            "私人 1.0.11 已固定 WKWebView 外框并解决电话键盘反复弹跳；1.0.12 建立 Keychain 会话和按账号隔离的云备份；1.0.13 将不可用的大陆手机号短信 OTP 改为“手机号 + 密码”，共享核心同步升级为 v888。连续原生语音、网页安全区和既有锁定规则继续保留。",
            "用户界面只输入中国大陆手机号和密码。原生桥把号码映射为内部邮箱后调用 Supabase password grant，密码不保存；删除或重装后重新登录即可读取同一 auth.uid() 下的云备份。首次登录不会静默覆盖本机数据。",
            "本阶段尚需在 Supabase 执行备份表迁移、创建首名已确认私人用户，并在 Mac 编译安装后做登录、上传、删除重装、恢复、断网和旧快照真机回归。完整大容量原生媒体存储与服务端单控制器租约仍属后续主线。",
        ],
    ),
    "AI开发项目_新聊天启动说明.docx": (
        "新聊天接手状态｜共享网页 v888／私人小手机 1.0.13（2026-08-11）",
        [
            "唯一仓库为 C:\\Users\\pc\\Documents\\小手机\\phone-work，分支 main。共享网页版本 v888，私人小手机版本 1.0.13 (13)。公开审核中的 North 工程保持不动，公开 North 与私人小手机不能同时控制同一台设备。",
            "用户已确认账号目标：App 显示手机号和密码，删除重装后登录恢复即可，不要求短信验证码。不要再引导用户配置 Twilio；Phone Provider 保持关闭。内部 Auth 邮箱按 smallphone.86手机号@example.com 生成，只在 Supabase 和原生桥内部使用。",
            "账号安全边界：密码不保存；Keychain 只存 token、到期时间、用户 ID 和遮罩手机号；云表按 auth.uid() RLS；旧采集时间禁止覆盖新备份；首次登录有云数据时必须由用户选择保留本机或恢复云端。",
            "下一步依次完成三项：一，在 Supabase SQL Editor 执行 202608110001_private_phone_accounts.sql；二，在 Authentication 手动创建已确认的内部邮箱用户和 8 位以上密码；三，在 Mac 安装 1.0.13，先登录并上传，再删除重装验证同手机号密码可以恢复。",
            "必须保留已经真机确认的网页适配、电话输入框固定和连续话筒修复。账号改动不得顺手调整锁定、同步、共同生活、时间轴、页面安全区或公开 North。",
        ],
    ),
}


def main():
    for name, (heading, paragraphs) in SECTIONS.items():
        append_section(name, heading, paragraphs)
        document = Document(DOCS / name)
        assert sum(p.text.strip() == heading for p in document.paragraphs) == 1
        assert not any(
            "codex-file-citation" in p.text or "PLACEHOLDER" in p.text
            for p in document.paragraphs
        )
        with ZipFile(DOCS / name) as archive:
            assert archive.testzip() is None
        print(f"verified {name}: {len(document.paragraphs)} paragraphs")


if __name__ == "__main__":
    main()
