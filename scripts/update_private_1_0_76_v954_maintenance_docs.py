from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, size=10.5):
    run.font.name = "等线"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(size)


def append_record(stem, heading, paragraphs):
    docx_path = DOCS / f"{stem}.docx"
    document = Document(docx_path)
    if not any(paragraph.text.strip() == heading for paragraph in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, 16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(docx_path)
    with ZipFile(docx_path) as archive:
        assert archive.testzip() is None

    txt_path = DOCS / f"{stem}.txt"
    text = txt_path.read_text(encoding="utf-8")
    if heading not in text:
        txt_path.write_text(text + "\n\n" + heading + "\n" + "\n".join(paragraphs) + "\n", encoding="utf-8")


records = {
    "AI开发项目_Bug记录模板": (
        "v954 私人版交付记录｜1.0.76 iOS 音乐导入与小号报备修复（2026-08-16）",
        [
            "现象与根因：iOS 文件选择器把音频和视频写在同一个 accept 过滤器里，部分系统会把已有 MP3/M4A 显示为不可选择；微信小号首次加好友时，请求历史可能只有 system 消息，部分兼容接口因此返回 HTTP 400；切回主号后的报备又依赖模型请求成功，失败后没有可靠交付。私人状态问句还会被“刚刚”等实时词误判为联网查询，旧兜底则直接拼接双方原话，形成机械复述。",
            "处理：音乐导入拆成独立的“选择音乐文件”和“选择录屏视频”，音频明确接受 MP3、M4A、AAC、WAV、FLAC、OGG、OPUS；好友通过自动消息以 user 角色进入首轮请求，并拦截兼容接口的客服式问候。小号事件保存来信、角色回复与处理动作；切回主号时即使模型失败，也会自然概括并交付一次，且不会逐句复述聊天记录。",
            "防误伤：普通明确联网请求仍可联网；“刚刚有人给你发消息吗”“刚才谁加你微信了”等私人状态问题固定留在本地角色记忆。小号仍按陌生身份隔离，主号只接收该次联系的自然报备，同一事件只交付一次。",
            "版本：网页 v954；私人 iOS 1.0.76 (76)；原生桥 18。PhoneWeb.bundle 必须由当前共享源码清单重建，不能从旧 ZIP 覆盖。Windows 自动测试通过后，仍需在 Mac 编译签名，并在真实 iPhone 上验收 Files 选择器、首句回复及小号切回主号报备。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "发布补充规范｜小号跨账号事件必须可交付、不可机械复述（1.0.76 起）",
        [
            "小号联系事件不得只依赖模型成功才标记完成。必须先持久化结构化事实，再生成角色口吻报备；模型失败时使用同一结构化事实自然交付，完成后才推进游标。",
            "报备内容禁止输出“某某：……；我：……”式原始记录、逐句转录、接口错误或客服式问候。只说明谁联系、主要话题和角色如何处理，保持角色口吻和关系上下文。",
            "涉及联系人、好友申请、刚才聊天等本地私人状态的问题，除非用户明确要求搜索互联网，否则不得触发联网。iOS 文件选择器必须按媒体类型拆分，避免混合 accept 造成系统文件灰显。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "当前交付基线｜网页 v954／私人 iOS 1.0.76 (76)（2026-08-16）",
        [
            "网页与私人 App 使用同一套微信统一自然系统。小号保持陌生身份隔离；切回主号后，对该次联系进行一次自然报备。报备失败不丢事件，成功后不重复，同一角色不会把原始双方聊天按标签机械念出。",
            "音乐库在 iOS 上分别选择音频文件和录屏视频，MP3/M4A 等常用音频不再因为混合过滤器而灰显。此变更只修复文件入口，不改变既有歌曲保存、播放或视频提取逻辑。",
            "网页核心 v954、私人 iOS 1.0.76 (76)、原生桥 18。Windows 负责自动回归与包结构检查；Mac 仍负责 Xcode 全 Target 编译、Apple 签名和真实设备验收。",
        ],
    ),
    "AI开发项目_新聊天启动说明": (
        "2026-08-16 当前版本补充｜网页 v954／私人 iOS 1.0.76 (76)",
        [
            "当前正式修改集中在 iOS 音乐文件可选择、小号首句兼容、切回主号必达报备、私人状态问句不误联网和报备自然化。不要重新尝试 iOS 系统黑条方案，也不要恢复微信自然模式开关。",
            "后续排查小号问题时先检查：事件是否保存 incoming/reply、切号时是否触发报告、可见消息是否带事件时间、游标是否只在真实交付后推进。禁止用泛化回复、旧话题续写或原始聊天转录冒充报备。",
            "继续使用 phone-work/main 单一直线，保护现有未提交现场。私人安装包每次只交最新一个完整 ZIP，不夹带旧安装说明、旧 ZIP、预览、缓存或临时脚本。",
        ],
    ),
}


for stem, (heading, paragraphs) in records.items():
    append_record(stem, heading, paragraphs)
