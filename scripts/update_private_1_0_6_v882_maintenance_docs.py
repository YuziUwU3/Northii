from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, name="等线", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def append_section(name, heading, paragraphs):
    path = DOCS / name
    document = Document(path)
    if any(p.text.strip() == heading for p in document.paragraphs):
        return
    document.add_page_break()
    title = document.add_heading(heading, level=1)
    for run in title.runs:
        set_font(run, size=16)
    for text in paragraphs:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.25
        set_font(paragraph.add_run(text))
    document.save(path)


SECTIONS = {
    "AI开发项目_Bug修改规范.docx": (
        "新增强制规范｜原生语音提交、键盘避让与安全区不得只验表象（v882／私人 1.0.6 起）",
        [
            "语音识别显示临时字幕不等于消息已经发送。原生桥必须提供可结束的 final 事件，网页必须验证 final 会进入用户消息、持久化记录和角色回复链；测试不得只断言字幕或权限存在。连续识别若系统长期不返回 final，必须以明确的静默窗口提交当前完整转写，并防止同一句重复提交。",
            "WKWebView 的键盘避让只允许一个权威来源。禁止同时叠加原生键盘 frame、visualViewport 差值、强制 window.scrollTo 和 CSS fixed 位移；出现输入框过高或消失时，先撤掉重复补偿并恢复 iOS/WKWebView 自带避让，再做逐页真机验证。",
            "Apple 主屏幕安全区测试必须验证计算结果有非零兜底，不能只检查选择器里出现 safe-area-inset。已安装网页可能返回 0；顶部返回键、独立 App 顶栏和底部输入区必须分别有页面级类、明确 fallback 和回归断言。",
            "全屏外壳只能维护一套高度模型。已经验证稳定的 100% shell 不得被可选适配开关改写为 100dvh，否则容易同时产生上方挤压和底部黑块。适配开关只能增加安全区内边距，不能替换根视口高度。",
            "每个安装候选必须同时记录共享网页版本和私人原生版本。Windows 测试通过只能证明静态链路与网页回归，Mac 编译、权限弹窗、键盘、语音入库和真机安全区仍须单独验收。",
        ],
    ),
    "AI开发项目_Bug记录模板.docx": (
        "v882／私人 1.0.6 Bug 记录｜字幕未发送、输入框二次上移与网页安全区零值（2026-08-11）",
        [
            "现象：私人 App 的 iOS 原生语音能在屏幕显示字幕，但后台通话记录没有用户说的话，角色也无法读取；通话和音乐输入框被键盘推到屏幕中部或直接看不到。苹果主屏幕网页版开启适配后，微信、购物、音乐、抖音、外卖等页面顶部返回键仍与系统状态栏重叠，底部还出现黑块。",
            "语音根因：原生 SFSpeechRecognizer 持续返回 partial 转写，现有网页通话逻辑只在 isFinal=true 时调用 hfHeard。partial 只更新字幕，不会生成带 _call 的用户消息，不会 save，也不会触发 callAI，所以视觉上识别成功但业务链完全没有收到。",
            "键盘根因与失败方案：私人 1.0.5 同时使用 iOS/WKWebView 自带键盘避让、原生 keyboard frame、visualViewport 差值、CSS fixed 偏移和强制滚动。多套坐标系统重复补偿，导致原生 App 输入框二次上移。该方案已经由真机证明失败，本次完整移除自定义键盘 frame 桥和对应 CSS，不得再次照搬。",
            "网页安全区根因与失败方案：v878 的适配开关把 v877 已验证的 100% 外壳改写为 100dvh，同时回归测试只确认选择器包含 env(safe-area-inset-*)，没有验证已安装网页中 env 可能为 0。结果是根外壳变短产生底部黑块，而顶部返回键仍没有得到真实偏移。",
            "修复：私人语音每次 partial 更新后重置 1.15 秒静默计时；停顿后发出同一会话的 isFinal=true，再走既有 hfHeard→用户通话消息→save→角色回复链并结束本轮识别。输入框恢复 WKWebView/iOS 单一键盘避让。v882 保留 100% shell，仅通过 CSS 变量为顶部安全区提供 max(env(...),47px)、底部提供 max(env(...),34px)，并继续逐页使用明确页面类。",
            "隔离与验证：公开审核中的 North App 工程未修改。共享网页升级为 v882，私人小手机升级为 1.0.6 (6)。全量 node --test tests\\*.test.mjs 为 447/447 通过，git diff --check 无错误（只有 Windows LF→CRLF 提示）；Windows 无法替代 Mac 编译和 iPhone 真机验证。",
        ],
    ),
    "AI开发项目_项目说明文档.docx": (
        "共享网页 v882／私人小手机 1.0.6｜语音真实提交与双入口安全区修复（2026-08-11）",
        [
            "本轮继续遵守一份共享网页业务核心、私人原生能力层和公开审核 App 隔离的架构。共享网页提升到 v882；私人小手机安装版本提升到 1.0.6 (6)；当前审核中的公开 North App 工程不改动。",
            "私人 App 的原生语音不再停留在字幕层。SFSpeechRecognizer 的 partial 结果用于实时字幕，连续 1.15 秒没有新内容时由原生桥提交 final，网页沿用既有 hfHeard 数据路径，将用户发言写入通话历史并让角色读取后回复。",
            "私人 App 输入框撤销 1.0.5 的第二套原生键盘位移，重新只依赖 WKWebView/iOS 自带避让。网页版 Apple 主屏幕模式继续是显式开关，但只增加顶部和底部安全区，不再更改根外壳高度；Safari、Android 和私人原生 App 不套用该网页开关。",
            "本轮没有声称手机号账号、Keychain 云恢复、大容量原生媒体存储、服务端控制器租约和长时间真机稳定性已经完成。这些仍是私人真实 App 主线，不得因网页装入 App 就误判完成。",
            "自动化证据：全量 447/447 通过，覆盖版本缓存、原生桥、语音 final 链、键盘重复补偿禁用、Apple 主屏幕安全区和所有既有业务功能。下一道硬门槛是 Mac 编译 1.0.6 并在真实 iPhone 分别验证网页与原生 App。",
        ],
    ),
    "AI开发项目_新聊天启动说明.docx": (
        "新聊天接手状态｜共享网页 v882／私人小手机 1.0.6（2026-08-11）",
        [
            "唯一仓库为 C:\\Users\\pc\\Documents\\小手机\\phone-work，分支 main。当前共享网页版本 v882，私人小手机版本 1.0.6 (6)。公开审核中的 North App 工程保持不动，不能让公开 North 与私人小手机同时控制同一台设备。",
            "本轮已修复三件事：原生语音停顿 1.15 秒后提交 final 并进入真实通话记录/角色回复；删除导致通话和音乐输入框二次上移的原生键盘 frame 补偿；网页版 Apple 主屏幕适配恢复稳定 100% shell，并为顶部/底部安全区提供非零 fallback。",
            "1.0.5 的键盘 frame+visualViewport+CSS fixed 组合已由真机证明失败，禁止恢复。v878 只断言 safe-area 选择器存在且用 100dvh 覆盖 100% shell 的方案也已证明不完整，禁止重复。",
            "接手后先做 Mac 编译和三项真机验证：一，通话说一句完整话，停顿后后台通话记录出现该句且角色回复；二，通话和音乐键盘弹出时输入框紧贴键盘而不在屏幕中部；三，网页版开启苹果主屏幕适配后微信、购物、音乐、抖音、外卖返回键可点且底部无黑块。",
            "后续主线仍包括手机号登录、Keychain、云端主数据与重装恢复、大容量原生媒体存储、单控制器租约、永久锁定/同步/跨日/断网/后台长稳测试。当前全量自动化 447/447 通过，但不得把 Windows 自动化写成 Mac 编译或真机验收完成。",
        ],
    ),
}


def main():
    for name, (heading, paragraphs) in SECTIONS.items():
        append_section(name, heading, paragraphs)
        document = Document(DOCS / name)
        assert any(p.text.strip() == heading for p in document.paragraphs)
        print(f"verified {name}: {len(document.paragraphs)} paragraphs")


if __name__ == "__main__":
    main()
