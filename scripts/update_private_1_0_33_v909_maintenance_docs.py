from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, name="等线", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def append(name, heading, paragraphs):
    path = DOCS / name
    document = Document(path)
    if not any(p.text.strip() == heading for p in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, size=16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(path)
    with ZipFile(path) as archive:
        assert archive.testzip() is None


append("AI开发项目_Bug记录模板.docx", "v909／私人 1.0.33 Bug 记录｜App 快照误判与通知先确认后落盘（2026-08-13）", [
    "现象：查看当前软件的后台通知偶尔说没有看清；后台能看到角色通知，但点进聊天没有对应消息。网页发布后还可能长期停留在旧构建。",
    "根因：Screen Time 的真实采集时间和 App 列表位于 snapshot.screenTime，服务器却只读取外层 capturedAt／generatedAt，外层为空时把新鲜子快照误判为无效。客户端收到服务端消息后先 ACK，再尝试写入 IndexedDB；页面中断、临时角色状态或持久化失败时，服务端行已经消费，客户端消息却未落盘。发布版本只改了部分资源参数，HTML 构建标记、脚本构建守卫、Service Worker 缓存仓、index 和 repair 入口仍有旧号，也会把新版挡回旧缓存。",
    "修复：App 查看统一读取 snapshot.screenTime 及其 generatedAt／capturedAt；设备接力读取所有子快照中的最新时间。服务端消息取消普通文本去重，按唯一 outbox ID 入聊；先 save 并等待 persistWechatMessagesNow 成功，随后才 ACK，失败则保留服务端行重试。v909 的 HTML、app.js、sw.js、index.html、repair.html 和私人 PhoneWeb.bundle 原子统一。",
    "补充：宠物睡眠锚点再次依据实际背景下移到粉色窝垫；私人 iOS 26 增加 AlarmKit 系统闹钟同步。Windows Node 全量回归 506/506 通过，phone-role-push 已部署。Mac 五 Target 编译、AlarmKit 授权和真机后台响铃仍需验收。",
])

append("AI开发项目_Bug修改规范.docx", "新增强制规范｜子快照时间、消息确认顺序与版本原子发布（v909 起）", [
    "设备快照如果包含 screenTime、wellness、location 等子快照，必须使用对应子快照自己的 generatedAt／capturedAt 判断新鲜度；外层时间只能作为兜底，禁止因外层为空否定已有真实数据。",
    "服务端消息队列必须遵守先本地持久化、后服务端确认。save 调用不等于 IndexedDB 已完成；只有持久化 Promise 成功后才能 ACK。失败、角色临时不可用或页面中断时必须让服务端行保持可重试。",
    "网页发布号必须原子更新 APP_VER、HTML shell build、脚本 build guard、Service Worker BUILD 与缓存仓、注册 URL、index／repair 入口及私人内置资源；全量 cache-version 测试不通过时禁止发布。",
    "系统后台能力必须使用平台正式 API。iOS 26 闹钟使用 AlarmKit 并声明 NSAlarmKitUsageDescription；网页和 Android 保留原前台实现。没有 Mac 编译与真机证据时必须明确标为待验收，不得把源码完成写成实机通过。",
])

append("AI开发项目_项目说明文档.docx", "v909／私人 1.0.33｜后台消息可靠入聊与 iOS 系统闹钟（2026-08-13）", [
    "网页版升级为 v909，私人 iOS 安装版升级为 1.0.33 (33)，原生桥契约升级为 12。App 查看测试改用 Screen Time 子快照和真实采集时间，phone-role-push v909 已在线部署。",
    "角色后台消息现在以服务端 outbox ID 作为唯一身份，先写入角色聊天并等待本地持久化完成，再向服务器确认；如果中途失败，消息仍会在下一次拉取时重试。这样修复了只看到通知、聊天没有正文的问题。",
    "私人 iOS 26 版把小手机闹钟同步到 AlarmKit，支持一次、每日、关闭和删除，App 退到后台后由系统继续提醒。网页和 Android 的原前台闹钟保持不变。宠物睡眠坐标再次下移到左侧粉色窝垫。Windows 自动回归 506/506 通过，Mac 编译和真实 iPhone 验收仍是证据边界。",
])

append("AI开发项目_新聊天启动说明.docx", "新聊天接手状态｜v909／私人 1.0.33 已部署候选（2026-08-13）", [
    "当前基线：网页版 v909、私人 iOS 1.0.33 (33)、原生桥契约 12。phone-role-push v909 已部署；本轮没有新增数据库迁移。",
    "不要退回只读 snapshot.capturedAt 的 App 判断；真实软件列表与采集时间优先位于 snapshot.screenTime。不要把角色服务端消息 ACK 移到本地持久化之前，也不要恢复 24 小时纯文本去重覆盖唯一 outbox 消息。",
    "私人版新增 NativeAlarmService.swift，使用 AlarmKit 保存自身闹钟记录并与 AlarmManager.shared.alarms 比较一次闹钟是否已经响过。Info.plist 已加入 NSAlarmKitUsageDescription；网页／Android 不走该桥。",
    "下一步真机优先验收：立即查看当前软件；后台通知点进聊天确认正文存在；创建两分钟后的一次闹钟并回桌面等待系统响铃。随后再检查宠物夜间睡眠是否全部位于粉色窝垫。Windows 506/506 已通过，Mac 五 Target 编译和上述真机链路尚未完成。",
])
