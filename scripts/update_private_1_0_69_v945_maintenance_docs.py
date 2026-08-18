from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, size=10.5):
    run.font.name = "等线"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(size)


def append_docx(stem, heading, paragraphs):
    path = DOCS / f"{stem}.docx"
    document = Document(path)
    if not any(p.text.strip() == heading for p in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, 16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(path)
    with ZipFile(path) as archive:
        assert archive.testzip() is None


def append_txt(stem, heading, paragraphs):
    path = DOCS / f"{stem}.txt"
    content = path.read_text(encoding="utf-8")
    if heading not in content:
        path.write_text(
            content + "\n\n" + heading + "\n" + "\n".join(paragraphs) + "\n",
            encoding="utf-8",
        )


records = {
    "AI开发项目_Bug记录模板": (
        "v945／私人版 1.0.69 Bug 记录｜主屏、消息删除与大小号隔离（2026-08-15）",
        [
            "现象一：少数 401–649px 宽屏设备的顶部仪表组件仍受旧 348px 上限约束，视觉上缩向左侧；下方组件和右下四个 App 因沿用窄屏纵坐标而过度靠近。另有少数 WebView 指针流中断后没有继续派发 pointermove，造成主页横滑偶发卡住。修复只在 401–649px 宽屏取消仪表与 Dock 的旧宽度上限，并在高度不少于 820px 时拉开各区纵向间距；360／390／393px 规则保持不变。横滑补入非被动 touchmove 回退、touchend、失焦与后台清理，指针流正常时不重复接管。",
            "现象二：微信自己发出的文字和图片没有单条删除入口。修复在本人文字／图片消息菜单中加入删除，保存后立即从当前账号消息数组移除并持久化；图片若仍在识图等待中会终止迟到回复，无其他引用时由媒体回收清理。角色消息、撤回和已有菜单语义不变。",
            "现象三：小号原本应当是陌生人，但少数大号关系模块仍会进入小号系统上下文；后来加强隔离时又断开了旧版切回大号后的报备入口。修复后小号只读取小号资料、聊天、小号摘要与长期记忆，主号恋爱、极端关系、共同生活、好友来源和记忆清除事件全部隔离；切回大号时只对最新一段新增小号聊天报备一次，同一活动不重复。",
            "私人 App 顶部黑色系统区只用同色安全区背景融合，WKWebView 不进入状态栏或 Dynamic Island。iPhone 主屏幕 Web App 的底部系统区域未修改。AI 账户同时固定声明：内置 AI 只用于语音生成和影院字幕识别，不用于普通聊天、聊天识图或聊天生图。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "新增强制规范｜宽屏隔离、消息删除与大小号边界（v945 起）",
        [
            "主页稀有设备修复必须使用窄范围媒体条件，并提供 360／390／393px 坐标未变化证据。401–649px 宽屏调整时要同时检查顶部仪表、唱片、双人组件、八个 App、页点、Dock 和第二页；不得只修用户截图里的一个组件。",
            "主页手势必须保留 pointer 主路径，同时允许未被 pointer 接管的非被动 touchmove 作为回退；touchend、pointercancel、blur 和 visibilitychange 必须清理 pending pan、拖拽幽灵和占用状态。不得在 pointerdown 全局 preventDefault，也不得改变长按拖拽的稳定槽位与唯一性校验。",
            "本人消息删除必须账号内持久化，只允许本人文字与图片；图片删除还要取消未完成识图的迟到回复并执行无引用媒体回收。不得把角色消息、撤回、引用或图片查看入口误改成删除。",
            "小号必须是独立联系人：只读当前小号资料、聊天、摘要和小号记忆，不得读取大号称呼、关系、共同生活、好友来源、关系事件或主号记忆。允许的跨账号行为只有切回大号后的单向报备；同一段新增活动只报一次，小号有新聊天后才可再报，且提示必须禁止猜测两个账号是同一人。",
            "私人 App 状态栏融合只能绘制系统安全区背景，不能把网页伸入状态栏。内置 AI 功能白名单固定为语音生成和影院字幕识别；普通聊天、聊天识图和聊天生图继续使用用户自己的外置配置。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "v945／私人版 1.0.69｜主屏稳定与微信账号修复（2026-08-15）",
        [
            "当前正式版本：网页版 v945；私人 iOS 1.0.69 (69)；原生桥契约 18。",
            "401–649px 少数宽屏主页恢复完整仪表与 Dock 宽度，并在高屏设备拉开组件和右下四个 App；360／390／393px 与第二页既有结构保持原样。主页横滑加入缺失指针流时的触摸回退和生命周期清理，既有长按拖拽与布局持久化不变。",
            "微信支持删除本人文字和图片。大小号继续使用独立消息键、摘要、长期记忆和异步回复账号；小号按独立陌生联系人生成，切回大号后只对每段新增小号交流报备一次。",
            "私人 App 顶部系统安全区使用深色背景融合，网页仍在安全区下方。AI 账户明确内置 AI 只提供语音生成和影院字幕识别，普通聊天、聊天识图和聊天生图不走内置服务。",
        ],
    ),
    "AI开发项目_新聊天启动说明": (
        "新聊天接手状态｜v945／私人版 1.0.69（2026-08-15）",
        [
            "当前基线：网页版 v945、私人 iOS 1.0.69 (69)、原生桥契约 18。固定 main 单线工作，iPhone 主屏幕 Web App 的底部系统区域禁止继续修改。",
            "本版宽屏修复只允许命中 401–649px；360／390／393px 必须保持原坐标。主页手势同时保留 pointer 主路径与未接管 touch 回退，任何修改都要复测横滑、反向滑、长按拖拽、两页应用唯一性和第二页组件。",
            "小号必须继续作为独立陌生联系人，只读小号资料、聊天和小号记忆。切回大号后可对新增小号聊天报备一次；没有新活动不得重复，报备不能反向泄露大号身份给小号。",
            "本人文字／图片删除要持久化并阻止图片迟到识图回复。私人 App 顶部只铺安全区背景。内置 AI 白名单只有语音生成和影院字幕识别，普通聊天、聊天识图和聊天生图仍走外置配置。",
        ],
    ),
}


for stem, (heading, paragraphs) in records.items():
    append_docx(stem, heading, paragraphs)
    append_txt(stem, heading, paragraphs)
