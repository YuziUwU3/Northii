from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"

def set_font(run, name="等线", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)

def append(name, heading, paragraphs):
    path = DOCS / name
    document = Document(path)
    if not any(p.text.strip() == heading for p in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, size=16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(path)
    with ZipFile(path) as archive:
        assert archive.testzip() is None

append("AI开发项目_Bug记录模板.docx", "v905／私人 1.0.29 Bug 记录｜退后台中断与一分钟测试失效（2026-08-12）", [
    "现象：普通回复或真实设备读取生成期间退出小手机，WKWebView 会被冻结，原任务随页面中断；设置1分钟也仍受正式30至60分钟安静期限制，无法用于后台链路验收。",
    "根因：模型生成和设备读取只存在于页面内存；服务器只有定时主动消息，没有当前回合的持久任务。1分钟配置代表主动消息额外间隔，不是测试旁路。",
    "修复：用户消息与设备读取开始时预登记可取消的后台接管任务；前台完成立即取消，页面隐藏则用keepalive把任务提前到5秒后。服务器领取、重试、生成、写outbox并通过APNs发送。新增独立一分钟真实测试，不计日次数，不查看或锁定软件。用户新回复会由数据库触发器取消旧接管和App后续任务。",
    "同时完成：普通主动消息和App查看共用单一到期队列，50%分流；App查看只使用明确授权的稳定ID，每天0至5次，五分钟后只执行询问或锁定之一，完成后90分钟冷却。后台六项设备自动规则通过分钟任务主动唤醒并只使用新鲜快照。",
])
append("AI开发项目_Bug修改规范.docx", "新增强制规范｜iOS后台长任务必须服务端持久接管（v905起）", [
    "WKWebView、JavaScript定时器和前台fetch都不能作为iOS后台可靠执行器。任何要求退出页面后继续完成的生成、读取、延迟或通知，必须在退出前写入服务端持久任务；任务要有基线用户消息、幂等去重、超时重领、用户新回复取消和APNs最终送达。",
    "主动消息、App查看和后续锁定不得各自建独立随机定时器。它们必须共享同一到期认领和占用状态；查看和控制只准使用用户明确授权的稳定App ID，系统关键App与未授权App一律不进入候选。旧快照不得冒充当前数据。",
    "全局时间感知关闭必须在所有提示入口和服务端入口同时处理：不注入当前日期时间、星期、时段、持续时长或消息间隔，并清除事件上下文中的机器时间戳；开关变化要同步所有服务器角色资料。",
])
append("AI开发项目_项目说明文档.docx", "v905／私人 1.0.29｜后台接管、App感知与全局时间开关（2026-08-12）", [
    "普通微信回复和真实iPhone查看在页面退到后台后由服务器接管并通过角色逐句通知送达；前台正常完成会取消服务端任务，避免双回复。一分钟后台测试是独立旁路，不受正式安静期、日次数和App查看影响。",
    "新增可关闭的角色App感知：每次统一主动机会在普通联系与查看App之间各50%，每日查看上限0至5。服务器先请求新鲜授权App用量，角色自然联系；用户回复立即取消后续，否则五分钟内由角色选择再问一次或锁定该App，随后90分钟冷却。",
    "共同生活状态可分别手动编辑生活阶段、当前活动和地点。时间感知关闭后，微信、视频/语音、线下约会、共同生活及服务器消息均不知道日期、时间、星期、时段和间隔。睡眠步数、每日全部App用量、失联电量位置、5%电量、难过心率及手动解锁事件进入服务器分钟任务，普通后台和锁屏可继续；iOS上划强退仍不保证系统唤醒。",
])
append("AI开发项目_新聊天启动说明.docx", "新聊天接手状态｜v905／私人1.0.29后台接管候选（2026-08-12）", [
    "当前候选基线为网页v905、私人iOS 1.0.29 (29)。本版新增服务器持久后台任务、真正一分钟测试、统一主动/App查看队列、共同生活手动状态和全局时间感知关闭。",
    "发布前必须应用迁移202608120002_background_role_tasks.sql并部署phone-role-push。需真机验证：回复生成中退后台、设备读取中退后台、一分钟通知、普通主动与App查看互斥、五分钟取消/锁定、90分钟冷却，以及六项后台自动规则。",
    "iOS普通后台和锁屏可由静默推送唤醒；用户上划强退后系统可能拒绝后台唤醒，这一限制不得对用户承诺为百分之百可用，也不得用旧快照伪装成功。",
])
