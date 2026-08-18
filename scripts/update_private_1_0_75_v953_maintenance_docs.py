from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, size=10.5):
    run.font.name = "等线"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(size)


def append_record(stem, heading, paragraphs):
    docx_path = DOCS / f"{stem}.docx"
    document = Document(docx_path)
    if not any(paragraph.text.strip() == heading for paragraph in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, 16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(docx_path)
    with ZipFile(docx_path) as archive:
        assert archive.testzip() is None

    txt_path = DOCS / f"{stem}.txt"
    text = txt_path.read_text(encoding="utf-8")
    if heading not in text:
        txt_path.write_text(
            text + "\n\n" + heading + "\n" + "\n".join(paragraphs) + "\n",
            encoding="utf-8",
        )


records = {
    "AI开发项目_Bug记录模板": (
        "v953 私人版交付记录｜1.0.75 同步小号报备修复（2026-08-16）",
        [
            "现象与决定：网页 v953 的小号报备语义校验已经通过完整回归，但第七十四次私人安装包仍内置旧 PhoneWeb.bundle。用户明确要求把同一套 AI 修复打入私人 App，并继续遵守每次只交一个全新 ZIP 的死规则。",
            "处理：私人 App 升级到 1.0.75 (75)，原生桥契约仍为 18；通过 private-phone-web.manifest.json 和 stage-private-phone-web.mjs 从当前共享源码重建 PhoneWeb.bundle，不从旧 ZIP 解压覆盖。网页正式版本仍为 v953，没有另起网页分支。",
            "范围：本次只同步已经验证的小号报备交付校验、旧错误游标修复和首次直加好友 HTTP 400 本地开场兜底；不改主屏布局、iOS 系统栏、网页版本、账号隔离边界或其他原生能力。",
            "验证：Windows 完整回归 200 个测试文件、657 项通过、0 失败。交付 ZIP 必须只有完整 Xcode 工程、请在 Mac 编译前先读和本次第七十五次安装说明；必须检查只有一个 PhoneWeb.bundle/index.html、无嵌套 ZIP、无历史安装说明和临时文件。Mac 编译、Apple 签名及真机报备仍需在 Mac/iPhone 完成。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "发布补充规范｜共享网页修复进入私人 App 必须重建内置 Bundle（1.0.75 起）",
        [
            "网页核心修复只有经过清单重新生成 PhoneWeb.bundle 并升级私人构建号，才算进入私人 App。不得把网页已提交等同于私人安装包已更新，也不得从历史 ZIP 解压后只覆盖 app.js。",
            "私人包发布前必须同时核对 MARKETING_VERSION、CURRENT_PROJECT_VERSION、__SMALL_PHONE_PRIVATE_BUILD__、设置页可见版本和安装说明；完整回归通过后再从全新暂存目录压缩。",
            "交付目录只保留最新一个 ZIP；ZIP 内只允许完整 Xcode 工程、当前 Mac 编译说明和本次安装说明。历史说明、嵌套 ZIP、预览、缓存和受保护未提交现场一律不得进入。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "当前交付基线｜网页 v953／私人 iOS 1.0.75 (75)（2026-08-16）",
        [
            "私人 iOS 1.0.75 (75) 继续内置网页核心 v953，原生桥契约 18。统一自然系统、小号身份隔离与主号一次性报备保持同一路线；报备只有在可见内容明确包含指定小号名称及联系动作后才推进游标。",
            "本版把网页已验证的小号报备交付校验、旧误消费事件修复和首次直加好友 HTTP 400 本地开场兜底同步进私人 PhoneWeb.bundle。普通旧话题续写、泛化、否认、空回复或跳出角色内容不会误记为报备成功。",
            "私人交付仍是可独立编译的完整 Xcode 工程。Windows 已通过 657 项自动测试；Mac 编译、Apple 签名和真实 iPhone 的小号聊天后切回主号报备验收仍需另行完成。",
        ],
    ),
    "AI开发项目_新聊天启动说明": (
        "2026-08-16 当前版本补充｜私人 iOS 1.0.75 (75)",
        [
            "当前网页正式版本为 v953，私人 iOS 正式安装版本为 1.0.75 (75)，原生桥契约为 18。私人版已通过清单重建 PhoneWeb.bundle，同步 v953 小号报备交付校验，不是旧 ZIP 覆盖包。",
            "继续使用 phone-work/main 单一直线。先只读检查 Git 状态并保护既有未提交现场；私人包每次只交最新一个完整 ZIP，不得夹带旧安装说明、旧 ZIP、预览、缓存或临时脚本。",
            "小号报备回归重点：小号保持陌生身份；切回主号后报备必须具名且包含联系动作；旧话题、泛化、否认和空回复不得推进游标；同一事件只报一次，新活动可再报一次。",
        ],
    ),
}


for stem, (heading, paragraphs) in records.items():
    append_record(stem, heading, paragraphs)
