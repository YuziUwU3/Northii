from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, name="等线", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def append(name, heading, paragraphs):
    path = DOCS / name
    document = Document(path)
    if not any(p.text.strip() == heading for p in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, size=16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(path)
    with ZipFile(path) as archive:
        assert archive.testzip() is None


append("AI开发项目_Bug记录模板.docx", "v921／私人版 1.0.45 Bug 记录｜主屏幕通话字幕恢复（2026-08-14）", [
    "现象：普通电话和普通视频主屏幕底部的大字幕被缩小，字幕整体下沉并贴近输入框。用户明确要求只恢复主屏幕大字幕；共享屏幕小框和后台画中画字幕不改。",
    "根因证据：对照 v850（caac487），主屏幕 .csline 固定为 18px，.callsub 使用 align-items:center。v918 将主屏幕容器改为 align-items:flex-end，并按文本长度给主屏幕附加 compact／dense 类，把字号降为 15px／12px，因此位置和字号都发生回归。",
    "修复：主屏幕 .callsub 恢复 v850 的垂直居中，主屏幕 .csline 恢复固定 18px，并删除网页主屏幕的长文本缩小分类。保留 v918 以后已经确认的整句淡入动画、自动换行和长文本展示范围。原生 PiP／共享小框仍保留独立的 14／11／9.5 字号策略。",
    "风险边界：不修改角色声音、免提识别、TTS、屏幕共享、字幕内容、字幕淡入关键帧、输入框和控制按钮。新增回归测试同时锁定主屏幕 v850 字号／位置与 PiP 独立缩放，防止以后再次混用。",
])

append("AI开发项目_Bug修改规范.docx", "新增强制规范｜主屏幕字幕与共享小框独立维护（v921 起）", [
    "主屏幕通话字幕和共享／PiP 小框字幕是两个不同展示面。修改字号、位置、换行或长文本适配前必须明确目标展示面；不得把小框为容纳长文本采用的缩小规则套到主屏幕。",
    "恢复历史样式时必须逐项对照字号、行高、容器 bottom、垂直对齐和动画，不得只看其中一个属性。若用户要求恢复字号和位置但保留当前动画，应只改布局属性，不能顺带回退动画关键帧。",
    "回归测试必须同时证明：主屏幕保持固定 18px 和垂直居中；共享／PiP 小框仍能按长文字独立缩放。任何一侧变化都必须由用户明确授权。",
])

append("AI开发项目_项目说明文档.docx", "v921／私人版 1.0.45｜主屏幕通话字幕恢复（2026-08-14）", [
    "当前版本：网页 v921；私人 iOS 1.0.45 (45)；原生桥契约 18。PhoneWeb.bundle 由同一份 v921 网页核心重新生成。",
    "普通电话和普通视频主屏幕字幕恢复 v850 的固定 18px 与容器垂直居中，不再因总文字长度自动缩成 15px 或 12px。字幕仍使用当前整句从透明到实心的淡入动画，并保留自动换行。",
    "共享屏幕小框和后台 PiP 没有跟随主屏幕改动，仍使用其独立的长文本字号策略。本轮没有修改 v920 的角色声音、识别隔离、前后台播放器分流或屏幕共享逻辑。",
])

append("AI开发项目_新聊天启动说明.docx", "新聊天接手状态｜v921／私人版 1.0.45（2026-08-14）", [
    "当前代码基线：网页 v921、私人 iOS 1.0.45 (45)、原生桥契约 18、PhoneWeb.bundle v921。",
    "v921 只修普通电话／普通视频主屏幕大字幕：固定 18px、垂直居中，恢复 v850；当前整句淡入动画继续保留。共享屏幕小框和原生 PiP 的长文本缩放没有修改。",
    "真机验收时用一条短双语字幕和一条较长双语字幕检查主屏幕：字号应保持一致，整体不贴输入框；再打开共享／PiP 小框确认长文字仍可缩小适配。角色声音仍按 v920 链路单独验收，不要因字幕布局再次改声音代码。",
])
