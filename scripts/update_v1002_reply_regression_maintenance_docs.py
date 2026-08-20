from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, size=10.5):
    run.font.name = "等线"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(size)


def append_record(stem, heading, rows):
    docx_path = DOCS / f"{stem}.docx"
    document = Document(docx_path)
    if not any(paragraph.text.strip() == heading for paragraph in document.paragraphs):
        title = document.add_heading(heading, level=1)
        title.paragraph_format.page_break_before = True
        for run in title.runs:
            set_font(run, 16)
        for row in rows:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(row))
        document.save(docx_path)
    with ZipFile(docx_path) as archive:
        assert archive.testzip() is None

    txt_path = DOCS / f"{stem}.txt"
    text = txt_path.read_text(encoding="utf-8")
    if heading not in text:
        txt_path.write_text(
            text + "\n\n" + heading + "\n" + "\n".join(rows) + "\n",
            encoding="utf-8",
        )


RECORDS = {
    "AI开发项目_Bug记录模板": (
        "v1002 回复回归热修｜自动重试占用普通聊天与朋友圈（2026-08-20）",
        [
            "用户证据与最后正常基线：用户确认未改前的 v1000 普通角色回复和朋友圈回复完整正常。安装后在共同生活读取屏幕时间之后，出现普通角色回复需数分钟、朋友圈长期显示回复中、模型测试偶发连接超时、重新打开 App 后仍在回复中的连续现象。Git main、HEAD 与 origin/main 均为 v1000 提交 42a7b1587ce53070d82d38320c79d1941c603a80；v1001／v1002 只存在于未提交工作区。",
            "代码根因：v1002 新增的 phoneFactRetry 把失败结果连同 retryAt 持久化。cohabPhoneAutonomyTick 每 15 秒检查，时间到后重新取得全局真实读取通道并调用 cohabPhoneDeliverFact；该函数一次最多进行主请求加两次安全重说，而 chatAPI 单次超时上限为 190 秒，因此一轮失败最多可占用约 570 秒，并可跨 App 重启继续。普通微信 aiReply 在该通道期间主动静默，朋友圈评论又会同时争用同一模型线路，形成用户看到的分钟级延迟和近似无回复。",
            "第二个回归：replyGenerationRun 只把“运行期间读取代次发生变化”识别为正常接管。若用户点击回复前读取通道已经占用，aiReply 会立即静默返回，但代次没有变化，包装器会再请求一次并错误显示“模型未回复”，把正常读取接管误报成 API 故障。",
            "测试遗漏：原 v1002 测试只静态断言失败结果存在自动重试，没有限制重试时长、跨重启次数、全局通道占用，也没有验证 v1000 普通聊天和朋友圈在读取失败后的可用性。850／850 因此不能证明这条真机并发链安全。",
            "修复：真实读取结果仍保存在 phoneFactRetry，但状态改为 manual，不再由 15 秒循环自动重放模型请求；旧存档中的 retryAt 启动后原位迁移为 manual，保留真实数据且停止继续占用线路。replyGenerationRun 在第一次调用前、第一次返回后、重试前和显示错误前都检查真实读取通道；正常接管保持静默，不再二次请求或误报模型错误。朋友圈中断状态在冷启动时改为真实失败与手动重试，不写入假回复。",
            "风险隔离：没有修改用户 API 地址、Key、模型名、chatAPI 190 秒上限、统一角色 Prompt、普通微信正文生成、角色人设、世界书或朋友圈真实回复 Prompt。首次补丁因同一文件在补丁中重复声明而被编辑器整体拒绝，未落盘；拆为单一文件块后成功。没有采用裁剪统一 Prompt 的高风险方案。",
            "验证：app.js 与私人 PhoneWeb.bundle/app.js 语法检查通过；伴生、普通手动回复、朋友圈和共同生活专项 79／79；Windows 完整 Node 回归 850／850。未提升版本、未打包、未提交、未推送。Mac 编译、签名和真实 iPhone 上的普通回复时延、朋友圈评论、屏幕时间完成态及旧 retryAt 迁移仍未验证。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "持久模型重试不得占用旧稳定回复链（v1002 回复回归后）",
        [
            "用户指出明确历史正常版本时，必须把该版本作为最后正常基线，逐行比较新增状态、定时器、全局锁和模型调用次数。不能因为新功能测试通过，就默认普通聊天、朋友圈和模型测试没有被并发链影响。",
            "真实数据可以持久保存，但模型生成失败不得由短周期定时器跨重启无限自动重放，更不得在重放期间长期持有会阻止普通回复的全局读取通道。任何自动重试都必须有总次数、总时长、并发所有权和旧功能可用性测试；没有这些保护时只能保存为人工重试状态。",
            "回复包装器必须同时识别两种正常读取接管：运行前通道已经占用，以及运行期间读取代次发生变化。两种情况都不得二次请求、不得显示模型空回复或 API 故障提示。",
            "朋友圈回复失败或 App 挂起后只显示真实失败和重试入口，禁止补写角色假回复。回归至少覆盖：旧 retryAt 存档迁移、读取通道预先占用、读取中途接管、朋友圈冷启动 pending 修复、普通回复与朋友圈不被后台重放拖慢。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "v1002 未提交回复回归热修（2026-08-20）",
        [
            "v1000 提交 42a7b1587ce53070d82d38320c79d1941c603a80 是用户确认的最后正常回复基线，也是当前 main／HEAD／origin/main。v1001／v1002 与本热修均在未提交工作区，不得写成正式发布完成。",
            "本热修只停止共同生活真实读取失败后的跨重启自动模型重放，并修复读取已占用时手动回复误报模型失败；真实结果继续保留，朋友圈仍只接受真实模型回复。共享 app.js 与私人 PhoneWeb.bundle 已同步本热修相关代码，Windows 完整回归 850／850。Mac 编译和真实 iPhone 验证仍待完成。",
        ],
    ),
    "AI开发项目_新聊天启动说明": (
        "v1002 回复回归热修接手说明（2026-08-20）",
        [
            "真实 Git 基线仍是 v1000：main、HEAD、origin/main 均为 42a7b1587ce53070d82d38320c79d1941c603a80。工作区包含未提交的 v1001／v1002 与回复回归热修；不得整体回退、覆盖用户未跟踪资料、提交、推送或打包，除非用户明确授权。",
            "禁止恢复 phoneFactRetry 的 retryAt 自动模型重放。旧 retryAt 只迁移为 manual 并保留真实数据；真实读取预先占用或中途接管时，手动回复必须静默结束，不能二次请求或误报 API。朋友圈中断只变为真实失败与手动重试，不能生成假回复。",
            "Windows 当前结果为专项 79／79、完整 850／850。下一步必须在 Mac 从全新工程编译，并在真实 iPhone 先验证普通微信回复时延、朋友圈评论、屏幕时间完成态和重开后的旧 retryAt 迁移，再继续处理其他功能。",
        ],
    ),
}


for stem, (heading, rows) in RECORDS.items():
    append_record(stem, heading, rows)
