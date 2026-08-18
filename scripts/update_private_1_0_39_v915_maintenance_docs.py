from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, name="等线", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def append(name, heading, paragraphs):
    path = DOCS / name
    document = Document(path)
    if not any(p.text.strip() == heading for p in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, size=16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(path)
    with ZipFile(path) as archive:
        assert archive.testzip() is None


append("AI开发项目_Bug记录模板.docx", "v915／私人 1.0.39 Bug 记录｜外部 App 共享帧被前台帧覆盖与实时观察连续抢话（2026-08-13）", [
    "现象：用户切到其他 App 后，角色偶尔仍描述小手机自己的通话界面；实时共享理解只要画面变化就连续说话，没有给用户解释和回答的空档。原生 PiP 角色名居中且容器偏高，字幕空间不够。公开在线音乐中部分受限曲目无法播放。",
    "根因：Broadcast Upload 扩展与小手机前台 WebView 共用同一个 latest.jpg，用户切回小手机时新前台帧可能覆盖刚才的外部 App 帧。实时理解只有固定定时器和画面差异判断，没有角色观察状态机。PiP 内容视图高度与上下约束偏宽。公开目录返回结果里含受限或当前不可流式播放曲目。",
    "修复：ReplayKit 在宿主进入后台时另写一次性 background-latest 帧，网页请求实时帧时优先冻结该外部 App 帧，确认消费后删除，普通 latest 链路不变。实时共享理解新增 observing、waiting、answering、ended 状态；角色可静默继续、只问一个问题并暂停、听完用户回答后决定继续或结束。口头让角色看画面继续走原不限次数链路。PiP 内容高改为 144，名字上移，字幕四行；公开音乐搜索过滤不可用、受流限制曲目并增加一次刷新重试。",
    "边界：角色自主观察只在用户已经通过系统面板授权并开启屏幕共享、且偏好设置实时共享理解开关开启时运行；不能自行开启系统录屏或绕过权限。原生 PiP 最终黑色系统背板由 iOS 合成控制，应用只能把自己的内容根视图设为透明，不能保证系统层完全透出。",
])

append("AI开发项目_Bug修改规范.docx", "新增强制规范｜实时共享自主观察与后台帧归属（v915 起）", [
    "ReplayKit 共享外部 App 画面时，宿主前台帧与后台外部 App 帧必须分槽保存；不得用单个 latest 文件判断画面归属。后台帧只能一次性冻结、消费并回执删除，避免下一轮重复描述旧画面，同时不得破坏普通共享帧和手动识图链路。",
    "角色实时观察不能把每一帧都变成台词。必须显式维护继续静默、暂停提问、等待用户回答、结束本轮四类状态；等待期间不得继续取帧，收到用户回答后只能继续或结束，不能立刻连问第二个问题。口头‘看一下’属于用户主动请求，仍不限次数且优先于自主等待拦截。",
    "任何自主查看都必须同时满足：用户已在 iOS 系统面板授权共享、当前共享真实 active、实时共享理解开关开启、本次视频仍在自动识别上限内。角色不得自行启动系统录屏，也不得用提示词伪装已经看到未提供的画面。",
])

append("AI开发项目_项目说明文档.docx", "v915／私人 1.0.39｜后台共享帧交接与角色自主观察（2026-08-13）", [
    "当前版本：网页 v915；私人 iOS 1.0.39 (39)；原生桥契约 16。ReplayKit 后台扩展会把其他 App 的最新帧与小手机前台帧分开交接，实时识别优先消费后台外部 App 帧，减少切回小手机后只看到自身界面的情况。",
    "实时共享理解改为角色自主观察：角色看到画面后可以安静继续，不需要每帧说话；遇到在意内容时只问一个问题并暂停，用户用语音或打字回答后，角色按性格决定继续观察或结束本轮。设置里的每次视频识别次数上限仍有效；用户口头要求看画面仍不限次数。关闭开关后恢复原流程。",
    "系统 PiP 内容高由 176 压缩为 144，名字与状态靠上，字幕最多四行；应用内容根视图设为透明，但系统悬浮窗最终黑色背板由 iOS 决定。公开音乐搜索过滤 Audius 不可用和受流限制项目，播放错误会进行一次刷新重试；本地音乐、导入歌单和一起听未改。",
    "网易云、QQ 音乐、腾讯视频、优酷、爱奇艺的个人会员登录与受保护媒体播放没有作为通用功能接入：缺少面向此类第三方 App 的公开消费者会员播放授权，且受 FairPlay 等 DRM 保护的视频会在 ReplayKit 录屏或镜像时变黑。只有获得平台正式商务授权和允许的 SDK 后才能另行开发。",
])

append("AI开发项目_新聊天启动说明.docx", "新聊天接手状态｜v915／私人 1.0.39（2026-08-13）", [
    "当前基线：网页 v915、私人 iOS 1.0.39 (39)、原生桥契约 16。不要回退 ReplayKit 外部 App 后台帧独立槽、一次性消费和宿主前后台标记；这是解决角色总看到小手机自身界面的关键。",
    "实时共享理解使用 observing／waiting／answering／ended 状态。角色提问后必须停取帧；用户回答后才可继续或结束。不要把静默继续强行修正成必须说台词，也不要把普通口头‘你看一下’纳入等待拦截。关闭实时共享理解开关时必须完全停用自主状态机。",
    "PiP 名字上移、内容高度 144、字幕四行；透明只指应用自己的 PiP 内容根视图，不能承诺 iOS 系统黑色背板一定透明。在线音乐只保留合法公开 Audius 目录和本地来源；不得抓取商业会员平台、嵌入个人账号密码或绕过 DRM。",
    "下一步真机只做三件事：一，安装第三十九次包，开视频和实时共享理解后开始系统共享；二，切到另一个普通 App，确认角色描述的是外部 App，若提问则回答并观察它是否停下后再继续；三，缩小系统通话窗，确认名字上移、框更扁、字幕空间增加，并测试一首公开在线歌曲。",
])
