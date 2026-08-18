from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, name="等线", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def append(name, heading, paragraphs):
    path = DOCS / name
    document = Document(path)
    if not any(p.text.strip() == heading for p in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, size=16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(path)
    with ZipFile(path) as archive:
        assert archive.testzip() is None


append("AI开发项目_Bug记录模板.docx", "v917／私人 1.0.41 Bug 记录｜通话同声、字幕一致与线上群聊（2026-08-14）", [
    "现象：网页大通话字幕与原生共享小框的动画和追加方式不完全一致；角色播音时用户话筒会暂停，外部视频声、角色声音和用户说话无法稳定并行，回声也会增加识别错字。苹果兼容开启后宠物返回键和角色微信登录顶部键仍偏高。共同生活里角色口头答应发微信或群聊，却只在共同生活现场说出来。真人群没有真正的群主解散能力。",
    "根因：网页与原生各维护一份字幕时序，PiP 的每次字幕更新还会误走启动路径并重新激活音频会话；角色播音前后通过停止和重建识音规避回声，导致插话丢失。共同生活回复标签只有线下动作和手机查看，没有线上消息投递协议。真人群只有本地清空记录，缺少验证群主身份并级联删除云端数据的 RPC。",
    "修复：网页和原生共享字幕正文、说话人、逐字时长、延迟、位移、缩放及两套贝塞尔曲线；PiP 更新只更新画面。原生输入启用 voice processing，角色播音期间持续识别，最终用户插话进入最多三条队列。共同生活新增真实线上投递标签，仅允许角色微信和角色参加的 AI 群，群名歧义必须追问。新增 phone_friend_group_disband，以账号密钥和 owner_id 双重校验后原子删除群、成员、邀请、消息、回执及撤回标记。",
    "验证边界：Windows 的 555 项自动测试覆盖网页、SQL 契约和原生静态集成，但不能替代 Mac Xcode 六 Target 编译、签名及真实 iPhone 的 AVAudioSession、Speech、PiP、ReplayKit 和第三方媒体策略。解散群聊迁移 202608140001_phone_friend_group_disband.sql 已在现用项目部署并登记。",
])

append("AI开发项目_Bug修改规范.docx", "新增强制规范｜通话同声、线上代发与群聊权限（v917 起）", [
    "网页大通话和原生 PiP 字幕必须读取同一份协议参数。正文、说话人、逐字时长、逐字间隔、最大延迟、整行动画时长、位移、缩放和动画曲线任一改变时，必须同时更新协议与双端回归测试；字幕更新不得重新启动 PiP 或重置全局音频会话。",
    "连续通话不得为了播放角色声音而停止用户话筒。用户在角色播音期间产生的最终识别结果必须保留并按序处理；原生应使用 voice processing 降低外放回声，但不得宣称能强制覆盖第三方 App 的暂停、DRM 或系统策略。",
    "共同生活角色只能以角色身份写入该角色的微信或其实际参加的 AI 群。目标群不唯一时必须先追问，绝不能随机投递；不得冒充用户向真人小手机好友或真人群发言。AI 群的 AI 发言不得出现编辑入口。真人群解散必须客户端仅群主展示、客户端再次校验，并由服务端验证密钥与 owner_id 后原子级联删除。",
])

append("AI开发项目_项目说明文档.docx", "v917／私人 1.0.41｜通话同声、线上消息与群主解散（2026-08-14）", [
    "当前版本：网页 v917；私人 iOS 1.0.41 (41)；原生桥契约 17。PhoneWeb.bundle 已从共享网页核心重新生成。AI 账户首屏使用已确认红字：内置配置为新手便利服务并收取人工服务费，自己注册外置配置通常更省钱，两种方式自选且不强制。",
    "通话字幕协议覆盖网页与原生 PiP 的正文、说话人、逐字渐显、追加、位移、缩放和曲线。原生 voice processing 与 playAndRecord/voiceChat/mixWithOthers 会话允许外部媒体、角色播放和用户话筒并行；角色播音期间最终识别结果排队，不再通过暂停和重建识音丢弃插话。",
    "苹果兼容开关仅在开启时下移宠物页返回按钮和角色微信登录页顶部按钮。共同生活的角色线上消息会真实进入该角色微信或其参加的 AI 群并通知，目标歧义时先追问。AI 群消息菜单没有 AI 发言编辑入口。真人小手机群聊由群主调用服务端原子解散 RPC，其他成员同步后移除该群。",
    "云端状态：supabase/migrations/202608140001_phone_friend_group_disband.sql 已在现用项目部署并登记。Windows 已通过 node --check 和 555/555 自动测试；Mac 六 Target 编译、签名和真实 iPhone 同声、字幕、PiP 仍需验收。",
])

append("AI开发项目_新聊天启动说明.docx", "新聊天接手状态｜v917／私人 1.0.41（2026-08-14）", [
    "当前正式代码：网页 v917、私人 iOS 1.0.41 (41)、原生桥契约 17、PhoneWeb.bundle v917。安装工程包为小手机私人版_第四十一次安装_v917_2026-08-14.zip。Windows 555/555 自动测试通过；Windows 不能冒充 Mac 编译或 iPhone 真机验收。",
    "不要恢复角色播音前暂停识音或 760ms 后重建的旧方案。PiP 的 call.pip.update 只能更新名字、状态和字幕，不能调用 start 或重复激活音频会话。字幕动画参数只维护 CALL_SUBTITLE_MOTION 和原生同名字段。用户在角色播音期间的最终识别必须进入 _callHFPending 队列。",
    "共同生活线上代发只能去角色微信或角色参加的 AI 群，歧义群名先追问，不能冒充用户进入真人群。真人群解散依赖 202608140001_phone_friend_group_disband.sql，部署前界面会提示更新云端补丁。AI 群中的 AI 发言不可编辑。苹果兼容关闭和 Android 不得套用新增顶部偏移。",
    "下一步按顺序执行：Mac 编译全部六 Target；真机同时播放外部视频、让角色说话并由用户插话；核对网页大通话与 PiP 字幕；双账号验证只有群主能解散及其他成员同步消失。群解散迁移已部署，不要重复执行。",
])


for docx_path in DOCS.glob("*.docx"):
    document = Document(docx_path)
    in_v917 = False
    changed = False
    for paragraph in document.paragraphs:
        if "v917" in paragraph.text:
            in_v917 = True
        if in_v917:
            for run in paragraph.runs:
                if "553" in run.text:
                    run.text = run.text.replace("553", "555")
                    changed = True
                replacements = {
                    "解散群聊上线前必须先部署 202608140001_phone_friend_group_disband.sql。": "解散群聊迁移 202608140001_phone_friend_group_disband.sql 已在现用项目部署并登记。",
                    "发布前置：部署 supabase/migrations/202608140001_phone_friend_group_disband.sql。": "云端状态：supabase/migrations/202608140001_phone_friend_group_disband.sql 已在现用项目部署并登记。",
                    "下一步按顺序执行：部署群解散迁移；": "下一步按顺序执行：",
                    "部署前界面会提示更新云端补丁。": "迁移已部署，不要重复执行。",
                }
                for old, new in replacements.items():
                    if old in run.text:
                        run.text = run.text.replace(old, new)
                        changed = True
    if changed:
        document.save(docx_path)
    with ZipFile(docx_path) as archive:
        assert archive.testzip() is None
