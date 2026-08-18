from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, size=10.5):
    run.font.name = "等线"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(size)


def append_record(stem, heading, paragraphs):
    docx_path = DOCS / f"{stem}.docx"
    document = Document(docx_path)
    if not any(paragraph.text.strip() == heading for paragraph in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, 16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(docx_path)
    with ZipFile(docx_path) as archive:
        assert archive.testzip() is None

    txt_path = DOCS / f"{stem}.txt"
    text = txt_path.read_text(encoding="utf-8")
    if heading not in text:
        txt_path.write_text(
            text + "\n\n" + heading + "\n" + "\n".join(paragraphs) + "\n",
            encoding="utf-8",
        )


records = {
    "AI开发项目_Bug记录模板": (
        "v966 主屏点按迟钝与组件间歇消失修复记录（2026-08-17）",
        [
            "用户真机补充证据为：主屏横向往返滑动始终顺畅，但点击应用非常迟钝；组件有时出现、有时随滑动消失。该证据否定了 v965 仅以大面积玻璃合成层解释全部现象的判断，也说明手机发热最多是放大因素，并非唯一根因。",
            "点按根因是每个主屏项目从 pointerdown 起同时启用 touch-action:none、自定义 12 像素滑动判定、长按拖拽和原生 click。普通手指轻微位移会先被自定义滑动分支取消并阻止默认行为，使点击丢失或明显延迟。v966 把普通点按、横向翻页和纵向滚动交还浏览器原生手势；位移只取消尚未成立的长按，真正进入拖拽后才阻止默认行为和接管手势。原长按排序能力保留。",
            "组件消失根因是后续绝对定位主屏页使用 content-visibility:auto。移动 Safari、WKWebView 及部分 Android 浏览器在横向滚动容器中可能错误跳过页面绘制，使组件的占位或背景仍在而内容层被卸载。v966 让三个主屏页始终参与绘制，并减少滚动期间无意义的页码重绘；没有改变主题、图标、尺寸、布局、角色、人设、世界书、记忆、媒体共享或后台主动联系。",
            "Edge 移动端模拟分别按 iPhone 与 Android 视口验证：带轻微位移的设置点按在 1200 毫秒内打开；第 0、1 页往返后三个页面 computed content-visibility 均为 visible；第二页 12 个应用保持 80.5×88 几何框；没有残留拖拽、编辑状态或脚本错误。Windows 全量 Node 回归 682/682 通过，PhoneWeb.bundle 已由共享清单重建并对齐 v966。没有在 Mac 编译、签名，也未冒充真实 iPhone 长时间点击、发热或内存压力验收通过。版本为网页 v966、私人 iOS 1.0.88 (88)、原生桥 23。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "发布补充规范｜主屏点按与翻页不得同时由两套手势接管（1.0.88 起）",
        [
            "遇到“滑动顺畅但点击迟钝”时，应优先检查 touch-action、pointer/touch 监听、位移阈值、preventDefault、pointer capture 与 click 回退之间的竞争，不能继续只调合成层或把问题归因于发热。普通触摸应只有一个滚动所有者；长按成立前不得接管原生翻页。",
            "绝对定位组件位于横向滚动页时，不得未经真机证据使用 content-visibility:auto 或类似屏外卸载优化。若出现背景仍在但组件内容消失，应同时核对 computed content-visibility、几何框、命中状态和拖拽遮罩清理；页面隐藏、失焦及 pointercancel 都必须恢复透明度、pointer-events 与编辑类名。",
            "修复主屏交互时必须保留长按拖拽、主题、图标、布局和页面翻页，并对角色人设、世界书、记忆、主动功能、媒体会话和后台通知做全量回归。浏览器移动端模拟只证明 DOM、绘制与点按链路在该环境成立，不能替代 Mac 编译和真实 iPhone 的热压力、内存压力及连续使用验收。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "当前交付基线｜网页 v966／私人 iOS 1.0.88 (88)／原生桥 23（2026-08-17）",
        [
            "v966 的主屏普通触摸由系统原生处理点按、横向翻页和纵向滚动；轻微位移只取消待定长按，不再模拟 scrollLeft 或提前 preventDefault。长按真正成立后才进入 home-drag-active，原有跨页排序、Dock 与组件布局能力保持。",
            "全部主屏页保持 content-visibility:visible，避免移动浏览器把横向滚动容器中的绝对定位应用和组件误判为屏外内容后反复卸载。滚动页码只在实际页号变化时刷新，取消或页面隐藏会清理拖拽幽灵、透明度、pointer-events 和编辑状态。",
            "本轮未更改角色基础人设、世界书、真实事件与记忆、角色自主能力、挂断拒接失联事实、微信与放映室媒体、后台主动联系、主题颜色、图标和业务数据。当前为网页 v966、私人 iOS 1.0.88 (88)、原生桥 23；Windows 全量自动回归 682/682 通过，iPhone/Android 浏览器模拟通过。仍须在 Mac 编译签名，并在真实 iPhone 验证持续点按、长按拖拽、反复翻页、发热和内存压力下的组件稳定性。",
        ],
    ),
}


for stem, (heading, paragraphs) in records.items():
    append_record(stem, heading, paragraphs)
