from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, name="等线", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def append_section(name, heading, paragraphs):
    path = DOCS / name
    document = Document(path)
    if any(p.text.strip() == heading for p in document.paragraphs):
        return
    document.add_page_break()
    title = document.add_heading(heading, level=1)
    for run in title.runs:
        set_font(run, size=16)
    for text in paragraphs:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.25
        set_font(paragraph.add_run(text))
    document.save(path)


SECTIONS = {
    "AI开发项目_Bug修改规范.docx": (
        "新增强制规范｜视频通话画面事件必须与旧对话隔离（v902／私人 1.0.26 起）",
        [
            "视频通话取得新摄像头画面后，该帧是独立的新事件。角色生成时不得继续携带旧用户话语作为本轮对话，不得复读或补答上一句话；第一句可听台词必须点明画面描述中的具体人、物、文字、颜色、动作或环境细节。",
            "定时自动识别受分钟间隔和单次视频自动次数上限约束；用户口头说“你看一下”“看看这个”等主动要求时，必须重新截取当下画面且不受自动次数上限限制。主动请求若遇到正在识别，应排队执行而不是丢弃。",
            "识别过程不得占用通话中央字幕或显示技术系统提示。右下角只允许显示无文字旋转圆圈；成功后消失，失败时短暂显示“失败”。原始画面不得写入聊天或长期存档。",
        ],
    ),
    "AI开发项目_Bug记录模板.docx": (
        "v902／私人 1.0.26 Bug 记录｜画面已识别却复读上一句话（2026-08-12）",
        [
            "现象：视频通话设置 60 秒自动识别后，视觉接口可能已经返回画面描述，但角色继续回答用户上一句话；中央字幕还显示“正在让角色看画面”“画面已看清”等系统提示，用户无法确认画面是否真正送达。",
            "根因：画面描述作为 system note 追加到包含旧用户消息的普通通话历史中，旧消息仍会争夺当前任务；通话忙时画面事件与普通回复共用一个可覆盖的 pending 槽。自动次数上限还错误地限制了用户主动口令。",
            "修复：画面回复使用独立高优先队列和空对话历史，只保留当前画面事件；回复必须包含画面具体词，首次不合格自动重说，仍不合格则用真实视觉描述兜底。主动口令每次重新截帧且不限次数，自动识别才累计上限。",
            "界面：删除中央识别提示；右下角识别时仅转圈，成功消失，失败短暂显示“失败”。间隔设置由秒改为分钟，旧秒数自动换算，默认 60 秒迁移为 1 分钟。专项测试通过后仍需真机核验摄像头、视觉接口和角色语音完整链路。",
        ],
    ),
    "AI开发项目_项目说明文档.docx": (
        "v902／私人 1.0.26｜视频画面真实回应与分钟设置（2026-08-12）",
        [
            "网页版和私人 App 共用的视频通话摄像头链路升级为画面独立事件：视觉模型返回的当前画面不会再和历史最后一句用户话语混在同一轮，角色必须先说出画面中的具体细节，再按人设评价、关心或追问。",
            "用户主动口头让角色看画面时不限次数，每次都会重新截取当下画面；偏好设置的单次次数只限制定时自动识别。自动识别间隔从秒改为整分钟，旧设置自动迁移，0 仍表示关闭定时、只允许主动口令。",
            "通话中央区域不再显示任何识别系统词。最右下角仅在识别期间显示小圆圈，成功后消失；识别失败时短暂显示“失败”。画面仍只作为瞬时单帧发送，不写入聊天或存档。",
        ],
    ),
    "AI开发项目_新聊天启动说明.docx": (
        "新聊天接手状态｜v902／私人 1.0.26 视频画面链路（2026-08-12）",
        [
            "唯一仓库仍为 C:\\Users\\pc\\Documents\\小手机\\phone-work，分支 main。网页核心 v902，私人 iOS 1.0.26 (26)。",
            "视频通话自动识别按分钟配置并受单次自动次数上限限制；用户口头说“你看一下／看看这个”时必须重新取得当前画面且不限次数。画面回复必须隔离旧用户消息，并在第一句说出具体可见细节。",
            "界面识别状态固定在右下角：工作时只有旋转小圆圈，成功消失，失败显示“失败”；中央字幕不得出现识图系统说明。Mac 编译和真实 iPhone 的摄像头、前后翻转、分钟定时、主动不限次、失败提示及语音回应仍需实机验收。",
        ],
    ),
}


def main():
    for name, (heading, paragraphs) in SECTIONS.items():
        append_section(name, heading, paragraphs)
        document = Document(DOCS / name)
        assert sum(p.text.strip() == heading for p in document.paragraphs) == 1
        with ZipFile(DOCS / name) as archive:
            assert archive.testzip() is None
        print(f"verified {name}: {len(document.paragraphs)} paragraphs")


if __name__ == "__main__":
    main()
