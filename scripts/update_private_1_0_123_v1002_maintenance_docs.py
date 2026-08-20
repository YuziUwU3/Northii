from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def font(run, size=10.5):
    run.font.name = "等线"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(size)


def append(stem, heading, rows):
    docx = DOCS / f"{stem}.docx"
    document = Document(docx)
    if not any(p.text.strip() == heading for p in document.paragraphs):
        title = document.add_heading(heading, level=1)
        title.paragraph_format.page_break_before = True
        for run in title.runs:
            font(run, 16)
        for row in rows:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            font(paragraph.add_run(row))
        document.save(docx)
    with ZipFile(docx) as archive:
        assert archive.testzip() is None
    txt = DOCS / f"{stem}.txt"
    text = txt.read_text(encoding="utf-8")
    if heading not in text:
        txt.write_text(text + "\n\n" + heading + "\n" + "\n".join(rows) + "\n", encoding="utf-8")


records = {
    "AI开发项目_Bug记录模板": (
        "v1002 聊天入口分工与屏幕时间完成态修复（2026-08-20）",
        [
            "聊天入口回归：v1001 误把聊天右上角三个点和角色头像都指向新增聊天详情页。v1002 恢复旧分工：三个点进入角色资料页（朋友圈、权限、角色功能等），角色头像进入聊天详情页，消息气泡进入消息菜单。新增互斥断言，防止两个入口再次合并。",
            "屏幕时间悬浮条根因：共同生活读取已经完成，但完成横幅仍等待角色模型生成现场反应；模型超时、返回空内容或未通过真实数据校验时，横幅会长期停留，真实结果也没有独立重试状态。",
            "修复：读取完成横幅在 450 毫秒后立即关闭，不再等待模型；真实结果单独持久化，角色回复失败时只用同一真实结果重试模型，不重复读取、不编造替代回复。回复属于共同生活而用户停留在微信时，明确提示进入共同生活查看。",
            "版本与验证：网页 v1002，私人 iOS 1.0.123 (123)，原生桥 25。Windows 全量自动回归 850/850 通过；Mac 编译、签名和真实 iPhone 验证仍未完成。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "入口语义与异步完成状态规范（v1002／1.0.123 起）",
        [
            "同一界面中的头像、右上角更多按钮和消息气泡是三个独立入口。修改其中一个时必须分别断言各自目标，禁止用包含某个路由的宽泛测试掩盖入口合并。",
            "设备读取完成与角色回复生成是两个阶段。完成横幅只能反映真实读取阶段，读取结束后必须按固定短延迟关闭，不能被网络模型请求占用。",
            "模型回复失败时保存真实读取结果并退避重试；不得重新读取制造重复命令，不得写入固定假回复。跨页面生成的现场回复必须告诉用户真实落点，不能把共同生活消息伪装成微信消息。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "当前补丁基线｜网页 v1002／私人 iOS 1.0.123 (123)／原生桥 25（2026-08-20）",
        [
            "v1002 在 v1001 主动伴生与微信修复基础上恢复聊天入口分工，并修复共同生活屏幕时间读取完成横幅不结束、回复失败后结果丢失的问题。",
            "三个点进入角色资料；角色头像进入聊天详情；消息气泡进入消息操作。屏幕时间读取完成后先收起横幅，再独立生成真实角色反应；失败只重试真实模型。Windows 自动回归 850/850，真机仍待验。",
        ],
    ),
    "AI开发项目_新聊天启动说明": (
        "v1002 接手补充｜三个聊天入口不可再次合并（2026-08-20）",
        [
            "当前版本：网页 v1002、私人 iOS 1.0.123 (123)、原生桥 25。三个点、角色头像、消息气泡必须分别进入角色资料、聊天详情、消息菜单。",
            "共同生活屏幕时间完成横幅必须先关闭，再等待角色回复。模型失败时真实结果保存在 phoneFactRetry 中并自动重试，不得补写假回复。Windows 850/850；Mac 和真实 iPhone 未验证。",
        ],
    ),
}


for stem, (heading, rows) in records.items():
    append(stem, heading, rows)
