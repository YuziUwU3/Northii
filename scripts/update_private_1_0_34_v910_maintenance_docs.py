from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, name="等线", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def append(name, heading, paragraphs):
    path = DOCS / name
    document = Document(path)
    if not any(p.text.strip() == heading for p in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, size=16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(path)
    with ZipFile(path) as archive:
        assert archive.testzip() is None


append("AI开发项目_Bug记录模板.docx", "v910／私人 1.0.34 Bug 记录｜屏幕共享与后台 App 查看新鲜快照（2026-08-13）", [
    "现象：后台 App 查看偶尔只能收到类似普通主动消息，无法确认是否真实看到了当前软件；视频通话只能看摄像头，切到其他 App 后看不到角色字幕。角色也不能发起屏幕共享请求。",
    "根因：服务器此前可能拿请求前的旧 Screen Time 快照立即比较，旧基线与当前快照没有新增量时就误判失败；主动消息主要参考后台 outbox，真实普通聊天上下文和用户沉默时长优先级不足。iOS 原工程没有 ReplayKit Broadcast Upload 扩展、视频通话 PiP 容器和后台通话音频接力。",
    "修复：App 查看改为 start→request_snapshot→awaiting 三阶段，命令发出后只接受采集时间不早于 requestedAt 的新快照；没有新增量会再次请求，不再拿旧数据冒充成功。二选一动作在第一次查看成功时固定保存为 followupChoice，五分钟后严格执行继续询问或锁定。主动联系优先读取普通真实聊天并计算长时间未回复，再参考后台主动消息去重。",
    "屏幕共享：新增 PhoneScreenBroadcast 第六 Target，通过 ReplayKit 只在用户确认共享时把最新一帧写入 App Group；主 App 读取单张最新帧交给原有视觉识别。摄像头与屏幕共享共用分钟间隔和每通自动次数，口头让角色看仍不限次数。角色只能发起请求，必须经过应用同意弹框和 iOS 系统确认；拒绝、未完成确认、开始和结束都会作为通话事实交给角色。",
    "验证：Windows Node 全量回归 512/512 通过；网页 v910、私人 1.0.34 (34)、原生桥契约 13、六 Target 工程引用与版本静态校验通过。Windows 无法编译 Xcode，仍需 Mac 编译六 Target，并在真实 iPhone 验证系统共享倒计时、PiP 字幕、麦克风与后台音频。",
])

append("AI开发项目_Bug修改规范.docx", "新增强制规范｜屏幕共享授权、新鲜快照与通话后台连续性（v910 起）", [
    "屏幕共享属于高敏感权限。角色可以提出请求，但不得直接开启；每次必须先显示清楚的角色请求弹框，再由用户主动同意，并继续遵守 iOS／浏览器的系统级共享确认。拒绝、取消或超时不得伪装成已共享，模型也不得声称看到了尚未取得的画面。",
    "ReplayKit 扩展只能保留本次共享需要的最新一帧，使用原子覆盖写入 App Group；共享结束必须删除临时帧并清空 active 状态。不得后台常驻录屏，不得绕过系统红色共享指示，也不得把屏幕帧混入普通聊天或遥测上传。",
    "设备 App 查看必须记录 requestedAt，并只接受采集时间大于等于 requestedAt 的 Screen Time 子快照。旧快照、没有可识别名称或没有新增使用量时必须报告等待／无增量并重试，禁止把旧数据写成当前事实。基线要兼容旧数值格式和新 {used,name} 格式。",
    "摄像头与共享屏幕必须是两个独立媒体源；识别入口统一但取帧时共享屏幕优先。自动识别严格遵守分钟间隔、0 关闭和每通次数；用户口头明确要求看画面时不受自动次数限制。每次识别必须使用当次新帧，不能复用上一句用户话语生成回答。",
    "通话退后台需要使用平台正式能力：iOS 视频通话 PiP 承载角色名、原文字幕和翻译，AVAudioSession 保持通话音频。网页只能使用 getDisplayMedia 且每次由浏览器授权；移动网页不能承诺跨 App 共享。无 Mac 编译与真机证据时必须标注待验收。",
])

append("AI开发项目_项目说明文档.docx", "v910／私人 1.0.34｜视频通话屏幕共享、PiP 字幕与后台查看修复（2026-08-13）", [
    "当前版本：网页版 v910；私人 iOS 1.0.34 (34)；原生桥契约 13。私人 Xcode 工程由主 App、Report、Monitor、Shield、通知服务和 PhoneScreenBroadcast 共六个 Target 组成，内置 PhoneWeb.bundle 已重新生成并与 v910 对齐。",
    "视频通话左下角新增小型屏幕共享按钮。私人 iOS 使用 ReplayKit 系统共享面板；电脑网页使用 getDisplayMedia。共享开始后，画面识别改取共享屏幕，摄像头仍可独立开关；二者共用偏好设置里的分钟间隔和每通自动次数，口头要求角色查看仍不限次数。角色知道当前是普通摄像头还是屏幕共享，也知道共享结束。",
    "角色可按自身性格输出屏幕共享请求标记，App 显示带头像、原因、拒绝和同意的请求弹框。用户同意后仍需通过 iOS／浏览器系统确认；20 秒未真正开始会被视为未完成确认，角色不得假装看到。切到其他 App 时，iOS PiP 显示角色状态、外语原文字幕和中文翻译，原生音频接力用于维持通话声音。",
    "后台 App 查看以请求后的新鲜 Screen Time 快照为准，缺少增量时会重新请求；五分钟跟进动作在首次成功时固定为继续询问或锁定，避免模型话术与执行相撞。主动消息以普通真实聊天上下文为第一优先，并按角色性格感知长时间未回复，不再只沿用旧主动消息。",
    "自动化证据：Windows 全量 512/512 通过；PBX 六 Target、版本和资源静态核验通过。phone-role-push 需要部署本版。Mac 编译、签名、Broadcast Upload 扩展嵌入、系统共享倒计时、PiP、后台麦克风／音频以及真实 Screen Time 新鲜快照仍需真机验收。",
])

append("AI开发项目_新聊天启动说明.docx", "新聊天接手状态｜v910／私人 1.0.34（2026-08-13）", [
    "当前基线：网页 v910、私人 iOS 1.0.34 (34)、原生桥契约 13、Xcode 六 Target、PhoneWeb.bundle v910。Windows 全量自动化测试 512/512 已通过；不要回退摄像头口头识图、0 关闭自动识图、通知入聊、时间感知全局开关、共同生活输入和既有宠物逻辑。",
    "屏幕共享不是静默远控。角色只能提出请求；用户必须在角色弹框同意，并完成 iOS／浏览器系统确认。共享帧只保留 App Group 中最新一张，结束即删除。识别共享屏幕和识别摄像头使用同一视觉链路，但共享屏幕优先；口头查看不限次数。",
    "后台 App 查看必须等待 requestedAt 之后的新 Screen Time 快照；没有新增量就重试，不得把旧快照写成成功。主动消息先读普通真实聊天和沉默时长，再用最近后台消息防重复。followupChoice 一旦确定不可由后续模型文字改变。",
    "下一步真机只做三个明确操作：一，在 Mac 编译全部六个 Target 并安装真实 iPhone；二，视频通话中分别由用户和角色发起共享，切到其他 App，确认系统共享、PiP 双语字幕、麦克风和口头识图；三，立即 App 查看与一分钟后台测试各跑一次，核对新快照时间、通知正文和点入聊天后的对应消息。任何一步失败先收集 Xcode／设备日志再修改。",
])
