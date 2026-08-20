from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, size=10.5):
    run.font.name = "等线"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(size)


def append_record(stem, heading, rows):
    docx_path = DOCS / f"{stem}.docx"
    document = Document(docx_path)
    if not any(paragraph.text.strip() == heading for paragraph in document.paragraphs):
        title = document.add_heading(heading, level=1)
        title.paragraph_format.page_break_before = True
        for run in title.runs:
            set_font(run, 16)
        for row in rows:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(row))
        document.save(docx_path)
    with ZipFile(docx_path) as archive:
        assert archive.testzip() is None

    txt_path = DOCS / f"{stem}.txt"
    text = txt_path.read_text(encoding="utf-8")
    if heading not in text:
        txt_path.write_text(
            text + "\n\n" + heading + "\n" + "\n".join(rows) + "\n",
            encoding="utf-8",
        )


RECORDS = {
    "AI开发项目_Bug记录模板": (
        "v1006／私人 1.0.127 后台模型、通话与朋友圈回复修复记录（2026-08-20）",
        [
            "真实根因：v996 迁移到独立伴生云后，只配置了 APNs 和系统密钥，没有任何 OPENAI、DashScope 或 MiniMax 模型提供方；后台任务因此在生成阶段返回 no-provider，根本没有走到 APNs。旧“当前模型”提示只核验任务入库，没有把私人 App 当前角色所用的 API 地址、Key 和模型名同步到服务器，也没有核对最终投递结果。",
            "第二个独立根因：私人 App 的隔离包装把 roleBackgroundDispatchNow(true) 改成了无参数调用，丢失 keepalive；退到后台后 iOS 可能取消加速派发请求。旧界面又把入队成功误写成即将送达，使生成失败、APNs 失败和用户看不到通知无法区分。",
            "后台修复：开启后台主动联系时，把当前角色实际使用的 HTTPS API 线路放入受 owner secret 保护的角色资料；服务器优先执行该线路，拒绝 localhost、私网和非 HTTPS 地址。关闭功能时清除模型线路。立即测试保留任务 UUID，通过 task_status 核对生成、outbox 与 APNs 结果，并显示 no-provider、鉴权、余额、限流、无令牌、APNs 未配置或 Apple 拒绝等真实原因；隔离包装完整转发所有参数。phone-role-push 已部署到独立云，安全在线探针已确认新 task_status 分支生效。",
            "通话修复：电话提示不再无界携带完整长期记忆与全部历史；只选相关记忆，历史最多 48 条、12000 字符，超长单条压缩。保留原 chatAPI 与通话结构，避免私人 App 因巨大请求长时间无字幕、无声音后超时。",
            "朋友圈修复：评论保存准确角色目标，重试仍回复原评论；评论线程、近期聊天和长期记忆均设上限并按内容选取。只进行一次真实 chatAPI 请求，真实结果只追加一次；失败只留下重试状态，禁止生成假评论。",
            "验证与边界：JavaScript 语法、后台任务、私人隔离、朋友圈原函数运行时、通话提示预算及 Windows 完整 Node 回归 857／857 通过；独立云函数已部署且在线安全探针通过。Mac 编译、签名、真实 iPhone 的电话声音／字幕、朋友圈真实模型时延、APNs 可见通知仍未完成，不能写成真机通过。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "v1006 后台送达、通话与朋友圈回复回归规则",
        [
            "伴生云迁移或新建项目时，必须分别清点数据库／RPC、定时任务、APNs、设备令牌和模型提供方。只看到 APNs 密钥或预检通过，绝不能推断角色已经可以在服务器生成消息。",
            "后台测试的入队成功不等于送达成功。必须保留任务 ID，继续核对模型生成、outbox、APNs 状态与错误；界面只能报告实际到达的阶段，不能用倒计时或本地假消息代替回执。",
            "私人能力隔离包装必须用 rest 参数完整转发调用参数，并以测试固定 keepalive 等语义参数；禁止包装后静默改变函数契约。",
            "凡界面声称使用“当前模型”，必须真的同步角色当前线路并执行服务器端测试。同步地址只允许 HTTPS 且拒绝本机和私网；关闭后台能力时清除线路，并向用户明确披露同步范围。",
            "通话和朋友圈模型请求必须设置历史、线程、单条内容与总提示预算，并按当前内容选取长期记忆。朋友圈失败只能显示重试，禁止任何本地假回复或固定兜底冒充角色。",
            "Windows 自动测试只能证明代码与回归基线；Mac 编译、Apple 签名、后台挂起、电话音频／字幕和通知展示必须在真实 iPhone 单独验收。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "当前候选基线｜网页 v1006／私人 iOS 1.0.127 (127)／原生桥 25（2026-08-20）",
        [
            "v1006 修复独立云缺少模型线路、私人包装丢失 keepalive 和立即测试只认入队的三段后台故障；服务器现在使用当前角色线路并返回真实任务与 APNs 结果。",
            "电话改为相关记忆与有界历史；朋友圈评论保留准确回复目标、有界上下文和一次真实模型请求，失败不伪造。Windows 完整回归 857／857；phone-role-push 已部署并通过安全在线探针。",
            "仍待完成：在 Mac 全新解压 SmallPhone_v1006_CallMomentsBackgroundRepair，完成 Xcode 编译／签名，并在真实 iPhone 验证普通聊天、电话声音与字幕、朋友圈评论／重试，以及 App 退到后台后的真实 APNs 通知和返回 App 后的结果提示。",
        ],
    ),
    "AI开发项目_新聊天启动说明": (
        "v1006 接手说明｜后台模型、通话与朋友圈回复修复候选（2026-08-20）",
        [
            "当前候选：网页 v1006、私人 iOS 1.0.127 (127)、原生桥 25。Windows 完整回归 857／857；独立云 phone-role-push 已部署。不得把提交、ZIP、云函数部署或 Windows 测试写成 Mac／真机通过。",
            "Mac 必须完整解压 SmallPhone_v1006_CallMomentsBackgroundRepair，全新打开 PhoneCompanionTest.xcodeproj，执行 Product → Clean Build Folder 后再编译安装，不覆盖旧工程。",
            "真机按顺序验收：普通微信回复；电话声音与字幕且不连接超时；朋友圈真实评论回复与失败重试不伪造；设置 → 聊天与媒体立即模拟后台主动消息，确认 App 已退到桌面后收到通知，再返回 App 查看服务器模型和 APNs 的真实结果。",
        ],
    ),
}


for stem, (heading, rows) in RECORDS.items():
    append_record(stem, heading, rows)
