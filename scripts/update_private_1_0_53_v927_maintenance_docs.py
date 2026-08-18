from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, size=10.5):
    run.font.name = "等线"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(size)


def append_docx(stem, heading, paragraphs):
    path = DOCS / f"{stem}.docx"
    document = Document(path)
    if not any(p.text.strip() == heading for p in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, 16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(path)
    with ZipFile(path) as archive:
        assert archive.testzip() is None


def append_txt(stem, heading, paragraphs):
    path = DOCS / f"{stem}.txt"
    content = path.read_text(encoding="utf-8")
    if heading not in content:
        path.write_text(
            content + "\n\n" + heading + "\n" + "\n".join(paragraphs) + "\n",
            encoding="utf-8",
        )


records = {
    "AI开发项目_Bug记录模板": (
        "v927／私人版 1.0.53 Bug 记录｜苹果首屏文字、安全区与美化导入（2026-08-14）",
        [
            "现象一：只有 iPhone 浏览器添加到主屏幕后的玻璃第一页会把应用名称上下截断；第二页、安卓和私人 App 正常。根因是第一页绝对定位槽只有 75px，而 66px 图标、间距与文字行盒总高度超过槽位，WebKit 会压缩文字 flex 项。修复为八个首屏应用槽统一预留 90px，并把名称行盒固定为不可收缩的 20px；图标坐标和主题资源不变。",
            "现象二：iPhone 添加到主屏幕后的页面底部出现黑条。根因是 standalone 模式的可见窗口与百分比外壳高度存在差异，黑色 html/body 背景从底部露出。修复仅在真实 iPhone／iPad 主屏幕环境且用户开启“苹果兼容适配”时，把 phone/screen 固定铺满可见视口；安卓、普通 Safari 和私人 App 不命中。",
            "现象三：导出或导入美化后，接收方第一页组件、应用和 Dock 排列被覆盖，表现为空槽、图标散开或页面错位。根因是美化包错误包含 widgets、appLayout、homeLayout、appDock 和 homeReferenceAppSlots 等设备布局字段，与界面文字所承诺的‘只处理外观’不一致。修复后新包不再导出这些字段，导入旧包时也忽略它们，只恢复壁纸、图标、头像、气泡、组件外观与音乐外观。",
            "隔离范围：此轮没有修改玻璃组件尺寸关系、第二页网格、安卓主屏布局、私人 App 安全区，也没有修改电话声音、字幕、识别、屏幕共享、微信聊天或 v910 屏保。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "新增强制规范｜WebKit 行盒、安全区与美化包设备字段隔离（v927 起）",
        [
            "绝对定位应用槽必须容纳图标、间距和完整名称行盒，名称行盒必须禁止 flex 收缩。不能以 overflow:visible 代替真实高度，因为 iOS WebKit 仍可能先压缩文字再溢出。",
            "安全区修复必须按运行环境隔离。iPhone／iPad 添加主屏幕产生的外壳高度修复只能由 standalone 检测与苹果兼容开关共同启用；不得用全局 100vh／100dvh 规则改变安卓或私人 WKWebView。",
            "美化包属于可移植外观数据，不得携带设备布局。widgets、appLayout、homeLayout、appDock、homeReferenceAppSlots 等字段既不能导出，也不能从旧美化包导入；完整备份仍保留全部布局数据。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "v927／私人版 1.0.53｜苹果首屏文字、安全区与美化导入修复（2026-08-14）",
        [
            "当前正式版本：网页版 v927；私人 iOS 1.0.53 (53)；原生桥契约 18。",
            "玻璃第一页八个应用槽为 66px 图标和完整名称预留 90px 高度，名称行盒固定 20px，不再被 iPhone 主屏 WebApp 压缩。第二页、安卓和私人 App 继续沿用原布局。",
            "苹果兼容适配开启后，只有真实 iPhone／iPad 主屏幕版会把外壳铺满可见视口并覆盖底部安全区；普通 Safari、安卓和私人 App 不受影响。",
            "美化包只迁移外观，不再迁移页面、组件、Dock 和应用槽位。完整备份功能仍会保存全部布局。v910 屏保保持原版。",
        ],
    ),
    "AI开发项目_新聊天启动说明": (
        "新聊天接手状态｜v927／私人版 1.0.53（2026-08-14）",
        [
            "当前基线：网页版 v927、私人 iOS 1.0.53 (53)、原生桥契约 18。v910 屏保保持原版。",
            "iPhone 添加到主屏幕后的第一页名称依赖 90px 应用槽和不可收缩的 20px 行盒；不要再把首屏槽位压回 75px。",
            "底部安全区铺满只允许在 appleHomeCompatBrowserEnvironment() 为真且苹果兼容开关开启时生效；安卓、普通 Safari和私人 App 必须保持原外壳。",
            "美化包不得读写 widgets、appLayout、homeLayout、appDock 或 homeReferenceAppSlots。需要迁移完整布局时只能使用完整备份。",
        ],
    ),
}


for stem, (heading, paragraphs) in records.items():
    append_docx(stem, heading, paragraphs)
    append_txt(stem, heading, paragraphs)


drag_heading = "\u0076\u0039\u0032\u0037 \u8865\u5145\uff5c\u73bb\u7483\u9996\u9875\u957f\u6309\u62d6\u52a8\u4e0d\u518d\u8df3\u4f4d"
drag_paragraphs = [
    "\u6839\u56e0\uff1a\u73bb\u7483\u9996\u9875\u5e38\u6001\u4fdd\u7559 38px \u9876\u90e8\u5e03\u5c40\u95f4\u8ddd\uff0c\u4f46\u957f\u6309\u8fdb\u5165\u7f16\u8f91\u72b6\u6001\u65f6\u4f1a\u547d\u4e2d\u65e7\u89c4\u5219\uff0c\u5c06\u8be5\u95f4\u8ddd\u7acb\u5373\u6e05\u4e3a 0\uff0c\u5bfc\u81f4\u6574\u9875\u4e0a\u79fb 38px\u3001\u624b\u6307\u4e0e\u5e94\u7528\u5750\u6807\u9519\u4f4d\uff0c\u968f\u540e\u53c8\u56de\u5f39\u3002",
    "\u4fee\u590d\uff1a\u73bb\u7483\u4e3b\u9898\u5728\u7f16\u8f91\u72b6\u6001\u7ee7\u7eed\u4fdd\u7559 38px \u95f4\u8ddd\uff1b\u62d6\u52a8\u5e7d\u7075\u4e0d\u518d\u989d\u5916\u7f29\u653e\uff1b\u8865\u5165 iOS touchmove/touchend \u975e\u88ab\u52a8\u5904\u7406\uff0c\u957f\u6309\u540e\u53ef\u7ee7\u7eed\u8ddf\u624b\u79fb\u52a8\u5e76\u843d\u5230\u65b0\u69fd\u4f4d\u3002",
    "\u771f\u5b9e\u6d4f\u89c8\u5668\u51e0\u4f55\u9a8c\u8bc1\uff1a390\u00d7844 iPhone \u4e3b\u5c4f\u5e55\u6a21\u5f0f\u4e0b\uff0c\u9996\u9875 8 \u4e2a\u540d\u79f0\u5747\u4e3a\u5b8c\u6574 20px \u884c\u76d2\uff0c\u5e95\u90e8\u5916\u58f3\u95f4\u9699\u4e3a 0\uff0c\u957f\u6309\u524d\u540e\u62d6\u52a8\u5750\u6807\u5dee\u4e3a 0\uff0c\u5e76\u6210\u529f\u6539\u53d8\u69fd\u4f4d\u3002\u5b89\u5353\u4e3b\u5c4f\u5e55\u73af\u5883\u4e0d\u4f1a\u542f\u7528\u82f9\u679c\u5b89\u5168\u533a\u7c7b\u3002",
]

for stem in records:
    append_docx(stem, drag_heading, drag_paragraphs)
    append_txt(stem, drag_heading, drag_paragraphs)
