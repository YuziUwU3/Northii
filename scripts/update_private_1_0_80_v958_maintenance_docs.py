from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, size=10.5):
    run.font.name = "等线"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(size)


def append_record(stem, heading, paragraphs):
    docx_path = DOCS / f"{stem}.docx"
    document = Document(docx_path)
    if not any(paragraph.text.strip() == heading for paragraph in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, 16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(docx_path)
    with ZipFile(docx_path) as archive:
        assert archive.testzip() is None

    txt_path = DOCS / f"{stem}.txt"
    text = txt_path.read_text(encoding="utf-8")
    if heading not in text:
        txt_path.write_text(
            text + "\n\n" + heading + "\n" + "\n".join(paragraphs) + "\n",
            encoding="utf-8",
        )


records = {
    "AI开发项目_Bug记录模板": (
        "v958 主动消息重复与 App 关闭后不推送修复记录（2026-08-17）",
        [
            "现象一：自然模式融合后，随机主动消息会再次识别并延续用户最后一句，即使该句已经由角色正常回复。根因是主动记忆检索仍以最后一条用户文本为种子，服务端提示又把最近聊天设为第一优先上下文；现有去重只比较角色输出，无法识别“已经结束的用户回合”。",
            "处理一：网页端和云端均显式计算最近用户消息及其后的角色回复。已完成回合被标记为对话结束，随机主动联系只能从独立新事件开场，禁止回答、复述或改写最后一句；尚未回复的用户消息存在时，定时主动联系保持安静。主动记忆检索同时排除与最后一句相同或包含关系的候选。",
            "现象二：线下约会、共同生活面对面等临时状态会让角色资料同步 enabled=false，云端把它当成用户永久关闭后台联系；App 关闭后没有新的完整同步，因而持续不推送。根因是临时暂停与持久偏好共用同一个 enabled 字段。",
            "处理二：enabled 只表达用户持久开启状态，临时线下、通话、睡眠等状态改由 automation_config.suspended 表达。普通定时任务在暂停时顺延十分钟且不关闭资料；生成前后均复核暂停状态。明确的 reply_handoff、device_handoff、one_minute_test 与 app_watch_test 不被误删，面对面期间的普通 App 查看后续仍暂停。",
            "影响核对：保留每日额度、睡眠、通话、随机静默、去重、通知头像、回复接力、设备接力、一分钟测试和 App 查看测试。phone-role-push 已部署到现有 Supabase 项目；Windows 自动回归 668/668 通过。未在 Mac 编译、签名，也未冒充真实 iPhone 前台、后台或上划强退通知通过。版本为网页 v958、私人 iOS 1.0.80 (80)、原生桥 18。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "发布补充规范｜主动消息必须区分已结束回合与临时暂停（1.0.80 起）",
        [
            "随机主动消息不得仅凭最近用户文本生成。必须先判断该用户消息之后是否已有正常角色回复：已有回复代表回合结束，只能把历史当事实背景并从独立新事件开场；没有回复时，定时主动任务必须静默，交给正常回复或明确接力任务处理。输出去重不能替代回合边界判断。",
            "后台主动联系的持久偏好与临时场景暂停必须使用不同字段。enabled 只能由用户的长期开关决定；线下约会、共同生活面对面、通话、睡眠等状态写入 suspended，并允许任务顺延。任何临时状态都不得把云端角色资料永久同步为 disabled。",
            "修复后台推送时必须逐项核对明确任务：回复接力、设备接力、一分钟测试、App 查看测试；并核对普通任务的每日额度、睡眠、通话、随机静默、去重和通知头像。部署成功与 Windows 回归不能写成 Mac 编译或真实 iPhone 后台／强退推送通过。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "当前交付基线｜网页 v958／私人 iOS 1.0.80 (80)（2026-08-17）",
        [
            "主动消息系统现在有明确的对话边界：最近用户消息已有角色回复时，本轮视为结束，随机主动消息不得继续回答、复述或改写该句；最近用户消息尚未回复时，随机任务保持安静。服务端与前台队列使用同一边界语义，主动记忆检索不再把最后一句当作必选主题。",
            "后台联系把用户持久开关与临时场景暂停分离。角色资料 enabled 保留用户的长期选择，automation_config.suspended 负责线下约会、共同生活面对面、通话和睡眠等短暂停顿；普通定时任务暂停时顺延而不关闭资料，明确接力和测试任务仍按原规则保留。",
            "当前版本为网页 v958、私人 iOS 1.0.80 (80)、原生桥 18。phone-role-push 已部署，Windows 自动回归 668/668 通过。仍必须在 Mac 用 Xcode 编译五个 Target，并在真实 iPhone 分别验证前台、后台、锁屏与上划强退后的 APNs 行为。",
        ],
    ),
}


for stem, (heading, paragraphs) in records.items():
    append_record(stem, heading, paragraphs)


theme_records = {
    "AI开发项目_Bug记录模板": (
        "v958 屏保返回箭头主题同步补充（2026-08-17）",
        [
            "现象：主页顶部“回到屏保”小箭头位于 home 容器外，粉白、蓝白、灰白透明玻璃主题切换时它仍使用纯黑主题的深灰按钮，视觉上没有跟随主题。",
            "处理：只为 north-pack-pink、north-pack-blue、north-pack-gray 根类增加 lockpull 按钮与箭头描边配色；north-pack-black 不增加覆盖，继续沿用原来的深灰按钮。尺寸、位置、动画、可见条件和 lockShow 点击路径完全不改。",
            "验证：新增静态回归，确认三套浅色主题分别命中自己的按钮和箭头颜色，纯黑主题没有新增 lockpull 覆盖；屏保其他 lock 元素继续保持 v910 样式。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "界面补充规范｜跨层固定控件必须跟随主题根类",
        [
            "主题相关控件若位于页面主题容器之外，不能依赖容器后代选择器。应使用已经由运行时维护的根节点主题类做最小覆盖，并逐个列出受支持主题。",
            "为跨层控件补主题时不得顺带改动尺寸、定位、动画、点击函数、显示条件或其他同名前缀元素。要求保留原主题时，应以“不增加该主题覆盖”作为硬门槛并写入测试。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "v958 界面补充｜屏保返回箭头随透明玻璃主题",
        [
            "主页顶部“回到屏保”小箭头现在同步粉白、蓝白、灰白三套透明玻璃图标主题；纯黑主题继续使用原来的深灰按钮。此补充只改变按钮与箭头颜色，不改变屏保、主页布局、顶部安全区、动画或返回行为。",
        ],
    ),
}


for stem, (heading, paragraphs) in theme_records.items():
    append_record(stem, heading, paragraphs)


bubble_records = {
    "AI开发项目_Bug记录模板": (
        "v958 第二页姓名签名气泡主题亮色补充（2026-08-17）",
        [
            "现象：第二页头像卡下方的姓名／签名气泡固定使用深灰半透明背景。切换粉白、蓝白、灰白透明玻璃主题并使用亮色壁纸后，气泡仍整体偏暗灰，与主题不一致。",
            "处理：只为 north-pack-pink、north-pack-blue、north-pack-gray 的 glass-second-portrait-copy 增加更明亮的对应主题渐变、边框、阴影和深色文字；north-pack-black 不增加覆盖，继续保留原暗灰气泡。头像、文字内容、卡片尺寸、布局、拖动与编辑入口均不改变。",
            "验证：第二页布局回归明确要求三套浅色主题各有独立亮色气泡，并确认纯黑主题不存在该覆盖。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "界面补充规范｜亮色主题卡片必须同时校验文字对比度",
        [
            "把暗色玻璃改为亮色主题时，必须同时指定可读的深色文字、边框和阴影，不能只提高背景透明度后继续使用白字。主题覆盖范围应限制到目标子卡片，纯黑主题要求保持时不得新增对应覆盖。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "v958 界面补充｜第二页姓名签名气泡使用亮色主题玻璃",
        [
            "第二页头像卡的姓名／签名气泡在粉白、蓝白、灰白透明玻璃主题中使用更明亮的对应颜色和深色文字，以适配亮色壁纸；纯黑主题继续使用原来的暗灰气泡。此调整不改变第二页结构、头像、文字编辑或应用拖动。",
        ],
    ),
}


for stem, (heading, paragraphs) in bubble_records.items():
    append_record(stem, heading, paragraphs)


status_bar_records = {
    "AI开发项目_Bug记录模板": (
        "v958 私人 App 顶部系统安全区主题同步记录（2026-08-17）",
        [
            "现象：私人 App 为避免网页进入 Dynamic Island 和系统状态栏，WKWebView 固定从顶部安全区下方开始；该系统安全区此前固定为纯黑，因此在浅色壁纸与粉白、蓝白、灰白或线条白色主题下形成一条突兀的黑色区域。",
            "处理：保留 WKWebView 现有安全区边界，只新增 appearance.statusBar 原生桥消息。网页根据当前主题发送 black、pink、blue、gray 或 white；SwiftUI 外壳仅重绘顶部系统安全区，并用相应明暗模式保证时间、信号和电量图标可读。纯黑使用纯黑，粉蓝使用很浅对应色，灰色使用浅灰，白色使用纯白。主题值写入 UserDefaults，冷启动时可恢复。",
            "风险隔离：没有恢复已撤回的底部动态采样方案，没有让 WKWebView 忽略顶部安全区，也没有改变底部 home indicator、安全区高度、网页布局、屏保箭头或第二页卡片结构。Windows 只能完成源码与自动回归；仍需在 Mac 编译并用真实 iPhone 验证状态栏图标明暗和 Dynamic Island 机型布局。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "原生外壳补充规范｜系统安全区换色不得移动网页边界",
        [
            "需要按网页主题重绘 iOS 系统安全区时，只允许改变原生安全区背景及状态栏内容明暗，不得让 WKWebView 延伸到顶部状态栏或重新加入底部动态颜色采样。主题消息必须限制为固定白名单并提供纯黑回退；浅色背景必须配深色系统图标，深色背景必须配浅色系统图标。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "v958 界面补充｜私人 App 顶部系统安全区随主题变色",
        [
            "私人 App 顶部系统安全区保留在 WKWebView 之外，以避开 Dynamic Island；其背景现在随纯黑、粉白、蓝白、灰白和线条白色配色变化，并同步状态栏图标明暗。网页位置、底部安全区和各页面功能保持原样。此能力属于私人 iOS 外壳，普通网页版不会伪装成可控制浏览器或系统状态栏。",
        ],
    ),
}


for stem, (heading, paragraphs) in status_bar_records.items():
    append_record(stem, heading, paragraphs)


music_disc_records = {
    "AI开发项目_Bug记录模板": (
        "v959 音乐唱片独立四色与一起听按钮联动记录（2026-08-17）",
        [
            "需求：音乐播放页唱片需要脱离主屏透明玻璃主题，允许用户在音乐设置内单独选择黑、白、蓝、粉四色；头像下方原粉色一起听圆形按钮应跟随唱片颜色。白色唱片对应的白色按钮必须使用深色外圈和深色图标，避免与周围浅色线条融在一起。",
            "处理：在 S.music 中新增受白名单约束的 discColor，默认黑色并持久化。播放页根节点只增加 music-disc-black／white／blue／pink 状态类，四色规则只命中 .music-vinyl 与 .music-headphone-action；主屏 .home-vinyl-card、音乐进度条、发送按钮和其他主题颜色均不改动。白色按钮固定深色边框与图标。",
            "数据与发布：discColor 写入整包和分首音乐备份，导入时只接受四个合法值。网页 Service Worker 热修订键更新，使 v959 线上缓存重新拉取本次正式 app.js 与样式；私人 PhoneWeb.bundle 从共享清单重新生成。Windows 自动回归通过后仍需真实 Android 浏览器与 Mac／iPhone 验证视觉和点击行为。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "音乐界面补充规范｜应用内独立配色不得污染全局主题",
        [
            "音乐播放器内部的唱片配色必须由音乐自己的持久状态控制，不得复用主屏 north-pack-* 根类。选择器作用域必须同时包含音乐播放页根类和目标控件，禁止改写主屏唱片、全局强调色、进度条、发送按钮或其他同色控件。",
            "浅色控件与相邻浅色装饰可能重叠时，必须单独提供深色边界与图标对比度。新增可选颜色必须同时验证默认值、非法值回退、保存恢复、整包／分首备份导入，以及四种颜色之间连续切换。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "v959 界面补充｜音乐播放页唱片可独立选择黑白蓝粉",
        [
            "音乐设置的播放页外观新增黑、白、蓝、粉四色唱片选项。选择只影响完整播放器中的唱片和头像下方的一起听圆形按钮，不会改变主屏主题、主屏唱片、音乐背景、封面、进度条或其他按钮。白色圆形按钮使用深色外圈和深色图标保持清晰。",
            "所选颜色保存在 S.music.discColor，并随音乐整包和分首备份导出、导入恢复。私人 App 内置 PhoneWeb.bundle 与网页共享同一实现；Windows 验证不能代替 Android 真机渲染、Mac 编译或 iPhone 真机视觉验收。",
        ],
    ),
}


for stem, (heading, paragraphs) in music_disc_records.items():
    append_record(stem, heading, paragraphs)


v959_visible_ui_records = {
    "AI开发项目_Bug记录模板": (
        "v959 唱片入口可见性、顶部主题条与第二页透明气泡复核（2026-08-17）",
        [
            "用户真机反馈：播放页外观只看见背景和导出项目，没有找到唱片选色；顶部黑色状态区域切换主题后没有变色；第二页姓名／签名主题气泡虽然变亮，但填充过实，亮壁纸透不出来。仅确认源码存在选择器不能视为交付完成。",
            "处理：播放页外观列表新增独立且带当前颜色摘要的“唱片颜色”设置行，点击进入黑／白／蓝／粉选择；原有四色按钮仍保留。网页根节点新增统一 shell 主题状态并同步 theme-color，网页状态条直接使用浅粉、浅蓝、浅灰、纯白或原纯黑；私人 App 在主题切换和 native-ready 时强制重发 appearance.statusBar，避免桥初始化时序或旧缓存使顶部仍黑。",
            "第二页姓名／签名气泡保留粉、蓝、灰白主题和可读深色文字，但明显降低渐变填充 alpha 并提高背景模糊；线条淡粉、线条白色和默认暗色气泡也有对应透明规则。尺寸、点击编辑、头像、主屏布局、音乐背景、主屏唱片和其他音乐按钮均未改动。Windows 自动测试只能核对结构、作用域、缓存版本和资源包一致性，仍需真实 Android、Mac 与 iPhone 视觉验收。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "发布补充规范｜用户可见入口与主题视觉必须按实际交付层复核",
        [
            "新增设置不能只在模板中出现，还要给用户一个文字明确、可直接点击、能显示当前值的入口，并验证网页、私人 PhoneWeb.bundle 和最终 ZIP 三层都包含它。主题安全区修改必须同时覆盖网页状态条、原生安全区、冷启动桥时序和主题切换重发；静态源码存在不等于真机已生效。",
            "玻璃主题的亮色适配不得用接近不透明的浅色块代替玻璃。应同时验证背景可透出、文字对比度、边框和模糊强度，并覆盖图标包主题与线条配色两条状态路径。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "v959 界面复核｜唱片颜色入口、顶部主题条与透明气泡",
        [
            "音乐播放页外观列表直接显示“唱片颜色”一行和当前颜色，进入后可选黑、白、蓝、粉。网页顶部状态条和私人 App 顶部系统安全区使用同一主题映射，切换主题会立即同步；纯黑保持黑色，粉蓝灰使用浅色，白色为纯白。",
            "第二页姓名／签名气泡仍随粉白、蓝白、灰白和线条配色变化，但填充已降低到透明玻璃范围，亮壁纸可明显透出。以上调整不改变主屏布局、气泡编辑、音乐背景、主屏唱片或其他音乐控制。",
        ],
    ),
}


for stem, (heading, paragraphs) in v959_visible_ui_records.items():
    append_record(stem, heading, paragraphs)


for docx_path in DOCS.glob("AI开发项目_*.docx"):
    document = Document(docx_path)
    changed = False
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            updated = (
                run.text.replace("667/667", "668/668")
                .replace("v958 音乐唱片独立四色与一起听按钮联动记录", "v959 音乐唱片独立四色与一起听按钮联动记录")
                .replace("v958 界面补充｜音乐播放页唱片可独立选择黑白蓝粉", "v959 界面补充｜音乐播放页唱片可独立选择黑白蓝粉")
                .replace("使 v958 线上缓存重新拉取本次正式 app.js 与样式", "使 v959 线上缓存重新拉取本次正式 app.js 与样式")
            )
            if updated != run.text:
                run.text = updated
                changed = True
    if changed:
        document.save(docx_path)
    with ZipFile(docx_path) as archive:
        assert archive.testzip() is None

for txt_path in DOCS.glob("AI开发项目_*.txt"):
    text = txt_path.read_text(encoding="utf-8")
    updated = (
        text.replace("667/667", "668/668")
        .replace("v958 音乐唱片独立四色与一起听按钮联动记录", "v959 音乐唱片独立四色与一起听按钮联动记录")
        .replace("v958 界面补充｜音乐播放页唱片可独立选择黑白蓝粉", "v959 界面补充｜音乐播放页唱片可独立选择黑白蓝粉")
        .replace("使 v958 线上缓存重新拉取本次正式 app.js 与样式", "使 v959 线上缓存重新拉取本次正式 app.js 与样式")
    )
    if updated != text:
        txt_path.write_text(updated, encoding="utf-8")
