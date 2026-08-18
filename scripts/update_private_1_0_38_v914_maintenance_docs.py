from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, name="等线", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def append(name, heading, paragraphs):
    path = DOCS / name
    document = Document(path)
    if not any(p.text.strip() == heading for p in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, size=16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(path)
    with ZipFile(path) as archive:
        assert archive.testzip() is None


append("AI开发项目_Bug记录模板.docx", "v914／私人 1.0.38 Bug 记录｜透明对象误解、开关假开启与空字幕条（2026-08-13）", [
    "现象：上一版把展开通话页字幕容器做成半透明玻璃框，用户实际要求的是退出小手机后覆盖在 iPhone 主屏幕或其他 App 上方的系统悬浮通话小窗透明；完整通话页出现无字幕时仍占位的黑条。实时共享理解提示已开启，但设置页开关外观没有打开；左下共享按钮在等待系统面板确认期间没有清晰状态。",
    "根因：透明对象被错误理解为网页完整通话层；callsub 即使内容为空仍有背景、边框和最小高度。设置开关在活动通话期间只调用 renderCall，没有重绘当前设置页面，因此状态已保存但当前 DOM 仍旧。原生共享开始请求只表示系统面板已打开，并不等于用户已点开始直播，网页缺少 pending 状态。",
    "修复：完整通话页恢复原实心无框字幕布局，并在空内容时 display:none；自然渐显上移动画保留。透明度仅应用在 AVPictureInPictureVideoCallViewController 的内容根视图，背景 alpha 调到 0.38，并隐藏空字幕。设置开关直接同步当前元素的 on 类与 aria-checked；共享按钮增加等待系统确认态并在原生 active 回调或轮询结束后清除。",
    "新增：音乐页接入 Audius 公开在线曲库搜索。用户无需创建或登录音乐账号；搜索结果可直接在线播放或收藏，收藏结果进入现有 S.music.songs，并继续使用原有邀请角色一起听、上下文和聊天记录链路。不引入密钥，不替换本地音乐与导入歌单能力。",
    "验证：JavaScript 语法检查通过；通话、共享、原生基础、后台识图、音乐与 v914 新增回归共 30/30 通过。iOS PiP 是否把透明内容与主屏幕真实混合由系统合成决定，Windows 不能完成该真机验收；Mac 六 Target 编译和真实 iPhone PiP 透明度、搜歌播放仍需实机确认。",
])

append("AI开发项目_Bug修改规范.docx", "新增强制规范｜区分 App 内通话层与系统悬浮 PiP（v914 起）", [
    "收到“通话框透明、透视后台”需求时，必须先明确目标是 App 内网页通话页、App 内缩小浮层，还是退出 App 后由 iOS 管理的系统 PiP。不得用网页背景透明冒充主屏幕透视，也不得把一个对象的样式要求误施加到另一个对象。",
    "任何开关若在不重绘当前页面的分支中修改状态，必须同步修改触发元素的视觉类、aria-checked 和必要的禁用状态；不能只弹“已开启”提示。原生系统面板已打开、用户已确认、扩展已 active 是三个不同阶段，按钮必须显示 pending，只有真实 active 事件才能显示开启。",
    "空字幕容器不得保留背景、边框或最小高度造成黑条。字幕动画可复用，但无文本时必须隐藏。PiP 透明属于系统合成能力，代码可以设置非不透明根视图和 alpha，但发布说明必须保留真机验证边界，禁止承诺系统一定透出其他 App。",
    "在线音乐只能接入合法公开曲库或用户自己的合法来源；不得抓取受保护商业平台或绕过登录、DRM。在线目录故障不能破坏本地音乐、导入歌单或一起听核心链路。",
])

append("AI开发项目_项目说明文档.docx", "v914／私人 1.0.38｜系统悬浮通话透视与免登录在线音乐（2026-08-13）", [
    "当前版本：网页 v914；私人 iOS 1.0.38 (38)；原生桥契约 16。完整通话页恢复原来的无玻璃框字幕样式，空字幕不再显示黑条；只有退出小手机后显示在 iPhone 主屏幕或其他 App 上方的原生 PiP 内容背景使用轻透明设置。",
    "实时共享理解开关会立即同步真实外观。用户点左下共享按钮后，iOS 系统面板尚未确认时按钮进入等待态；只有 Broadcast Upload 扩展报告 active 后，才显示真实开启并通知角色共享已开始。既有共享状态独立队列、口头识图不限次数、自动次数上限和后台帧链路不变。",
    "音乐页新增公开在线搜索，不需要用户登录音乐账号。搜索结果可播放或收藏，收藏后与本地歌曲一样可邀请角色一起听；角色仍能读到当前歌名、歌手、一起听聊天和普通微信上下文。本地上传、直链、歌单导入导出继续保留。",
    "验收边界：网页功能可在 Windows 静态和自动化环境验证；iOS 系统 PiP 的最终透明合成、系统面板、后台音频、ReplayKit 和六 Target 签名必须在 Mac／真实 iPhone 验收。在线曲库受网络、地区和公开曲库收录影响，故障时不得影响本地音乐。",
])

append("AI开发项目_新聊天启动说明.docx", "新聊天接手状态｜v914／私人 1.0.38（2026-08-13）", [
    "当前基线：网页 v914、私人 iOS 1.0.38 (38)、原生桥契约 16。不要把完整通话页改透明；用户要的是系统 PiP 悬浮窗尽量透出 iPhone 主屏幕或其他 App。完整通话页应无玻璃字幕框，空字幕必须完全隐藏。",
    "实时共享理解开关已修复当前 DOM 假关闭；共享按钮具有等待系统确认态。不要回退共享状态高优先队列、原生轮询补偿、口头识图不限次数、自动上限、帧变化检测和有限后台任务成对释放。",
    "音乐新增 Audius 公开曲库，无用户登录。在线歌曲保存为普通 S.music.songs 项并复用原一起听链路；不得加入登录要求、密钥或盗版抓取，也不要删除本地上传、直链和导入歌单。",
    "下一步真机只做三件事：一，安装第三十八次包并在偏好设置反复开关实时共享理解，确认外观立即变化；二，视频通话点共享并完成系统开始直播，缩到 iPhone 主屏幕确认 PiP 样式、字幕和麦克风；三，音乐搜索一首公开曲目，播放、收藏并邀请角色一起听。失败时记录 Xcode、iPhone 时间点和实际系统画面，不盲改稳定链路。",
])
