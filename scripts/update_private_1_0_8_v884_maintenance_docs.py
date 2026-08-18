from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, name="等线", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def append_section(name, heading, paragraphs):
    path = DOCS / name
    document = Document(path)
    if any(p.text.strip() == heading for p in document.paragraphs):
        return
    document.add_page_break()
    title = document.add_heading(heading, level=1)
    for run in title.runs:
        set_font(run, size=16)
    for text in paragraphs:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.25
        set_font(paragraph.add_run(text))
    document.save(path)


SECTIONS = {
    "AI开发项目_Bug修改规范.docx": (
        "新增强制规范｜播放与录音必须显式交接音频会话（v884／私人 1.0.8 起）",
        [
            "原生连续语音不能只在每句结束后轮换 recognitionTask。只要角色语音播放与用户录音共用 AVAudioSession，就必须建立明确的播放/录音状态机：角色播放前 pause 并释放录音会话，全部播放结束后 resume；网页层不得仅凭 active 标志猜测原生麦克风仍可用。",
            "iOS 出现“黄色录音点还在但只有退后台才恢复”时，必须把后台恢复视为音频引擎被系统重建的证据。下一版应真正新建 AVAudioEngine，不得继续把同一个 engine.reset 当成等价替代。恢复失败必须保留 paused 状态并有限重试，不能让 JS active 与原生实际状态分叉。",
            "WKWebView 键盘适配禁止用 focus 前后切换两个绝对 bottom 值。输入条必须从未聚焦到聚焦始终锚定同一可见底边，相关按钮在同一坐标系中重新排布；出现一次上下弹跳即判定失败。",
            "带 !important 的页面级顶栏必须用同等或更高优先级补入安全区。共同生活微信顶栏与迷你通话属于独立验收对象：分别检查返回键、标题、状态胶囊、头像和卡片高度，不能用通用 nav/callscreen 断言推断覆盖。",
        ],
    ),
    "AI开发项目_Bug记录模板.docx": (
        "v884／私人 1.0.8 Bug 记录｜v883 语音再次中断、微信顶栏压缩与输入框跳动（2026-08-11）",
        [
            "v883 真机失败事实：第一句可识别并能交给角色，但角色语音播放后话筒自动失效；再次点击无效，只有把 App 退到后台再回来才能再识别一句。App 通话输入框聚焦时会先上下弹跳。网页版仅微信角色聊天仍被系统状态栏挤压，返回键位置过高，共同生活状态胶囊和迷你通话头像消失。",
            "纠正上一版诊断：v883 的原生代际轮换只覆盖“用户一句话结束”，没有覆盖“角色开始播放语音”。识别任务仍占用 playAndRecord/voiceChat 音频会话，角色 TTS 播放会中断输入节点；JS client.__active 没收到 end，后续 start 又因 active 直接返回。复用同一 AVAudioEngine.reset 也没有达到退后台时系统重建引擎的效果。",
            "布局根因：原生 callinput 默认 bottom:150px，focus-within 又切为 12px，键盘聚焦必然跳动。网页版 .cohab-wx-nav 用 !important 固定 56px，压过安全区总高度；全屏 .callscreen .cav 的高优先级 margin-top 又覆盖迷你通话头像的 margin:0，使头像被 58px 裁剪区隐藏。",
            "v884 修复：JS 在角色真实播放前调用 speech.pause，整轮播放结束后调用 speech.resume；原生 pause 清理任务、tap 和 AVAudioSession，resume 每次新建 AVAudioEngine，失败保留 paused 并有限重试。原生输入条始终 bottom:12px，挂断区固定在其上方。网页用最终高优先级选择器补足共同生活微信顶栏高度，并强制迷你通话头像 margin:0、卡片高 58px。",
            "隔离边界：公开审核中的 North App 未修改；网页语音识别未修改；锁定、同步、共同生活时间轴和其他 App 布局未修改。Windows 自动化只能证明静态链路和既有回归，最终仍需 Mac 编译和真实 iPhone 连续多轮“我说→角色说→我再说”验收。",
        ],
    ),
    "AI开发项目_项目说明文档.docx": (
        "共享网页 v884／私人小手机 1.0.8｜音频会话交接与微信最终安全区（2026-08-11）",
        [
            "架构方向不变：公开 North、私人小手机原生 App、安卓/浏览器网页共享业务核心并彼此隔离；公开 North 不控制本人设备，审核中的公开工程保持不动。",
            "私人小手机 1.0.8 将免提通话明确分成录音、角色播放、恢复录音三个状态。播放前原生释放录音会话，播放结束后建立全新的音频引擎与识别任务；同一个 JS speech client 保留，避免字幕可见但角色收不到、只认一句或必须退后台恢复。",
            "App 通话输入条不再在键盘聚焦前后改变底距；网页版 v884 只补微信共同生活顶栏和迷你通话最终选择器，保证返回键、状态胶囊、头像与紧凑高度同时恢复。普通浏览器、安卓和其他已稳定 App 不受本轮规则影响。",
            "后续主线仍未完成：手机号登录与 Keychain、云端主数据和重装恢复、大容量原生媒体存储、单控制器租约，以及锁定/同步/跨日/断网/后台长时间真机稳定性。本轮只清除阻塞这些工作的通话和微信适配回归。",
        ],
    ),
    "AI开发项目_新聊天启动说明.docx": (
        "新聊天接手状态｜共享网页 v884／私人小手机 1.0.8（2026-08-11）",
        [
            "唯一仓库为 C:\\Users\\pc\\Documents\\小手机\\phone-work，分支 main。共享网页版本 v884，私人小手机版本 1.0.8 (8)。公开审核中的 North 工程保持不动，公开 North 与私人小手机不能同时控制同一台设备。",
            "必须保留的最新诊断：v883 已正确推送但真机仍失败，原因不是缓存。原生识别任务虽然会在用户句尾轮换，却没有在角色 TTS 播放前释放录音会话；角色一开口就中断麦克风，而 JS active 阻止重新 start。v884 改为显式 pause/resume，并在 resume 时新建 AVAudioEngine。",
            "网页版本轮只改微信：共同生活顶栏的 !important 高度必须包含安全区，迷你通话头像必须覆盖全屏通话 margin，卡片固定 58px。原生通话输入条始终贴同一底边，禁止恢复 focus 前后 150px→12px 的跳变。网页话筒、音乐聊天、锁定同步和其他 App 不得顺手改动。",
            "真机验收按三步：一，安装 1.0.8 (8) 并确认共享核心 v884；二，免提连续完成至少三轮“用户说一句→角色语音回复→用户再说一句”，每句都要进入字幕、后台记录和角色上下文，全程不退后台；三，检查 App 输入框打开无跳动，以及网页版微信返回键、状态胶囊、迷你通话头像与卡片高度。",
            "验收通过后才回到主线：手机号账号/Keychain/云端恢复、大容量原生存储、单控制器租约、永久锁定与可靠同步长稳测试。不得再把网页装进 App 误判为完整原生基础能力已经完成。",
        ],
    ),
}


def main():
    for name, (heading, paragraphs) in SECTIONS.items():
        append_section(name, heading, paragraphs)
        document = Document(DOCS / name)
        assert any(p.text.strip() == heading for p in document.paragraphs)
        assert sum(p.text.strip() == heading for p in document.paragraphs) == 1
        assert not any(
            "codex-file-citation" in p.text or "PLACEHOLDER" in p.text
            for p in document.paragraphs
        )
        with ZipFile(DOCS / name) as archive:
            assert archive.testzip() is None
        print(f"verified {name}: {len(document.paragraphs)} paragraphs")


if __name__ == "__main__":
    main()
