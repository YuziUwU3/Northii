from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, name="等线", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def append(name, heading, paragraphs):
    path = DOCS / name
    document = Document(path)
    if not any(p.text.strip() == heading for p in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, size=16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(path)
    with ZipFile(path) as archive:
        assert archive.testzip() is None


append("AI开发项目_Bug记录模板.docx", "v919／私人 1.0.43 Bug 记录｜普通通话角色无声与微信输入栏遮挡（2026-08-14）", [
    "现象：未开启屏幕共享的普通语音电话和普通视频通话中，用户说话能显示，但角色不再出声；v918 新增语音门禁后还可能只显示用户临时识别，角色收不到最终一句。微信输入框第一次点击时 iPhone 蓝色光标偶尔落到框下方，需再次点击；打开功能或表情面板后，面板位于输入栏上方并可能把输入栏顶出可见区域。",
    "版本证据与根因：逐行对比 v917 提交 3903e23 与 v918 提交 6db66af，普通通话角色播放链唯一被直接删除的保障是 CallPictureInPictureController.playAudio() 中 stopAudio() 后的 activateCallAudio()。通话开始虽曾激活会话，但识音、系统切换或其他音频操作可能改变会话状态；AVAudioPlayer 仍创建却可能无声。v918 网页新增 voiceActivity、置信度和近期回声硬门禁，可能在已经出现中间字幕后丢弃最终结果。输入栏问题来自 iOS 对小于 16px 表单字号的焦点缩放，以及 DOM 中面板排在输入栏前方。",
    "修复：只恢复 v917 已验证的 playAudio() 播放前 activateCallAudio()，不改字幕、屏幕共享、角色生成和 TTS 内容。免提 onresult、busy 排队、speakWait 和 stop 路径按 v917 恢复，原生 voice processing 和输入 tap 保留，不新增猜测阈值。微信 textarea 使用 16px、明确盒模型和最小高度；#panel 以 flex order 排在输入栏后，打开面板先 blur 收起键盘。",
    "风险与验证：每次修改前必须先列影响面、历史正确基线、最小改动点和回归测试。本轮新增普通语音／视频播放前激活、v917 识别路径、字幕动画不变、输入栏面板层级和 16px 字号测试；完整 Windows 回归、Xcode 工程静态校验和 PhoneWeb.bundle 对齐通过后方可打包。Windows 不能代替 Mac 编译及真实 iPhone 的普通电话、普通视频、连续角色发声、静音误发和键盘真机验收。",
])

append("AI开发项目_Bug修改规范.docx", "新增强制规范｜改动前影响面与已验证回退优先（v919 起）", [
    "每次代码改动前必须先明确四项：会影响哪些已经正常的旧功能；是否存在用户确认正常的历史版本；能否只改一处或一条链路；需要新增哪些回归测试。没有完成这四项时不得为了新需求顺手重构相邻功能。",
    "遇到新改动破坏旧功能时，优先逐行对比最后正常提交和首个异常提交。若能定位唯一被删除或替换的保障，应先恢复历史已验证实现，不得同时调整字幕、提示词、共享、识别、TTS 或其他无关模块。普通通话角色出声的硬约束是每段原生 AVAudioPlayer 播放前确认 playAndRecord/voiceChat/mixWithOthers 会话已激活。",
    "语音降噪或回声门禁若造成真人最终结果被丢弃，应先回退到用户确认正常的上一版处理，再以独立实验和真机证据迭代。不能因中间字幕能显示就认定最终消息已送达角色。字幕动画有独立基准，音频和识别修复不得触碰其 DOM、CSS、PiP UILabel 或运动参数。",
    "iPhone 聊天输入控件不得使用小于 16px 的字号，避免 WebKit 焦点缩放和光标错位；可展开面板必须位于输入栏之后且先收起软键盘，不能覆盖或挤走发送区。界面修复必须与通话核心修复分文件、分测试验证。",
])

append("AI开发项目_项目说明文档.docx", "v919／私人 1.0.43｜普通通话声音恢复与微信输入栏修复（2026-08-14）", [
    "当前版本：网页 v919；私人 iOS 1.0.43 (43)；原生桥契约 18。PhoneWeb.bundle 由唯一共享网页核心重新生成。本版不修改字幕动画和屏幕共享功能。",
    "普通语音和普通视频通话的角色音频恢复 v917 行为：每段音频在创建 AVAudioPlayer 前重新调用 activateCallAudio()，会话仍为 playAndRecord、voiceChat、defaultToSpeaker、allowBluetoothHFP 和 mixWithOthers。该恢复针对未开启屏幕共享的普通通话，同样覆盖首次和连续回复。",
    "v918 新增网页硬门禁可能在已经显示用户中间字幕后丢掉最终识别，导致角色没有收到输入。v919 将免提最终结果、忙碌期排队、角色播音起点和停止清理恢复到 v917 已验证路径；原生 voice processing、Speech 连续识别和共享画面元数据保持不变。当前字幕动画保持 v918 状态，不参与本轮修改。",
    "微信 textarea 使用 16px、明确最小高度和盒模型；功能／表情面板通过 flex order 固定在输入栏之后，打开前先收起键盘。这样输入框第一次聚焦不触发 iOS 页面缩放，表情区打开时输入栏仍可见可点。",
])

append("AI开发项目_新聊天启动说明.docx", "新聊天接手状态｜v919／私人 1.0.43（2026-08-14）", [
    "当前正式基线：网页 v919、私人 iOS 1.0.43 (43)、原生桥契约 18、PhoneWeb.bundle v919。安装工程包为小手机私人版_第四十三次安装_v919_2026-08-14.zip。Windows 自动回归结果以本节发布记录为准，不能冒充 Mac 编译或真机验收。",
    "v919 只针对明确回归：普通语音／视频角色每段播放前重新激活 iOS 通话音频会话；免提网页识别回到 v917 路径；微信输入框使用 16px，面板排在输入栏之后并先收键盘。字幕动画、屏幕共享、角色生成、TTS 文本和宠物功能没有修改。",
    "以后修改前先做四项风险检查：旧功能影响面、最后正常版本、最小改动点、回归测试。遇到回归优先用提交对比恢复已验证代码，不在同一次修复里顺手重构相邻模块。",
    "真机先不打开屏幕共享：分别测试普通电话和普通视频的首次与连续角色出声；再测试免提真人话语和静音环境；最后检查微信输入框第一次光标、功能面板、表情面板和发送。字幕只确认保持当前效果，不再调整。",
])
