from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, name="等线", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def append(name, heading, paragraphs):
    path = DOCS / name
    document = Document(path)
    if not any(p.text.strip() == heading for p in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, size=16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(path)
    with ZipFile(path) as archive:
        assert archive.testzip() is None


append("AI开发项目_Bug记录模板.docx", "v904／私人 1.0.28 Bug 记录｜0 仍自动识图与通知头像正文回归（2026-08-12）", [
    "现象一：用户把视频自动识图间隔设置为 0，当前通话仍每隔一小会识别画面。根因是打开摄像头时创建的定时器只在创建前读取设置；保存 0 后没有重新布置定时器，已经存在的回调也没有在执行边界复核开关。",
    "修复一：全新默认改为 0；设置保存时若间隔变化且摄像头正在运行，立即重新布置或清除定时器；定时回调与自动识图入口再次检查非零间隔。口头让角色看画面仍走 manual 分支，不受次数或 0 值限制。",
    "现象二：后台通知中心把同一角色消息折叠为‘多个通知’，看不到每句正文；头像退化成昵称首字。根因是 APNs 和通知扩展为所有气泡复用同一个 threadIdentifier，同时后台头像同步在 IndexedDB 图片尚未进入内存缓存时直接生成了首字占位缩略图。",
    "修复二：每个 outbox 气泡使用 roleId、outboxId 和 messageIndex 组成独立通知标识；卡片保留角色昵称和当前句正文。头像同步在缓存缺失时主动从大容量图片存储读取真实头像，通知扩展等待通信意图捐赠完成后再更新内容。苹果强制通信头像为圆形并叠加 App 来源图标，无法改成方形或删除。完整自动化回归 483／483 通过。",
])
append("AI开发项目_Bug修改规范.docx", "新增强制规范｜关闭定时能力必须撤销存量任务并在执行边界复核（v904 起）", [
    "任何以 0 表示关闭的定时功能，不得只在创建定时器前判断。设置保存后必须立即撤销或重新布置已存在的任务，定时回调和最终副作用入口还要再次复核当前开关，防止已经排队的旧回调越权执行。手动触发和自动触发必须保留独立语义，关闭自动不得误伤用户主动操作。",
    "iOS 通知分组由 threadIdentifier 控制。产品要求每句话独立显示时，不得复用角色级 threadIdentifier；通信通知头像必须在上传前等待本地大容量图片读取，不能在缓存暂未命中时静默上传首字占位图。系统强制的圆形裁剪和 App 来源图标不得承诺可自定义。",
])
append("AI开发项目_项目说明文档.docx", "v904／私人 1.0.28｜识图关闭与后台通知修复（2026-08-12）", [
    "视频摄像头自动识图间隔 0 现在会立即清除当前通话定时器，并在旧定时回调和自动识图入口双重拦截；新用户默认 0。用户口头说‘你看一下’仍每次重新取得当前画面且不限次数。",
    "后台主动消息的每个气泡作为独立通知卡显示，直接保留角色昵称与当前句正文。角色头像同步会在内存缓存缺失时重新读取 IndexedDB 或私人本机大容量图片，再上传真实缩略图；通知扩展等待通信意图处理结束后提交卡片。苹果系统仍会把通信头像裁成圆形，并显示 App 来源小图标。",
])
append("AI开发项目_新聊天启动说明.docx", "新聊天接手状态｜v904／私人 1.0.28 通知与识图关闭修复（2026-08-12）", [
    "当前候选基线为网页 v904、私人 iOS 1.0.28 (28)。视频自动识图间隔 0 必须同时阻止新定时器、已存在定时器和排队旧回调；口头识图仍不限次数。",
    "后台角色消息每个气泡使用独立通知标识并显示该句正文。真实头像上传前必须等待大容量图片读取，通知扩展必须等待通信意图捐赠完成。苹果系统的圆形头像和右下角 App 来源图标不可自定义。",
    "Windows 完整自动化回归 483／483 通过。后续需在 Mac 编译五个 Target，并在真实 iPhone 验证 0 间隔、口头识图、逐句通知正文和真实头像。用户已补充但尚未进入本版的下一项：共同约会地点允许手动输入；关闭时间感知后所有线上、线下场景均不得注入或推断时间。",
])
