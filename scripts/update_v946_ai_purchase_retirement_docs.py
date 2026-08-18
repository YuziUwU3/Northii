from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, size=10.5):
    run.font.name = "等线"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(size)


def append_docx(stem, heading, paragraphs):
    path = DOCS / f"{stem}.docx"
    document = Document(path)
    if not any(p.text.strip() == heading for p in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, 16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(path)
    with ZipFile(path) as archive:
        assert archive.testzip() is None


def append_txt(stem, heading, paragraphs):
    path = DOCS / f"{stem}.txt"
    content = path.read_text(encoding="utf-8")
    if heading not in content:
        path.write_text(
            content + "\n\n" + heading + "\n" + "\n".join(paragraphs) + "\n",
            encoding="utf-8",
        )


records = {
    "AI开发项目_Bug记录模板": (
        "v946｜内置 AI 新购买渠道下线（2026-08-16）",
        [
            "用户决定立即关闭网页版内置 AI 的新购买渠道和套餐，但老用户余额、已有克隆音色、内置语音、流水和历史订单必须继续可用。原页面同时存在套餐卡片、付款二维码、音色克隆申请、付款凭证上传、低余额充值提示和旧版使用说明；若只隐藏套餐，旧入口或旧客户端仍可能继续创建订单。",
            "修复以最小范围删除前端套餐、支付方式、收款码、音色申请及凭证上传代码，历史订单改为只读；低余额提示与使用说明改为现有余额继续使用或按公告退款。普通聊天、聊天识图和生图继续明确使用用户自行配置的外置接口。",
            "服务端账户响应固定返回空 plans；purchase_create 与 purchase_submit 均在写库和上传前返回 410 purchase-channel-closed，阻止旧客户端绕过新版页面。历史账户、私有音色、流水、管理员历史审核与退款记录保留。仓库中的三张付款／联系方式二维码及 Service Worker、私人资源清单引用一并删除。",
            "验证：JavaScript 语法检查与购买下线、历史订单、内置图片边界、私有音色和内置 TTS 专项测试通过；完整 Windows 回归与线上 Edge Function 部署结果在发布记录中继续补充。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "新增强制规范｜收费入口下线必须前后端同时封口（v946 起）",
        [
            "关闭收费功能时，不得只隐藏套餐卡片。必须同时检查套餐数据、支付按钮、付款二维码或链接、新订单创建、凭证上传、联系方式、低余额引导、使用说明、Service Worker 缓存、安装包资源清单和服务端写入口。",
            "服务端应在任何数据库写入、文件上传或付款处理前返回明确的已下线状态，避免旧客户端继续提交。账户响应不得再返回已下线套餐。历史余额、已有权益、流水、订单和退款核对数据必须保留只读收尾路径。",
            "老用户使用与新购买必须分开验证：既要证明新订单／新凭证无法产生，也要证明既有余额可读取、内置语音仍可调用、已有专属音色仍能选择、历史订单和流水仍可查看。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "v946｜关闭内置 AI 新购买，保留老用户使用（2026-08-16）",
        [
            "当前网页版核心升级为 v946；私人 iOS 仍为 1.0.69 (69)，内置共享网页核心同步为 v946，原生桥契约 18 不变。",
            "AI账户不再展示点数套餐、支付渠道、付款二维码、新订单、付款凭证上传或新的音色克隆申请。页面只保留老用户已有余额、已绑定音色、历史订单、流水、内置语音及当前仍开放的影院字幕功能。普通聊天、聊天识图和生图继续使用用户自行填写的外置接口。",
            "phone-ai 服务端不再返回套餐，并对 purchase_create 与 purchase_submit 固定返回 410 purchase-channel-closed；因此旧网页或旧缓存也不能继续创建或提交新购买。三张旧付款／联系方式二维码已经从网页仓库与安装包清单中移除。",
            "本次没有清空任何用户点数、专属音色、流水或历史订单，也没有删除管理员处理历史退款所需的数据。",
        ],
    ),
    "AI开发项目_新聊天启动说明": (
        "新聊天接手状态｜v946 内置 AI 新购买已关闭（2026-08-16）",
        [
            "当前基线：网页版 v946；私人 iOS 1.0.69 (69)，内置网页核心 v946；原生桥契约 18。",
            "内置 AI 的新点数购买、新音色克隆申请、付款二维码和付款凭证上传均已从前端移除；服务端 purchase_create／purchase_submit 固定返回 410。不得恢复旧套餐、二维码、付款链接或上传入口。",
            "老用户余额、已有专属音色、内置语音、流水和历史订单必须继续保留。历史订单只读；退款与历史核对继续由管理员数据处理。普通聊天、识图和生图仍走用户自己的外置接口。",
            "后续若正式关闭影院字幕或全部内置 AI，必须另行处理其前端开关、服务端能力、余额退款和历史记录；不得把本次仅关闭新购买误写成所有老用户功能已经立即停用。",
        ],
    ),
}


for stem, (heading, paragraphs) in records.items():
    append_docx(stem, heading, paragraphs)
    append_txt(stem, heading, paragraphs)
