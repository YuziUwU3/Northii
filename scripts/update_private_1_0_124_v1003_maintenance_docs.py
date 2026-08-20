from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, size=10.5):
    run.font.name = "等线"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(size)


def append_record(stem, heading, rows):
    docx_path = DOCS / f"{stem}.docx"
    document = Document(docx_path)
    if not any(paragraph.text.strip() == heading for paragraph in document.paragraphs):
        title = document.add_heading(heading, level=1)
        title.paragraph_format.page_break_before = True
        for run in title.runs:
            set_font(run, 16)
        for row in rows:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(row))
        document.save(docx_path)
    with ZipFile(docx_path) as archive:
        assert archive.testzip() is None

    txt_path = DOCS / f"{stem}.txt"
    text = txt_path.read_text(encoding="utf-8")
    if heading not in text:
        txt_path.write_text(
            text + "\n\n" + heading + "\n" + "\n".join(rows) + "\n",
            encoding="utf-8",
        )


RECORDS = {
    "AI开发项目_Bug记录模板": (
        "v1003／私人 1.0.124 回复链路回归修复发布记录（2026-08-20）",
        [
            "发布范围：在用户确认 v1000 为最后正常回复基线后，只发布已经完成根因对照的回复回归热修。网页版本提升为 v1003，私人 iOS 提升为 1.0.124 (124)，原生桥保持 25；不得把版本提升写成 Mac 或真机已经验证。",
            "修复内容：phoneFactRetry 继续保存真实读取结果，但不再由 15 秒循环跨重启自动调用模型；旧 retryAt 原位迁移为 manual。replyGenerationRun 同时识别读取通道预先占用与运行中代次变化，两种正常接管都不重试、不误报模型或 API。朋友圈冷启动 pending 状态只转为真实失败和重试入口，不生成假回复。",
            "保留能力：v1001 的主动伴生、每日必查、低电、失联、心率关心、手动解锁、朋友圈真实评论／置顶／配图、聊天详情、拍一拍删除和精确远控，以及 v1002 的三个入口分工和屏幕时间 450 毫秒完成横幅全部保留。API 配置、统一角色 Prompt、普通正文生成和角色资料没有修改。",
            "验证与交付：共享源码和唯一 PhoneWeb.bundle 由 manifest 重建；JavaScript 语法、版本一致性、专项与 Windows 完整 Node 回归为 850／850。ZIP 必须只含全新 SmallPhone_v1003_ReplyPipelineRegressionRepair 工程、请在Mac编译前先读.md 和第一百二十四次安装说明，不夹带旧说明、预览、临时脚本或嵌套 ZIP。Mac 编译、签名和真实 iPhone 仍未完成。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "v1003 回复回归发布门槛",
        [
            "涉及模型重试、全局读取通道或持久 pending 状态的发布，除功能测试外必须验证旧正常聊天和朋友圈在失败、挂起、重启后的可用性。只断言“有重试”不算覆盖。",
            "发布身份必须原子更新网页 APP_VER、HTML build、Service Worker build／cache／注册 URL、index／repair、私人 Bundle Info、十二处 Xcode marketing/build、原生可见构建号、测试断言和安装说明。",
            "私人包必须从 manifest 重建唯一 PhoneWeb.bundle，并从全新暂存目录压缩；只允许一个当前安装说明。Windows 850／850 不能替代 Mac 编译和真机回复时延验收。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "当前候选基线｜网页 v1003／私人 iOS 1.0.124 (124)／原生桥 25（2026-08-20）",
        [
            "v1003 在 v1001／v1002 功能基础上停止共同生活真实读取失败后的跨重启自动模型重放，修复读取已占用时普通回复误报 API，并把朋友圈挂起回复恢复为真实失败与重试。",
            "Windows 完整回归 850／850；PhoneWeb.bundle 必须从共享 manifest 重建。Mac 编译、Apple 签名、普通微信回复速度、朋友圈评论、屏幕时间完成态和旧 retryAt 真机迁移尚未验证。",
        ],
    ),
    "AI开发项目_新聊天启动说明": (
        "v1003 接手说明｜回复链路回归修复候选（2026-08-20）",
        [
            "当前候选：网页 v1003、私人 iOS 1.0.124 (124)、原生桥 25。Git 提交前最后正常历史基线仍是 v1000；v1003 是包含 v1001／v1002 功能与回复回归修复的新候选。",
            "Mac 必须完整解压 SmallPhone_v1003_ReplyPipelineRegressionRepair，全新打开 PhoneCompanionTest.xcodeproj。先测普通回复和朋友圈，再做共同生活屏幕时间读取并立刻复测两者；重开 App 后确认旧自动重试不再运行。",
            "Windows 完整回归 850／850。Mac 编译、签名和真实 iPhone 未完成；不得把 ZIP、提交或推送等同于真机通过。",
        ],
    ),
}


for stem, (heading, rows) in RECORDS.items():
    append_record(stem, heading, rows)
