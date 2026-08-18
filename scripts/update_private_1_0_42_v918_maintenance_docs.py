from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, name="等线", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def append(name, heading, paragraphs):
    path = DOCS / name
    document = Document(path)
    if not any(p.text.strip() == heading for p in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, size=16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(path)
    with ZipFile(path) as archive:
        assert archive.testzip() is None


append("AI开发项目_Bug记录模板.docx", "v918／私人 1.0.42 Bug 记录｜字幕回归、杂音误发与宠物回窝（2026-08-14）", [
    "现象：v917 后网页大通话、原生共享小框以及用户说话识别都变成生硬逐字显示，原先小框整句从透明到实心的跳出动画被破坏；长翻译会超出可见范围。没有说话时偶尔会把外部视频声、角色外放或杂音识别成用户话语并发给角色。角色说话时外部视频可能暂停。语音电话没有共享屏幕入口。睡觉宠物停在窝外地毯上。角色资料仍暴露一分钟后台真实测试和立即查看软件测试。",
    "根因：历史对照确认 v700 使用整句节点替换和 300ms 淡入上浮；后续为逐字同步改成逐字符动画，错误覆盖了原本正确的小框行为。v917 为保留角色播音期间的用户插话，移除了忙碌期拦截并把所有最终识别积压到 _callHFPending，没有要求真人声活动、中间识别、置信度或回声排除，因此角色外放、外部媒体和杂音也可能被延迟发送。角色每句播放时重复激活全局音频会话，以及共享状态切换时停止媒体，增加了第三方视频暂停风险。宠物睡眠坐标仍落在地毯。",
    "修复：网页和 PiP 都恢复 v700 的整句 300ms 透明到实心、上浮 8px 和轻微缩放动画，用户与角色完全同路；翻译按长度降为 15px 或 12px 并允许换行。原生识别加入输入 RMS 真人声活动、Apple Speech 置信度、角色近期文本回声匹配与忙碌期三条件门槛；真实插话仍保留，静音、低置信度单字和角色回声不入队。音频会话只在通话开始配置，角色 AVAudioPlayer 发声链路保留；移除共享状态切换的媒体停止。语音和视频电话都可由用户授权 ReplayKit 共享。所有宠物按同时睡觉数量使用粉色窝内独立槽位和物种缩放。两个真实测试入口及网页函数已删除。",
    "验证边界：Windows 自动测试、JavaScript 语法、Xcode 工程静态校验可证明契约和回归保护，但不能代替 Mac 六 Target 编译与真实 iPhone。Apple 官方能力允许 playAndRecord 同时输入输出、mixWithOthers 与其他 App 混音、ReplayKit 录屏；但第三方 App 可选择不可混音会话，受保护视频也可拒绝捕获，本 App 无权强迫其继续播放。真机必须同时验证外部视频声、角色声、用户插话、静音不误发、长字幕和系统共享。",
])

append("AI开发项目_Bug修改规范.docx", "新增强制规范｜通话字幕基准、杂音门禁与外部媒体边界（v918 起）", [
    "通话字幕的已确认视觉基准为 v700：每次文本更新以整句话整体从透明到实心并上浮，不得拆成逐字符动画。用户识别、角色原文、翻译、网页大通话和原生 PiP 必须遵守同一运动语义；长文本必须优先缩小并换行，禁止横向裁掉。改字幕前必须先对照历史正确版本和双端回归测试，不能为了新增协议破坏既有小框动画。",
    "角色播音期间允许真实用户插话，不等于保留所有识别结果。进入待处理队列前必须结合真人声活动、中间识别、置信度、近期角色文本回声和重复结果；静音、低置信度单字、角色外放复述不得发送。不得通过暂停识音来消除回声，也不得删除角色 AVAudioPlayer 或既有 TTS 发声链路。",
    "连续通话只在开始时配置 playAndRecord/voiceChat/mixWithOthers 音频会话，不得为每句角色语音重新激活全局会话。语音与视频电话都必须提供用户确认后的 ReplayKit 共享入口。不得承诺强制第三方 App 混音或捕获受保护内容；真机验收应分别记录本 App 行为与第三方媒体自身策略。",
    "宠物睡眠坐标必须以背景图左侧粉色窝的实际内垫为边界；所有物种共用槽位算法，同时睡觉时按数量和物种缩放且不得重叠。新增功能不得改坏角色发声、用户话筒、电话挂断、摄像头、既有宠物成长与其他已稳定链路。",
])

append("AI开发项目_项目说明文档.docx", "v918／私人 1.0.42｜整句字幕、语音门禁、双通话共享与宠物回窝（2026-08-14）", [
    "当前版本：网页 v918；私人 iOS 1.0.42 (42)；原生桥契约 18。PhoneWeb.bundle 从唯一共享网页核心重新生成。角色资料的一分钟后台真实测试和立即查看软件测试已删除。",
    "通话字幕恢复 v700 整句动画：用户或角色的当前整句话每次更新都以 300ms 淡入、上浮 8px、轻微缩放进入；长翻译按长度缩小并换行。原生 PiP 使用多行 UILabel 承载同样的整句动画，不再使用逐字绘制视图，也不会因字幕更新重启 PiP 或重新激活音频会话。",
    "v917 杂音回归的直接原因是忙碌期最终识别不再拦截且无条件积压。v918 原生桥返回 Speech 置信度和实时输入声活动，网页叠加中间识别门槛、近期角色回声匹配、低置信度与重复过滤。角色播音期间识音保持连续，真实插话继续排队；角色 AVAudioPlayer 保留。音频会话为 playAndRecord、voiceChat、mixWithOthers，并只在通话开始配置。",
    "语音和视频电话都显示共享屏幕按钮，用户授权后由 ReplayKit 提供外部画面并允许角色按既有节奏理解。Apple 官方能力支持应用内同时录音和播放及可混音会话；能否让某个第三方视频继续播放、能否捕获受保护内容仍由对方音频会话、播放器和系统策略决定。所有睡觉宠物使用左侧粉色窝内的独立槽位，按数量和物种缩放，避免叠放。",
])

append("AI开发项目_新聊天启动说明.docx", "新聊天接手状态｜v918／私人 1.0.42（2026-08-14）", [
    "当前正式基线：网页 v918、私人 iOS 1.0.42 (42)、原生桥契约 18、PhoneWeb.bundle v918。安装工程包为小手机私人版_第四十二次安装_v918_2026-08-14.zip。Windows 自动回归结果以本节后续发布记录为准；Windows 不能冒充 Mac 编译或 iPhone 真机验收。",
    "字幕视觉基准固定为 v700 的整句淡入跳出，不得恢复逐字符 DOM 或逐字 PiP 绘制。用户识别与角色字幕必须同样处理，长翻译缩字并换行。不得删除角色发声链路；playAndRecord/voiceChat/mixWithOthers 只在通话开始配置，不得每句重新激活。",
    "杂音根因是 v917 为插话取消忙碌期拦截后无条件积压最终识别。当前必须保留输入 RMS 声活动、Speech 置信度、中间识别、近期角色回声和重复过滤；真实用户插话仍允许排队。语音与视频电话都可以在用户确认后共享屏幕。宠物睡眠必须进入左侧粉色窝独立槽位。角色资料两个真实测试入口已删除。",
    "下一步真机顺序：Mac 编译六 Target；播放外部视频并让角色连续说话，同时用户插话；静音观察不得产生消息；核对网页大通话与 PiP 的用户、角色整句动画和长翻译；分别从语音与视频电话启动共享；让一至四只不同动物同时睡觉并检查都在粉色窝内且不重叠。第三方媒体拒绝混音或受保护内容黑屏必须按系统边界单独记录。",
])
