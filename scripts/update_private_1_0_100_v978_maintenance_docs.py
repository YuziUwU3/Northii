from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, size=10.5):
    run.font.name = "等线"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(size)


def append_record(stem, heading, paragraphs):
    docx_path = DOCS / f"{stem}.docx"
    document = Document(docx_path)
    if not any(paragraph.text.strip() == heading for paragraph in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, 16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(docx_path)
    with ZipFile(docx_path) as archive:
        assert archive.testzip() is None

    txt_path = DOCS / f"{stem}.txt"
    text = txt_path.read_text(encoding="utf-8")
    if heading not in text:
        txt_path.write_text(
            text + "\n\n" + heading + "\n" + "\n".join(paragraphs) + "\n",
            encoding="utf-8",
        )


records = {
    "AI开发项目_Bug记录模板": (
        "v978 全锁实执行、通话记录状态与组件透明度修复记录（2026-08-18）",
        [
            "实际现象一：用户本轮明确要求全部锁定或解除全锁，模型回复却只输出带“仅内置”的部分控制标签，角色口头说已全锁但真实 iPhone 没有执行。根因不是云端通知，而是确定性全量识别只检查角色回复文本；角色回复没有再次说出“全部”时，执行层无法知道本轮用户要求的完整范围。v978 将当前真实用户消息显式传入控制落地层；只有用户语义是确定的全锁或全解、且角色输出同方向执行标签时才升级为内外完整动作。角色拒绝、相反动作、问句、否定句和条件句均不触发。",
            "实际现象二：通话记录展开一通后删除或编辑其中一句，页面完整重绘会重建 details，导致展开状态丢失；通话记录未注册到通用滚动恢复表，旧 clKeep 又只在重绘后立即写一次 scrollTop，布局变化后仍会跳到顶部。v978 用稳定的联系人加通话 session 键保存每一通的开合状态，并在当前帧和连续两个动画帧恢复距底部或原滚动位置；删除和编辑不再自动折叠。",
            "实际现象三：组件透明度数值能够保存，但四套图标包的固定玻璃规则选择器比自定义规则更具体，导致背景仍用固定透明度；私人 App 第二页低成本合成分支还写死 56% 实体底色。v978 仅在存在用户自定义值时，以更高优先级覆盖图标包规则，并让私人 App 第二页实体底色读取同一个透明度变量；未自定义设备继续使用原兼容默认底色，避免把部分设备的透明卡片问题误当成黑色主题专属问题。",
            "新增回归覆盖用户全锁请求加仅内置标签、拒绝与反方向保护、通话记录开合与滚动恢复、四主题自定义选择器优先级及私人 App 透明度变量。网页版本为 v978，私人 iOS 为 1.0.100 (100)，原生桥保持 23。Windows 自动测试通过后仍须 Mac 编译、签名和真实 iPhone 验收；不得宣称已通过 Family Controls 回执或真机 WebKit 渲染。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "发布补充规范｜全量控制依据、重绘状态与视觉偏好优先级（1.0.100 起）",
        [
            "全量锁定或解除的目标范围必须以当前真实用户请求与角色本轮执行决定共同确定，不能只看角色回复是否重复“全部”。只有当前用户语义为确定执行、角色给出同方向控制标签时，才能把错误的仅内置或仅外置标签升级为完整范围；拒绝、反方向、问句、否定和条件句必须保持不执行。不得读取更早的用户消息代替当前轮。",
            "可展开长列表在删除、编辑等局部操作触发整页重绘前，必须用稳定业务键保存开合状态，并保存滚动位置或距底部距离；重绘后要等布局稳定再恢复。不能只给 details 写固定 open，也不能只在 render 返回后同步写一次 scrollTop。",
            "用户可调颜色或透明度必须验证最终 CSS 层叠结果，而不只是检查值已保存或变量已生成。主题固定规则、设备兼容回退和自定义规则同时存在时，自定义规则只在用户主动设置后取得足够优先级；设备默认兼容底色仍须保留。原生低成本合成分支不得再用固定 alpha 覆盖用户透明度变量。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "当前交付基线｜网页 v978／私人 iOS 1.0.100 (100)／原生桥 23（2026-08-18）",
        [
            "v978 修复用户明确要求全锁或解除全锁时被模型“仅内置”标签降级的问题。当前用户请求会与角色同方向执行标签共同进入确定性落地；完整动作覆盖小手机已授权内置 App 和真实 iPhone 当前全部已选 App。拒绝、相反动作、问句、否定句和条件句不会误执行。",
            "通话记录按联系人和通话 session 保存展开状态。删除或编辑单句后，页面会恢复原开合状态，并在布局稳定后恢复原阅读位置或距底部距离。",
            "四套透明玻璃图标主题的组件颜色和透明度偏好现在覆盖固定主题值；私人 App 第二页实体玻璃底色使用同一透明度变量。未设置自定义值时继续使用既有设备兼容默认值。",
            "当前版本为网页 v978、私人 iOS 1.0.100 (100)、原生桥 23。共享网页与 PhoneWeb.bundle 已同步并完成 Windows 自动回归；仍须在 Mac 编译签名，并在真实 iPhone 验证全部锁定与解除回执、通话记录删除后的状态，以及四主题透明度 0／默认／100 的视觉变化。",
        ],
    ),
}


for stem, (heading, paragraphs) in records.items():
    append_record(stem, heading, paragraphs)
