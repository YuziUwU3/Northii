from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, name="等线", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def append_section(name, heading, paragraphs):
    path = DOCS / name
    document = Document(path)
    if any(p.text.strip() == heading for p in document.paragraphs):
        return
    document.add_page_break()
    title = document.add_heading(heading, level=1)
    for run in title.runs:
        set_font(run, size=16)
    for text in paragraphs:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.25
        set_font(paragraph.add_run(text))
    document.save(path)


SECTIONS = {
    "AI开发项目_Bug修改规范.docx": (
        "新增强制规范｜连续原生语音必须由原生层管理任务生命周期（v883／私人 1.0.7 起）",
        [
            "SFSpeechRecognizer 显示并提交第一句，不等于连续语音已经可用。真机必须连续说至少三句，并确认每句都进入通话记录和角色上下文；退出后台再回来只能作为故障证据，不能作为恢复方案。",
            "普通一句话结束时，原生层不得结束整个 JS 会话并依赖网页重新 start。必须在原生层清理旧 recognitionTask、移除 input tap、endAudio、audioEngine.reset、停用再激活 AVAudioSession，然后创建下一代任务；旧代回调必须用 generation 标记丢弃。只有用户明确关闭话筒或真正的终止错误才发送 end。",
            "WKWebView 键盘弹出后已经缩小可见区域。原生专用输入框只能贴新的可见底边，禁止保留网页时代的固定 150px 底距，也禁止重新引入 keyboard frame 与 visualViewport 双重补偿。缩小通话卡必须显式 bottom:auto 和最大高度，不能只写 inset:auto 后假设所有 WebKit 版本都会清掉旧边。",
            "页面安全区断言必须命中最终运行时 DOM。若后加载的 presentation layer 覆盖 renderShop、renderFood 或 renderDouyin，测试和 CSS 必须检查最终 commerce-top、dy-topbar、dynav 等选择器；只验证 app.js 中已被覆盖的占位类属于误检。底部 fallback 不得强制制造 34px 黑块，优先使用系统真实 env 值。",
        ],
    ),
    "AI开发项目_Bug记录模板.docx": (
        "v883／私人 1.0.7 Bug 记录｜语音只认一句、原生输入框错位与运行时安全区漏覆盖（2026-08-11）",
        [
            "真机补充现象：私人 1.0.6 每次退出后台再回来只能识别第一句话，之后黄色录音点仍在但无字幕；再次退出后台才又恢复一句。通话打字框被键盘推到屏幕中部，音乐聊天框可能完全看不到；缩小通话呈现异常高的空卡。网页版开启苹果主屏幕适配后，微信、购物、音乐、抖音、外卖仍有返回键过高和底部黑块。",
            "对上一条 v882 记录的纠正：v882 确实已经成功提交并推送，故障不是缓存或推送失败。1.0.6 只修通了字幕→final→角色读取这一条链，但原生 finishCurrentSession 在每句话后停止并取消任务、清空会话，且没有 audioEngine.reset；网页重启与旧原生任务形成竞态。后台切换由 iOS 强制重置音频栈，所以只能暂时再认一句。",
            "原生布局根因：callinput 仍硬编码 bottom:150px，而 WKWebView 已自行避让键盘；两者叠加后输入框被抬高。callscreen.mini 只用 inset:auto，没有显式 bottom:auto 和高度上限。音乐聊天也缺少原生键盘聚焦时贴底规则。",
            "网页安全区根因：商城、外卖和新版抖音由 commerce-ui.js 在 app.js 之后替换渲染，v879/v880 的 shop-nav、food-nav、dy-safe-nav 断言只命中了不会出现在最终页面的旧实现；音乐首页实际使用 music-app-head，却只改了播放器 music-topbar。底部 max(env,34px) 在已由系统避让的主屏幕 WebApp 中额外制造黑块。",
            "修复：私人原生语音保留同一 JS 会话，在原生层按代际轮换识别任务，每轮完整清理并 audioEngine.reset；取消普通句尾 end。原生输入框聚焦时贴 WKWebView 可见底边，缩小通话显式清边和限高。网页 v883 直接覆盖最终 commerce-top、music-app-head、dy-topbar、dynav，并把底部安全区改为系统真实 env 值。",
            "隔离与验证：公开审核中的 North App 工程未修改；网页话筒未修改。共享网页为 v883，私人小手机为 1.0.7 (7)。Windows 全量 node --test tests\\*.test.mjs 为 447/447 通过；Windows 仍不能替代 Mac 编译与连续三句真机验收。",
        ],
    ),
    "AI开发项目_项目说明文档.docx": (
        "共享网页 v883／私人小手机 1.0.7｜连续原生语音与最终 DOM 安全区（2026-08-11）",
        [
            "架构方向不变：公开 North、私人小手机原生 App、安卓/浏览器网页共享同一业务核心并彼此隔离；公开审核中的 North 不参与控制本人设备，也未在本轮修改。",
            "私人小手机 1.0.7 的语音任务生命周期改由原生层连续维护。每句停顿提交真实用户发言后，原生层立即清理、重置并建立下一轮识别，网页层不需要通过退出后台或重复创建会话才能继续。原生通话与音乐输入框按 WKWebView 的实际可见区域贴住键盘，缩小通话卡限制为紧凑高度。",
            "网页版 v883 仍保留 v877 已验证的 100% 外壳，只修苹果主屏幕模式。安全区规则以 commerce-ui.js 最终渲染的商城、外卖、抖音类名和音乐首页类名为准，底部不再强制增加 34px；普通 Safari、安卓和私人原生 App 不套用该网页开关。",
            "本轮没有完成手机号登录、Keychain、云端主数据与重装恢复、大容量原生媒体存储、服务端单控制器租约，以及锁定/同步/跨日/断网/后台长时间真机稳定性。这些仍是把小手机做成完整私人 App 的后续主线。",
            "自动化证据：全量 447/447 通过，覆盖共享缓存版本、最终页面安全区选择器、原生语音代际轮换、音频引擎 reset、原生键盘贴底、缩小通话限高和既有业务回归。下一门槛是在 Mac 编译 1.0.7，并在真实 iPhone 连续说至少三句。",
        ],
    ),
    "AI开发项目_新聊天启动说明.docx": (
        "新聊天接手状态｜共享网页 v883／私人小手机 1.0.7（2026-08-11）",
        [
            "唯一仓库为 C:\\Users\\pc\\Documents\\小手机\\phone-work，分支 main。共享网页版本 v883，私人小手机版本 1.0.7 (7)。公开审核中的 North 工程保持不动，公开 North 与私人小手机不能同时控制同一台设备。",
            "已确认 v882 推送成功，不能再把真机未变化归因于推送。1.0.6 的真实失败是只识别第一句：每句后原生任务被结束但音频引擎未完整重置；退出后台只是 iOS 临时重置音频栈。1.0.7 已改为原生内部连续轮换任务并丢弃旧代回调。",
            "原生 App 只改原生连续语音和原生布局，网页话筒不动。网页版只改苹果主屏幕适配：最终 commerce-top、music-app-head、dy-topbar、dynav 获得顶部安全区，底部只使用真实 env 值。禁止恢复 100dvh 根外壳、screen.height、keyboard frame+visualViewport 双重补偿。",
            "Mac 真机下一步依次做三项：一，编译安装 1.0.7 (7)，确认设置里共享核心显示 v883；二，保持通话页面连续说三句，每句字幕、后台通话记录和角色回应都存在，不退出后台；三，分别验证原生键盘输入框/缩小通话，以及网页版微信、购物、音乐、抖音、外卖返回键和底部黑块。",
            "长期未完成主线：手机号登录与 Keychain、云端主数据和删除 App 后恢复、大容量原生存储、单控制器租约、永久锁定和可靠同步真机长稳测试。不得因为网页已装入原生 App 就误判这些基础能力完成。当前 Windows 全量自动化为 447/447。",
        ],
    ),
}


def main():
    for name, (heading, paragraphs) in SECTIONS.items():
        append_section(name, heading, paragraphs)
        document = Document(DOCS / name)
        assert any(p.text.strip() == heading for p in document.paragraphs)
        print(f"verified {name}: {len(document.paragraphs)} paragraphs")


if __name__ == "__main__":
    main()
