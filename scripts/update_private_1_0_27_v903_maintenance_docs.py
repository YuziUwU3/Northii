from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, name="等线", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def append(name, heading, paragraphs):
    path = DOCS / name
    document = Document(path)
    if not any(p.text.strip() == heading for p in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, size=16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(path)
    with ZipFile(path) as archive:
        assert archive.testzip() is None


append("AI开发项目_Bug记录模板.docx", "v903／私人 1.0.27 Bug 记录｜识图可用但角色回答报上游异常（2026-08-12）", [
    "现象：用户已验证识图测试正常，普通聊天也正常，但视频通话取得画面后中央显示“上游聊天服务暂时异常”。v901 口头让角色看画面曾可正常回应，v902 出现回归。",
    "根因：v902 为防止自动识别复读而清空历史，把定时取得的当前画面只作为 system 消息发送；最终请求没有 user 消息且以 system 结尾。部分 OpenAI 兼容聊天接口拒绝这种消息结构，视觉识别本身并未失败。",
    "修复：只调整定时自动识图后的角色回答封装。自动取得的当前画面作为本轮唯一 user 消息发送，旧对话仍完全隔离；所有自动识图重试也保证以当前画面收尾。v901 已验证可用的口头让角色看画面链路保持原状。角色回答失败时不再污染中央字幕，只在右下角显示失败。",
])
append("AI开发项目_Bug修改规范.docx", "新增强制规范｜隔离历史不得删除当前 user 事件（v903 起）", [
    "定时自动识图为防复读而隔离聊天历史时，必须保留一个明确的当前 user 画面事件。不得向兼容聊天接口发送纯 system 请求，也不得以 system 消息作为需要生成回答的最终消息。已经真机验证正常的口头识图链路不得随自动链路一起重构。",
    "视觉识别成功与角色回答成功必须分开诊断。视觉接口正常、聊天接口异常时不得误报成识图失败；修复消息封装不得顺手改动普通聊天、相机采集或视觉路由。",
])
append("AI开发项目_项目说明文档.docx", "v903／私人 1.0.27｜视频识图上游兼容修复（2026-08-12）", [
    "v903 是 v902 的窄范围回归修复。定时自动取得的视频画面仍与旧聊天完全隔离，但会作为本轮唯一 user 事件交给角色聊天接口，参考放映厅的自动画面接话结构并继续防止复读；v901 已验证的口头识图结构保持不变。",
    "视觉识别、普通聊天、摄像头、前后翻转、分钟间隔和口头识图不限次数均不变；角色回答阶段失败只在右下角显示失败，不再覆盖通话中央字幕。设备授权同时移除迁移码、备用恢复码及其登录输入入口，邀请码保留，设备恢复和浏览器合并只允许本人扫脸或指纹。",
])
append("AI开发项目_新聊天启动说明.docx", "新聊天接手状态｜v903／私人 1.0.27 识图回答兼容（2026-08-12）", [
    "v903 修复 v902 定时自动识图纯 system 回答请求导致的兼容接口上游异常。自动取得的当前画面必须作为唯一 user 事件，旧对话不得重新加入；口头让角色看沿用 v901 链路。",
    "后续真机重点验证：同一视觉配置下图片测试正常、视频口头让角色看能说出具体画面、不会复读上一句话；失败只在右下角显示。",
    "授权规则：第一次正常邀请码仍保留；设置中的生成迁移码、生成备用恢复码和登录页代码输入入口已删除，旧代码兑换接口会拒绝。恢复设备和 Safari／Edge 合并只允许扫脸或指纹。",
])

append("AI开发项目_Bug记录模板.docx", "v903 授权收紧记录｜停用迁移码与备用恢复码（2026-08-12）", [
    "需求：保留第一次正常邀请码以及扫脸／指纹恢复；删除设置中的生成迁移码、生成备用恢复码按钮，并删除进入小手机时的代码恢复输入入口。",
    "实现：客户端不再导出迁移码、备用恢复码、本机身份恢复函数；登录页和设置页不再提供相关控件。旧版本仍可能缓存旧界面，因此 phone-license 云端对 transfer_create、transfer_redeem、recovery_create、recovery_redeem、local_identity_restore 统一返回 biometric-required。管理员放回授权不再生成恢复码，并保留已绑定的生物通行密钥。",
    "边界：正常邀请码没有删除；人脸／指纹绑定、恢复与 Safari／Edge 生物验证合并继续保留。聊天、角色和本机存档未改动。Windows 全量测试 483／483 通过。",
])
append("AI开发项目_Bug修改规范.docx", "新增强制规范｜设备恢复只允许生物验证（v903 起）", [
    "设备恢复和跨浏览器授权合并只允许本人扫脸或验证指纹。不得重新增加迁移码、备用恢复码、管理员恢复码或依靠本机保存身份绕过生物验证的入口。",
    "停用授权能力必须同时处理 UI、客户端导出和云端动作，不能只隐藏按钮；第一次正常邀请码属于首次授权入口，必须与恢复码区分，不得误删。",
])
append("AI开发项目_项目说明文档.docx", "v903 授权补充｜仅人脸或指纹恢复设备（2026-08-12）", [
    "第一次正常邀请码继续保留。设置中的生成迁移码、生成备用恢复码以及进入小手机时的代码恢复输入框已删除；Safari／Edge 合并只保留扫脸／指纹方式。",
    "客户端已移除代码生成与兑换能力，云端拒绝旧版本的迁移码、备用恢复码和本机身份自动恢复动作。管理员可以放回授权状态，但不会签发号码，用户仍必须使用已绑定的人脸或指纹。",
])
append("AI开发项目_新聊天启动说明.docx", "新聊天接手状态｜v903 设备恢复仅限生物验证（2026-08-12）", [
    "不得恢复迁移码或备用恢复码功能。第一次正常邀请码保留；设备恢复与 Safari／Edge 合并只允许人脸或指纹。",
    "相关 UI、客户端函数和云端兑换动作均已关闭；旧页面调用会收到 biometric-required。后续修改授权时必须同时验证邀请码和 WebAuthn 通行密钥未被误伤。",
])
