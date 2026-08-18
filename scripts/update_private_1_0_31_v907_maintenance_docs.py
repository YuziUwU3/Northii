from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, name="等线", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def append(name, heading, paragraphs):
    path = DOCS / name
    document = Document(path)
    if not any(p.text.strip() == heading for p in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, size=16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(path)
    with ZipFile(path) as archive:
        assert archive.testzip() is None


append("AI开发项目_Bug记录模板.docx", "v907／私人 1.0.31 Bug 记录｜后台接力被取消与网页停在901（2026-08-13）", [
    "现象：App感知立即测试在每日次数为0时无法开始；一分钟测试已入库却没有通知；角色回复或设备查看途中退出后任务和消息消失。多名用户的网页版长期停在v901。",
    "根因：立即测试错误复用了正式随机查看的每日配额门禁；服务器把人工测试和显式后台接力也按瞬时profile.enabled判断并取消；reply_handoff与device_handoff互相覆盖，且前端在回复真正持久化前取消兜底。服务工作线程的授权文件完整性检查仍要求v903已删除的redeemTransfer迁移码函数，导致v902以后安装阶段持续失败，旧v901接管页面。",
    "修复：立即测试只要求App感知开关和真实权限，不计每日次数；四类显式任务不受瞬时enabled影响，同角色仅替换完全相同种类任务；回复先写恢复尾、持久化聊天，再取消服务器接力，设备查看退后台后保留服务器兜底。服务工作线程改为校验restorePasskey与supportsPasskey，彻底移除redeemTransfer依赖。迁移202608130002已单独应用并登记，phone-role-push已部署。",
    "验证：Windows Node完整自动回归499/499通过，三份JavaScript语法检查通过；远端迁移列表已确认202608130002本地与远端一致。Windows无Xcode、Swift与LibreOffice，五Target编译、真机APNs可见性、App读取和文档分页仍需Mac／iPhone验收。",
])
append("AI开发项目_Bug修改规范.docx", "新增强制规范｜显式后台任务、离线升级与发布（v907起）", [
    "用户点击的一分钟测试、立即App查看、正在回复接力和正在设备读取接力属于显式任务。不得用页面离开后可能瞬变的普通主动消息enabled状态取消；不同任务种类不得互相覆盖。只有同种旧任务或已确认持久化成功的本地结果可以取消对应兜底。",
    "任何本地回复完成路径必须遵守：写恢复尾与聊天持久存储成功在前，取消服务器接力在后。页面隐藏或上划冻结发生在任意一步时，至少保留一个可恢复来源。",
    "删除授权能力后必须同步修改Service Worker完整性校验，并增加一条‘当前授权文件能通过安装检查’测试。只改版本号而不验证新Service Worker可安装，视为未完成发布。",
])
append("AI开发项目_项目说明文档.docx", "v907／私人 1.0.31｜后台可靠性、聊天与宠物成长（2026-08-13）", [
    "后台链路修复三项核心失败：App立即查看在正式每日次数0时仍可测试；一分钟测试不再被瞬时资料状态取消；聊天回复和设备查看退后台后保留独立服务器接力，且本地消息先持久化再撤销兜底。网页Service Worker升级校验已修复，旧v901可获取并安装v907。",
    "微信气泡设置新增独立的本人／角色语音翻译框背景与文字颜色。点击角色消息打开操作菜单时可直接删除该条角色消息，删除后同步持久化聊天。",
    "宠物每满30天进化一个阶段，不再由照顾分数加速；成长手册可退化回最初奶团，也可恢复自然阶段。睡眠坐标统一回到左侧窝垫，今日三餐的‘让TA照顾’会关闭手册并在房间展示照顾动作。",
    "版本为网页v907、私人iOS 1.0.31 (31)。共享PhoneWeb.bundle已重新生成，迁移202608130002已应用，phone-role-push已部署，Windows自动回归499/499通过。",
])
append("AI开发项目_新聊天启动说明.docx", "新聊天接手状态｜v907／私人 1.0.31 已发布候选（2026-08-13）", [
    "当前基线：网页v907、私人iOS 1.0.31 (31)，原生桥契约11。迁移202608130002_background_handoff_reliability.sql已经单独应用并登记，phone-role-push已经部署，不要重复执行全部历史迁移。",
    "本版最重要根因是v902以后Service Worker仍要求已删除的redeemTransfer，造成所有人停在v901；现已改为生物识别恢复函数检查。若仍显示901，先等待旧Service Worker下一次更新检查或从修复页打开，不要再次增加迁移码逻辑。",
    "真机优先验收：每日次数0且App感知开关开时立即测试；一分钟退出后通知；回复生成和设备查看中途退出后的通知与聊天落盘；旧901自动升级；宠物窝内睡眠、每月成长与退化；角色消息单条删除和语音翻译框配色。Windows自动回归499/499已通过，Mac五Target编译仍未执行。",
])
