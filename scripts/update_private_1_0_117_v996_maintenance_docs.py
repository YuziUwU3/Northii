from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, size=10.5):
    run.font.name = "等线"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(size)


def append_record(stem, heading, paragraphs):
    docx_path = DOCS / f"{stem}.docx"
    document = Document(docx_path)
    if not any(paragraph.text.strip() == heading for paragraph in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, 16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(docx_path)
    with ZipFile(docx_path) as archive:
        assert archive.testzip() is None

    txt_path = DOCS / f"{stem}.txt"
    text = txt_path.read_text(encoding="utf-8")
    if heading not in text:
        txt_path.write_text(
            text + "\n\n" + heading + "\n" + "\n".join(paragraphs) + "\n",
            encoding="utf-8",
        )


records = {
    "AI开发项目_Bug记录模板": (
        "v996 七项修复与独立伴生云迁移记录（2026-08-20）",
        [
            "问题一（共同生活摘要消失）：总结入口在异步 save 尚未真正落盘时就提示成功，iOS 随后挂起会丢失刚生成的角色记忆。修复为等待 saveNowAsync 持久化成功后再返回成功；失败不再假报完成。",
            "问题二（跨日与睡眠时间错误）：通话睡眠缺少开始、醒来和跨日持续时间锚点，模型会把睡前话题当成醒来后的当前事件；精确时间回答还可能在模型生成期间变旧。现持久化睡眠起止，计算真实时长和跨越天数，醒来提示明确标注旧话题，并在消息实际送达前校正精确时间。",
            "问题三（启动显示任务布置失败）：自动任务生成与手动任务页共用可见错误通道，启动网络波动会直接打扰用户。现自动路径静默、无效 JSON 使用安全后备任务、失败按持久化指数退避；只有用户主动进入任务页时保留明确错误。",
            "问题四（视频摄像头逐句卡顿）：每次角色 TTS 回复都会重置原生音频会话，摄像头开启时造成媒体会话反复拆建。现摄像头视频活动期间复用同一原生媒体会话，不随每句识别结果重建；摄像头关闭后的既有音频恢复逻辑保持。",
            "问题五（角色时区）：新增跟随设备、北京时间、常用国外时区和 Intl 支持的全球时区选择。角色日期、时间、星期、昼夜、日程、共同生活自动状态和后台推送时区统一读取角色所选时区；旧用户默认跟随设备，不改变原行为。",
            "问题六（“我去刷抖音”后静默）：查岗解析器把“刷”当作刷新/读取关键词，普通自用陈述进入早期查岗通道，查询无变化后又没有回复。现先识别“我去刷/看看某 App”等自用语句并保持正常聊天；“你查一下我的抖音”等明确查看指令仍进入真实读取。",
            "问题七（后台提醒和手机号绑定失效）：旧 Supabase 项目鉴权与函数请求超时，而邀请码新项目只部署 phone-license；锁定/查看仍可用是因为走 iPhone 原生直连，后台 APNs、角色任务和手机号账号走的是失效旧云。按用户授权新建完全独立项目 small-phone-companion-cloud（qvuahlqimcfgeoetosnl），未修改旧项目和邀请码项目。安装 20 个前置迁移及 202608200001 定时端点迁移，部署 phone-role-push 与 phone-companion-push，配置 APNs，核验 cron 只指向独立项目。",
            "手机号首次绑定根因与修复：原界面只有登录，新独立项目又为空，因此用户没有可登录账号。原生桥 25 新增 account.password.signup，界面分为“创建并绑定”和“登录并恢复”；独立项目启用自动确认，不发短信/邮件。云端已有备份时仍先展示备份并要求用户选择，本机在确认恢复前不改变，空账号才建立第一份备份。",
            "验证：云端 Auth、角色状态 RPC、伴生函数、备份 RPC、迁移记录、定时任务目标与旧项目隔离均已探测成功；Windows 全量自动回归 810/810 通过，三份维护文档共 423 页均已渲染检查。Windows 无法替代 Mac 编译、签名和真实 iPhone 的摄像头、跨日、后台/锁屏/上划退出 APNs 验收。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "时间连续性、静默自动任务与独立云边界规范（v996／1.0.117 起）",
        [
            "任何显示“已总结、已保存、已绑定、已备份”的成功反馈，都必须发生在对应持久化完成之后；iOS 可能在下一瞬间挂起，延迟 save 不能作为成功依据。云端已有备份时禁止自动上传或下载覆盖，必须先展示来源与时间并由用户选择。",
            "角色时间不能只在提示词里写当前小时。跨日事件必须同时保留绝对时间、角色时区、日历日期和持续时长；睡眠/醒来等状态转换要明确区分睡前旧话题与醒来后的新事件。精确报时时在实际送达前再校正。",
            "启动、恢复前台、定时扫描等自动任务不得复用手动页面的可见失败提示。自动失败应静默并持久化退避，避免每次启动重复请求；只有用户主动进入对应功能时才显示明确错误和可操作重试。",
            "视频通话开启摄像头时，TTS、识别和画面抓帧必须共享稳定的原生媒体会话。不得因每句角色回复暂停、拆除或重建摄像头/音频会话；摄像头关闭后才允许按原流程恢复普通音频会话。",
            "意图路由必须先区分“用户自己使用 App”和“要求角色查看/控制 App”。词语“刷、看看”不能单独触发查岗；主语、指令对象和动作方向必须同时成立。早期工具通道即使无数据变化也不能吞掉普通聊天回复。",
            "邀请码、旧业务和伴生后台必须保持三条独立云边界：邀请码备用项目只能部署 phone-license；伴生同步、APNs、角色后台任务和私人账号只使用独立伴生项目；其他既有网页业务仍保留旧服务。端点迁移后要同时核对客户端、原生桥、Edge Function、cron、迁移记录和允许域名，禁止把一个项目健康误判成全部链路健康。",
            "新建云项目时必须记录项目 ref、区域、迁移清单、函数列表和无密钥运行说明。不得把 PAT、数据库密码、APNs 私钥或用户凭据写入仓库。若建项密码未安全保存，应明确要求后台重置，不能猜测或伪造。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "当前交付基线｜网页 v996／私人 iOS 1.0.117 (117)／原生桥 25（2026-08-20）",
        [
            "本版集中修复七类问题：共同生活总结等待持久化；通话睡眠具备真实跨日时长和醒来事件锚点；启动自动任务失败静默退避；视频摄像头识别期间复用原生媒体会话；角色可选择跟随设备、北京时间或全球时区；“我去刷抖音”等自用陈述不再误触发查岗；后台通知、伴生同步和手机号账号迁到独立云。",
            "独立云项目为 small-phone-companion-cloud（qvuahlqimcfgeoetosnl，Singapore）。它独占私人账号、伴生 RPC、后台角色任务、APNs 和两个 Edge Functions。旧业务项目 lkhlyfpssmrjkkzhuzag 没有被迁移或覆盖；邀请码项目 lovbzibismsjqvjujilz 继续只承担 phone-license。可复现的 cron 端点迁移为 202608200001。",
            "手机号账号支持首次“创建并绑定”和既有“登录并恢复”。登录令牌仅保存在 iPhone Keychain，不返回网页脚本；账号内部以手机号映射的私有邮箱登录，不发送验证码。空云端才自动建立第一份备份；已有备份必须由用户选择保留本机上传或查看后恢复。",
            "当前版本为网页 v996、私人 iOS 1.0.117 (117)、原生桥 25。PhoneWeb.bundle 从共享清单重建；Windows 全量回归和文档渲染通过后，仍须在 Mac 打开全新 SmallPhone_v996_SevenFixesIndependentCloud，完成编译签名，并在真实 iPhone 验证手机号首次创建、跨夜醒来、国外时区、摄像头连续识别和前台/后台/锁屏/上划退出通知。",
        ],
    ),
}


for stem, (heading, paragraphs) in records.items():
    append_record(stem, heading, paragraphs)


old_verification = (
    "验证：云端 Auth、角色状态 RPC、伴生函数、备份 RPC、迁移记录、定时任务目标与旧项目隔离均已探测成功；"
    "相关本地回归 115/115 通过，版本/包定向回归 68/68 通过。最终全量数量以发布记录为准。"
    "Windows 无法替代 Mac 编译、签名和真实 iPhone 的摄像头、跨日、后台/锁屏/上划退出 APNs 验收。"
)
new_verification = records["AI开发项目_Bug记录模板"][1][-1]
bug_docx = DOCS / "AI开发项目_Bug记录模板.docx"
bug_document = Document(bug_docx)
changed = False
for paragraph in bug_document.paragraphs:
    if paragraph.text == old_verification:
        paragraph.text = new_verification
        for run in paragraph.runs:
            set_font(run)
        changed = True
if changed:
    bug_document.save(bug_docx)
with ZipFile(bug_docx) as archive:
    assert archive.testzip() is None

bug_txt = DOCS / "AI开发项目_Bug记录模板.txt"
bug_text = bug_txt.read_text(encoding="utf-8")
if old_verification in bug_text:
    bug_txt.write_text(
        bug_text.replace(old_verification, new_verification),
        encoding="utf-8",
    )


isolation_records = {
    "AI开发项目_Bug记录模板": (
        "v996 网页版与私人 App 伴生功能硬隔离修正（2026-08-20）",
        [
            "现象与风险：伴生设备、后台 APNs、手机号云账号和真实 iPhone 原生控制本应只属于私人 App，但网页与私人 App 共用 app.js。虽然部分入口已有原生桥判断，情侣空间伴生页、角色后台开关及若干定时拉取仍可能在网页版出现或尝试连接独立伴生云；旧本地数据还可能让角色提示词带入真实设备缓存。",
            "根因：过去把“共享业务源码”和“共享全部功能入口”混为一谈。privateNativeAppOn 只覆盖局部原生能力，没有形成界面、提示词、定时器、RPC 和直接 fetch 的统一硬门槛，因此单靠页面密码或原生桥缺失不能证明网页版完全隔离。",
            "修复：新增 privateCompanionAppOn 作为唯一产品边界。网页版隐藏情侣空间第三页、全部伴生快捷入口、关闭 App 后主动联系、软件使用感知和次数设置；底层 companionRpc、配对、同步、APNs、角色后台队列、定时拉取、状态核验和直接 keepalive 请求均在非私人 App 环境返回，不访问独立伴生云。角色设备提示、日常伴生背景和设备动作标签在网页版返回空内容或无操作。网页版原有普通主动消息、聊天、电话、查岗和其他网页功能不被删除。",
            "验证：新增 web-private-companion-isolation 回归，覆盖界面隐藏、后台网络拦截和角色事实隔离；同步修正安全读取测试以识别私人包装层。Windows 全量自动回归 814/814 通过，PhoneWeb.bundle 已从同一共享清单重建。Mac 编译、签名和真实 iPhone 私人伴生功能仍需真机验证，网页版还需发布后做一次无原生桥网络观察。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "网页版与私人 App 伴生能力隔离规范（v996／1.0.117 起）",
        [
            "共享 app.js 只代表共享可复用业务实现，不代表网页版可以获得私人能力。真实 iPhone 伴生、APNs、手机号私人账号、后台角色云队列、Family Controls、Device Activity、HealthKit 和原生控制一律以 privateCompanionAppOn/privateNativeAppOn 为硬门槛。",
            "私人能力必须同时隔离四层：网页版不渲染入口；直接调用返回明确的私人 App 限制；定时器、恢复前台和旧数据不能触发网络；角色提示词与动作解析不得读取或复述伴生缓存。只隐藏按钮而保留后台请求，或只拦 RPC 而泄露提示词，都不算隔离完成。",
            "共享文件中的独立伴生云 URL 可以作为私人构建所需常量保留，但网页版运行时不得向它发出请求。新增任何伴生端点时必须同时检查 companionRpc 之外的直接 fetch、keepalive、Edge Function 和启动/前台定时器，并加入网页版无原生桥回归。",
            "网页版的普通主动消息、网页聊天、网页电话、既有查岗和其他公开功能与私人后台伴生不是同一能力。做隔离时只关闭私人链路，不得借机删除或改变网页自己的正常功能。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "v996 产品边界补充｜网页版不使用伴生功能（2026-08-20）",
        [
            "网页版与私人 App 继续复用同一份业务源码，但产品能力已硬分层。网页版当前不显示或使用伴生设备、真实 iPhone 原生控制、手机号私人云账号、APNs 后台通知、角色服务器后台队列和软件使用感知；这些能力只在私人小手机 App 的可信原生桥存在时启用。",
            "网页版仍保留自己的普通聊天、电话、主动消息、情侣空间日常、既有查岗和其他网页功能。此次修正不是删除网页版功能，而是阻止网页接触私人伴生云和真实设备事实。即使浏览器保留旧伴生数据，角色提示词、自动背景和动作标签也不会在网页版使用它。",
            "当前发布身份不另起版本：网页 v996、私人 iOS 1.0.117 (117)、原生桥 25。PhoneWeb.bundle 已重新从共享清单生成；Windows 全量回归 814/814 通过。Mac 编译和真实 iPhone 验证仍未完成，且不得把 Windows 静态隔离测试写成真机伴生已通过。",
        ],
    ),
}


for stem, (heading, paragraphs) in isolation_records.items():
    append_record(stem, heading, paragraphs)


def remove_redundant_break_before_heading(stem, heading):
    path = DOCS / f"{stem}.docx"
    document = Document(path)
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip() != heading:
            continue
        paragraph.paragraph_format.page_break_before = True
        if index:
            previous = document.paragraphs[index - 1]
            breaks = previous._element.xpath('.//w:br[@w:type="page"]')
            if not previous.text.strip() and breaks:
                previous._element.getparent().remove(previous._element)
        document.save(path)
        with ZipFile(path) as archive:
            assert archive.testzip() is None
        return


for stem, (heading, _) in isolation_records.items():
    remove_redundant_break_before_heading(stem, heading)
