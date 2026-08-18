from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def find_doc(fragment: str) -> Path:
    matches = [path for path in DOCS.glob("*.docx") if fragment in path.name]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one document for {fragment!r}, got {matches}")
    return matches[0]


def set_chinese_font(run, name="等线", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def append_section(path: Path, title: str, paragraphs: list[str]):
    document = Document(path)
    if any(paragraph.text.strip() == title for paragraph in document.paragraphs):
        return
    document.add_page_break()
    heading = document.add_heading(title, level=1)
    for run in heading.runs:
        set_chinese_font(run, "等线", 16)
    for text in paragraphs:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.25
        run = paragraph.add_run(text)
        set_chinese_font(run)
    document.save(path)


append_section(
    find_doc("Bug修改规范"),
    "新增强制规范｜共同生活必须按长期约会完整链路验收（v864 起）",
    [
        "共同生活不是独立演示页，而是一次性线下约会的长期形态。两者必须复用同一角色设定、角色所选模型路线、线上长期记忆、旧约会记忆和准确时间；微信、一次性约会、共同生活的可见记录仍分别保存，禁止复制气泡或改变旧旁白格式。",
        "「约会中同步到线上 · 含共同生活」是唯一互通开关。开启时只把相关内容作为隐藏上下文，关闭时按平行独立世界处理。共同生活启用期间，微信静默间隔判断必须知道角色处于生活现场，不能再生成与现场矛盾的久未联系提示。",
        "生活状态必须由结构化标签解析并持久化，模型每轮必须得到当前准确日期时间、共同生活持续时长、当前状态持续时长和最近回家时间。开门只代表到门口；明确到家或进门后才能切回在家。",
        "状态横幅、顶部标签或 pending 队列都不等于角色已经回复。线上从外出切回在家后，必须在共同生活记录中真实落下一段角色旁白或动作与一段角色台词；生成失败时保留待承接任务并允许重试，禁止硬编码伪造成功。",
        "发布前必须同时验证角色模型路线、开关开闭两种上下文、可见格式隔离、跨日和离开时长、微信到家到共同生活真实消息的完整链路，以及一次性约会旧回归。静态预览、源码断言和测试总数不得替代正式代码链路验证。",
    ],
)

append_section(
    find_doc("Bug记录模板"),
    "v864 Bug 记录｜共同生活半独立、到家不同步与无真实回复（2026-08-10）",
    [
        "现象：角色在微信里已经说到家、开门、进来了，共同生活仍显示外出；进入共同生活后「让TA回应」可能没有生成内容。界面看似有共同生活状态，但角色模型路线、微信上下文和旧约会记忆没有形成完整闭环。",
        "错误验收：曾把静态共同生活预览、状态横幅出现和自动化测试通过描述成已经与微信和旧约会完整互通。这一判断不成立，也造成了对完成度的错误陈述。",
        "根因：状态解析、微信提示词、共同生活提示词、模型路由和可见消息写入分散实现；线上到家只改变状态，没有可靠的到家承接事务。共同生活回复还误用了部分一次性约会修复路径，角色所选模型没有显式放行。",
        "修复：共同生活改用独立的长期生活提示词和修复链，但复用旧约会的角色防漂移、格式校验、逐字显示与记忆能力；加入线上与共同生活的双向隐藏上下文、旧约会记忆、角色模型路线、准确时间与持续时长。",
        "到家闭环：非在家到在家的状态迁移会写入 pendingArrival，并生成真正的共同生活旁白和台词。微信原句只作隐藏承接背景，不复制到线下记录；失败时保留任务并可重试。",
        "验证要求：必须看到微信明确到家后状态切换、共同生活真实消息落库且可继续当面对话；同时验证同步开关关闭后完全隔离、旧约会格式不变、模型路线正确、跨日时间正确。",
    ],
)

append_section(
    find_doc("项目说明文档"),
    "v864｜共同生活升级为长期约会上下文（2026-08-10）",
    [
        "共同生活现在按「长期线下约会」实现：一次性约会仍保留原入口和原格式；共同生活不会因退出软件结束，并和同一角色的微信、电话、线上长期记忆及旧约会重要记忆保持语义连续。",
        "互通通过既有「约会中同步到线上 · 含共同生活」开关控制。开启时不同页面只交换隐藏背景，不复制可见消息；关闭时线上和线下按平行独立世界处理，因此不会破坏微信气泡或旧约会旁白格式。",
        "共同生活每轮使用角色当前选择的模型路线，主路线无有效内容且备用路线已配置时可安全回退。一次性约会原有模型默认和修复行为不变。",
        "角色可通过结构化生活状态自主选择在家、上班、回家路上或临时外出。模型能读取当前准确时间、共同生活开始和持续时长、当前状态开始和持续时长、最近回家时间，避免把数小时前或前一天说成刚刚。",
        "微信里明确到家或进门会触发真正的线下承接：共同生活记录新增角色动作旁白和自然台词，然后才能继续当面对话。只说开门不会提前判定已经在家；只显示回家横幅也不算承接完成。",
    ],
)

append_section(
    find_doc("新聊天启动说明"),
    "新聊天接手状态｜v864 共同生活长期上下文与到家承接（2026-08-10）",
    [
        "唯一仓库仍为 C:\\Users\\pc\\Documents\\小手机\\phone-work，唯一分支 main，远端 origin/main。禁止创建分支、worktree、强推或使用旧目录；一次只允许一个聊天修改和发布。",
        "v864 把共同生活从半独立状态页改成旧线下约会的长期形态。共享角色设定、角色所选模型、线上长期记忆、旧约会相关记忆和准确时间；三类可见消息仍分别保存，不得互相复制或改动旧格式。",
        "既有「约会中同步到线上 · 含共同生活」控制隐藏上下文互通。开启时微信能知道当前生活状态，共同生活能接续相关微信背景；关闭时完全隔离。",
        "线上从外出切回在家时必须生成真实的共同生活旁白和台词。状态标签、横幅、静态预览和测试总数都不能单独证明成功；必须验证消息已经写入共同生活且下一轮可继续回复。",
        "后续修改必须先跑共同生活上下文与到家承接专项、旧线下约会回归和完整 Node 回归，再做新的本地浏览器空数据链路验证；任何未完成环节都不得表述为已经接好或已经发布成功。",
    ],
)
