from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, size=10.5):
    run.font.name = "等线"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(size)


def append_record(stem, heading, paragraphs):
    docx_path = DOCS / f"{stem}.docx"
    document = Document(docx_path)
    old_heading = heading.replace("2026-08-18", "2026-08-17")
    renamed = False
    if old_heading != heading:
        for paragraph in document.paragraphs:
            if paragraph.text.strip() == old_heading and paragraph.runs:
                paragraph.runs[0].text = heading
                for run in paragraph.runs[1:]:
                    run.text = ""
                renamed = True
    if not any(paragraph.text.strip() == heading for paragraph in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, 16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        renamed = True
    if renamed:
        document.save(docx_path)
    with ZipFile(docx_path) as archive:
        assert archive.testzip() is None

    txt_path = DOCS / f"{stem}.txt"
    text = txt_path.read_text(encoding="utf-8")
    original_text = text
    if old_heading != heading:
        text = text.replace(old_heading, heading)
    if heading not in text:
        txt_path.write_text(
            text + "\n\n" + heading + "\n" + "\n".join(paragraphs) + "\n",
            encoding="utf-8",
        )
    elif text != original_text:
        txt_path.write_text(text, encoding="utf-8")


records = {
    "AI开发项目_Bug记录模板": (
        "v967 回复显示、设备卡片与网页数据隔离修复记录（2026-08-18）",
        [
            "本轮接续了 app.js 中五处尚未接通的局部修改。模型可能已经返回角色正文，但正文被包在 JSON、choices、output、message/content 或 Markdown 代码围栏中；旧链路只在部分校验处看到原始字符串，后续解析、重试和存储仍可能把它判成空回复。v967 在微信、线下约会、共同生活、最终可见性诊断及回复拆分入口统一提取可见正文。",
            "仅缺少【第三人称旁白】或混入未包裹动作时属于格式问题，改用简短格式修复提示；只有拒绝模板、OOC、AI／模型说明等不安全漂移才继续拦截。修复输出、重复修复和最终重试也先解包再判断，避免已有正常台词被清空成“模型没有生成”。空 JSON 或畸形对象不会作为原始 JSON 气泡显示。",
            "第二页卡片透明依据用户补充按设备／WKWebView 合成差异处理，不归因于黑色主题。私人 App 已有的低成本无 backdrop-filter 分支在部分设备上只剩极低透明底色，v967 仅在该原生稳定分支为黑、粉、蓝、灰白及自定义色卡片补足实色背景和说明文字底色；网页和正常支持模糊的设备不被全局改色。",
            "公开网页的数据隔离只切断真实 iPhone 伴生缓存、设备快照和原生控制上下文的读取／注入。网页原有的抖音、微信、朋友圈、电话短信、浏览器、钱包、X、机票查岗，以及角色权限、内置软件锁定、每日限额等功能继续存在，没有用删除功能代替隔离。",
            "网页版本升级为 v967，私人 iOS 为 1.0.89 (89)，原生桥保持 23；PhoneWeb.bundle 由共享清单重新生成。Windows JavaScript 语法检查和全量 Node 回归 686/686 通过。当前环境没有 Mac/Xcode 与真实 iPhone，未宣称编译、签名、真机卡片合成或真实设备数据边界已经验收通过。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "发布补充规范｜模型可见正文、设备透明兼容与网页数据隔离（1.0.89 起）",
        [
            "接入兼容模型时，校验、格式修复、重复检查、气泡拆分、存储和无回复诊断必须消费同一份已解包的可见正文。常见 JSON／代码围栏可以兼容提取，但空对象、无可见字段或畸形 JSON 不得直接显示为角色台词。格式不合规与不安全拒绝／OOC 必须分流，不能用同一清空规则把可恢复正文丢掉。",
            "玻璃卡片在部分设备透明时，必须先核对该设备是否走了禁用 backdrop-filter 的稳定分支、背景实际 alpha、文字底色和合成层，不能仅凭用户当前主题把根因写成黑色主题。兼容修复应覆盖实际受影响的所有色板，并限定在有证据的原生／设备分支，避免全局改变网页版或正常设备的视觉。",
            "网页版与私人 App 共享源码时，应通过运行环境门禁隔离真实手机缓存、设备权限、原生控制提示和主动上下文；不得删除网页原有查岗应用、角色授权、内置锁定和每日限额等业务能力来实现隔离。导入旧状态后也必须保证公开网页不能把残留伴生快照当成实时 iPhone 数据。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "当前交付基线｜网页 v967／私人 iOS 1.0.89 (89)／原生桥 23（2026-08-18）",
        [
            "v967 统一处理模型回复的常见外层封装，并把可恢复的格式问题与真正不安全的拒绝／跳出角色分开。线下约会、共同生活与微信回复会在校验、修复、拆分和最终显示前使用同一份可见正文，降低模型已有回复却被误报为空的概率。",
            "私人 App 第二页卡片的兼容底色只作用于原生无背景模糊稳定分支，覆盖黑、粉、蓝、灰白和自定义色；问题按不同设备合成行为处理，不定义成黑色主题专属问题。正常网页与支持原玻璃合成的环境继续使用既有主题效果。",
            "公开网页只隔离真实 iPhone 伴生数据和原生控制上下文，经典网页查岗、角色权限、内置软件锁定、每日使用限额及既有应用均保留。私人 App 继续通过原生桥 23 使用获授权的真实设备能力。",
            "当前版本为网页 v967、私人 iOS 1.0.89 (89)、原生桥 23。PhoneWeb.bundle 已从共享源码重建，Windows 全量 Node 回归 686/686 通过；仍须在 Mac 编译签名，并在真实 iPhone 验证模型返回封装、四套主题第二页卡片，以及网页／私人 App 数据边界。",
        ],
    ),
}


for stem, (heading, paragraphs) in records.items():
    append_record(stem, heading, paragraphs)
