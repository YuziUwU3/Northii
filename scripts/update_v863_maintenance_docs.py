from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


SECTIONS = {
    "AI开发项目_Bug记录模板.docx": (
        "v863 Bug 记录｜黄色快照冲突、假回执与后台通知（2026-08-10）",
        [
            "现象：伴生页大量外置 App 卡片变黄，显示“锁定回执与回执后的 iPhone 快照不一致”；反复点同步仍然变黄。点击锁定还会收到一条可见通知，但真实 App 没有锁上。用户原始目标是更快、更灵敏、无需手动同步，不接受简单退回旧轮询。",
            "含义澄清：黄色冲突不是“快照过期”。它表示命令记录已经是 completed，但同一命令之后的 iPhone 状态仍为未锁定。手动同步只会再次上传未锁定，因此不会自动消失。真正的过期只指 pending 命令在规定时间内没有收到设备回执。页面从 v863 起明确写“锁定未生效……不是快照过期”。",
            "根因一：phone-companion-push 把后台唤醒做成 alert 推送，包含标题、声音和角标，所以用户会看到一条像成功提示的通知；但 APNs 只代表唤醒尝试，不代表 iPhone 已执行锁定。根因二：原生端写 ManagedSettings 后立即把命令回执为成功，随后才单独上传快照，没有从系统设置读回验证。根因三：完整同步与后台快车道可能并发，后台入口在命令通道忙时返回 0，却仍上传旧控制快照；服务端也没有快照顺序保护，旧快照可覆盖新快照。",
            "修复：APNs 改为仅 content-available 的静默 background 推送，priority=5，不再携带 alert、sound、badge，并用 collapse id 合并重复唤醒。原生完整同步先处理命令，再读取耗时的 DeviceActivity 数据；所有命令入口串行等待，忙超时则明确失败，禁止上传假快照。锁定/解锁前检查 Screen Time 控制权限，写入 ManagedSettingsStore 后短暂等待并从同一 store 读回，只有真实配置与目标一致才算成功。快照锁定状态直接读取 ManagedSettingsStore，不再把 UserDefaults 镜像当成系统真相。",
            "服务端一致性：新增 snapshotSequence 单调序号；phone_companion_push_snapshot 只接受比现有序号更新的 v863 快照。为安装过渡，旧客户端在服务端尚未收到任何带序号快照前仍可上传；v863 首次上传后，旧快照永久不能覆盖新版。新增 phone_companion_complete_command，在同一数据库事务中锁定 pending 命令、保存已验证的控制快照、更新 last_sync_at 并完成命令回执，任何一步失败都不产生 completed 假回执。",
            "验证：app.js 语法检查通过；伴生专项 38/38 通过；完整 Node 回归 381/381 通过；git diff --check 通过。Windows 无法完成 Xcode 编译和真机 ManagedSettings 生效验证，发布后仍必须在 Mac Xcode 替换 v863 两个 Swift 文件、安装原 iPhone，并实测后台锁定、两分钟保持、解锁和失败提示。",
            "部署边界：GitHub main 中的 Edge Function 与 SQL migration 只是源码。只有在 Supabase 实际执行 202608100001_phone_companion_atomic_snapshot.sql，并实际部署 phone-companion-push 后，服务端修复才生效；二者必须独立核验，不能与 GitHub 推送混为一谈。",
        ],
    ),
    "AI开发项目_Bug修改规范.docx": (
        "新增强制规范｜伴生命令必须以系统读回和原子快照为准（v863 起）",
        [
            "APNs、网页排队成功、Edge Function 返回 pushed=true 都只能证明“已请求唤醒”，不得在界面或维护记录中表述为设备执行成功。静默后台唤醒不得夹带 alert、sound 或 badge 来制造成功感。",
            "锁定或解锁 ManagedSettings 后必须从同一个 ManagedSettingsStore 读回目标 token。只有读回状态与命令目标一致，才能把命令标为 completed；授权缺失、token 不存在、读回不一致或网络提交失败都不得生成成功回执。",
            "成功命令的 completed 状态、acknowledged_at、post-command snapshot 和 last_sync_at 必须在同一数据库事务中提交。禁止先成功回执、再异步补快照；禁止用网页收到命令或本地 UserDefaults 镜像代替系统状态。",
            "所有设备快照必须携带单调 snapshotSequence。服务端不得允许旧序号覆盖新序号。完整同步、前台同步和 APNs 快车道必须共享串行命令通道；若通道仍忙，应等待或明确失败，禁止假装没有命令并上传过时控制快照。",
            "产品文案必须区分三种状态：pending 超时才叫“命令过期”；动态数据较早叫“快照较早”；completed 回执与设备状态矛盾叫“锁定/解锁未生效”。使用时长百分比不得伪造最小 2%，0 使用应显示 0%。",
            "伴生修复至少验证：静默 APNs 头和 payload、命令先于重型数据刷新、并发命令不上传旧快照、ManagedSettings 读回、原子 RPC、单调序号、冲突文案、旧客户端过渡，以及完整项目回归。Windows 静态测试不能替代 Mac Xcode 编译和真实 iPhone 验收。",
        ],
    ),
    "AI开发项目_新聊天启动说明.docx": (
        "新聊天接手状态｜v863 伴生静默快锁与原子快照（2026-08-10）",
        [
            "v863 以已发布的 v862 main 提交 2c81ce07543b31fada7f219b08d6504b18eed340 为开工基线。唯一仓库仍为 C:\\Users\\pc\\Documents\\小手机\\phone-work，唯一分支 main，远端 origin/main；禁止创建分支、worktree、强推或使用旧目录。",
            "本次修复的是伴生真实锁定闭环，不是简单回退。黄色卡片不是快照过期，而是成功回执与回执后 iPhone 状态矛盾。v863 改为静默 APNs、命令优先、ManagedSettings 系统读回、原子完成命令与快照、单调 snapshotSequence，并把冲突文案改成“锁定未生效……不是快照过期”。",
            "数据库新增 supabase/migrations/202608100001_phone_companion_atomic_snapshot.sql；Edge Function 源码修改为 background/priority 5 的静默推送。务必分别核验：源码是否进入 GitHub main、SQL 是否实际在 Supabase 执行、phone-companion-push 是否实际重新部署。缺少后两项时网页版本虽可变为 v863，服务端锁定闭环仍未上线。",
            "原生端必须用 v863 包内 CompanionSyncView.swift 和 PhoneCompanionTestApp.swift 替换 Mac Xcode 主 App Target 的同名文件并安装到原 iPhone。旧 App 在首个 v863 序号快照到来前仍可过渡同步；v863 首次成功同步后，旧快照不能再覆盖新版。不要从多任务界面强制划掉 App，否则 iOS 通常不会交付后台唤醒。",
            "Windows 自动验证：伴生专项 38/38、完整 Node 回归 381/381。尚未完成的唯一关键验收是 Mac Xcode 编译和真实 iPhone：后台锁定应无可见通知、数秒内真实锁上、页面收到原子确认；保持两分钟后状态不丢；再后台解锁；权限被关闭时必须显示失败而不是 completed。",
        ],
    ),
    "AI开发项目_项目说明文档.docx": (
        "v863｜伴生静默快锁、系统读回与原子快照（2026-08-10）",
        [
            "v863 延续 v862 全部功能，专门修复真实 iPhone 外置 App 锁定链路。目标仍是用户最初要求：不必退出伴生 App或手点刷新，后台可被静默唤醒，并尽快执行真实锁定/解锁；没有退回旧的慢轮询方案。",
            "新的成功定义：网页排队和 APNs 推送都不是成功；原生 App 必须检查 Screen Time 控制授权，把目标 token 写入 ManagedSettingsStore，再从同一 store 读回确认。成功命令与读回后的控制快照由 phone_companion_complete_command 在一个事务中提交，网页不会先看到 completed、后看到相反快照。",
            "快照模型从 v863 起增加 snapshotSequence。服务器只允许更新序号前进，防止完整同步、后台唤醒和重试之间的旧结果倒灌。动态时长、位置、健康数据可以随时间变旧，但已经由真实 post-command snapshot 确认的锁定事实不会因为动态数据较早而自动过期。",
            "后台通知改为 APNs background push（content-available=1、priority=5），不显示横幅、声音或角标。Apple 后台推送属于尽力交付，系统仍可能节流；用户强制划掉 App 后通常停止后台交付。因此服务端继续保留 pending 队列，下次系统唤醒或打开 App 时幂等补执行。",
            "页面状态定义：蓝/红等正常状态来自真实设备快照；黄色 conflict 表示锁定或解锁未生效，不是过期。pending 超过命令期限才叫过期。使用占比允许显示 0%，不再为了进度条可见而伪造 2%。",
            "v863 Windows 验证为伴生专项 38/38、完整 Node 回归 381/381。服务端上线需要独立执行 SQL migration 和部署 Edge Function；原生上线需要 Mac Xcode 编译安装。只有 GitHub main、Supabase、iPhone 三层都完成并真机锁定/解锁通过，才可称为完整修复。",
        ],
    ),
}


def append_section(path: Path, heading: str, paragraphs: list[str]) -> None:
    document = Document(path)
    if any(p.text.strip() == heading for p in document.paragraphs):
        return
    document.add_page_break()
    document.add_heading(heading, level=1)
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(path)


def main() -> None:
    for name, (heading, paragraphs) in SECTIONS.items():
        path = DOCS / name
        append_section(path, heading, paragraphs)
        document = Document(path)
        assert any(
            paragraph.text.strip() == heading
            for paragraph in document.paragraphs
        ), f"missing v863 heading in {name}"
        assert len(document.sections) >= 1
        print(f"verified {name}: {len(document.paragraphs)} paragraphs")


if __name__ == "__main__":
    main()
