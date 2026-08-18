from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, size=10.5):
    run.font.name = "等线"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(size)


def append_docx(stem, heading, paragraphs):
    path = DOCS / f"{stem}.docx"
    document = Document(path)
    if not any(p.text.strip() == heading for p in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, 16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(path)
    with ZipFile(path) as archive:
        assert archive.testzip() is None


def append_txt(stem, heading, paragraphs):
    path = DOCS / f"{stem}.txt"
    content = path.read_text(encoding="utf-8")
    if heading not in content:
        path.write_text(
            content + "\n\n" + heading + "\n" + "\n".join(paragraphs) + "\n",
            encoding="utf-8",
        )


records = {
    "AI开发项目_Bug记录模板": (
        "v926／私人版 1.0.52 Bug 记录｜首屏重叠、旧通知跳页、线下约会与 v910 屏保恢复（2026-08-14）",
        [
            "现象一：玻璃主屏第一页偶尔严重卡顿、图标消失或无法点击。根因不是设备性能，而是布局修复把缺失应用全部塞进只有八个绝对定位槽位的第一页；第九个及后续应用没有坐标，透明层互相重叠并拦截触摸。修复为第一页严格保留八个槽位，溢出和新发现应用进入第二页；窄屏组件宽度改以主屏容器计算，不再使用可能与内嵌 WebView 不一致的 100vw。",
            "现象二：打开几分钟后可能无操作跳进角色微信。根因不是久未打开规则，而是原生端把通知路由持久保存，后续任意一次 WKWebView didFinish 都会无条件消费旧路由。修复为原生和网页双层校验 source=notificationTap、唯一 nonce 与两分钟新鲜度，并且每个 nonce 只执行一次；过期路由直接清除。",
            "现象三：线下约会或共同生活点击后回到主屏。根因是历史存档中的 msgs、memory、history 等字段可能缺失或类型错误，弹层关闭后 renderOff 抛错，视觉上像被送回主页。修复为每次进入前统一规范化旧存档，并把角色入口改成语义化 button 的单一 click 路径。",
            "现象四：锁定应用、自定义上传图标会变宽，名称可能被截断。根因是玻璃参考页后置 CSS 的 width:auto 覆盖固定槽位，图标又允许 flex-shrink；上传图片也没有统一正方形裁切。修复为应用外层固定槽位、图标禁止收缩且固定正方形，自定义图片在保存前中心裁切为 256×256，名称使用固定高度行盒。",
            "屏保决定：不再继续修改玻璃版屏保。移除 glass-theme.css 中全部 lock* 覆盖，屏保结构、时间、日期、星期、通知、箭头和解锁层完整沿用 v910 原版；玻璃主题只作用于解锁后的主屏。",
            "验证：专项自动化 28 项通过，实际 390×844 页面中第一页八个图标均为 66×66，锁定与自定义图标不改变外层几何；旧存档可进入线下约会；屏保计算样式恢复 v910 的无边框、无玻璃卡片布局。完整 Windows 回归与原生资源清单在发版前再次执行；Mac Xcode 编译、签名和真实 iPhone／Android 触摸验收仍需安装后完成。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "新增强制规范｜固定槽位、持久通知路由与屏保版本隔离（v926 起）",
        [
            "绝对定位的主屏参考页必须先定义槽位容量。布局修复、数据迁移和新增应用不得把超过容量的项目继续写入第一页；溢出项必须进入下一页，任何没有坐标类的绝对定位应用都视为阻断点击的高风险缺陷。",
            "原生持久化导航不得在普通页面加载完成时无条件执行。通知跳转必须同时验证真实点击来源、唯一 nonce、短时新鲜度和单次消费；后台消息生成只能写消息与通知，不能主动切换用户当前页面。",
            "历史存档入口必须在关闭当前页面或弹层前完成数据形状校验。数组、布尔值和枚举字段需要在每次进入时兼容旧值，禁止先关闭界面再让渲染异常把用户留在主页。",
            "应用槽位、图标容器、内部图片和名称行盒必须分层定尺寸。锁标、自定义图片和主题资源只能改变内容，不得改变外层宽高；自定义图片统一中心正方形裁切，图标在 flex 布局中禁止收缩。",
            "屏保与解锁后主题必须版本隔离。恢复指定旧版屏保时，应删除新主题对全部 lock* 选择器的覆盖，而不是模拟外观；以后玻璃主题不得重新覆盖 v910 屏保布局。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "v926／私人版 1.0.52｜首屏稳定、线下约会与 v910 屏保恢复（2026-08-14）",
        [
            "当前正式版本：网页版 v926；私人 iOS 1.0.52 (52)；原生桥契约 18。网页版与私人 PhoneWeb.bundle 使用同一份 v926 业务源码。",
            "玻璃主屏保持黑、粉白、蓝白、灰白四套主题和既定参考排版。第一页固定八个应用槽，新增或溢出应用进入第二页；锁定、自定义图标和不同屏宽不再改变应用正方形尺寸。自定义图标保存前会中心裁切，AI 账户四套资源继续使用人头与脑内网络球图案并提升缓存版本。",
            "线下约会／共同生活入口兼容旧存档缺字段；原生角色通知只有用户刚刚真实点击时才允许打开聊天，普通刷新、重新进入前台和久未打开消息均不导航。",
            "屏保已完整恢复 v910 原版，玻璃主题不再影响屏保时间、年月日、星期、通知、顶部箭头或解锁层。此轮没有修改电话／视频声音、字幕动画、语音识别、屏幕共享和微信聊天生成逻辑。",
        ],
    ),
    "AI开发项目_新聊天启动说明": (
        "新聊天接手状态｜v926／私人版 1.0.52（2026-08-14）",
        [
            "当前基线：网页版 v926、私人 iOS 1.0.52 (52)、原生桥契约 18。屏保必须保持 v910 原版；不得再让玻璃主题覆盖任何 lock* 样式。",
            "玻璃第一页只有八个固定应用槽，超过八个必须进入第二页。若再次出现第一页卡顿、图标消失或无法点击，先检查是否有无槽位的绝对定位应用和透明重叠层，不要先改滚动动画或缩小图标。",
            "角色聊天只能由用户主动点击微信、角色或刚刚的真实系统通知进入。原生 pendingRolePushRoute 必须有 notificationTap 来源、nonce、两分钟新鲜度并单次消费；普通 didFinish 绝不能执行旧路由。",
            "线下约会入口每次都必须通过 offData 规范化旧存档。自定义应用图标固定中心裁切为正方形，锁标与名称不得改变应用外层几何。后续修改仍须避开电话声音、字幕、识别和屏幕共享，除非用户重新明确要求。",
        ],
    ),
}


for stem, (heading, paragraphs) in records.items():
    append_docx(stem, heading, paragraphs)
    append_txt(stem, heading, paragraphs)
