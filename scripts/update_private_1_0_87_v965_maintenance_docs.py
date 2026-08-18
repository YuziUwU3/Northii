from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, size=10.5):
    run.font.name = "等线"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(size)


def append_record(stem, heading, paragraphs):
    docx_path = DOCS / f"{stem}.docx"
    document = Document(docx_path)
    if not any(paragraph.text.strip() == heading for paragraph in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, 16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(docx_path)
    with ZipFile(docx_path) as archive:
        assert archive.testzip() is None

    txt_path = DOCS / f"{stem}.txt"
    text = txt_path.read_text(encoding="utf-8")
    if heading not in text:
        txt_path.write_text(
            text + "\n\n" + heading + "\n" + "\n".join(paragraphs) + "\n",
            encoding="utf-8",
        )


records = {
    "AI开发项目_Bug记录模板": (
        "v965 角色自主、外部共享音频、首页合成与上下班主动通知记录（2026-08-17）",
        [
            "现象与决定一：自然模式融合后，极端依恋、数值心情、上下位、知识滑块、关系意图和随机当前活动仍会在多条提示链中替角色决定情绪与行为，压过基础人设、世界书、真实事件和记忆。处理为统一优先级：基础人设与说话习惯、世界书真实规则、当前真实事件与相关记忆、角色自主判断与自然表达、可用功能。旧控制器迁移删除并在运行时停用；挂断、拒接、失联、共同经历和实际功能动作仍作为真实事实进入上下文。电话、礼物、约会、游戏、音乐、放映室等能力继续向角色说明，由角色自行决定何时主动使用。",
            "现象与处理二：私人 App 共享小手机外部屏幕并播放抖音等媒体时，角色每次发声会暂停视频。根因是角色音频播放前沿用普通通话的识别暂停、音频会话停用和重建流程，外部媒体与免提识别共用的 AVAudioSession 被切断。修复沿用 v964 放映室成功路径：外部屏幕共享被标记为 shared media，角色语音复用已经激活的音频会话，不执行普通通话暂停和停用；未共享媒体的普通通话保留旧保护。",
            "现象与处理三：私人 App 首页滑动后偶发大型组件内容消失成黑色矩形，随后所有应用点不动；手机发热时更容易出现。根因收窄为 WKWebView 在大面积 backdrop-filter、连续滚动及内存／热压力叠加时丢失合成内容层并阻塞命中测试。只对私人 App 首页大型玻璃组件、Dock 和第二页大玻璃卡使用低风险合成路径，关闭实时 backdrop-filter；主题颜色、透明背景、尺寸、布局、图标、拖动以及网页／Android 原路径不改。热量是放大因素，不是唯一根因。",
            "现象与处理四：删除随机当前活动后，角色仍需要知道自己真实去上班或下班并主动通知。前端从角色已启用的作息提取工作地、上班和下班时间、工作日及请假日期，写入独立 roleSchedule；云端在时间边界生成每日一次的真实上班／下班事件。角色按人设、世界书、记忆和关系自然通知，可自主决定是否附带位置、电话或邀请；周末和请假跳过，不编造同事、会议、加班，也不恢复随机活动控制器。",
            "字幕边界：受保护外部平台不会向小手机暴露完整字幕轨。私人屏幕共享只把当前帧中已经清晰可见的烧录字幕交给视觉模型，禁止猜测尚未出现、已经消失或被遮挡的台词；这不是完整字幕文件提取。Windows 全量 Node 回归 679/679 通过，PhoneWeb.bundle 已由共享清单重建并与 v965 对齐。没有在 Mac 编译、签名，也未冒充真实 iPhone 连续播放、首页压力测试、字幕识别或后台通知通过。版本为网页 v965、私人 iOS 1.0.87 (87)、原生桥 23。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "发布补充规范｜角色自主不能丢失事实，媒体会话与大型玻璃必须分层处理（1.0.87 起）",
        [
            "角色自然化不能等同于删除功能知识和事件事实。基础人设、世界书、当前真实事件与记忆必须先于功能提示；系统只说明角色拥有哪些能力和边界，不替角色决定情绪、关系姿态和使用时机。挂断、拒接、失联、通话、礼物、约会、共同观看等实际动作必须继续写入上下文，否则功能会变成角色不知道自己做过什么的空壳。",
            "角色语音与外部媒体并存时，必须区分普通通话和 shared media。普通通话可在角色发声前暂停／重建识别；放映室或系统屏幕共享已承载外部媒体时，不得停用或重设当前 AVAudioSession。每次修改必须分别回归普通通话、放映室、微信外部屏幕共享和后台原生播放。",
            "WKWebView 出现组件占位仍在但内容层消失、滑动后全页失去点击时，应先检查大面积 backdrop-filter 的合成层，而不是直接归因于手机发热或数据丢失。修复应优先限定到私人原生 App 和大型容器，保留主题、布局、图标与网页／Android 路径，并对拖动和点击命中做回归。",
            "上班、下班等主动事件必须来自用户已配置的真实作息，并独立于随机当前活动。服务端要按账号时区、工作日、请假和每日幂等键判断；提示只提供真实事件，不编造未配置的同事、会议、加班或路线。Windows 测试与函数部署均不能写成真实 iPhone APNs 已验收。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "当前交付基线｜网页 v965／私人 iOS 1.0.87 (87)／原生桥 23（2026-08-17）",
        [
            "角色生成链现在按基础人设与说话习惯、世界书真实设定与规则、当前真实事件和相关记忆、角色自主判断与自然表达、可用功能的顺序组织。极端依恋、数值心情、上下位、知识滑块和随机当前活动不再替角色做决定；挂断拒接失联、共同经历和所有实际执行动作继续被角色感知。角色可以自主发起电话、礼物、约会、游戏、音乐和共同观看，不要求用户先提出。",
            "私人 App 的微信外部屏幕共享复用放映室已经验证的媒体音频会话保护，角色发声不再主动暂停抖音等外部媒体；普通通话仍保留原识别暂停与重建。受保护平台只支持逐帧识别当前可见画面和烧录字幕，不提供完整字幕轨抓取或版权绕过。",
            "私人 App 首页的大面积玻璃组件改用低风险合成路径，降低滑动、发热或内存压力下内容层消失并造成全页点不动的概率，不改变主题、布局、图标、拖动或网页／Android 行为。角色真实作息同步为独立云端计划：工作日上班和下班边界由角色按人设、世界书和记忆自然通知，周末与请假不触发。",
            "当前版本为网页 v965、私人 iOS 1.0.87 (87)、原生桥 23。私人 PhoneWeb.bundle 已从唯一共享源码重建，Windows 全量自动回归 679/679 通过。仍必须在 Mac 编译签名，并在真实 iPhone 验证外部媒体连续播放、首页高温／滚动稳定、当前帧字幕识别和后台上班／下班 APNs。",
        ],
    ),
}


for stem, (heading, paragraphs) in records.items():
    append_record(stem, heading, paragraphs)
