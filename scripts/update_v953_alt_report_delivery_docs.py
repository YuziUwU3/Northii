from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, size=10.5):
    run.font.name = "等线"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(size)


def append_docx(stem, heading, paragraphs):
    path = DOCS / f"{stem}.docx"
    document = Document(path)
    if not any(p.text.strip() == heading for p in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, 16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(path)
    with ZipFile(path) as archive:
        assert archive.testzip() is None


def append_txt(stem, heading, paragraphs):
    path = DOCS / f"{stem}.txt"
    content = path.read_text(encoding="utf-8")
    if heading not in content:
        path.write_text(
            content + "\n\n" + heading + "\n" + "\n".join(paragraphs) + "\n",
            encoding="utf-8",
        )


records = {
    "AI开发项目_Bug记录模板": (
        "v953 补充 Bug 记录｜小号报备被旧话题误判为成功（2026-08-16）",
        [
            "现象：小号与角色发生新聊天后切回大号，角色可能继续回答大号切换前的旧话题，没有主动说明小号联系人及聊天内容；之后即使询问，角色也可能否认或不记得。关闭旧自然模式时曾能报备，统一自然系统启用后更容易暴露此问题。",
            "根因：账号切换和报备事件实际已经触发，但异步报备队列只依据本轮是否产生任意可见角色消息判断成功。模型带着大号旧历史生成普通续聊时，回调仍推进 _altReportAt、_altReportDeliveredAt 和事件 reportedAt，导致没有完成的报备被提前消费。旧修复还把 _altReportDeliveredAt 本身当作交付证据，无法自动修复已经误消费的事件。",
            "修复：每个报备事件加入稳定时间戳，并对模型结果执行语义交付校验。具名小号必须同时出现该联系人名称和加好友、联系、发消息或聊天动作；延续旧话题、否认、不记得及泛泛的“有人联系我”均不算成功。首次结果不合格时使用隔离旧历史的报备纠正请求；仍不合格或接口失败时，使用仅依据事件快照的事实型本地兜底。只有带事件标记的真实可见消息才推进游标。",
            "旧数据修复：历史 _altReportAt 或 _altReportDeliveredAt 不能再单独证明已报备。若大号消息中找不到对应事件标记或有效报备语义，就自动清除错误交付游标并重新排队，因此已被旧版本误消费的小号活动也能恢复一次报备。",
            "边界：小号仍是独立陌生联系人，只读取小号账号资料、聊天、摘要和记忆；大号只接收单向报备，不向小号泄露大号身份，也不猜测两个账号是同一人。同一事件只报一次，小号有新活动才生成新事件。本次仅修改网页核心，私人 App 安装包未改。",
            "验证：账号隔离、自然系统、首句生成、好友恢复及完整项目回归通过；完整结果为 200 个测试文件、657 项通过、0 失败。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "新增强制规范｜跨账号报备必须验证可见语义后再推进游标（v953 起）",
        [
            "报备队列的“接口调用成功”“生成了角色消息”与“报备交付成功”必须分开。不得因为本轮出现任意角色气泡就写入 reportedAt、_altReportAt 或 _altReportDeliveredAt；必须先验证最终可见内容确实说明了指定小号联系人及其联系行为。",
            "每个待报备事件必须拥有稳定事件标记，并把标记写入最终可见的报备消息。恢复旧数据时只能以该标记或严格的具名报备语义作为证据；历史游标和 deliveredAt 不得自证成功。没有证据时必须恢复待处理状态并重新排队。",
            "模型延续大号旧话题、否认联系、不记得、发服务型助手问候或只说泛泛的“有人联系我”，均视为报备失败。纠正请求必须隔离大号旧历史，只携带角色稳定设定和本次报备事件；纠正仍失败时允许使用事件快照生成简短事实型兜底，但不得编造事件之外的态度、关系或内容。",
            "跨账号回归测试至少覆盖：旧话题续聊不得通过；不具名泛化不得通过；具名联系人加联系动作可以通过；失败不推进游标；成功只推进一次；新活动可再次报备；旧错误游标可修复；小号继续保持独立陌生人边界。",
        ],
    ),
    "AI开发项目_项目说明文档": (
        "v953 补充｜小号切回大号报备交付校验（2026-08-16）",
        [
            "微信统一自然系统保留大小号原有功能：小号作为独立陌生联系人生成，角色不得读取大号关系和共同记忆；切回大号后，角色对每段新增小号交流主动报备一次。",
            "报备现在使用持久事件队列和可见交付校验。普通续聊、否认或不具名的泛化内容不会消费事件；系统会隔离旧话题重试，并在模型仍未完成时使用当前事件快照生成简短事实报备。已被旧版本错误标记为完成但没有真实报备消息的事件会自动修复并重排。",
            "本次修复范围仅为网页 app.js 和回归测试，未修改私人 iOS App、原生桥、主屏布局、系统安全区或受保护未提交现场。",
        ],
    ),
}


for stem, (heading, paragraphs) in records.items():
    append_docx(stem, heading, paragraphs)
    append_txt(stem, heading, paragraphs)
