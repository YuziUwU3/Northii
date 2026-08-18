from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, name="等线", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def append_section(name, heading, paragraphs):
    path = DOCS / name
    document = Document(path)
    if any(p.text.strip() == heading for p in document.paragraphs):
        return
    document.add_page_break()
    title = document.add_heading(heading, level=1)
    for run in title.runs:
        set_font(run, size=16)
    for text in paragraphs:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.25
        set_font(paragraph.add_run(text))
    document.save(path)


SECTIONS = {
    "AI开发项目_Bug修改规范.docx": (
        "新增强制规范｜网页能力迁入私人 App 必须逐项原生化验收（私人 1.0.5 起）",
        [
            "WKWebView 不是 Safari。凡是麦克风、语音识别、键盘、通知角标、文件、长期存储和后台能力，必须建立原生权限、原生桥、失败状态和真机回归；禁止因为网页版可用就推断私人 App 可用。",
            "键盘避让不得把原生键盘高度和 visualViewport 已经缩小的高度重复相加。原生端提供键盘遮挡量，网页端先扣除浏览器已经覆盖的可视区域，只补足缺口；通话输入栏和音乐聊天栏必须逐页真机验证。",
            "推送 payload 中的固定 badge=1 不是未读计数。私人 App 在没有真实未读账本前必须由通知扩展丢弃该占位角标，并在启动、回到前台和前台收通知时清零，不能让假红点永久存在。",
            "设备管理入口属于私人 App 原生能力，只能放在设置页固定位置；禁止作为覆盖所有网页页面的悬浮层，以免遮挡通话、音乐、购物等独立 App。",
            "Windows 源码断言只能证明桥接和权限配置存在，不能证明 iOS 录音、转写、键盘位置或角标已真实生效。每次安装候选包都必须在真机重新授权并逐项验收后才能写成完成。",
        ],
    ),
    "AI开发项目_Bug记录模板.docx": (
        "私人 1.0.5 候选 Bug 记录｜假红点、话筒失效与键盘遮住输入栏（2026-08-11）",
        [
            "现象：私人小手机安装后长期显示红色角标 1；网页原本可用的话筒转字幕在 App 内不可用；通话输入框被键盘推得过高，音乐页弹出键盘后看不到聊天输入框；原生设备管理悬浮按钮还会覆盖每个内置页面。",
            "根因一：共享角色推送仍携带公开 North 使用的固定 badge=1，私人 App 没有未读计数账本，也没有在生命周期内清角标。根因二：WKWebView 没有稳定提供 Safari 的 Web Speech 能力，工程同时缺少语音识别说明、原生 Speech 桥和媒体权限代理。",
            "根因三：输入布局仍按浏览器假设处理，没有读取原生键盘 frame；初版候选若直接把整段键盘高度加到 CSS bottom，还会与 visualViewport 已经完成的避让重复计算，把输入栏再次推走。根因四：设备管理按钮在 SwiftUI 根视图全局叠加，天然会出现在所有网页页面之上。",
            "修复候选：私人通知扩展清除固定 badge，主 App 在启动、激活和前台通知时清零；加入 SFSpeechRecognizer、AVAudioEngine、麦克风/语音权限和版本 3 原生桥；键盘事件只把浏览器尚未避让的缺口写给通话和音乐输入栏；设备管理入口移到设置页右上角。",
            "隔离边界：所有新增网页分支都以 __SMALL_PHONE_PRIVATE__ 为条件；公开 North 保持 v881 和原推送行为，不改审核版本。私人版本提升为 1.0.5 (5)。",
            "当前验证：JS 语法、私人 App 基础、音乐键盘、音乐视觉和 iOS 视口专项通过。Windows 无法编译 Xcode 或代替真机，因此麦克风授权、转写结果、角标消失、通话与音乐输入栏位置仍必须在 iPhone 安装 1.0.5 (5) 后确认。",
        ],
    ),
    "AI开发项目_项目说明文档.docx": (
        "私人小手机 1.0.5 候选｜原生语音、键盘与角标生命周期（2026-08-11）",
        [
            "私人小手机继续采用一份共享网页业务核心和独立原生能力层。公开 North 保持 v881 不动；私人安装包单独提升到 1.0.5 (5)，新增能力只有检测到私人原生桥时才启用。",
            "私人 App 的话筒转字幕改为原生 Speech + AVAudioEngine，经版本 3 原生桥模拟现有 SpeechRecognition 接口，因此现有通话、持续话筒等业务仍复用同一套网页逻辑。拒绝授权或系统不可用必须返回真实错误。",
            "原生层监听键盘 frame，网页层结合 visualViewport 计算尚未避让的差值，分别固定通话输入栏和音乐聊天栏，避免输入框被键盘遮住或被重复推到屏幕中部。",
            "设备管理入口从覆盖全 App 的悬浮按钮迁到设置页右上角。角色推送的固定 badge 占位值在私人通知扩展中被丢弃，App 启动、回前台和前台收通知时清理旧角标。",
            "手机号账号、Keychain、云端恢复、完整大容量原生媒体存储、服务端控制器租约以及永久锁定/同步长稳测试仍是后续任务；本候选包不把这些项目误报为已完成。",
        ],
    ),
    "AI开发项目_新聊天启动说明.docx": (
        "新聊天接手状态｜私人小手机 1.0.5 原生输入候选（2026-08-11）",
        [
            "唯一仓库为 C:\\Users\\pc\\Documents\\小手机\\phone-work，分支 main。公开 North 仍为 v881 且审核版保持不动；私人小手机候选版本为 1.0.5 (5)。",
            "本候选处理私人 App 的永久假红点、WKWebView 话筒转字幕失效、通话/音乐键盘遮挡和全局设备管理悬浮按钮。核心实现位于 PhoneNativeBridge.swift、LocalPhoneWebView.swift、PhoneCompanionTestApp.swift、RoleNotificationService/NotificationService.swift 及共享 app.js/小手机.html 的私人条件分支。",
            "真机安装后必须依次验证：系统首次询问麦克风和语音识别权限；通话转字幕可写入输入；音乐聊天键盘出现时输入栏仍可见可发送；回到桌面无永久红点；真实 iPhone 管理入口只在设置页右上角。",
            "若输入栏仍异常，先记录 iPhone 型号、系统版本、页面和键盘前后截图，检查 native keyboard payload 与 visualViewport 差值，禁止再次直接叠加整段键盘高度。",
            "后续主线不变：浏览器安全区逐页回归、完整原生持久存储、手机号/Keychain/云恢复、控制器租约隔离，以及锁定、同步、断网、后台、跨日和重装长稳验收。",
        ],
    ),
}


def main():
    for name, (heading, paragraphs) in SECTIONS.items():
        append_section(name, heading, paragraphs)
        document = Document(DOCS / name)
        assert any(p.text.strip() == heading for p in document.paragraphs)
        print(f"verified {name}: {len(document.paragraphs)} paragraphs")


if __name__ == "__main__":
    main()
