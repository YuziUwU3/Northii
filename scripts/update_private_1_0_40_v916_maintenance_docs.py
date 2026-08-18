from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, name="等线", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def append(name, heading, paragraphs):
    path = DOCS / name
    document = Document(path)
    if not any(p.text.strip() == heading for p in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, size=16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(path)
    with ZipFile(path) as archive:
        assert archive.testzip() is None


append("AI开发项目_Bug记录模板.docx", "v916／私人 1.0.40 Bug 记录｜固定扫描、作息失效与通话音频争抢（2026-08-13）", [
    "现象：实时共享理解仍按固定间隔反复取帧，角色容易连续说话；角色让用户切换页面后，用户说‘打开了’不会立即读新画面。共同生活作息修改后状态不按时间推进；宠物长期未喂会自动变瘦。播放抖音等外部媒体时，系统通话音频会暂停外部声音或丢失麦克风。中文主流歌曲在原 Audius 目录中经常搜不到。",
    "根因：v915 自主观察仍保留默认八秒重调度，模型没有合法选择时还会兜底成继续；语音完成时冻结的 ReplayKit 帧没有直接交给‘切好了’分支。共同生活状态只把作息写进提示词，没有钟点到状态的确定性同步器。宠物仍保留两天未喂的视觉缩放。AVAudioSession 的 playAndRecord 默认不可与其他 App 音频混合。音乐只查 Audius 公共目录，中文商业目录覆盖不足。",
    "修复：自主观察只有角色明确输出继续及下一次秒数时才再取帧；无合法选择直接结束，提问与等待切换时停取帧，‘打开了／切好了／给你看’优先消费说话当刻冻结帧。作息按工作日时段自动推进在家、上班与回家路上，手动共同生活状态保留到下一时段边界。宠物增长只进不退，只有显式退回奶团期可变小。原生识音和通话播放会话加入 mixWithOthers；字幕逐字渐显且原生最终提交延长到1.65秒。搜索并行接入 Apple 中国区官方试听与 Audius 完整公开曲目。",
    "边界：角色不能绕过用户的 ReplayKit 系统授权主动开启共享；自动观察仍受开关和每次视频次数限制，口头看画面不限次数。mixWithOthers 允许系统混音，但第三方 App 若自身检测录音后主动暂停或受 DRM／ReplayKit 限制，应用不能强制覆盖。Apple 目录只提供官方预览并附官方页面，不等于网易云、QQ 音乐或 Apple Music 会员完整播放授权。",
])

append("AI开发项目_Bug修改规范.docx", "新增强制规范｜角色驱动共享观察与确定性状态推进（v916 起）", [
    "实时共享理解不得由固定四秒、八秒或其他默认循环持续扫描。第一次画面只供角色做决定；只有角色明确选择‘继续’并给出下一次时间才可调度。模型没有输出合法决定时必须停止，不能兜底继续；选择提问或等待切换后必须停止取帧。",
    "用户在角色要求切页后说‘打开了／切好了／给你看’，必须优先消费该次语音最终结果绑定的 ReplayKit 冻结帧，不能等用户回到小手机后再抓前台自身画面。视觉提示必须忽略小手机自己的悬浮通话框，不能把角色自己的名字、字幕和头像认成第三方内容。",
    "时间感知开启且作息开关开启时，共同生活状态必须由真实钟点确定性推进；不能只把作息写进模型提示词。用户手动状态在当前作息槽内优先，直到下一时段边界才由作息接管。时间感知关闭后不得自动推进。宠物自然成长阶段只能单调增加，任何饥饿、患病和旧字段迁移都不得让体型或成长阶段自动倒退。",
])

append("AI开发项目_项目说明文档.docx", "v916／私人 1.0.40｜自主共享观察、作息推进与混合音频（2026-08-13）", [
    "当前版本：网页 v916；私人 iOS 1.0.40 (40)；原生桥契约 17。私人 PhoneWeb.bundle 已由共享网页核心重新生成。主题玻璃质感与图标放大仍是本地预览项，不在本次正式推送内。",
    "实时共享理解由角色控制：共享开始后取得第一帧，角色可静默继续并自选3到90秒后的下一次、针对在意内容只问一个问题并暂停、要求用户切换页面后等待、或结束本轮观察。未给合法决定不会继续扫描；口头‘你看一下’仍走不限次数的手动链路。视觉模型明确忽略小手机自身悬浮通话窗。",
    "共同生活作息已与确定性状态机相连，保存新作息会立即按当前时段同步状态；手动生活状态、活动和地点保留到下一作息槽。宠物不再因未喂自动变小，超过三天仍可生病并需要药物，只有‘退回奶团期’会缩小。视频通话字幕逐字渐显，原生语音等待更稳定文本后提交。",
    "原生麦克风和角色语音改为 AVAudioSession playAndRecord + mixWithOthers，减少外部视频／音乐被系统会话强制暂停；最终能否同时播放还取决于第三方 App 自身策略。音乐搜索同时查询 Apple 中国区官方试听与 Audius 公开完整曲目，可搜索中文主流歌名并跳转官方页；不保存或伪装成会员完整音源。",
])

append("AI开发项目_新聊天启动说明.docx", "新聊天接手状态｜v916／私人 1.0.40（2026-08-13）", [
    "当前正式基线：网页 v916、私人 iOS 1.0.40 (40)、原生桥契约 17、PhoneWeb.bundle v916。不要回退外部 App 后台帧独立槽、冻结令牌、序列号和一次性回执；不要恢复固定四秒／八秒自主扫描。",
    "自主共享状态为 observing／waiting／waiting-screen／answering／ended。只有角色合法选择继续才安排下一帧；提问和等待切换都停取帧。用户说‘打开了／切好了／给你看’时优先使用语音事件绑定的新帧。口头要求看画面不限次数；开关关闭后整个自主状态机失效。",
    "作息自动推进只在时间感知和角色作息均开启时运行；手动状态保留到下一作息槽。宠物成长只进不退，生病不缩小。原生音频已启用 mixWithOthers，但真机必须同时验证外部视频声音、用户麦克风和角色 TTS；第三方 App 自行暂停不应误判为小手机仍未配置混音。",
    "下一轮真机只做三件事：一，共享另一个 App，让角色自行决定继续／提问／等待切换／结束，确认没有固定循环抢话；二，角色要求切页后说‘打开了’，确认立即描述新页且忽略悬浮通话窗；三，播放外部视频并交替说话、听角色回复，同时搜索‘薛之谦 丑八怪’确认出现 Apple 官方试听。",
])
