from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def set_font(run, name="等线", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def append(name, heading, paragraphs):
    path = DOCS / name
    document = Document(path)
    if not any(p.text.strip() == heading for p in document.paragraphs):
        document.add_page_break()
        title = document.add_heading(heading, level=1)
        for run in title.runs:
            set_font(run, size=16)
        for text in paragraphs:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            set_font(paragraph.add_run(text))
        document.save(path)
    with ZipFile(path) as archive:
        assert archive.testzip() is None


append("AI开发项目_Bug记录模板.docx", "v913／私人 1.0.37 Bug 记录｜共享开始不知情与复读旧话（2026-08-13）", [
    "现象：用户主动打开屏幕共享后，角色有时不知道共享已经开始，并会继续回答或复读共享前的上一句话；展开通话框字幕切换生硬，样式比缩小悬浮框厚重。",
    "根因：共享开始事件和普通通话回复共用单个覆盖式待处理槽，旧模型请求完成后仍可播放；iOS 切换系统共享面板时可能错过一次 WKWebView 状态回调。共享事件还携带普通聊天历史，模型容易把最后一句当作本轮问题。",
    "修复：共享开始／结束使用独立高优先状态队列和代次门控；新状态到达立即停止旧音频、清空旧字幕并使旧生成作废。共享状态事件隔离普通聊天历史，禁止复读旧话和在没有画面描述时谎称看到具体内容。私人版主动轮询原生 screenShare.status，补回切后台时错过的状态事件。",
    "新增：默认关闭的实时共享理解测试开关。ReplayKit 帧序列约每 4 秒最多触发一次，网页将帧缩成 12×12 亮度签名，只有明显变化才识图和回应；静止画面不重复，自动次数共用单次视频上限，口头让角色看仍不限次数。原生帧令牌使用有限后台任务并在角色回复完成后释放。",
    "验证：Windows Node 全量回归 528/528 通过；JavaScript 语法、版本与六 Target 静态引用通过。Windows 不能完成 Xcode 编译与 iOS 后台调度证明，真实 iPhone 的长时间切其他 App、锁屏、发热、流量和接口响应时间仍需真机验收。",
])

append("AI开发项目_Bug修改规范.docx", "新增强制规范｜共享状态事件与准实时画面理解（v913 起）", [
    "通话中的媒体状态改变必须使用独立队列和生成代次，禁止与普通用户话语共用覆盖式 pending 槽。新共享状态到达后，旧生成、旧字幕和旧音频不得继续播放；状态事件不得携带上一轮普通聊天历史。",
    "共享开始只证明 ReplayKit 已处于活动状态，不证明模型已经理解某个具体画面。没有当轮真实画面描述时，角色只能确认共享已开始；不得编造软件、文字、物品或动作。",
    "准实时理解必须默认关闭、可随时关闭，并在关闭后恢复原分钟定时与口头识图链路。变化检测应在本地低分辨率完成，静止画面不得重复计费；自动调用严格遵守每次视频上限，口头明确要求查看不受该上限影响。",
    "iOS WKWebView 定时器在后台不可靠。任何承诺切其他 App 后继续的共享功能必须由 ReplayKit 原生帧序列提供触发，并对每轮网络识别申请有限后台任务、完成或失败时成对释放。仍不得承诺无限后台执行或逐帧视频模型效果。",
])

append("AI开发项目_项目说明文档.docx", "v913／私人 1.0.37｜共享状态可靠通知与实时共享理解（2026-08-13）", [
    "当前版本：网页 v913；私人 iOS 1.0.37 (37)；原生桥契约 16。共享业务核心仍只有一份，私人 App 内置 PhoneWeb.bundle 与网页版本原子同步。",
    "屏幕共享开始和结束现在是独立的通话状态事件。私人版除原生每半秒状态推送外，网页在用户点共享按钮后还会主动查询状态，避免系统面板切换期间丢事件。角色会先自然确认知道共享已开始，但没有真实帧描述时不会声称看到了具体内容。",
    "偏好设置新增默认关闭的实时共享理解测试开关。开启后仅在屏幕共享期间工作：约每 4 秒检查一轮，只有明显画面变化才调用视觉模型并让角色按人设回应。摄像头原分钟识别不变；共享实时模式开启时不再叠加分钟定时，避免双重识别。",
    "全屏通话字幕使用与原生 PiP 相同的 0.3 秒、8px 渐显上移动画。字幕框保留原文字宽度，缩短高度并使用轻透明模糊背景；PiP 宽度保持 360，高度从 202 缩至 176，名字和内容略向上。",
    "验证边界：528/528 Windows 自动测试通过。Mac 六 Target 编译、签名、Broadcast Upload 扩展、真实 iPhone 后台帧触发、语音持续、PiP 透明度及长时间耗电仍待实机验证。",
])

append("AI开发项目_新聊天启动说明.docx", "新聊天接手状态｜v913／私人 1.0.37（2026-08-13）", [
    "当前基线：网页 v913、私人 iOS 1.0.37 (37)、原生桥契约 16、Xcode 六 Target、PhoneWeb.bundle v913。Windows 全量自动化测试 528/528 已通过。",
    "不要回退共享状态独立队列、旧生成作废、原生状态补偿轮询、口头识图不限次数、0 关闭分钟识图和单次视频自动次数上限。实时共享理解默认关闭，只有用户主动开启才运行；关闭后恢复旧链路。",
    "实时共享理解不是逐帧视频模型：ReplayKit 持续提供系统画面，客户端约每 4 秒做变化检测，明显变化才识别。静止画面不重复；每轮有限后台任务必须在角色回复完成或失败后释放。",
    "下一步真机只做三个操作：一，Mac 编译安装第三十七次包；二，主动开启共享，确认角色知道共享开始且不复读上一句话；三，打开实时共享理解并切到其他 App，改变明显画面，确认数秒后回应、静止不刷屏、达到次数上限后停止自动理解。失败先记录 Xcode 和真机时间点，不盲改旧稳定链路。",
])
