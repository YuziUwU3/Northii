from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
EXPECTED = {
    "AI开发项目_项目说明文档.docx": [
        "v865｜主屏时间隐藏保持原布局",
        "visibility:hidden 和 pointer-events:none",
        "387/387 通过",
    ],
    "AI开发项目_Bug记录模板.docx": [
        "v865 Bug 记录｜关闭主屏时间后整体上移并露出底部黑块",
        "home-clock-hidden",
        "部分手机比例下",
    ],
    "AI开发项目_Bug修改规范.docx": [
        "视觉开关不得意外删除布局槽位（v865 起）",
        "禁止通过条件模板、display:none 或删除 DOM 节点实现",
        "页面底部背景",
    ],
    "AI开发项目_新聊天启动说明.docx": [
        "v865 主屏时间隐藏保持原布局",
        "f6d3a10db1277bd1fccf7dd39a0ac842a8add150",
        "完整 Node 回归 387/387",
    ],
}


for name, markers in EXPECTED.items():
    path = DOCS / name
    with ZipFile(path) as archive:
        assert archive.testzip() is None, f"corrupt zip member in {name}"
        assert "word/document.xml" in archive.namelist(), f"missing document.xml in {name}"
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for marker in markers:
        assert marker in text, f"missing marker {marker!r} in {name}"
    assert text.count(markers[0]) == 1, f"duplicate v865 section in {name}"
    print(f"verified {name}: {len(document.paragraphs)} paragraphs, {len(document.tables)} tables")
