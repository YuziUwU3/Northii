from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1] / "docs" / "maintenance"
MARKER = "2026-08-12 v896／私人 1.0.21 全量读取完成回执与后台表结构修复"

ENTRIES = {
    "AI开发项目_项目说明文档.docx": [
        "网页核心升级为 v896，私人 iOS 候选升级为 1.0.21 (21)，原生桥契约仍为 11。用户明确要求读取全部时，该范围成为本轮不可降级的事实，角色模型的提前猜测不能把任务缩成位置、电量或任一单项。",
        "原生 device.snapshot 在电量、Screen Time、HealthKit 和位置各读取器均已返回成功、不可用或超时后，才写入 readFinishedAt、readComplete=true 与 readOutcomes。网页端同时核对 readSessionId、requestedFocus=全部、完成时间新鲜度和 readComplete，拿不到同一轮完成回执就禁止角色说结果。",
        "完成回执通过后，角色使用确定性汇总逐项说明九类结果：电量、屏幕总时长、逐 App、步数、睡眠、心率、心电、HRV 和位置；可用项说真实值，不可用项说本轮真实失败原因，不再用模型自由生成的单项结论代替全量结果。",
        "后台核验 column controller_user_id does not exist 的真实原因不是按钮或账号，而是远端已存在依赖这些字段的 004 统一控制器函数，却漏装了更早的 002 表结构迁移。已实际执行 002，登记远端迁移历史，并核验 phone_companion_links 的 controller_user_id、controller_kind、controller_instance_id、controller_claimed_at 四列和两组控制器 RPC 均存在。",
        "Windows 全套 475 项测试通过，包含聊天、通话、原生完成回执、九类汇总和控制器表结构回归。Mac 五 Target 编译、签名以及真实 iPhone 的 HealthKit、DeviceActivityReport、后台／强退通知送达仍需最终验收。",
    ],
    "AI开发项目_Bug记录模板.docx": [
        "现象：用户说读取全部，顶部已经开始逐项读取，但角色在原生任务完成前先说没有读到，最后有时只报告位置；单独读取心率却能成功。后台资料页同时提示 column controller_user_id of relation phone_companion_links does not exist。",
        "v895 未成功的读取根因：普通 AI 回复先生成，guardUnverifiedRolePhoneReply 又先于用户读取意图处理。模型若提前只声称位置，守卫会把用户原本的全部请求降级为位置单项；即使进入全量分支，旧代码还把 capturedAt 新鲜当成读取成功，没有等待原生所有异步读取器结束。单项心率成功因此与全量提前开口并不矛盾。",
        "v895 未成功的后台根因：上轮只确认并部署了后续 003–006 和统一控制器函数，却没有核对其前置依赖 002 是否在线。函数可出现在 schema cache 中，但运行时一访问 phone_companion_links.controller_user_id 就失败；本地有 002 SQL 文件不能证明远端表已经增加字段。",
        "修复：聊天和通话先以用户原话识别本机真实读取意图，命中后立刻压制普通模型回复；原生端最后写完成回执，网页只接受同一轮全部请求的新鲜完成回执，再用固定九类汇总逐项回复。远端实际应用 002 并修复迁移历史，确认四个 controller 字段和新旧参数兼容 RPC 均存在。",
        "验证：node --test tests/*.test.mjs 共 475／475 通过；supabase migration list --linked 显示 202608110002–006 均为远端已应用。Windows 无法替代 Xcode 编译和真实 iPhone 权限、Report 扩展执行及 APNs 后台／强退送达验证。",
    ],
    "AI开发项目_Bug修改规范.docx": [
        "用户请求的读取范围必须在调用原生桥之前确定并贯穿同一 readSessionId。不得从角色模型生成的回复反推或缩小范围；模型输出只能在真实读取完成后负责语气，不能决定读取了哪些项目。",
        "异步全量读取的成功条件必须是显式完成回执，而不是请求已发出、capturedAt 更新、任一字段有值或进度横幅走完。完成回执只能在每个请求项已有成功／不可用／超时终态后写入，消费者必须核对 requestedFocus、readSessionId、readFinishedAt 和 readComplete。",
        "数据库函数存在不等于依赖表结构完整。部署或修复 RPC 时必须按迁移依赖顺序核对：先查 migration history，再查被引用表的真实列，最后执行带认证上下文的函数。不得只因后续函数已创建就把前置迁移视为已完成。",
        "全量读取验收应一次触发后检查九类逐项结果；不能要求用户逐个项目反复测试来掩盖调度错误。未产生数据、未授权和超时必须作为各自的真实结果保留，不能把缺失项静默省略。",
    ],
    "AI开发项目_新聊天启动说明.docx": [
        "当前候选基线：网页 v896、私人 iOS 1.0.21 (21)、原生桥契约 11。远端 202608110002–006 已核验为应用；002 补齐 phone_companion_links 四个 controller 字段，解决后台核验缺列错误。",
        "读取全部的关键不再是逐项手工测试：用户只需发一次读取全部，顶部依次显示各真实类别；角色必须等本轮 readComplete 完成回执后，再逐项报告电量、屏幕总时长、逐 App、步数、睡眠、心率、心电、HRV 和位置。",
        "若单项可读但全量只回一个项目，优先检查用户请求是否在普通模型守卫中被降级、requestedFocus 是否仍为全部、readSessionId 与 readFinishedAt 是否属于同一轮；不要再次只改提示词或延长横幅动画。",
        "后台出现 relation column does not exist 时，直接核对远端前置迁移和表列，不能只看统一函数是否存在。当前结构缺口已服务器端即时修复；客户端 1.0.21 仍用于全量完成回执修复。",
        "Windows 自动测试当前 475／475 通过。Mac 只从 SmallPhone_v896_AllReadCompletionPushSchemaFix 全新目录打开工程，编译五 Target，并在真机验证一次全量读取、HealthKit、Report 回传、前后台／强退通知、头像、主动来电和手动解锁告警。",
    ],
}


def main() -> None:
    for filename, items in ENTRIES.items():
        path = ROOT / filename
        document = Document(path)
        changed = not any(MARKER in paragraph.text for paragraph in document.paragraphs)
        if changed:
            document.add_heading(MARKER, level=1)
            for item in items:
                document.add_paragraph(item, style="List Bullet")
            document.save(path)
        with ZipFile(path) as archive:
            assert archive.testzip() is None
            assert "[Content_Types].xml" in archive.namelist()
        verified = Document(path)
        assert sum(MARKER in paragraph.text for paragraph in verified.paragraphs) == 1
        print(f"{'updated' if changed else 'unchanged'}: {path.name}")


if __name__ == "__main__":
    main()
