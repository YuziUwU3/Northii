from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
EXPECTED = {
    "AI开发项目_项目说明文档.docx": (
        "v870｜主辅模型统一 API 路线与就近保存（2026-08-10）",
        ["同时保存两组配置", "标题右上角", "旧数据迁移", "408/408"],
    ),
    "AI开发项目_Bug记录模板.docx": (
        "v870 Bug 记录｜API 路线只保存主模型且保存入口过远（2026-08-10）",
        ["S.settings.aux", "统一的捕获、应用、回填", "不能把旧路线缺失 aux", "408/408"],
    ),
    "AI开发项目_Bug修改规范.docx": (
        "新增强制规范｜一条 API 路线必须原子保存主辅模型（v870 起）",
        ["历史字段缺失", "S.settings.chat 与 S.settings.aux", "完整配置单元", "不得写入真实 API Key"],
    ),
    "AI开发项目_新聊天启动说明.docx": (
        "新聊天接手状态｜v870 主辅模型统一 API 路线（2026-08-10）",
        ["chatRouteApply", "aux:{base,key,model}", "两个标题按钮只保存各自一半", "408/408"],
    ),
}


for filename, (marker, required) in EXPECTED.items():
    document = Document(DOCS / filename)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    text = "\n".join(paragraphs)
    assert paragraphs.count(marker) == 1, f"duplicate or missing v870 section in {filename}"
    assert paragraphs[-1], f"blank ending in {filename}"
    for phrase in required:
        assert phrase in text, f"missing {phrase!r} in {filename}"
    print(filename, f"paragraphs={len(paragraphs)}", f"tables={len(document.tables)}")

print("v870 maintenance documents verified")
