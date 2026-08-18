from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
EXPECTED = {
    "AI开发项目_项目说明文档.docx": (
        "v869｜共同生活限额与每日必查（2026-08-10）",
        ["共同生活限额|准确App名", "睡眠与今日步数", "automationRuns", "408/408"],
    ),
    "AI开发项目_Bug记录模板.docx": (
        "v869 Bug 记录｜共同生活缺少限额与每日必查交接（2026-08-10）",
        ["没有共享的必查接管函数", "dailyKind/dailyDay", "心率仅保留", "五分钟后再试"],
    ),
    "AI开发项目_Bug修改规范.docx": (
        "新增强制规范｜跨通道每日必查必须显式交接（v869 起）",
        ["防重复变成漏执行", "睡眠与步数", "事实哈希和逻辑日", "1到720"],
    ),
    "AI开发项目_新聊天启动说明.docx": (
        "新聊天接手状态｜v869 共同生活限额与每日必查（2026-08-10）",
        ["cohabDailyRequiredMaybe", "scope:'both'", "rolePhoneInspectionAcquire", "Supabase Edge Function"],
    ),
}


for filename, (marker, required) in EXPECTED.items():
    document = Document(DOCS / filename)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    text = "\n".join(paragraphs)
    assert paragraphs.count(marker) == 1, f"duplicate or missing v869 section in {filename}"
    assert paragraphs[-1], f"blank ending in {filename}"
    for phrase in required:
        assert phrase in text, f"missing {phrase!r} in {filename}"
    print(filename, f"paragraphs={len(paragraphs)}", f"tables={len(document.tables)}")

route_notes = [
    "每次请求重新从主模型开始",
    "路线真正改变时提示三秒",
    "连续同路线不得提示",
    "_cohabActualModelRoute 只记录当前页面最后实际使用路线",
]
all_text = []
for path in DOCS.glob("*.docx"):
    all_text.append("\n".join(paragraph.text for paragraph in Document(path).paragraphs))
for note in route_notes:
    assert sum(note in text for text in all_text) == 1, f"duplicate or missing route note: {note}"

print("v869 maintenance documents verified")
