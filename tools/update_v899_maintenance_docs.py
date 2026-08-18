from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1] / "docs" / "maintenance"
MARKER = "2026-08-12 v899／私人 1.0.24：角色自定读取回复与线下错误诊断"

ENTRIES = {
    "AI开发项目_项目说明文档.docx": [
        "网页核心升级为 v899，私人 iOS 升级为 1.0.24 (24)，原生桥契约仍为 11。读取中的普通模型输出继续由全局读取状态和生成代次拦截。",
        "读取完成后的可见话术不再由程序固定汇总。程序只向角色模型提供本轮真实事实；角色自行决定重点、语气、句数以及一条或多条发送。安全校验失败时最多重新生成两次，仍失败则不发送。",
        "线下回应失败新增原因分类：余额或额度、401 密钥、403 权限或地区、404 接口或模型、429 限流或配额、408/504 网络超时、连接失败，以及接口成功但没有可显示角色内容。失败不改动原对话。",
        "Windows 全套 478 项自动测试全部通过。Mac 五 Target 编译、签名以及真机微信／电话／共同生活／线下失败提示仍需最终验收。",
    ],
    "AI开发项目_Bug记录模板.docx": [
        "现象：读取门禁正确拦住普通回复后，手动回复包装器却显示“模型没有返回可见消息，请再点一下”；读取完成路径有时又用固定逐项清单替代角色话术。线下模型失败只显示“暂时没有生成”，无法判断余额、网络还是接口问题。",
        "真实原因：replyGenerationRun 仅以本轮是否新增可见气泡判断成功，把读取门禁的故意静默误判成模型空输出，因此触发第二次请求和通用提示。另有两个完成路径把不合格模型输出替换成 rolePhoneInspectionExactSummary 固定模板。线下 catch 丢失了原始 HTTP、鉴权、限流和网络错误类别。",
        "前次方案为何不完整：v898 解决了读取与普通回复的竞态，但没有让手动回复包装器理解“读取代次变化即为正常接管”；同时为了保证数字齐全保留了固定模板兜底，因此虽然不再提前说话，仍可能出现突兀提示或系统式清单。",
        "修复：手动回复运行前记录读取代次，模型返回后若代次变化则静默结束，不重试、不提示；删除固定汇总函数，两条完成路径只允许模型生成并做最多两次安全修复；线下错误保留并分类输出真实原因。",
        "验证：node --test tests/*.test.mjs 共 478 项，478 项通过；新增线下错误分类、读取静默不重试、完成回复仅模型生成及角色自主分条测试。",
    ],
    "AI开发项目_Bug修改规范.docx": [
        "异步门禁主动取消输出时，外层重试器必须区分“正常接管”与“模型空输出”。应使用任务代次、取消令牌或明确状态，不能仅依据是否新增可见气泡自动重试。",
        "角色查看真实数据后的可见回复必须由角色模型生成。程序可提供事实、安全校验和有限重试，但不得用固定清单或固定人格话术替换模型输出。",
        "错误提示必须保留可行动的原始类别。余额／配额、鉴权、权限、限流、接口或模型配置、网络超时、连接失败、空内容不得全部压缩成同一句“暂时没有生成”。",
        "任何失败提示都不得修改、删除或伪造原对话；未知错误可显示经过净化的接口原因，但不能泄露密钥。",
    ],
    "AI开发项目_新聊天启动说明.docx": [
        "当前候选基线：网页 v899、私人 iOS 1.0.24 (24)、原生桥契约 11。远端迁移 202608110002–006 已核验应用；本轮没有把 APP 专属真实数据能力推送到普通网页版。",
        "正确查看时序：识别真实读取后立即变更代次并静默拦截普通输出；顶部横幅继续逐项展示；同一 readSessionId 完成后，把成功事实交给角色模型；角色自行决定一段或多句，程序不代写固定汇总。",
        "若看到“模型没有生成／再点一次”，先检查 replyGenerationRun 是否把读取代次变化误判为空输出。若看到固定逐项系统清单，检查是否重新引入 rolePhoneInspectionExactSummary 或其他模板兜底。",
        "线下失败排查先看界面给出的类别：余额／额度、401、403、404、429、超时、网络连接或空输出。空输出只表示接口有响应但没有合格可见内容，不可擅自断言余额不足。",
        "Windows 自动测试当前 478／478 通过。Mac 只从 SmallPhone_v899_ModelOnlyInspectionAndOfflineDiagnostics 新目录打开工程，编译五 Target 并进行真机验收。",
    ],
}


def main() -> None:
    for filename, items in ENTRIES.items():
        path = ROOT / filename
        document = Document(path)
        changed = not any(MARKER in paragraph.text for paragraph in document.paragraphs)
        if changed:
            document.add_heading(MARKER, level=1)
            for item in items:
                document.add_paragraph(item, style="List Bullet")
            document.save(path)
        with ZipFile(path) as archive:
            assert archive.testzip() is None
            assert "[Content_Types].xml" in archive.namelist()
        verified = Document(path)
        assert sum(MARKER in paragraph.text for paragraph in verified.paragraphs) == 1
        print(f"{'updated' if changed else 'unchanged'}: {path.name}")


if __name__ == "__main__":
    main()
