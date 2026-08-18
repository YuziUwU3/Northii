from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
EXPECTED = {
    "AI开发项目_项目说明文档.docx": ["v856｜聊天输入不中断", "375 项自动化测试全部通过"],
    "AI开发项目_Bug记录模板.docx": ["角色消息清空正在编辑的内容", "captureChatComposer()"],
    "AI开发项目_Bug修改规范.docx": ["异步来信不得替换编辑器", "至少 375 项测试全部通过"],
    "AI开发项目_新聊天启动说明.docx": ["v856 待发布包", "此次没有数据库迁移"],
}

for name, markers in EXPECTED.items():
    path = DOCS / name
    with ZipFile(path) as archive:
        assert archive.testzip() is None, f"corrupt zip member in {name}"
        assert "word/document.xml" in archive.namelist(), f"missing document.xml in {name}"
    document = Document(path)
    text = "\n".join(p.text for p in document.paragraphs)
    for marker in markers:
        assert marker in text, f"missing marker {marker!r} in {name}"
    assert len(document.paragraphs) >= 15, f"unexpectedly short document {name}"
    print(f"verified {name}: {len(document.paragraphs)} paragraphs, {len(document.tables)} tables")
