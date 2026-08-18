from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_release(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    document = Document(path)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)


append_release(
    "AI开发项目_项目说明文档.docx",
    "v856｜聊天输入不中断（2026-08-09）",
    [
        "当前前端版本统一升级为 v856。本次只修复微信角色聊天的输入稳定性，不改变 v855 的主动联系、记忆、密码、伴生 App、查岗或线下约会行为。",
        "根因是角色普通文字已经采用消息列表增量插入，但服务器主动推送、角色卡片、表情、日程提醒、查岗回应等分支仍会调用整页 render()。整页重绘会销毁并重新创建 cinput 输入框，因此用户尚未发送的文字、选区、光标和键盘焦点会被清空或打断。",
        "v856 把聊天消息列表提取为独立渲染单元。角色消息到达时只更新 chatbg 消息区，不替换输入栏；服务器推送、主动消息、特殊卡片、节日与日程消息、查岗消息均接入此安全刷新路径。手动回复按钮的生成状态改为原位更新，不再通过整页重绘刷新。",
        "同时保留兜底保护：任何确需整页渲染的旧路径都会在替换页面前保存输入草稿、selectionStart、selectionEnd、输入框高度、内部滚动位置及焦点，渲染后仅在仍是同一角色聊天时恢复，防止遗漏分支再次清空输入。",
        "验证结果：app.js 语法检查通过；新增 chat-composer-preservation.test.mjs；全项目 375 项自动化测试全部通过，0 失败。",
    ],
)

append_release(
    "AI开发项目_Bug记录模板.docx",
    "v856 Bug 记录｜角色消息清空正在编辑的内容（2026-08-09）",
    [
        "现象：用户正在角色聊天输入框编辑文字时，角色普通回复、主动消息或服务器推送到达，会突然清空输入内容，打断光标和键盘状态。特殊卡片、表情、日程提醒与查岗回应更容易复现。",
        "根因：消息到达后的界面刷新策略不一致。普通文字使用 appendChatMessageHTML() 增量插入，而其他分支调用 render() 重建整个 #app；cinput 属于被替换的 DOM，因此浏览器无法保留未发送内容和输入法会话。",
        "修复：新增 chatMessageListHTML() 与 refreshChatMessages()，后台来信只重绘 #chatbg；服务器拉取、特殊角色消息、查岗与日程路径全部切换到消息区刷新。replyGenerationRefresh() 改为直接更新按钮状态。",
        "兜底：captureChatComposer() 与 restoreChatComposer() 保护同一聊天内的草稿、光标选区、焦点、高度和滚动位置；切换到其他角色或其他页面时不会错误注入旧草稿。",
        "防复发：新增专项测试模拟整页替换前后输入框实例变化，验证文字、选区、焦点、高度与滚动位置全部恢复；并断言服务器推送和特殊消息不得调用整页 render()。全量 375 项测试通过。",
    ],
)

append_release(
    "AI开发项目_Bug修改规范.docx",
    "新增强制规范｜异步来信不得替换编辑器（v856 起）",
    [
        "任何角色消息、服务器轮询、通知回填、后台同步、识图完成、语音完成或生成状态变化，只要用户正在 INPUT、TEXTAREA 或 contenteditable 中编辑，就不得通过整页 innerHTML 替换来刷新状态。",
        "聊天来信必须优先更新消息容器本身。若消息类型需要重新生成卡片，应重绘消息列表容器，不得连带重建输入栏；生成中、排队中等按钮状态必须原位更新。",
        "确需整页刷新时，必须按页面身份键保存并恢复草稿、选区、焦点、输入框尺寸和滚动位置；恢复前要确认刷新后仍是同一页面、同一角色，禁止把草稿带到其他会话。",
        "涉及输入、消息投递或后台同步的修改，必须同时覆盖普通文字、语音、图片、表情、位置、卡片、主动消息、服务器推送、查岗和日程提醒，并运行专项测试与全量测试。不得只验证一种消息格式。",
        "发布门槛更新为至少 375 项测试全部通过；若总数减少，必须记录删除原因，禁止通过删除断言或放宽断言规避输入中断回归。",
    ],
)

append_release(
    "AI开发项目_新聊天启动说明.docx",
    "新聊天接手状态｜v856 待发布包（2026-08-09）",
    [
        r"工作目录：C:\Users\pc\Documents\小手机\phone-release-web-v835-avatar。当前版本 v856，变更基于已发布的 v856。",
        "本次修复：角色来信不再替换聊天输入框；普通回复、主动消息、服务器推送、特殊卡片、表情、节日/日程和查岗消息只刷新消息区。旧路径整页刷新时仍有草稿与焦点兜底恢复。",
        "新增测试：tests/chat-composer-preservation.test.mjs；chat-date-boundary.test.mjs 已同步适配消息列表渲染函数。全项目 375 项测试通过，0 失败。",
        "版本缓存已同步：app.js、sw.js、index.html、repair.html、中文主页面及相关版本测试均为 856。此次没有数据库迁移，也没有 Edge Function 代码变化，只需提交到 main 并确认 GitHub Pages 构建成功。",
        "发布后验收：在角色聊天输入半句但不要发送，分别让角色普通回复、主动联系和服务器推送到达；新消息应正常出现，输入文字、光标和键盘必须保持不变。",
    ],
)

print("Updated v856 maintenance documents")
