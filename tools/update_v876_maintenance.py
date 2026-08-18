from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_release(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    document = Document(path)
    if any(paragraph.text.strip() == title for paragraph in document.paragraphs):
        print(f"Skipped existing section: {title}")
        return
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"Updated {filename}")


append_release(
    "AI开发项目_项目说明文档.docx",
    "v876｜North App Store 审核支持与隐私页面（2026-08-10）",
    [
        "分发决定：iOS 原生伴生 App North 继续保持独立，不把网页小手机整体嵌入原生 App；当前目标为 App Store Connect 的非公开 App（Unlisted App）审核路线。安卓继续使用网页小手机。非公开分发仍需 Apple 审核，不得描述成免审核或绕过 Family Controls entitlement。",
        "新增公开页面 north-support.html 与 north-privacy.html。支持页包含系统授权、App 选择、短时配对、同步和锁定故障恢复说明；隐私页披露屏幕使用时间、步数、睡眠、可选位置、设备配对、命令回执、Apple 系统服务与 Supabase 后端的真实数据边界，并明确不读取微信、抖音等第三方 App 的聊天、私信或页面内容。",
        "正式审核地址在发布后为 https://fenglina35-dotcom.github.io/phone/north-support.html 与 https://fenglina35-dotcom.github.io/phone/north-privacy.html。两个页面互相链接，支持页面使用项目 Issues 作为公开问题入口，并明确禁止公开提交配对码、设备密钥、位置或健康数据。",
        "Service Worker 原先会把全部 navigate 请求替换为缓存的小手机主页面。v876 为 north-support.html 与 north-privacy.html 增加最小排除规则，让它们按普通静态页面直接读取，避免审核员打开支持或隐私链接却看到小手机主页。小手机其他离线外壳、聊天数据与缓存恢复逻辑保持不变。",
        "发布缓存统一升级为 v876，版本标签为“v876 · North 审核支持页面”。专项页面、隐私边界、秘密泄漏防护、Service Worker 排除和版本对齐测试已加入；完整 Node 回归通过后才允许提交。",
    ],
)

append_release(
    "AI开发项目_Bug记录模板.docx",
    "v876 Bug 记录｜公开审核页面会被 Service Worker 错换成小手机主页（2026-08-10）",
    [
        "现象：准备 App Store Connect 支持 URL 与隐私政策 URL 时发现，即使仓库存在独立 HTML，已经安装过小手机 Service Worker 的浏览器打开这些地址，也会进入缓存的小手机主页面。Apple 审核员因此可能无法读取支持和隐私内容。该问题在正式推送前由静态检查发现，尚未作为线上失败交付。",
        "根因：sw.js 对所有 request.mode === 'navigate' 的同源请求统一返回 currentCore(cache, 'html')，没有区分应用外壳导航与独立公开文档。新增文件本身不能改变已注册 Service Worker 的导航接管范围。",
        "修复：在通用 navigate 处理之前识别 north-support.html 与 north-privacy.html，并直接 return，不调用 event.respondWith，让浏览器按普通网络请求读取真实静态文件。只排除这两个固定文件，不放宽 app.js、授权组件、缓存外壳或其他资源策略。",
        "验证：新增 north-app-store-pages.test.mjs，检查支持与隐私必要内容、第三方 App 内容边界、无 service_role/owner_secret/pair_secret/JWT 泄漏，并断言公开页面排除规则位于通用 navigate 处理之前。版本、语法和完整 tests/*.test.mjs 同步验证。",
        "失败方案记录：仅新增 HTML、但不检查 Service Worker 导航拦截的方案被判定不完整，未进入提交；不能用清缓存或要求审核员换浏览器代替代码修复。",
    ],
)

append_release(
    "AI开发项目_Bug修改规范.docx",
    "新增强制规范｜公开审核文档必须绕开应用外壳导航回退（v876 起）",
    [
        "在 PWA 根作用域新增支持页、隐私政策、审核说明、删除账户说明或法律文档时，不能只验证文件存在和直连 200。必须检查已注册 Service Worker 对 navigate 的处理，确保公开文档不会被缓存的应用 Shell 替换。",
        "排除规则应使用固定文件名或明确白名单，并位于通用 navigate fallback 之前。禁止为了公开文档直接关闭整个 Service Worker、扩大为任意路径绕过缓存，或破坏小手机已有离线启动与数据保护。",
        "审核与隐私页面不得包含真实 API Key、service_role、owner secret、pair secret、JWT、推送令牌、精确位置或健康样本。公开反馈入口必须提醒用户不要提交敏感信息；涉及删除或隐私请求时应转入私密处理渠道。",
        "回归至少覆盖：两份公开页面互链、必要披露、第三方 App 内容边界、秘密模式扫描、Service Worker 排除顺序、原有 app Shell 版本对齐和完整 Node 测试。",
    ],
)

append_release(
    "AI开发项目_新聊天启动说明.docx",
    "新聊天接手状态｜v876 North 非公开 App 审核资料（2026-08-10）",
    [
        "固定仓库 C:\\Users\\pc\\Documents\\小手机\\phone-work，固定分支 main，远端 origin/main。禁止新分支、worktree、旧目录、重置、覆盖和强推。当前协作约定为 Codex 完成测试并提交；网络推送由用户手动执行。",
        "North 是独立 iOS 伴生 App，不嵌入完整网页小手机；安卓继续使用网页。App Store Connect 已创建 North 1.0，主 Bundle ID 为 com.qianyi.PhoneCompanionTest，构建 1.0 (1) 已上传并验证，截图已上传。当前路线为非公开 App（Unlisted App），仍需正常完成元数据、审核和之后的非公开分发申请。",
        "v876 新增 north-support.html 与 north-privacy.html。发布后 URL 分别为 https://fenglina35-dotcom.github.io/phone/north-support.html 和 https://fenglina35-dotcom.github.io/phone/north-privacy.html；在用户推送并等待 GitHub Pages 部署前，不得声称线上已经可访问。",
        "sw.js 必须保留 north-(support|privacy).html 的 navigate 排除规则，否则已注册 PWA 会把审核文档替换成小手机主页。后续新增审核公开文档时也要加入白名单与专项测试，不能让审核员通过清缓存绕过。",
        "版本基线为 v876 · North 审核支持页面。继续提交 App Store 元数据时要如实说明：North 只处理主动授权的屏幕使用时间汇总、选定 App 管理、步数、睡眠、可选位置与设备回执，不读取第三方 App 内部内容，也不把服务器接收命令冒充设备已执行。",
    ],
)

def replace_once(filename: str, old: str, new: str) -> None:
    path = DOCS / filename
    document = Document(path)
    matches = [paragraph for paragraph in document.paragraphs if old in paragraph.text]
    if not matches:
        if any(new in paragraph.text for paragraph in document.paragraphs):
            return
        raise AssertionError(f"Missing verification marker in {filename}")
    if len(matches) != 1:
        raise AssertionError(f"Duplicate verification marker in {filename}")
    matches[0].text = matches[0].text.replace(old, new)
    document.save(path)


replace_once(
    "AI开发项目_项目说明文档.docx",
    "完整 Node 回归通过后才允许提交。",
    "完整 Node 回归 422/422 通过。",
)
replace_once(
    "AI开发项目_Bug记录模板.docx",
    "版本、语法和完整 tests/*.test.mjs 同步验证。",
    "版本、语法和完整 tests/*.test.mjs 同步验证，最终 422/422 通过。",
)

print("Updated v876 maintenance documents")
