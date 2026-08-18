from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
EXPECTED = {
    "AI开发项目_项目说明文档.docx": [
        "v860｜伴生推送上线与固定仓库交接",
        "不再是 HTTP 404",
        "GitHub 源码状态与 Supabase 部署状态必须分开描述",
    ],
    "AI开发项目_Bug记录模板.docx": [
        "v860 Bug 记录｜未发布部署记录迁入固定 main",
        "旧本地提交 2eb2ecb",
        "预期 HTTP 400 与 invalid-request",
    ],
    "AI开发项目_Bug修改规范.docx": [
        "固定仓库单一 main 发布（v860 起）",
        "禁止创建新分支、禁止创建 worktree",
        "实际核验线上页面是否读取到新版本",
    ],
    "AI开发项目_新聊天启动说明.docx": [
        "v860 固定仓库与伴生推送已核验",
        "不要重复部署函数",
        "下一步仅需在 Mac Xcode",
    ],
}


for name, markers in EXPECTED.items():
    path = DOCS / name
    with ZipFile(path) as archive:
        assert archive.testzip() is None, f"corrupt zip member in {name}"
        assert "word/document.xml" in archive.namelist(), f"missing document.xml in {name}"
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for marker in markers:
        assert marker in text, f"missing marker {marker!r} in {name}"
    assert text.count(markers[0]) == 1, f"duplicate v860 section in {name}"
    print(f"verified {name}: {len(document.paragraphs)} paragraphs, {len(document.tables)} tables")
