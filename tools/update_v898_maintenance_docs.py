from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1] / "docs" / "maintenance"
MARKER = "2026-08-12 v898／私人 1.0.23 全渠道读取完成后才回复"

ENTRIES = {
    "AI开发项目_项目说明文档.docx": [
        "网页核心升级为 v898，私人 iOS 候选版升级为 1.0.23 (23)，原生桥契约仍为 11。微信聊天、电话和共同生活共用全局真实读取状态，不再由各场景分别决定是否放行普通回复。",
        "读取请求一进入队列就变更生成代次。普通回复必须在调用模型前、模型返回后和实际发送前复核读取状态与代次；只要期间开始过真实读取，旧代次内容即使已经生成也永久作废。",
        "所有原生读取项目完成后，成功事实才作为隐藏上下文交给角色生成一次自然回复。完成回复是唯一允许越过读取锁的出口；电话必须等待完成回复真正说完才释放通道。",
        "顶部横幅继续逐项显示读取进度。不可用、未授权、超时和技术诊断保留在内部 readOutcomes，不得进入角色气泡、电话语音或共同生活对白。",
        "Windows 全套 476 项自动测试全部通过。Mac 五 Target 编译、签名及微信／电话／共同生活三个真机读取场景仍需最终验收。",
    ],
    "AI开发项目_Bug记录模板.docx": [
        "现象：用户说‘查一次所有的’后，角色在顶部仍显示读取中时先回复‘你在测试我’、位置或权限等普通话术；读取完成后又收到一次真实数据回复。",
        "真实原因：v897 的自然语言焦点解析没有把‘查一次所有的’归为‘全部数据’，早期 nativeOnly 门禁得到空焦点后退出。普通模型先生成；后置事实守卫虽然重新启动真实读取，却仍放行非数据句，形成两个独立回复通道。",
        "前次方案为何未成功：v897 只为已识别的同一用户消息设置 pending 键，并把几个固定入口提前；它没有全局冻结三个场景，也没有让已经在生成中的普通输出失效，所以继续补提示词无法封死竞态。",
        "修复：新增全局读取状态、稳定排队键和单调生成代次。聊天、电话、共同生活在模型前、模型后、发送前都检查；后置守卫若发现真实数据意图，只能排队读取并立即返回，不能保留任何旧文本。",
        "验证：node --test tests/*.test.mjs 共 476 项，476 项通过；另有针对‘查一次所有的’等自然说法、读取中旧输出作废和三场景独占通道的专项测试。",
    ],
    "AI开发项目_Bug修改规范.docx": [
        "等待原生异步读取的功能必须使用状态门禁，而不能只依赖提示词或在模型生成后删除文字。排队、执行、完成回复三个阶段需要明确，普通回复只能在无读取任务时运行。",
        "任何可能跨越异步边界的可见输出都必须携带生成代次，并在模型返回后和最终发送前复核。只要期间启动过读取，旧代次输出必须永久丢弃，不能在读取释放后补发。",
        "后置事实守卫只能作为安全兜底：发现未经本轮读取凭证支持的电量、位置、健康或屏幕数据时，应排队一次真实读取并返回空输出，不得一边读取一边发送剩余句子。",
        "读取完成回复必须是唯一例外且显式标记。它只能接收同一 readSessionId 的完成事实，不得包含 readErrors、missing、权限、超时或系统逐项清单。",
        "微信聊天、电话、共同生活新增或修改入口时必须共享同一锁与代次，不得复制出各自独立的 pending 逻辑。",
    ],
    "AI开发项目_新聊天启动说明.docx": [
        "当前候选基线：网页 v898、私人 iOS 1.0.23 (23)、原生桥契约 11。远端迁移 202608110002–006 已核验应用；本轮没有把 APP 专属真实数据能力推送到普通网页版。",
        "全渠道正确时序：用户请求查看真实 iPhone；立即排队并变更生成代次；普通回复全部冻结；顶部横幅逐项读取；同一 readSessionId 完成；成功事实进入隐藏上下文；只生成并发送一轮角色自然回复。",
        "v897 未成功的关键不是原生没读到，而是‘查一次所有的’焦点解析为空，并且同消息 pending 无法取消已经在生成的旧回复。以后不得只补固定说法，必须检查全局状态、生成代次和最终发送门。",
        "排查重复回复时依次检查 queueNativeInspection、rolePhoneInspectionEpoch、模型返回后的 stale gate、最终发送前 gate，以及完成回复是否被显式标记为 inspectionCompletion。",
        "Windows 自动测试当前 476／476 通过。Mac 只从 SmallPhone_v898_AllChannelInspectionGate 新目录打开工程，编译五 Target，并分别在微信、电话和共同生活验证读取期间零回复、完成后仅一轮自然回复。",
    ],
}


def main() -> None:
    for filename, items in ENTRIES.items():
        path = ROOT / filename
        document = Document(path)
        changed = not any(MARKER in paragraph.text for paragraph in document.paragraphs)
        if changed:
            document.add_heading(MARKER, level=1)
            for item in items:
                document.add_paragraph(item, style="List Bullet")
            document.save(path)
        with ZipFile(path) as archive:
            assert archive.testzip() is None
            assert "[Content_Types].xml" in archive.namelist()
        verified = Document(path)
        assert sum(MARKER in paragraph.text for paragraph in verified.paragraphs) == 1
        print(f"{'updated' if changed else 'unchanged'}: {path.name}")


if __name__ == "__main__":
    main()
