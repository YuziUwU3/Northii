from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_release(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    document = Document(path)
    if any(paragraph.text.strip() == title for paragraph in document.paragraphs):
        print(f"Skipped existing section: {title}")
        return
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"Updated {filename}")


def upsert_note_by_content(document_marker: str, legacy_markers: list[str], note_marker: str, text: str) -> None:
    for path in DOCS.glob("*.docx"):
        document = Document(path)
        current = "\n".join(paragraph.text for paragraph in document.paragraphs)
        if document_marker not in current:
            continue
        for paragraph in document.paragraphs:
            if note_marker in paragraph.text:
                if paragraph.text != text:
                    paragraph.text = text
                    document.save(path)
                    print(f"Updated note in {path.name}: {note_marker}")
                else:
                    print(f"Skipped existing note: {note_marker}")
                return
            if any(marker in paragraph.text for marker in legacy_markers):
                paragraph.text = text
                document.save(path)
                print(f"Corrected note in {path.name}: {note_marker}")
                return
        if note_marker in current:
            print(f"Skipped existing note: {note_marker}")
            return
        document.add_paragraph(text)
        document.save(path)
        print(f"Added note to {path.name}: {note_marker}")
        return
    raise RuntimeError(f"No maintenance document contains {document_marker!r}")


append_release(
    "AI开发项目_项目说明文档.docx",
    "v869｜共同生活限额与每日必查（2026-08-10）",
    [
        "共同生活里的角色现在可口头修改真实 App 每日限额。模型使用隐藏标签 [共同生活限额|准确App名|1到720的分钟数]；程序只接受已完成稳定 ID 关联、内置授权和限额权限的 App，并复用既有 companionDispatchRoleByText('limit') 路径强制同步同一份分钟数到小手机内置 App 与真实 iPhone。命令提交仍不等于成功，必须等待设备回执。",
        "每日必查由两次真实查看组成：一次 HealthKit 查看同时读取最近睡眠与今日步数；另一次读取 iPhone 今日总屏幕时长和全部逐 App 时长。心率保留为角色可自主查看项，但不属于每日必查；电量、位置及其他手机内容也不属于每日配额。",
        "共同生活与线上自动查看共用 automationRuns 当天完成标记。双方处于在家或一起外出的面对面现场时，共同生活接管到期的每日必查；线上通道保持静默。共同生活完成并实际生成回应后才写完成标记，失败后五分钟重试。同一执行占用继续阻止线上线下同时查看，事实签名与当日标记共同避免同一内容重复发言。",
        "共同生活自主查看仍不限次数，只能因实际查看结果发起主动消息，不得分享普通线上日常。睡眠与步数合并为一个查看目标，避免角色先看睡眠、再看步数时对同一份 HealthKit 摘要重复反应。角色独自上班、外出或回家路上时不运行共同生活查看。",
        "验证：app.js 语法和差异检查通过；共同生活、伴生自动化、设备控制、状态、作息等相关回归 63/63 通过；完整 Node 自动化回归 408/408 通过。未修改 Supabase Edge Function、数据库迁移或 iPhone 原生工程。",
    ],
)

append_release(
    "AI开发项目_Bug记录模板.docx",
    "v869 Bug 记录｜共同生活缺少限额与每日必查交接（2026-08-10）",
    [
        "现象：v867 已允许共同生活自主查看、锁定、解锁和微信登录，但角色不能在共同生活里口头修改真实 App 限额。v850 的睡眠和完整应用使用记录虽是线上每日必查，进入共同生活面对面现场后线上自动化会被正确静默，却没有把到期任务交给共同生活，因此整段共同生活期间可能完全漏查。旧睡眠必查还会顺带读取心率、电量和位置，与本次用户确认的必查范围不一致。",
        "根因：共同生活隐藏协议和标签解析器没有 limit 分支；cohabPhoneAutonomyTick 只运行自由自主决策，没有读取 companionAutomationCandidate 的必查状态。线上静默和线下执行分别实现，但没有共享的必查接管函数与完成标记。睡眠候选把可选健康/设备项目作为 extras 混入同一次必查，无法表达‘睡眠和步数必查、心率不是必查’的最新范围。",
        "修复：新增共同生活限额标签并复用既有已绑定双端统一限额路径；新增 companionRequiredDailyCandidate 与 cohabDailyRequiredMaybe，在共同生活同场时优先执行到期的睡眠+步数或完整屏幕时间查看，成功后调用 companionAutomationRecord。每日查看记录保存 dailyKind/dailyDay，既允许次日同值照常完成一次，又让当天普通重复查看继续被事实哈希挡住。",
        "范围纠正：HealthKit 每日必查只包含最近睡眠与今日步数，不读取心率、电量或位置；心率仅保留为授权后的自主查看目标。睡眠和步数映射到同一个查看键，防止拆成两个目标后对相同健康摘要重复发言。屏幕必查仍要求真实逐 App 快照、全部 App 明细和新鲜同步时间。",
        "失败与重试边界：设备数据未就绪、执行通道占用、页面隐藏或模型未产生有效共同生活回复时不写完成标记；五分钟后再试。APNs、网页入队和标签解析都不能冒充真实限额成功。相关 63/63 与完整 408/408 自动化测试通过。",
    ],
)

append_release(
    "AI开发项目_Bug修改规范.docx",
    "新增强制规范｜跨通道每日必查必须显式交接（v869 起）",
    [
        "当线上通道因共同生活、通话、影院或其他占用而静默时，任何标记为 requiredDaily 的任务都必须明确交给当前唯一可见通道，或保留为未完成等待后续执行。禁止只加静默条件而没有接管逻辑，否则会把防重复变成漏执行。",
        "线上与共同生活必须共用同一份每日完成标记；只有真实数据已读取且来源通道实际生成有效回复后才能落标。执行失败不得提前完成，重试必须有退避。任何时刻继续使用同一个查看/登录/远控占用，禁止两边同时查看或各发一次。",
        "每日健康必查的当前范围固定为一次读取睡眠与步数。心率、电量和位置是可选自主查看，不得作为睡眠必查的附带 extras。屏幕必查必须包含总屏幕时长和全部逐 App 时长；没有真实逐 App 新鲜快照时保持未完成，不得用 0、旧内置数据或推测补齐。",
        "跨日去重必须同时保存事实哈希和逻辑日：同一天相同事实不重复发言，下一天即使数值未变化也允许完成当天一次必查；完成后当天的普通重复查看仍应被同一事实签名拦截。相关健康子项合并查看时必须使用同一去重键，避免拆项造成重复反应。",
        "共同生活限额必须复用已验证的双端统一限额函数，只允许已关联且两端均授权的 App，分钟数限定 1到720。不得新建仅线下限额存储、不得把未绑定外置 App 设成独立限额，也不得把入队或 APNs 唤醒描述成设备已经执行成功。",
    ],
)

append_release(
    "AI开发项目_新聊天启动说明.docx",
    "新聊天接手状态｜v869 共同生活限额与每日必查（2026-08-10）",
    [
        "固定仓库 C:\\Users\\pc\\Documents\\小手机\\phone-work，固定分支 main，远端 origin/main。禁止新分支、worktree、旧目录和强推。开始前 clean、fetch、ff-only；发布前相关测试加完整回归，版本只在最新线上基础上加 1。",
        "共同生活真实手机入口集中在 cohabPhonePrompt、cohabApplyPhoneTags、cohabRunPhoneInspection、cohabDailyRequiredMaybe 与 cohabPhoneAutonomyTick。口头限额标签为 [共同生活限额|准确App名|分钟数]，内部必须走 companionDispatchRoleByText('limit', ..., scope:'both')；只对已绑定且授权的 App 生效。",
        "每日必查共两项：HealthKit 睡眠与步数合并一次；iPhone 总屏幕时长和全部逐 App 时长一次。companionRequiredDailyCandidate 决定到期项；共同生活同场时先执行必查，成功后 companionAutomationRecord 写共享完成标记。心率、电量、位置不是必查，仍可按角色意愿自主查看。",
        "通道边界：cohabOnlineQuiet 让线上在 home/together-away 静默；cohabDailyRequiredMaybe 负责线下接管。rolePhoneInspectionAcquire 是查看、微信登录和远控的单一占用；_phoneInspectionFacts 保存跨线上/线下事实哈希，并为每日项增加 dailyKind/dailyDay。禁止两边同时发或同事实重复说。",
        "稳定边界：不改变旧一次性约会旁白格式、共同生活独立总结、微信气泡、原生伴生 App、Supabase Edge Function、稳定外置 ID、回执与原子快照链路。发布前至少跑 cohabitation-phone-inspection、companion-automation、companion-device-control、共同生活状态/作息专项以及全量 Node 测试。",
    ],
)

upsert_note_by_content(
    "automationRuns",
    ["cohabPhoneDeliverFact 显式强制主模型"],
    "每次请求重新从主模型开始",
    "模型边界核对：普通共同生活对话和查看真实手机记录后的回应都遵守用户设置的共同生活回复线路。设置主模型时，每次请求重新从主模型开始；只有该次主模型报错或返回空内容时，才允许副模型临时重写一次。临时兜底不会写回模型设置，下一次对话仍先走主模型。实际从主模型切到副模型、或下一轮从副模型切回主模型时，会显示三秒提示；连续使用同一路线不重复提示。",
)
upsert_note_by_content(
    "dailyKind/dailyDay",
    ["原先调用通用 cohabReplyCore"],
    "本次只做诊断",
    "模型风格诊断：用户观察到查看手机后的语气偶尔像副模型。代码核对确认，主模型设置没有被永久改变；可能发生的是某一次主模型失败或空回复后，由副模型完成该次临时重写。route.fallback 只属于当前请求，不持久化到共同生活设置。保留既有主模型优先、失败时副模型单次兜底的行为；新增仅驻留当前页面内存的实际路线标记，路线真正改变时提示三秒，未切换不提示。",
)
upsert_note_by_content(
    "requiredDaily",
    ["只能给真实记录这一条调用链增加显式 forceMain"],
    "回退不得持久化",
    "模型路由规范确认：共同生活必须尊重用户选择的回复线路。选择主模型时允许在本次主请求失败或空回复后调用副模型重写一次，但回退不得持久化、不得修改角色或共同生活模型设置；下一次独立请求仍必须先尝试主模型。模型切换提示只能在实际路线与上一轮不同后显示三秒，首次建立基线和连续同路线不得提示，提示状态不得当作模型设置保存。",
)
upsert_note_by_content(
    "scope:'both'",
    ["cohabReplyCore(..., {forceMain:true})"],
    "不会写回模型设置",
    "模型交接：cohabRoleChat 每次根据 cohabSettings(d).replyRoute 与角色模型决定首选线路。用户设置主模型时先请求主模型；仅在该次失败或空回复且副模型可用时临时请求副模型一次，并把 fallback 只记在当前 route 对象，不会写回模型设置。_cohabActualModelRoute 只记录当前页面最后实际使用路线；cohabModelRouteNotice 仅在 main/aux 发生变化时 toast 三秒，后续新对话仍重新先请求主模型。",
)

print("Updated v869 maintenance documents")
