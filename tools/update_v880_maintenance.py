from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_release(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    document = Document(path)
    if any(paragraph.text.strip() == title for paragraph in document.paragraphs):
        print(f"Skipped existing section: {title}")
        return
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"Updated {filename}")


append_release(
    "AI开发项目_项目说明文档.docx",
    "v880｜持久锁定与当日同步（2026-08-11）",
    [
        "本次版本从 v879 升级为 v880 · 持久锁定与当日同步。目标是保证刷新、重连、跨日和每日限额监控重建都不能擅自解锁，同时让网页只接受 iPhone 当前本地日历日的递增使用量快照。",
        "原生 CompanionSyncView 新增持久锁账本，并统一合并手动锁、每日限额锁和持久锁三类 Token。初始化会恢复每日限额 shield 与持久账本；makeSnapshot() 使用有效锁集合；rebuildDailyLimitMonitoring() 不再清空 dailyLimitStore 或 limit.lockedTokens。明确解锁会同时移除三类锁并保存，系统读回失败时恢复三份先前状态。",
        "网页 Companion 状态升级到 schema 7，lockIntents 按外置稳定 ID 保存 desiredLocked，设备快照只更新 reportedLocked。schema 6 的已锁项目迁移为持久锁定意图；只有明确角色/面板锁定或解锁动作能更新意图。快照从已锁变成未锁时不再生成假的 manualUnlock，界面和角色上下文分别显示期望状态与设备回报。",
        "使用量快照新增 usageDay、timeZone 和 usageRevision。网页拒绝非当前日、旧 revision 和旧 generatedAt，控制回执快照不能冒充使用量刷新；同一天的总时长和逐 App 时长不允许被较小旧值覆盖。每日使用量主动查看只接受当前日期且 20 分钟内的新鲜使用量数据。轮询保持伴生页约 4 秒、其他前台页面约 20 秒，并继续服从 iOS 后台调度限制。",
        "缓存链已统一到 880：入口页、修复页、HTML shell、app.js、资源 query、Service Worker build 和预缓存版本全部对齐。Node 全量回归 tests/*.test.mjs 为 434/434；新增锁定账本测试覆盖 schema 迁移、快照无解锁权限、限额重建保锁、三存储显式解锁回滚、90 分钟限额已用超过 100 分钟仍保留意图，以及使用量日期/revision。",
        "本轮没有修改 Supabase Function 或 migration。Windows 仓库未包含 Monitor Extension 源码，也无法读取 Mac 实际 Xcode 文件或执行真机验收；把仓库 Swift 复制到 Mac 前仍必须先比较 /Users/zoushijie/Documents/AppleProjects/PhoneCompanionTest 中的实际文件，并核对 limit.lockedTokens 编码、编译、连续刷新、前后台、重启、跨日与真实回执。",
    ],
)

append_release(
    "AI开发项目_Bug记录模板.docx",
    "v880 Bug 记录｜刷新误解锁、假手动解锁与跨日旧使用量（2026-08-11）",
    [
        "现象：真实 iPhone App 已锁定后可能在刷新或每日限额重建后自行变成可用；网页会把旧状态 locked、新快照 unlocked 直接解释成手动解锁并触发角色反应；控制回执或旧报告还可能覆盖当天较新的总时长和逐 App 时长。",
        "根因：原生 rebuildDailyLimitMonitoring() 曾把 dailyLimitStore.shield.applications 设为 nil 并删除 limit.lockedTokens；初始化和 makeSnapshot() 又主要依赖手动锁 store。网页 schema 6 只有 app.locked 一个混合字段，快照差异被当成解锁授权。使用量缺少明确的本地日期、时区和单调 revision，且 lastSync 同时承担控制回执与使用量新鲜度。",
        "修复：原生增加独立持久锁账本，统一恢复/读取三类锁；限额重建保留 shield 和 Token；显式解锁成为同时清除三类锁的唯一代码入口并提供失败回滚。网页 schema 7 分离 desiredLocked/reportedLocked，删除由快照构造 manualUnlock 的逻辑。使用量按 usageDay/timeZone/revision 接受，并把 usageGeneratedAt 与普通 lastSync 分开。",
        "失败方案与禁止项：不能靠让快照永不过期掩盖问题；不能在重建限额时先清空全部 shield；不能用 app.locked=false 推导用户授权；不能只缩短轮询而不做日期与顺序保护；不能用控制回执时间触发每日使用量查看。",
        "自动验收：专项伴生测试 61/61，通过 Node 全量回归 434/434。尚未完成的外部验收是 Mac 工程差异比较、Monitor Extension 源码核对、Xcode 编译和真机连续刷新/前后台/重启/跨日/90 分钟限额已用超过 100 分钟/明确解锁回执；这些结果不得被文档写成已经通过。",
    ],
)

append_release(
    "AI开发项目_Bug修改规范.docx",
    "新增强制规范｜锁定意图与使用量快照必须分层（v880 起）",
    [
        "锁定意图是授权状态，设备快照是观察状态。desiredLocked 只能由明确锁定或明确解锁动作修改；reportedLocked 只能表示设备最近回报。刷新、重连、控制回执、跨日、限额重建、旧快照和报告不可用都不得清除锁定意图，也不得制造 manualUnlock 事件。",
        "原生有效锁定集合必须至少合并手动锁、每日限额锁和持久锁账本。新增任何监控重建或恢复路径时，不得清空仍有效的 shield/Token。显式解锁若涉及多个 store，必须先保存旧值、一次性移除并持久化、读回验证，任一失败恢复全部旧值。",
        "使用量 payload 必须包含 iPhone 本地 usageDay、timeZone、generatedAt 和单调 revision/sequence。网页只接受当前日且比已确认版本新的 payload；同一天不得用较小旧值静默覆盖较大已确认值。跨日可以开始新的当天统计，但绝不能重置 lockIntents。控制命令快照与使用量快照必须使用不同的新鲜度字段。",
        "涉及 Family Controls、ManagedSettings 或 DeviceActivity 的修改，Windows 静态测试不能代替 Mac/Xcode 和真机验证。复制仓库 Swift 前必须比较 Mac 实际工程文件；必须核对 Monitor Extension 对 limit.lockedTokens 的编码协议，并记录编译结果、授权状态、连续刷新、前后台、重启、跨日、限额阈值和真实回执。",
        "发布必须先有针对不变量的红灯测试，再实现修复，最后运行相关测试与全量测试。UI 和角色提示必须分别陈述期望锁定、设备回报和使用量数据新鲜度，不能把三者混成一个“锁定/不一致”结论。",
    ],
)

append_release(
    "AI开发项目_新聊天启动说明.docx",
    "新聊天接手状态｜v880 持久锁定与当日同步（2026-08-11）",
    [
        "固定仓库：C:\\Users\\pc\\Documents\\小手机\\phone-work；固定分支 main。只允许在该目录 fast-forward、修改、测试、提交和普通推送，不创建 worktree/支线，不强推。当前网页和缓存版本为 v880 · 持久锁定与当日同步。",
        "已实现：原生三类锁 Token 合并与恢复、持久锁账本、限额重建保锁、显式解锁三存储回滚；网页 schema 7 lockIntents、desiredLocked/reportedLocked 分离、快照无解锁权限；usageDay/timeZone/revision 防前日和旧值覆盖。Node 全量回归 434/434。没有改 Supabase Function 或 migration。",
        "仍需在 Mac 先比较仓库 CompanionSyncView.swift 与 /Users/zoushijie/Documents/AppleProjects/PhoneCompanionTest/PhoneCompanionTest.xcodeproj 的实际源码，不能覆盖用户手工改动。Windows 未找到 Monitor Extension 源码，必须在 Mac 核对 limit.lockedTokens 的 JSON 编码、App Group、named store 和清理时机后再编译。",
        "真机未完成项：手动锁定后连续刷新 10 次、前后台、关闭重开 North、重建限额、跨日、90 分钟限额且已用超过 100 分钟、角色明确解锁、用户手动解锁，以及同采样时刻的总时长/逐 App/步数/睡眠/电量日期与时间戳。只有拿到真实设备回执才可勾选。",
        "公开 North 与 Family Controls Distribution 仍以 Apple 实际审核状态为准。不要把正在审核的公开 North 改成完整小手机。私人 App 阶段尚未开始，必须等用户确认显示名并完成锁定真机验收后，再单独设计 Bundle ID、Keychain、共享核心和账号恢复。",
    ],
)

EXTRA_PARAGRAPHS = {
    "AI开发项目_项目说明文档.docx": [
        "v880 同时补齐跨渠道时间轴与逐页安全区：共同生活/一次线下消息写入真实 time；微信、电话、线下按真实时间合并并明确 user/assistant 身份，分别生成“当前输入前用户上一句”和“角色上一句”锚点。购物、外卖新增页面级 nav 类；抖音 feed 返回键/标签及其他页 nav、云程 head/scroll 均有独立安全区规则和回归断言。",
    ],
    "AI开发项目_Bug记录模板.docx": [
        "并行根因记录：v879 只加强了同居实时状态和提示规则，但底层共同生活消息没有 time；微信上下文只是 system 隐藏背景，线下历史仍单独作为主消息链，没有像电话那样按 role 与时间合并，因此“上一句话”会优先落到旧线下消息或发生说话人串位。安全区验收只检查通用/既有选择器，抖音首页返回键与云程顶栏的行内样式没有命中；购物/外卖没有页面级类名与专项断言，所以曾被误判为已覆盖。",
    ],
    "AI开发项目_Bug修改规范.docx": [
        "跨渠道消息强制规范：消息落库必须先保存真实 time，再按时间合并微信、电话、线下等渠道；上下文必须分别锚定当前输入前最新用户消息与最新角色消息，不能用隐藏 system 背景或单独主消息链推断“上一句话”。",
        "独立 App 安全区强制规范：每个独立 App 必须具有页面级安全区类名和逐页回归测试；返回键、顶栏、标签栏、滚动容器即使使用行内样式，也必须由专项选择器和断言覆盖，不能靠通用选择器推断。",
    ],
    "AI开发项目_新聊天启动说明.docx": [
        "已合并的另一组 v880 改动：共同生活/一次线下消息保存真实 time；微信/电话/线下按时间和 role 合成统一时间轴，并分别提供最新用户/角色锚点；购物、外卖、抖音、云程已增加逐页安全区类与回归测试。维护这些路径时必须保留该行为。",
    ],
}

for filename, paragraphs in EXTRA_PARAGRAPHS.items():
    path = DOCS / filename
    document = Document(path)
    existing = {paragraph.text.strip() for paragraph in document.paragraphs}
    changed = False
    for paragraph in paragraphs:
        if paragraph not in existing:
            document.add_paragraph(paragraph)
            changed = True
    if changed:
        document.save(path)
        print(f"Added combined v880 notes to {filename}")

for path in DOCS.glob("*.docx"):
    document = Document(path)
    in_v880 = False
    changed = False
    for paragraph in document.paragraphs:
        text = paragraph.text
        if "v880" in text and ("持久锁定与当日同步" in text or "刷新误解锁" in text or "锁定意图与使用量快照" in text):
            in_v880 = True
        if not in_v880:
            continue
        updated = text.replace("432/432", "434/434").replace("433/433", "434/434").replace("60/60", "61/61")
        if updated != text:
            for run in paragraph.runs:
                run.text = ""
            paragraph.add_run(updated)
            changed = True
    if changed:
        document.save(path)
        print(f"Finalized counts in {path.name}")

print("Updated v880 maintenance documents")
