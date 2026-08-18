from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1] / "docs" / "maintenance"
MARKER = "2026-08-11 v893／私人 1.0.18"


DOC_CONTENT = {
    "AI开发项目_项目说明文档.docx": (
        "逐项真实读取、明确解锁事件与统一后台推送",
        [
            "真实读取改为单次隔离会话：device.snapshot 返回 readSessionId、requestedFocus、capturedAt；角色只可使用本次原生快照，不得混用情侣空间旧缓存、虚拟位置、限额或内置网页计时。",
            "一键读取的可见进度逐项显示 iPhone 电量、总屏幕使用、每个真实 App 的实际使用时间、步数、睡眠、最新心率、心电/HRV 和位置。原生桥 25 秒结束悬挂请求；健康读取 12 秒超时；前台快照不再等待逐个地址反查。",
            "数值一致性守卫会拒绝模型引入快照中不存在的数字，尤其禁止把 120 分钟限额说成抖音实际使用。位置无真实授权时返回无数据，不再回退剧情位置。健康摘要新增最新 ECG 概要读取。",
            "手动解锁提醒改为必开规则。原生管理页和小手机控制页只在真实解锁成功后写入 explicit manualUnlock 事件；角色自己下达的解锁被排除；事件保留到角色成功发出反应后才标记 delivered。",
            "统一私人 App 可在手机号账号认证后直接认领云端伴生身份并登记本机 APNs 令牌，不再要求另配一个伴生页面。角色资料页显示通知令牌、角色资料、每分钟 cron 和下一次检查的真实状态。",
            "主动联系仍遵守会话后随机 30–60 分钟硬静默。间隔 1 分钟只代表静默期结束后的最小联系间隔，不能覆盖 30–60 分钟；睡眠、正在聊天、通话和线下场景仍禁止插入。",
            "服务端必须追加 202608110004_private_phone_unified_push.sql，并保留部署后的 phone-role-push Edge Function、APNs 密钥和每分钟 cron。未部署 004 时统一 App 无法完成 APNs 自登记，后台开关不得宣称已接通。",
            "Windows 自动化 463/463 通过，app.js 语法与 git diff 校验通过。Mac 五 Target 编译、签名、HealthKit ECG API、Family Controls 真机控制及前台/后台/强退 APNs 仍待验证。",
        ],
    ),
    "AI开发项目_Bug修改规范.docx": (
        "新增真实读取与后台开关强制规范",
        [
            "角色查看真实 iPhone 时必须建立独立 readSessionId，只允许消费本次快照。展示限额、旧缓存、虚拟位置或其他数据源会造成模型误读，必须在进入模型前隔离，而不是只靠提示词。",
            "真实读取必须有有界超时并显示当前字段。不得让反向地理编码、HealthKit 或原生桥 Promise 无限挂起；部分权限失败时，一键读取应返回其余真实项目并明确列出不可用项。",
            "对可核验数字必须在生成后做一致性校验。模型回答中出现原始事实没有的数字或把限额当用量时，直接替换为确定性真实摘要。",
            "解锁提醒不得由普通 locked→unlocked 快照差异推断。只有原生管理操作或 owner-command 的 executed 回执可生成 explicit manualUnlock；角色命令、旧快照和系统不明变化必须排除。",
            "界面开关必须有端到端状态。后台主动联系至少核验控制器认领、APNs 令牌、角色 profile、cron 和最近推送状态；保存失败应回滚开关并显示原因，不能只点亮 UI。",
            "统一 App 自登记后台能力必须使用已认证手机号账号、Keychain 设备密钥和服务端 cloudId 归属校验；不得把 owner/device secret、APNs token 或服务端密钥写入仓库。",
            "完成声明继续以 Mac 编译和真机为界。Windows 的 463/463 不能代替 Xcode、HealthKit、Screen Time、通知扩展和强退 APNs 验收。",
        ],
    ),
    "AI开发项目_Bug记录模板.docx": (
        "Bug 记录：角色读错来源、读取悬挂、解锁开关失效和后台主动联系未登记",
        [
            "现象：角色报告 78 步、睡眠 0、抖音 120 分钟，而同一时刻真实面板为 497 步、睡眠 9 小时 46 分钟；顶部长期停在泛化“正在连接真实手机”；手动解锁提醒和 1 分钟后台主动联系开关开启后没有消息。",
            "根因一：角色屏幕事实把真实 usedSec 与 limitMin 同时交给模型，120 分钟限额被误当用量；通用角色上下文还能看到旧伴生缓存。位置缺失时又回退虚拟 curLoc。",
            "根因二：本机快照在前台顺序反查最多 8 个足迹地名，HealthKit 和 JS→Swift Promise 没有总超时，任一系统请求不返回就会停在同一句泛化进度。",
            "根因三：companionState 和 companionApplyServerPayloadV7 会主动过滤所有 manualUnlock 事件；候选发送又会被上一条用户消息和普通主动开关阻断，因此界面开关实际上没有可送达事件。",
            "根因四：私人 App 的本机直连故意绕过旧配对流程，但 APNs 令牌登记只在 CompanionSyncService.isPaired 时运行。网页虽显示本机 linked，服务器 phone_companion_links 中却可能没有这台统一 App 的 device secret/token，导致后台 profile 无法真正推送。",
            "修复：新增隔离 readSession、逐项/逐 App 进度、25 秒桥超时、12 秒健康超时、禁用前台地址反查、ECG 摘要、真实来源隔离和数值一致性守卫；限额不再进入角色用量事实。",
            "修复：原生管理页记录明确解锁事件，owner 本机命令只在 stage=executed 后记录；事件合并保留 delivered 状态，手动解锁提醒为必开且可越过普通主动消息限制。",
            "修复：新增 202608110004 认证迁移，让统一私人 App 同时认领 owner/device 身份并登记 APNs；新增 phone_role_push_status 展示 token/profile/cron/nextDue；profile 每 10 分钟自愈同步。",
            "自动化结果：node --check app.js 通过；完整 node --test 为 463/463。新增覆盖逐项进度、限额与用量隔离、错误数字拒绝、原生手动解锁事件、统一 APNs 认领及后台诊断。",
            "未验证：Windows 无 Xcode。须在 Mac 应用迁移 003/004、部署 Edge Function并配置 APNs 后，编译五 Target，在真机复测读取 497 类实时值、ECG 权限、解锁即时反应和后台/强退通知。",
        ],
    ),
    "AI开发项目_新聊天启动说明.docx": (
        "下一次聊天接手基线",
        [
            "当前候选基线：网页 v893、私人 iOS 1.0.18 (18)、原生桥契约 11。先读总纲和三份长期文档，再检查 v893 未完成的 Mac/真机验收。",
            "最高优先级是本次 readSession 的真实来源一致性。角色不得看到 limitMin、旧缓存、虚拟 curLoc 或内置网页计时；任何新数字都必须能在本次原生快照中找到。",
            "后台部署必须同时包含 202608110003 和 202608110004，并重新部署 phone-role-push。进入角色资料页点击“重新核验”，必须同时看到 APNs、profile、cron 已接通和合理 nextDue。",
            "主动间隔填 1 分钟不能马上收到测试消息，因为最后一次聊天后仍有随机 30–60 分钟硬静默。验收时应先观察页面显示的下一次最早检查时间，再测试后台和强退。",
            "手动解锁验收分两条：在原生设备管理页亲自解锁，以及在情侣空间 owner 控制页解锁。两者都应产生一次角色反应；角色自己下达解锁不得触发；重复快照不得重复提醒。",
            "真机顺序：部署迁移/Edge/APNs → Xcode 编译五 Target → 覆盖安装保留数据 → 允许通知/健康/定位/Family Controls → 一键读取逐项核对 → 解锁事件 → 前台、后台、强退主动消息和来电通知。",
            "当前 Windows 回归为 463/463；这不等于真机通过。若 Xcode 编译失败，先记录具体 Swift/HealthKit 编译错误再改，禁止直接宣称已完成。",
        ],
    ),
}

SUPPLEMENT = {
    "AI开发项目_项目说明文档.docx": "规则开关独立生效：失联电量与位置、5% 低电量和难过时心率不再依赖另一个主动消息总开关；难过文字直接把新鲜心率事实并入当前回复，避免开关开启后被普通回复队列挡住。",
    "AI开发项目_Bug修改规范.docx": "任何独立自动查看开关开启后，不得再被无关的主动消息总开关二次拦截。需要接住当前用户消息的规则（如难过时心率）必须进入该次正常回复，而不是等待另一个主动调度窗口。",
    "AI开发项目_Bug记录模板.docx": "补充修复：三个可选自动查看规则改为各自开关即授权触发；难过时心率走当前消息回复链。明确手动解锁事件保留 24 小时，若当时正在通话、线下或睡眠，条件解除后仍可补发一次。",
    "AI开发项目_新聊天启动说明.docx": "验收所有开关时，应分别制造条件，不再额外开启主动消息总开关：低电量≤5%、失联≥3小时且满足限频、难过关键词+20分钟内新鲜心率都必须进入各自可见动作。",
}


def append_release(path: Path, heading: str, paragraphs: list[str]) -> bool:
    doc = Document(path)
    if any(MARKER in paragraph.text for paragraph in doc.paragraphs):
        return False
    spacer = doc.add_paragraph()
    spacer.add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading(MARKER, level=1)
    doc.add_heading(heading, level=2)
    for text in paragraphs:
        doc.add_paragraph(text, style="List Bullet")
    doc.save(path)
    return True


def main() -> None:
    for filename, (heading, paragraphs) in DOC_CONTENT.items():
        path = ROOT / filename
        changed = append_release(path, heading, paragraphs)
        doc = Document(path)
        supplement = SUPPLEMENT[filename]
        supplemented = False
        if not any(supplement in paragraph.text for paragraph in doc.paragraphs):
            doc.add_paragraph(supplement, style="List Bullet")
            doc.save(path)
            supplemented = True
        state = "updated" if changed or supplemented else "unchanged"
        print(f"{state}: {path.name}")


if __name__ == "__main__":
    main()
