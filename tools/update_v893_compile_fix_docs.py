from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1] / "docs" / "maintenance"
MARKER = "2026-08-11 v893／私人 1.0.18 Xcode 编译修正 1"

ENTRIES = {
    "AI开发项目_项目说明文档.docx": [
        "Mac 首次编译 v893 时出现红色错误：PhoneNativeBridge.swift 无法访问 registerPushTokenIfAvailable，因为该方法在 CompanionSyncView.swift 中被声明为 fileprivate。",
        "修正为模块内默认访问级别，使同一 PhoneCompanionTest Target 内的 Swift 文件可以共同调用；功能逻辑、桥契约和版本号均不变。",
        "Windows 静态回归已增加跨文件可访问性断言；黄色的 iOS 26、iOS 17 弃用提示与 Family Controls Development 提示不阻塞本次编译，留待真机编译通过后分项处理。",
    ],
    "AI开发项目_Bug修改规范.docx": [
        "Swift 方法若需要被同一 Target 的其他源文件调用，不得使用 fileprivate；fileprivate 只允许当前文件访问。跨文件内部能力应使用默认 internal，并通过测试同时检查声明端和调用端。",
        "收到 Xcode 截图时必须先区分红色编译错误和黄色警告。先最小化修复阻塞项，不得把仅有 Windows 静态测试写成 Xcode 已编译通过。",
    ],
    "AI开发项目_Bug记录模板.docx": [
        "现象：Xcode 在 PhoneNativeBridge 报错，registerPushTokenIfAvailable 因 fileprivate 保护级别不可访问；下方其余项目均为黄色警告。",
        "根因：统一 App 认领控制器会从 PhoneNativeBridge.swift 调用 CompanionSyncView.swift 中的推送令牌登记方法，但该方法误用了仅限本文件的 fileprivate。",
        "修复：移除 fileprivate，保留默认 internal；新增回归测试，确认声明不再是 fileprivate 且跨文件调用仍存在。",
        "验证边界：Windows 测试通过不等于 Mac 编译通过；须由 Mac 再次编译五个 Target，确认红色错误消失后继续真机验收。",
    ],
    "AI开发项目_新聊天启动说明.docx": [
        "v893／私人 1.0.18 已追加 Xcode 编译修正 1：registerPushTokenIfAvailable 已从 fileprivate 改为同模块可访问。后续只能使用带“编译修正版”字样的第十八次安装包。",
        "下一步先在 Mac 覆盖整个 XcodeProject 并重新编译；若仍有红色错误，记录完整错误和文件行号再修。黄色弃用与 Family Controls Development 警告不能误记为本次红色阻塞项。",
    ],
}


def append_entry(path: Path, items: list[str]) -> bool:
    doc = Document(path)
    if any(MARKER in paragraph.text for paragraph in doc.paragraphs):
        return False
    doc.add_heading(MARKER, level=1)
    for item in items:
        doc.add_paragraph(item, style="List Bullet")
    doc.save(path)
    return True


def main() -> None:
    for filename, items in ENTRIES.items():
        path = ROOT / filename
        changed = append_entry(path, items)
        with ZipFile(path) as archive:
            assert archive.testzip() is None
            assert "[Content_Types].xml" in archive.namelist()
        doc = Document(path)
        marker_count = sum(MARKER in paragraph.text for paragraph in doc.paragraphs)
        assert marker_count == 1
        print(
            f"{'updated' if changed else 'unchanged'}: {path.name}; "
            f"marker={marker_count}; paragraphs={len(doc.paragraphs)}"
        )


if __name__ == "__main__":
    main()
