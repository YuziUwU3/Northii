from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1] / "docs" / "maintenance"
MARKER = "2026-08-11 v895／私人 1.0.20 全量读取、后台链路与彻底清空"

ENTRIES = {
    "AI开发项目_项目说明文档.docx": [
        "网页核心升级为 v895，私人 iOS 候选升级为 1.0.20 (20)，原生桥契约仍为 11。角色读取全部时必须逐项返回电量、屏幕总时长、逐 App、步数、睡眠、心率、心电、HRV 和位置；每项只能使用本次真实读取值或明确失败原因。",
        "电量与位置成功是两条独立链路：电量直接来自 UIDevice，不需要系统权限；位置来自 CoreLocation。此前健康失败是网页角色权限没有同步打开原生 healthSyncEnabled；此前屏幕时间失败是主 App 只尝试了要求 approvedWithDataAccess 的 iOS 26 直接 API。",
        "1.0.20 在明确健康读取时启用 HealthKit；普通 Family Controls approved 授权下挂载既有 DeviceActivityReport，由 Report 扩展把本次当日总时长和逐 App 数据写入 App Group，主 App 只接受同一请求编号的新快照。正式用户直接 Screen Time 数据访问受欧盟设备与欧盟 Apple Account 限制，因此中国真机不得依赖该接口。",
        "后台主动联系远端实际缺少本地已有的迁移 003／004，导致统一控制器 RPC 不存在且全天随机调度仍是旧签名。已实际部署 003–006，增加旧 p_apns_env 参数兼容和服务端记忆清空 RPC，并重新部署 phone-role-push；cron 仍按分钟执行。",
        "彻底清空现在等待本地大聊天归档覆写，清除共同生活 home、服务端 recent_context／memory_context 和未消费 outbox。任一步失败都提示未彻底完成，不得假报成功。网页版大容量存档将 IndexedDB 升级到 v2，确保对象仓库存在并进行三次短退避重试。",
        "网页端只发布记忆清空与大容量存档修复；App 专属锁定、HealthKit、Screen Time 和 APNs 路径均以原生桥可用为硬门槛。Windows 472 项测试通过，Mac 五 Target 编译签名及真机读取、后台、强退通知仍需验收。",
    ],
    "AI开发项目_Bug记录模板.docx": [
        "现象：角色读取全部时只得到位置或电量，顶部没有逐项读取横幅；健康、屏幕时间和逐 App 未读到。主动消息页报 Could not find the function public.claim_private_phone_unified_controller；彻底清空后仍记得旧事；网页版大容量存档保存失败。",
        "真实读取根因：电量和位置本来就是 UIDevice 与 CoreLocation 独立接口。网页角色健康开关没有启用原生 healthSyncEnabled。Screen Time 主 App 只尝试 iOS 26 直接 API，该 API 在正式安装中要求 approvedWithDataAccess，且受欧盟设备和欧盟 Apple Account 限制；现有 Report 扩展虽能显示数据，但角色读取没有消费它。",
        "主动消息根因：本地 SQL 文件存在不等于远端已部署。远端缺少 003／004，统一控制器 RPC 不存在，调度函数签名也落后。第一次补旧参数兼容迁移时，用位置参数调用重载产生函数歧义而失败；改为完整具名参数后成功，这是与失败尝试实质不同的修正。",
        "记忆根因：旧清空同步返回，没有等待 __messages 大归档覆写；recovery hydration 随后可重新合并旧归档；共同生活 homes、服务端 recent_context／memory_context 和未消费 outbox 也未清。大存档根因：已有 IndexedDB v1 但缺对象仓库时不再触发 upgrade，写入又没有重试。",
        "修复：HealthKit 明确读取时启用；DeviceActivityReport 按请求编号回传；读取全部逐项报告。远端部署 003–006 与 Edge Function。清空改为异步等待本地和服务端全部完成。IndexedDB 升 v2 并重试。App 专属行为增加原生桥硬门槛，网页不获得真实设备功能。",
        "验证：node --test tests/*.test.mjs 共 472／472 通过；远端已核验两个统一控制器重载、主动联系状态 RPC、记忆清空 RPC、分钟 cron 和 Edge Function 部署。Windows 无法替代 Xcode 编译和真实 iPhone 权限、后台、强退通知验证。",
    ],
    "AI开发项目_Bug修改规范.docx": [
        "不要把电量、位置、HealthKit、Screen Time 视为同一权限或同一数据源。必须逐类确认框架、授权状态、数据产生条件、地区限制和本次读取凭证。一个类别成功不能推导其他类别已接通。",
        "数据库迁移文件存在于仓库不等于远端已应用。遇到 PostgREST schema cache / function not found，必须直接核验远端函数签名、迁移历史、cron 和 Edge Function；重载函数内部应使用完整具名参数，避免默认参数造成歧义。",
        "DeviceActivityReport 回传必须携带本次请求编号与生成时间。主 App 不得消费请求前的旧共享快照，也不得在报告未执行时编造零值。地区受限的直接 API 必须有符合官方能力边界的替代路径。",
        "彻底清空是跨存储事务：当前对象、大聊天归档、恢复快照、共同生活上下文、服务端记忆和待发送队列必须全部纳入。完成提示必须等待所有承诺范围成功；失败时保留待重试标记并阻止旧上下文再次同步。",
        "共享 app.js 中的 App 专属行为必须用原生桥能力硬门槛隔离，不能只依靠文案或调用者约定。网页发布前要用测试证明普通网页路径不会进入 HealthKit、Screen Time、锁定或 APNs 分支。",
    ],
    "AI开发项目_新聊天启动说明.docx": [
        "当前候选基线：网页 v895、私人 iOS 1.0.20 (20)、原生桥契约 11。远端迁移 003–006 和 phone-role-push 已部署；先核验远端真实状态，不要因为本地有 SQL 就假定线上存在。",
        "读取全部验收必须逐项观察顶部进度和最终九类结果。电量与位置是独立成功链路；健康需要 HealthKit 授权和真实记录；中国正式安装的屏幕／逐 App 应由 DeviceActivityReport 回传，不应等待 approvedWithDataAccess。",
        "彻底清空后必须同时检查聊天归档、恢复候选、共同生活 home、服务端 profile 上下文和未消费 outbox。任一残留都可能让角色再次说出旧事。",
        "网页版只验收记忆清空和大容量存档修复，不得加入 App 锁定、HealthKit、Screen Time 或 APNs。App 专属分支必须以 companionLocalNativeAvailable 为硬门槛。",
        "Windows 自动测试当前 472／472 通过。Mac 只从 SmallPhone_v895_FullReadMemoryPushFix 全新目录打开工程，编译五 Target，并在真机验证 Report 扩展、HealthKit、前后台／强退通知、头像、主动来电和手动解锁告警。",
    ],
}


def main() -> None:
    for filename, items in ENTRIES.items():
        path = ROOT / filename
        document = Document(path)
        changed = not any(MARKER in paragraph.text for paragraph in document.paragraphs)
        if changed:
            document.add_heading(MARKER, level=1)
            for item in items:
                document.add_paragraph(item, style="List Bullet")
            document.save(path)
        with ZipFile(path) as archive:
            assert archive.testzip() is None
            assert "[Content_Types].xml" in archive.namelist()
        verified = Document(path)
        assert sum(MARKER in paragraph.text for paragraph in verified.paragraphs) == 1
        print(f"{'updated' if changed else 'unchanged'}: {path.name}")


if __name__ == "__main__":
    main()
