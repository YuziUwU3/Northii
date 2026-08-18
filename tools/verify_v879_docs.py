from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"

EXPECTED = {
    "AI开发项目_项目说明文档.docx": (
        "v879｜共同生活上下文与苹果全局适配（2026-08-10）",
        ["持续同居", "共同生活记忆", "Apple standalone", "425/425"],
    ),
    "AI开发项目_Bug记录模板.docx": (
        "v879 Bug 记录｜共同生活上下文错位、主动消息越界与苹果独立 App 偏移（2026-08-10）",
        ["主动发线上消息", "summaries", "425/425"],
    ),
    "AI开发项目_Bug修改规范.docx": (
        "新增强制规范｜共同生活、线下和微信必须以实时状态统一判定（v879 起）",
        ["实时状态", "screen.height", "手动删除"],
    ),
    "AI开发项目_新聊天启动说明.docx": (
        "新聊天接手状态｜v879 共同生活上下文与苹果全局适配发布路线（2026-08-10）",
        ["phone-work", "git push origin HEAD:main", "425/425"],
    ),
}


for filename, (title, required) in EXPECTED.items():
    document = Document(DOCS / filename)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    text = "\n".join(paragraphs)
    assert paragraphs.count(title) == 1, f"missing or duplicate heading in {filename}"
    assert paragraphs[-1], f"blank ending in {filename}"
    for phrase in required:
        assert phrase in text, f"missing {phrase!r} in {filename}"
    print(filename, f"paragraphs={len(paragraphs)}", f"tables={len(document.tables)}")

print("v879 maintenance documents verified")
