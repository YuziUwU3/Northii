from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_release(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    document = Document(path)
    if any(paragraph.text.strip() == title for paragraph in document.paragraphs):
        print(f"Skipped existing section: {title}")
        return
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"Updated {filename}")


append_release(
    "AI开发项目_项目说明文档.docx",
    "v860｜伴生推送上线与固定仓库交接（2026-08-10）",
    [
        "v860 基于已发布的 v859 main（83600d314b3d640b097b54e9beaf869e3408032b）。本次不改变产品功能，只把此前因 GitHub 网络失败而未发布的伴生推送部署记录安全补入最新 main，并按用户新规则统一缓存与页面版本号。",
        "v857 的网页与原生伴生修复已随提交 6ce5181 进入 main。线上 GitHub Pages 当时已读取“v857 · 伴生极速回执与持久快照”，其命令快速通道、controlOnly 快照合并和持久确认状态继续由 v859、v860 完整保留。",
        "Supabase Edge Function phone-companion-push 已由用户在 Dashboard 手动部署。带项目公开凭据发送空请求后，线上返回预期 HTTP 400 与 {\"error\":\"invalid-request\"}，证明函数路由和代码已运行，不再是 HTTP 404；部署过程没有重建、替换或显示 APNs secrets。",
        "GitHub 源码状态与 Supabase 部署状态必须分开描述：supabase/functions/phone-companion-push/index.ts 进入 main 只代表源码已发布；本次已有独立的线上函数探测结果，才可记录为 Supabase 函数已部署。",
        "从 v860 起唯一仓库为 C:\\Users\\pc\\Documents\\小手机\\phone-work，唯一分支为 main，远端为 origin/main。禁止创建新分支或 worktree，禁止使用 phone-release-web*、voice-test、Documents\\GitHub\\phone 等旧目录；GitHub Desktop 只用于查看。",
        "Windows 侧继续使用完整 Node 自动化回归验证。剩余设备工作仅为在 Mac Xcode 安装 v857 原生替换包，并完成真实 iPhone 的后台唤醒时延、锁定、解锁和回执闭环验收。",
    ],
)

append_release(
    "AI开发项目_Bug记录模板.docx",
    "v860 Bug 记录｜未发布部署记录迁入固定 main（2026-08-10）",
    [
        "现象：伴生推送函数已在 Supabase 上线并通过探测，但记录该结果的旧本地提交 2eb2ecb 因连接 github.com:443 失败没有进入远端；与此同时远端 main 已发展到 v859，旧文档仍写着 phone-companion-push 返回 HTTP 404。",
        "风险：直接推送旧分支、复制整份旧 DOCX 或合并旧 worktree，可能覆盖 v858、v859 已发布的共同生活、礼物、聊天草稿和主动消息防复读记录，也违反固定仓库和单一 main 路线。",
        "处理：在干净的 C:\\Users\\pc\\Documents\\小手机\\phone-work/main 上重新 fetch 并以 --ff-only 对齐 origin/main；只把旧提交的部署完成事实追加到当前四份维护资料，没有从旧目录提交、合并或覆盖文件。",
        "线上事实：此前函数路径返回 HTTP 404；用户手动部署 phone-companion-push 后，带公开项目凭据的空请求返回预期 HTTP 400 与 invalid-request。网页排队命令现在具备请求 APNs 唤醒的服务端入口。",
        "发布边界：本次没有数据库迁移，也没有重新部署任何 Supabase Edge Function。GitHub 本次发布的是文档、版本缓存与既有函数源码状态；Supabase 函数已部署来自此前独立完成并核验的 Dashboard 操作。",
        "验证：app.js 语法检查、伴生专项测试和全项目测试必须在提交前重新执行。真实 iPhone 仍需在 Mac Xcode 安装 v857 的两个 Swift 文件后进行后台锁定、解锁和快照回执验收。",
    ],
)

append_release(
    "AI开发项目_Bug修改规范.docx",
    "新增强制规范｜固定仓库单一 main 发布（v860 起）",
    [
        "唯一允许修改和发布的仓库是 C:\\Users\\pc\\Documents\\小手机\\phone-work，唯一分支是 main，唯一远端目标是 origin/main。禁止创建新分支、禁止创建 worktree，禁止使用 phone-release-web*、voice-test、Documents\\GitHub\\phone 等旧目录。",
        "GitHub Desktop 只用于查看，不需要也不得要求用户在其中手动合并或发布。一次只能有一个聊天修改和发布；发现另一个聊天正在推送时必须等待。",
        "开工前必须确认固定仓库没有未提交内容、当前分支为 main，然后 fetch origin/main 并仅以 fast-forward 对齐。发现未提交内容、分支不一致或无法快进时立即停止，不得 reset、覆盖、强推或绕到旧目录。",
        "修改完成后必须运行相关测试，把当前线上版本号加 1，直接提交到 main，再普通推送 origin/main。推送前或推送时发现远端新增提交，必须安全纳入远端改动并重新测试，禁止 force push。",
        "发布后必须报告新版本号、origin/main 完整提交号，并实际核验线上页面是否读取到新版本。未读取到时只能报告等待部署或缓存状态，不得把本地/远端版本当作线上已生效。",
        "涉及 Supabase Edge Function 时，必须分别记录源码是否进入 GitHub main、函数是否实际部署到 Supabase。只有独立部署记录或线上函数探测可以证明 Supabase 已部署。",
    ],
)

append_release(
    "AI开发项目_新聊天启动说明.docx",
    "新聊天接手状态｜v860 固定仓库与伴生推送已核验（2026-08-10）",
    [
        "唯一仓库：C:\\Users\\pc\\Documents\\小手机\\phone-work；唯一分支：main；远端：origin/main。禁止新分支、worktree 和所有旧目录。每次开工都要重新确认干净状态并 fetch/fast-forward 最新 main。",
        "v860 以远端 v859 83600d314b3d640b097b54e9beaf869e3408032b 为基线，只迁入此前未推送的伴生函数部署记录、修正发布路线，并统一版本缓存；不改 v859 的主动消息防复读功能。",
        "伴生服务器状态：所需三个 RPC 已在线；phone-companion-push 已在 Supabase 部署，空请求探测返回预期 400 invalid-request。不要重复部署函数，不要重建或显示 APNs secrets。GitHub 源码发布与 Supabase 函数部署必须分别报告。",
        "原生替换包仍位于 C:\\Users\\pc\\Documents\\小手机\\PhoneCompanion\\v857伴生极速回执与持久快照_微信传输包_2026-08-09。下一步仅需在 Mac Xcode 替换 CompanionSyncView.swift 和 PhoneCompanionTestApp.swift，编译并安装到原 iPhone。",
        "真机验收：伴生 App 退到后台但不要强制划掉；网页锁定真实 App，确认数秒内执行并收到回执后的快照；等待两分钟确认状态保持；再解锁；最后确认完整同步仍更新时长、位置、足迹和健康数据。",
    ],
)

print("Updated v860 maintenance documents")
