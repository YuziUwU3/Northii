from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_release(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    document = Document(path)
    if any(paragraph.text.strip() == title for paragraph in document.paragraphs):
        print(f"Skipped existing section: {title}")
        return
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"Updated {filename}")


append_release(
    "AI开发项目_项目说明文档.docx",
    "v859｜主动联系防复读与单一发布路线（2026-08-09）",
    [
        "v859 基于已发布的 v858 main（cb32c667fc7ad285e05c3a2f47e3efb1c973421d）。完整保留 v858 的共同生活状态同步、礼物卡修复、聊天草稿保护、伴生命令优先级和原生替换包，再加入主动联系独立事件与极短消息防重复。",
        "客户端不再对少于 8 个规范化字符的主动消息直接放行。所有非空文本先做完全匹配，2—7 个字符也参与包含和最长公共子序列相似度判断；候选只取 assistant 文字或语音，因此普通角色回复、本地主动消息和服务器回填消息统一比较，用户自己的文字不会误触发拦截。",
        "服务端在写入 outbox 和发送 APNs 前，除比较近期服务器主动消息外，还从 recent_context 提取同一角色的普通回复。极短原句或一字改写会被判重；模型沿用原有一次重试，仍重复则保持安静，不增加固定话术、随机句库或硬编码兜底。",
        "本地与服务器提示词都把主动联系定义为与上一轮分开的独立新事件。最近聊天只用于理解已发生的事实、关系和用户交代，不是等待继续回答的当前回合，禁止隔一段时间后补答、复述或只改几个字重发上一条回复。",
        "发布路线从 v859 起收敛为一个：C:\\Users\\pc\\Documents\\小手机 的 main 为唯一长期工作目录与正式发布分支。临时 worktree 只允许在用户明确要求并行开发时创建，发布成功后必须核对干净、确认已合并并立即移除；旧本地分支只保留未合并或有独有工作的记录。",
        "验证结果：app.js 与 phone-role-push TypeScript 语法检查通过；合并高风险专项 67 项通过；全项目实际枚举 377 项自动化测试全部通过，0 失败。测试期间临时 ESM 声明已删除，没有进入发布提交。",
        "本次没有新数据库迁移。v854 的 message_min、message_max 和 upsert RPC 已存在；phone-role-push 源码必须随 v859 部署后，线上服务器主动推送才会获得跨普通回复的防复读能力。",
    ],
)

append_release(
    "AI开发项目_Bug记录模板.docx",
    "v859 Bug 记录｜主动消息极短复读与跨路径重复（2026-08-09）",
    [
        "现象：角色普通聊天已经回复“嗯，去吧宝宝。”，约二十分钟后的主动联系仍可能原句重发，或只改成“嗯，去吧宝贝。”。这不是用户再次提问后的正常回答，而是主动消息把已结束的上一轮误当成待续回合。",
        "根因：客户端 initiativeRecentlyRepeated() 对规范化长度少于 8 的文本直接返回 false；服务端相似匹配从较长文本才生效，并且只和 outbox 的历史主动消息比较，没有在 APNs 前和 recent_context 中的普通角色回复比较。两端提示词也没有明确最近聊天已经结束。",
        "修复：客户端对所有长度执行完全匹配，并让 2—7 字文本按长度参与一字差相似判断；服务端采用同一阈值，把近期 outbox 与 recent_context 中该角色的普通回复合并为候选。两端均声明主动联系是独立新事件。",
        "安全边界：用户消息不进入角色回复候选，用户再次提出同一问题后的普通回复不走主动消息静默规则。重复候选不得用固定台词替换；服务器最多重试一次，仍重复就保持安静。",
        "合并说明：v859 以远端 v858 为底稿解决版本号、app.js、四份二进制 DOCX 和维护脚本冲突，保留共同生活、礼物、草稿和伴生设备修改。没有用旧 v857 文件覆盖 v858。",
        "防复发：覆盖极短完全重复、一字改写、不同主题不误伤、用户文字不误伤、普通回复进入服务端候选、独立新事件提示词，以及共同生活、礼物、聊天草稿和伴生设备合并专项。专项 67 项、全量 377 项全部通过。",
    ],
)

append_release(
    "AI开发项目_Bug修改规范.docx",
    "新增强制规范｜主动联系防重复与发布路线收敛（v859 起）",
    [
        "主动联系必须被提示为与上一轮分开的独立新事件。最近聊天只能作为已经发生的事实、关系与用户交代的背景，不得被解释为等待角色继续回答的当前回合。",
        "主动消息防重复不得按固定最小长度整体豁免。所有非空文本必须参与完全匹配；2—7 个规范化字符必须参与包含或相似度判断，并采用与长度相符的一字差阈值。",
        "判重必须跨投递路径统一：客户端比较普通角色回复、本地主动消息与服务器回填消息；服务器在 outbox/APNs 前比较近期 outbox 与 recent_context 中的普通角色回复。用户消息不得进入角色回复候选。",
        "重复候选不得通过固定话术、随机句库或硬编码台词替换。允许模型按相同人格约束重试，仍无法生成真实新事件时保持安静；服务器必须在发送 APNs 前完成判重。",
        "项目长期只保留一个正式发布路线：主工作目录 C:\\Users\\pc\\Documents\\小手机，分支 main。不得从 GitHub Desktop 随意 Publish 旧分支或把旧分支合入当前页面。",
        "任务需要隔离时可临时使用 worktree，但创建前必须从最新 origin/main 开始；发布完成后必须确认工作区干净、分支已合入 main，再移除 worktree 和对应已合并本地分支。脏工作区或未合并提交不得强制删除。",
        "每次发布必须先 fetch 最新 main、解决并验证并行改动、统一唯一版本号、运行高风险专项和完整测试，再推送 HEAD:main。发布后记录远端提交号、线上版本和 Edge Function 部署状态。",
    ],
)

append_release(
    "AI开发项目_新聊天启动说明.docx",
    "新聊天接手状态｜v859 单一发布路线（2026-08-09）",
    [
        "唯一长期工作目录：C:\\Users\\pc\\Documents\\小手机；唯一正式发布分支：main。GitHub Desktop 中其他 codex/* 分支和 phone-* worktree 均视为历史或临时项，不得直接 Publish、切换或合并。",
        "当前发布版本 v859，基于远端 v858 cb32c667fc7ad285e05c3a2f47e3efb1c973421d。v858 的共同生活、礼物、聊天草稿、伴生命令与原生替换包全部保留；v859 新增主动消息极短文本和跨路径防复读。",
        "关键代码：app.js；supabase/functions/phone-role-push/index.ts；tests/proactive-contact.test.mjs；tests/role-server-push.test.mjs。缓存号、Service Worker、中文主页面和版本测试均为 859。",
        "验证：语法检查通过；合并高风险专项 67 项通过；全量 377 项通过，0 失败。普通回复“嗯，去吧宝宝。”之后，本地主动联系和服务器主动推送都不得原句或一字改写复读，不同新话题仍应正常出现。",
        "服务器：本次无数据库迁移；phone-role-push 必须部署本地 v859 源码。若没有 Supabase 管理凭据，只能确认代码已进入 main，不能宣称线上 Edge Function 已更新。",
        "后续开发固定流程：在主工作目录 fetch origin/main，确认干净后直接从 main 工作；完成、测试、提交并推送 main。除非用户明确要求并行任务，否则不要创建新 worktree 或新发布分支。",
    ],
)

print("Updated v859 maintenance documents")
