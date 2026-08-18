from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1] / "docs" / "maintenance"
TITLE = "2026-08-12 v900：语音模型防误配"
EXPECTED = {
    "AI开发项目_项目说明文档.docx": ["网页核心升级为 v900", "三个层级", "479 项"],
    "AI开发项目_Bug记录模板.docx": ["speech-2.8-hd", "/v1/chat/completions", "旧错误配置原样保留", "479 项通过"],
    "AI开发项目_Bug修改规范.docx": ["模型用途必须", "TTS 必须保持类型分离", "不得静默删除"],
    "AI开发项目_新聊天启动说明.docx": ["当前网页基线为 v900", "语音模型 / TTS", "479／479"],
}

for filename, required in EXPECTED.items():
    document = Document(ROOT / filename)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    text = "\n".join(paragraphs)
    assert paragraphs.count(TITLE) == 1, f"missing or duplicate v900 heading in {filename}"
    assert paragraphs[-1], f"blank ending in {filename}"
    for phrase in required:
        assert phrase in text, f"missing {phrase!r} in {filename}"
    print(filename, f"paragraphs={len(paragraphs)}", f"tables={len(document.tables)}")

print("v900 maintenance documents verified")
