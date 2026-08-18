from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
EXPECTED = {
    "AI开发项目_项目说明文档.docx": ["v854｜自然主动联系", "367 项自动化测试全部通过"],
    "AI开发项目_Bug记录模板.docx": ["v854 Bug 记录", "口头密码回归"],
    "AI开发项目_Bug修改规范.docx": ["新改动不得碰坏旧功能", "不得通过删测试"],
    "AI开发项目_新聊天启动说明.docx": ["v854 待发布整合包", "202608090001_phone_role_push_natural_messages.sql"],
}

for name, markers in EXPECTED.items():
    path = DOCS / name
    with ZipFile(path) as archive:
        assert archive.testzip() is None, f"corrupt zip member in {name}"
        assert "word/document.xml" in archive.namelist(), f"missing document.xml in {name}"
    document = Document(path)
    text = "\n".join(p.text for p in document.paragraphs)
    for marker in markers:
        assert marker in text, f"missing marker {marker!r} in {name}"
    assert len(document.paragraphs) >= 10, f"unexpectedly short document {name}"
    print(f"verified {name}: {len(document.paragraphs)} paragraphs, {len(document.tables)} tables")
