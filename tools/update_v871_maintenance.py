from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_release(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    document = Document(path)
    if any(paragraph.text.strip() == title for paragraph in document.paragraphs):
        print(f"Skipped existing section: {title}")
        return
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"Updated {filename}")


append_release(
    "AI开发项目_项目说明文档.docx",
    "v871｜共同生活模型与 API 路线分离（2026-08-10）",
    [
        "共同生活设置把“使用哪个模型”和“从哪条 API 路线请求”拆成两个互不替代的维度。界面分别提供对话模型、对话 API 路线、总结模型、总结 API 路线四个选择器；主模型／副模型不再被错误命名为主线路／备用线路。",
        "对话模型选择“跟随微信角色”时，读取角色当前主／副模型偏好；对话 API 路线选择“跟随微信当前路线”时，读取微信当前激活路线。用户也可固定路线一至路线四，固定值只随本次共同生活请求传入，不调用全局路线切换，不改变微信当前路线。",
        "总结模型与总结 API 路线独立配置。总结可跟随对话模型、强制主模型或副模型；API 路线可跟随对话路线、跟随微信当前路线或固定路线一至路线四。辅助模型回退到主模型时仍留在同一条已选 API 路线上。",
        "旧数据无损迁移：历史 replyRoute 自动迁移为 replyModel，历史 summaryRoute 自动迁移为 summaryModel；新增 replyApiRoute 默认跟随微信当前路线，summaryApiRoute 默认跟随对话路线。迁移不覆盖用户已经保存的新字段。",
        "验证结果：app.js 语法检查通过；迁移、独立解析、界面分区、固定路线不修改微信、上下文承接、会话内路线和回退链路等专项测试通过；完整 Node 自动化回归 412/412 通过。",
    ],
)

append_release(
    "AI开发项目_Bug记录模板.docx",
    "v871 Bug 记录｜共同生活把模型与 API 路线混成一个选择（2026-08-10）",
    [
        "现象：共同生活设置把“跟随角色设置／主线路／备用线路”放在同一个选择器里，用户无法判断它是在选主副模型，还是在选微信保存的 API 路线。共同生活也缺少独立固定 API 路线的入口。",
        "根因：replyRoute 与 summaryRoute 同时承担模型偏好和路由含义，调用 chatAPI 时只传 aux，没有传请求级 routeIndex，导致界面语义和实际请求路由无法一一对应。",
        "修复：字段拆为 replyModel、replyApiRoute、summaryModel、summaryApiRoute；chatAPI 新增可选 routeIndex，在请求局部读取指定路线的主／副配置。固定共同生活路线不调用 chatRouteApply，也不写回 S.settings.chatRouteActive。",
        "兼容：读取设置时只在新字段不存在时迁移历史 replyRoute／summaryRoute。固定路线的副模型不可用时，退回同一路线主模型，而不是漂移到微信当前路线。",
        "防复发：新增共同生活模型／路线专项测试，覆盖旧字段迁移、模型与 API 路线独立解析、四块设置界面以及固定路由不污染微信全局状态；完整回归 412/412 通过。",
    ],
)

append_release(
    "AI开发项目_Bug修改规范.docx",
    "新增强制规范｜模型选择与 API 路线必须正交（v871 起）",
    [
        "产品文案必须明确区分模型和 API 路线：主模型／副模型表示同一条路线内使用哪组模型配置；路线一至路线四表示请求所使用的接口、Key 与模型配置集合。禁止把主模型写成主线路、把副模型写成备用线路。",
        "子系统若允许固定 API 路线，必须通过请求级 routeIndex 或等价局部参数完成。禁止为了发送一次共同生活请求而修改 S.settings.chatRouteActive、调用 chatRouteApply 或永久切换微信当前路线。",
        "模型回退必须保持路由稳定：指定路线的副模型不可用时，只能退回指定路线的主模型；不得悄悄改用微信当前路线。总结模型与总结路线也必须分别配置和分别解析。",
        "旧字段迁移必须幂等且无损：仅当新字段缺失时读取旧字段，保存后以后以新字段为准。测试至少覆盖跟随角色、主副模型、跟随微信、固定路线、总结跟随、回退同路由和全局路线不变。",
        "任何测试、文档、日志与截图不得写入真实 API Key。涉及路线的回归必须同时验证用户可见标签和实际请求配置，不能只检查下拉框文字。",
    ],
)

append_release(
    "AI开发项目_新聊天启动说明.docx",
    "新聊天接手状态｜v871 共同生活模型与 API 路线分离（2026-08-10）",
    [
        "固定仓库 C:\\Users\\pc\\Documents\\小手机\\phone-work，固定分支 main，远端 origin/main。禁止新分支、worktree、旧目录和强推；开始前 clean、fetch、ff-only，发布前测试，线上最新版本基础上加 1。",
        "共同生活字段为 replyModel、replyApiRoute、summaryModel、summaryApiRoute。模型值为 follow／main／aux；API 路线值为 follow 或 0 至 3，总结路线另允许 reply 表示跟随对话路线。",
        "关键函数包括 chatRequestRoute、wechatAuxConfigured、cohabReplyAux、cohabReplyRouteIndex、cohabSummaryAux、cohabSummaryRouteIndex。chatAPI 的 opt.routeIndex 只决定本次请求读取哪条保存路线，不能修改微信当前路线。",
        "界面必须保留四个清晰区域：对话模型、对话 API 路线、总结模型、总结 API 路线。跟随角色设置同时继承角色的模型偏好和微信当前 API 路线，但两个选择器仍可分别改成固定值。",
        "v871 验证基线为完整 Node 412/412。后续修改至少运行 cohabitation-model-route-settings、cohabitation-inline-settings、cohabitation-context-handoff、in-session-api-routes、wechat-fallback-routing 以及完整 tests/*.test.mjs。",
    ],
)

append_release(
    "AI开发项目_项目说明文档.docx",
    "v871 补充｜共同生活查看手机只交付一次回复（2026-08-10）",
    [
        "共同生活里角色通过隐藏标签决定查看用户手机时，查看前的原始模型回合只负责启动查看，不再把其中的旁白或台词作为可见回复写入。界面等待查看进度完成后，只交付拿到真实后台记录的结果回合。",
        "外层共同生活发送流程会等待 cohabRunPhoneInspection 完成，期间不提前解除忙碌状态；结果回合已经自行写入共同生活记录并触发总结后，外层不再重复写入或重复触发总结。相同事实未变化时继续静默去重。",
        "新增回归验证了 schedule:false 不会再启动平行定时任务、原始回合被消费、结果查看被等待、无内容提示不会误报、总结不会执行两次。专项共同生活测试 22/22 通过。",
    ],
)

append_release(
    "AI开发项目_Bug记录模板.docx",
    "v871 Bug 记录｜共同生活查看手机后同时回复两次（2026-08-10）",
    [
        "现象：在线下共同生活聊天里，角色查看用户手机后台记录时，会先回复一次用户原话，查看完成后又根据真实记录回复一次，形成同一回合两次回答。微信线上链路没有这个问题。",
        "根因：cohabApplyPhoneTags 使用 setTimeout 独立启动查看，但 cohabReplyCore 同时把去掉标签后的原始文本作为普通回复返回；offAI 先落原始回复，随后 cohabPhoneDeliverFact 又落结果回复。",
        "修复：offAI 成为这类查看标签的唯一流程所有者。原始回合检测到查看标签后清空可见 items，并同步等待 cohabRunPhoneInspection；只有结果回合可写入。外层通过 inspection 标记跳过无内容提示和重复总结。",
        "防复发：测试必须同时断言没有平行定时查看、原始回合不落库、结果查看被 await、无内容提示被抑制以及总结只触发一次。",
    ],
)

append_release(
    "AI开发项目_Bug修改规范.docx",
    "新增强制规范｜异步查看动作必须单回合交付（v871 起）",
    [
        "模型回复若包含“先执行真实查看，再依据结果发言”的隐藏动作标签，动作前回合与结果回合只能有一个成为用户可见回复。默认由拿到真实事实的结果回合负责交付。",
        "禁止一边 setTimeout 启动后台动作，一边把原始模型正文照常落库。调用方必须明确拥有该动作、等待其完成，并用显式 inspection／deferred 标记跳过普通空回复提示、重复总结和重复通知。",
        "事实未变化导致结果回合静默时，不得回放动作前台词来填空；更不得编造一条新记录。线上微信、旧约会和远控链路不得因为共同生活去重修复而改变。",
    ],
)

append_release(
    "AI开发项目_新聊天启动说明.docx",
    "新聊天接手补充｜共同生活手机查看单回合规则（2026-08-10）",
    [
        "共同生活手动对话中的查看标签由 offAI 持有：cohabReplyCore 使用 inspectionOwner:'offAI' 解析后返回 inspection，不落查看前 items；offAI await cohabRunPhoneInspection 后才结束本轮。",
        "cohabApplyPhoneTags 的 opt.schedule=false 表示调用方将同步接管查看，不能再建立 700ms 平行任务。cohabPhoneDeliverFact 是唯一可见结果回复来源，并负责结果落库、事实提交与总结。",
        "回归入口为 tests/cohabitation-phone-inspection.test.mjs，必须保留“a co-living reply-tag turn produces only the completed inspection reaction”测试，并继续跑完整 tests/*.test.mjs。",
    ],
)

print("Updated v871 maintenance documents")
