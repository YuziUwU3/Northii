from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
EXPECTED = {
    "AI开发项目_项目说明文档.docx": (
        "v866｜共同生活记忆与 iOS 底部适配（2026-08-10）",
        ["S.cohabitation.homes[角色ID].summaries", "392/392", "-webkit-fill-available"],
    ),
    "AI开发项目_Bug记录模板.docx": (
        "v866 Bug 记录｜共同生活设置不可随手调整、记忆边界不清及部分 iOS 底部黑带（2026-08-10）",
        ["d.summaries", "WebKit standalone portrait", "visualViewport"],
    ),
    "AI开发项目_Bug修改规范.docx": (
        "新增强制规范｜长期场景设置、独立记忆库与设备定向视口修复（v866 起）",
        ["禁止复制进微信长期记忆", "分段不得为了达到条数而重复或编造", "禁止恢复全局 visualViewport"],
    ),
    "AI开发项目_新聊天启动说明.docx": (
        "新聊天接手状态｜v866 共同生活专属设置、独立记忆与 iOS 底部补齐（2026-08-10）",
        ["summaryMode", "不得把这些总结复制到微信记忆", "390×844"],
    ),
}


for filename, (marker, required) in EXPECTED.items():
    document = Document(DOCS / filename)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    text = "\n".join(paragraphs)
    assert paragraphs.count(marker) == 1, f"duplicate or missing v866 section in {filename}"
    assert paragraphs[-1], f"blank ending in {filename}"
    for phrase in required:
        assert phrase in text, f"missing {phrase!r} in {filename}"
    print(
        filename,
        f"paragraphs={len(paragraphs)}",
        f"tables={len(document.tables)}",
        f"last={paragraphs[-1][:50]}",
    )

print("v866 maintenance documents verified")
