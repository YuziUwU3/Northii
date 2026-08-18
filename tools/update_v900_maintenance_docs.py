from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1] / "docs" / "maintenance"
MARKER = "2026-08-12 v900：语音模型防误配"

ENTRIES = {
    "AI开发项目_项目说明文档.docx": [
        "网页核心升级为 v900；私人 iOS 原生代码未修改，仍为 1.0.24 (24)，原生桥契约仍为 11。聊天模型与 TTS 继续使用各自独立的配置和接口。",
        "主聊天和辅助模型的自由文本设置新增用途校验：speech-*、tts-* 及名称中作为独立段出现的 tts 标记被识别为语音合成专用模型，不能再作为文字聊天模型保存、测试或调用。",
        "防护覆盖三个层级：设置页就近保存和底部保存会拒绝误配；模型测试在联网前拒绝；chatAPI 在真实请求前再次拒绝。旧存档中的错误路线不会被自动删除，但启用或调用时会给出明确提示。",
        "正常聊天模型、API 四路线、辅助模型、共同生活固定路线和原有 TTS 生成链保持不变。Windows 全套 479 项自动测试全部通过。",
    ],
    "AI开发项目_Bug记录模板.docx": [
        "现象：部分用户把语音合成模型 speech-2.8-hd 填入主聊天或辅助模型。小手机随后把角色总人设和 Prompt 作为 messages 发送到 /v1/chat/completions；中转站旧行为可能仍把请求路由到 TTS 并产生费用，文字对话却不能正常工作。",
        "根本原因：聊天和 TTS 的运行代码本来分开，但聊天模型名是无类型自由文本。路线保存、就近保存、模型测试和 chatAPI 请求前都只检查地址、Key 与非空，没有检查模型用途，因此用户误填可以穿过整条文字聊天链。",
        "影响范围：误填主模型会影响微信等主模型功能；误填辅助模型会影响购物、外卖、X、查他手机等副模型功能。没有误填语音模型的用户不受影响，小手机也不会自动把 TTS 模型复制到聊天设置。",
        "修复：新增统一语音专用模型识别；保存和路线切换前拒绝并定位错误输入框；测试模型时在网络请求前拒绝；chatAPI 对当前主模型、辅助模型和固定路线做最终断言。旧错误配置原样保留，用户改正前只阻止调用，不静默改写模型名。",
        "验证：专项覆盖 speech-2.8-hd、speech-02-turbo、tts-1-hd、gpt-4o-mini-tts、正常 GPT/Claude 模型、主辅路线保存、旧路线启用、测试联网前拦截及会话内固定路线；node --test tests/*.test.mjs 共 479 项，479 项通过。",
    ],
    "AI开发项目_Bug修改规范.docx": [
        "模型用途必须在客户端和服务端双层校验。可自由填写模型名的界面不能只检查非空；至少在保存、测试和最终请求前确认该模型适用于当前接口。",
        "聊天与 TTS 必须保持类型分离。speech-*、tts-* 或明确带独立 tts 标记的语音合成模型不得进入 /chat/completions；语音模型只能进入 TTS 专用配置和接口。",
        "不能因为上游接口接受了错误模型就认定调用正确。中转服务可能兼容路由并计费；客户端必须在联网前阻止已知用途冲突，避免无效费用和 Prompt 被送往错误能力。",
        "发现旧存档误配时应阻止启用和调用并给出可行动提示，不得静默删除、替换或猜测用户想用的聊天模型。测试必须同时覆盖新保存、旧路线、主模型、辅助模型和最后请求门。",
    ],
    "AI开发项目_新聊天启动说明.docx": [
        "当前网页基线为 v900；私人 iOS 仍为 1.0.24 (24)，原生桥契约 11。本轮只修改网页共享核心的模型用途校验，没有修改 Swift、签名、HealthKit、Screen Time 或 APNs。",
        "聊天和 TTS 仍是两套独立路线。speech-2.8-hd 等语音合成模型只能填写在“语音模型 / TTS”；主聊天或辅助模型误填时，保存、测试和真实请求都会被拦截。",
        "若用户反馈角色不回复或购物、外卖、X、查他手机同时失败，先核对当前 API 路线的主模型和辅助模型是否含 speech 或独立 tts 标记，再检查地址、Key、余额和上游错误。",
        "禁止通过自动替换模型名修复旧数据。旧误配路线应保留供用户辨认，只阻止启用或调用并提示改到 TTS 栏。Windows 自动测试当前 479／479 通过。",
    ],
}


def main() -> None:
    for filename, items in ENTRIES.items():
        path = ROOT / filename
        document = Document(path)
        changed = not any(MARKER in paragraph.text for paragraph in document.paragraphs)
        if changed:
            document.add_heading(MARKER, level=1)
            for item in items:
                document.add_paragraph(item, style="List Bullet")
            document.save(path)
        with ZipFile(path) as archive:
            assert archive.testzip() is None
            assert "[Content_Types].xml" in archive.namelist()
        verified = Document(path)
        assert sum(MARKER in paragraph.text for paragraph in verified.paragraphs) == 1
        print(f"{'updated' if changed else 'unchanged'}: {path.name}")


if __name__ == "__main__":
    main()
