from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
EXPECTED = {
    "AI开发项目_项目说明文档.docx": ["v857｜伴生极速回执与持久快照", "376 项自动化测试通过"],
    "AI开发项目_Bug记录模板.docx": ["后台同步慢、快照过期与锁定回执等待", "phone-companion-push 仍为 HTTP 404"],
    "AI开发项目_Bug修改规范.docx": ["伴生命令与动态快照必须分层", "controlOnly"],
    "AI开发项目_新聊天启动说明.docx": ["v857 本地修复完成", "服务器函数待部署"],
}

for name, markers in EXPECTED.items():
    path = DOCS / name
    with ZipFile(path) as archive:
        assert archive.testzip() is None, f"corrupt zip member in {name}"
        assert "word/document.xml" in archive.namelist(), f"missing document.xml in {name}"
    document = Document(path)
    text = "\n".join(p.text for p in document.paragraphs)
    for marker in markers:
        assert marker in text, f"missing marker {marker!r} in {name}"
    assert len(document.paragraphs) >= 15, f"unexpectedly short document {name}"
    print(f"verified {name}: {len(document.paragraphs)} paragraphs, {len(document.tables)} tables")
