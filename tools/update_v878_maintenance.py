from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_release(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    document = Document(path)
    if any(paragraph.text.strip() == title for paragraph in document.paragraphs):
        print(f"Skipped existing section: {title}")
        return
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"Updated {filename}")


append_release(
    "AI开发项目_项目说明文档.docx",
    "v878｜苹果主屏幕安全区手动适配（2026-08-10）",
    [
        "本次在 v877 固定视口基线上增加“设置 → 偏好·工具 → 偏好 → 苹果主屏幕适配”开关。默认关闭，不改变现有 Safari 与安卓布局；用户只在 iPhone/iPad 添加到主屏幕后遇到灵动岛遮挡或底部安全区缺口时开启。",
        "运行范围采用双重门禁：必须识别为 iPhone/iPad，并且处于 standalone 主屏幕独立模式，同时用户已手动开启。安卓即使处于 PWA standalone 或保存了同名偏好，也不会获得 north-ios-home-safe 类；普通 iPhone Safari 同样不会命中。",
        "适配只使用标准 CSS 100dvh 与 env(safe-area-inset-top/bottom)：根外壳在该模式填满动态视口，主屏时间、通用导航避开顶部安全区，聊天输入栏避开底部安全区。未恢复 screen.height、物理屏幕高度比较、visualViewport/resize 写高或 -webkit-fill-available。",
        "专项测试覆盖 iPhone 主屏幕命中、iPhone Safari 不命中、Android standalone 不命中、iPad 桌面 UA 命中，并检查顶部与底部安全区规则；完整 Node 回归 424/424 通过。",
    ],
)

append_release(
    "AI开发项目_Bug记录模板.docx",
    "v878 Bug 记录｜iOS 主屏幕灵动岛遮挡与底部黑区（2026-08-10）",
    [
        "现象：同一页面在 iPhone Safari 中正常；添加到主屏幕后，首页时间进入状态栏/灵动岛区域，底部壁纸结束过早并露出黑区。安卓原本正常，但使用跨平台物理屏幕高度补齐会把安卓聊天页拉长，使输入框落到屏幕外。",
        "判断：viewport-fit=cover 会让 WebApp 使用边到边视口，重要内容应显式使用安全区；新系统把添加到主屏幕后的独立 WebApp 行为变得更常见，因此原本未处理的安全区问题暴露。问题不能按全平台统一高度修复。",
        "修复：新增用户可控的苹果主屏幕适配开关，并在 JS 中同时验证 Apple 设备和 standalone 环境。开启后只通过 100dvh 与 safe-area-inset 调整外壳、主屏头部、导航和输入栏；Safari 与安卓保持 v877 布局。",
        "失败方案记录：禁止再次使用 screen.height 或给所有 standalone 设备套相同高度；禁止仅凭触摸设备、display-mode 或屏幕尺寸推断 iOS；禁止让安全区开关在安卓生效。完整回归 424/424 通过。",
    ],
)

append_release(
    "AI开发项目_Bug修改规范.docx",
    "新增强制规范｜苹果安全区适配必须手动、独立且不影响安卓（v878 起）",
    [
        "涉及 iOS 主屏幕安全区的修复必须同时满足 Apple 设备、standalone 主屏幕模式和用户显式开启三项条件；普通 Safari、安卓浏览器与安卓 PWA 必须沿用原布局。",
        "允许使用 100dvh 和 env(safe-area-inset-*)；禁止使用 screen.height、物理屏幕与 innerHeight 差值、全局 visualViewport/resize 写高、-webkit-fill-available 或触摸设备泛化判断。",
        "适配开关必须持久化，关闭后立即移除 north-ios-home-safe；新增顶部或底部补偿时要同时检查主屏、普通导航页和聊天输入栏，且不得破坏 v877 的一屏外壳与内部滚动契约。",
        "每次发布至少自动验证 iPhone standalone、iPhone Safari、Android standalone、iPad 桌面 UA 四种判定，以及版本缓存对齐和完整回归。",
    ],
)

append_release(
    "AI开发项目_新聊天启动说明.docx",
    "新聊天接手状态｜v878 苹果主屏幕适配发布路线（2026-08-10）",
    [
        "固定仓库 C:\\Users\\pc\\Documents\\小手机\\phone-work，固定 main 与 origin/main。只在该目录修改、测试、提交并执行 git push origin HEAD:main；禁止建立新分支、worktree、旧版本目录、强推或覆盖远端。",
        "当前基线为 v878 · 苹果主屏幕安全区开关。设置页已有“苹果主屏幕适配”：默认关闭；仅 iPhone/iPad 添加到主屏幕后由用户开启，普通 Safari 和所有安卓环境不生效。",
        "不得恢复 v873 的 screen.height/north-standalone-shell，也不得把苹果补偿扩展给所有 standalone。v877 的根外壳最多一屏、chatbg 内滚动、inputbar 固定占位继续是不可破坏的稳定契约。",
        "发布前运行完整 tests/*.test.mjs，确认版本、缓存和 Service Worker 全部对齐；维护文档四件套继续追加。v878 完整 Node 回归为 424/424。",
    ],
)

print("Updated v878 maintenance documents")
