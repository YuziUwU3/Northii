from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1] / "docs" / "maintenance"
MARKER = "2026-08-12 v897／私人 1.0.22 读取完成后只自然回复一次"

ENTRIES = {
    "AI开发项目_项目说明文档.docx": [
        "网页核心升级为 v897，私人 iOS 候选升级为 1.0.22 (22)，原生桥契约仍为 11。聊天、通话和共同生活共用同一条规则：用户提出真实手机读取后，读取任务必须在普通角色模型之前取得本轮回复权。",
        "读取期间只允许顶部横幅逐项显示当前读取类别。所有请求读取器进入成功、不可用或超时终态并写入同一 readSessionId 的完成回执后，才把成功事实作为隐藏上下文交给角色生成一轮自然回复。",
        "v896 的固定九项汇总已撤销。不可用、未授权、超时等信息继续保存在内部 readOutcomes 中用于诊断，但不得变成角色气泡里的系统清单、技术提示或逐项失败报告；确定性兜底也必须是一段自然口吻。",
        "同一条用户消息使用 pending 键保护，读取任务尚未结束时不得再次进入普通聊天、通话或第二次原生读取回复，从而消除先答一轮、查完再答一轮的重复出口。",
        "Windows 全套 475 项测试通过，包含提前拦截、完成回执、自然回复守卫、系统诊断隐藏和并发去重。Mac 五 Target 编译、签名及真实 iPhone 一次全量读取仍需最终验收。",
    ],
    "AI开发项目_Bug记录模板.docx": [
        "现象：读取全部时，角色先正常回复一次；读取完成后又把电量、屏幕、逐 App、步数、睡眠、心率、心电、HRV 和位置逐条显示成多条角色消息，其中还夹带‘本次没有返回这项健康数据’等系统诊断。",
        "真实原因一：v896 为保证九类可核对，在 doSpyViewCore 和 cohabPhoneDeliverFact 中用 rolePhoneInspectionExactSummary 强制覆盖模型的自然回复。该汇总原本是内部验收报告，却直接进入聊天气泡，所以看起来像角色逐句念系统清单。",
        "真实原因二：微信和通话直到普通模型已经生成回复后才调用 nativeOnly 读取识别。虽然可清掉部分可见内容，但该轮模型链已运行，并把被压制文本继续当作 alreadySaid 传递，形成两个回复通道和前后重复。",
        "修复：在 wechatPrimaryReply 和通话 chatAPI 之前，用用户原话提前识别本机读取意图并立即返回；同一消息写入 pending 集合。完成后仅把成功事实交给角色生成一轮自然回复，并用数值一致性和技术措辞守卫检查；不安全时使用单段自然兜底。",
        "验证：node --test tests/*.test.mjs 共 475／475 通过。Windows 静态验证不能替代真实 iPhone 上的模型语气、HealthKit、DeviceActivityReport 和后台通知验收。",
    ],
    "AI开发项目_Bug修改规范.docx": [
        "任何需要等待外部／原生异步数据的用户请求，都必须在普通模型调用前取得回复权。不得先让模型生成，再依靠删除气泡、覆盖文本或 alreadySaid 参数补救；那会保留费用、上下文污染和重复出口。",
        "内部验收报告与角色可见回复必须分层。readOutcomes、权限、超时、缺字段和系统诊断只能用于守卫、日志和开发界面；角色气泡只能接收成功事实和符合人设的自然表达。",
        "禁止用硬编码逐项报告覆盖已通过事实校验的自然模型输出。确定性兜底只用于模型编造数值、遗漏已成功类别或泄漏技术提示时，且兜底本身必须是单段自然口吻。",
        "同一用户事件只能存在一个异步回复出口。需要用稳定事件键记录 pending，并在成功、失败和异常的 finally 路径中清理；聊天、通话和共同生活都必须遵守。",
    ],
    "AI开发项目_新聊天启动说明.docx": [
        "当前候选基线：网页 v897、私人 iOS 1.0.22 (22)、原生桥契约 11。远端 202608110002–006 已核验应用；本轮未把 APP 专属能力推送到普通网页版。",
        "读取全部的正确时序：用户消息先被原生读取入口拦截；顶部横幅逐项更新；同一 readSessionId 完成；成功事实进入隐藏上下文；角色只自然回复一轮。看到逐条九项角色消息、权限诊断或前后两轮回复即视为回归。",
        "v896 的失败方案是完成后强制 rolePhoneInspectionExactSummary，并在普通模型生成后才识别读取意图。以后不得恢复该硬覆盖，也不得把读取识别移回模型调用之后。",
        "排查重复回复时先检查 pending 键、早期 return 和所有调用出口；排查系统话术泄漏时检查 readErrors、missing、readOutcomes 是否进入可见 prompt 或 deterministic fallback。不要只改提示词。",
        "Windows 自动测试当前 475／475 通过。Mac 只从 SmallPhone_v897_NaturalSingleReadReply 全新目录打开工程，编译五 Target，并在真机验证一次全量读取只得到一轮自然回复。",
    ],
}


def main() -> None:
    for filename, items in ENTRIES.items():
        path = ROOT / filename
        document = Document(path)
        changed = not any(MARKER in paragraph.text for paragraph in document.paragraphs)
        if changed:
            document.add_heading(MARKER, level=1)
            for item in items:
                document.add_paragraph(item, style="List Bullet")
            document.save(path)
        with ZipFile(path) as archive:
            assert archive.testzip() is None
            assert "[Content_Types].xml" in archive.namelist()
        verified = Document(path)
        assert sum(MARKER in paragraph.text for paragraph in verified.paragraphs) == 1
        print(f"{'updated' if changed else 'unchanged'}: {path.name}")


if __name__ == "__main__":
    main()
