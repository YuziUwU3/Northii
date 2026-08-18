from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
EXPECTED = {
    "AI开发项目_项目说明文档.docx": (
        "私人「小手机」真实 App 总纲与基础工程（2026-08-11）",
        ["固定叫「小手机」", "controllerKind", "命令已发送", "native/private-small-phone"],
    ),
    "AI开发项目_Bug记录模板.docx": (
        "私人 App 转向风险记录｜双控制器、工程缺失与方向误解（2026-08-11）",
        ["单控制器租约", "0x8BADF00D", "尚未完成"],
    ),
    "AI开发项目_Bug修改规范.docx": (
        "新增强制规范｜私人 App 单一总纲与单设备单控制器（2026-08-11 起）",
        ["禁止同时控制同一台设备", "未知 action 默认拒绝", "禁止猜写"],
    ),
    "AI开发项目_新聊天启动说明.docx": (
        "新聊天最高优先级｜私人「小手机」真实 App 已启动（2026-08-11）",
        ["旧文中", "完整工程", "不得把 Node 测试通过"],
    ),
}


for filename, (heading, phrases) in EXPECTED.items():
    document = Document(DOCS / filename)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    text = "\n".join(paragraphs)
    assert paragraphs.count(heading) == 1, filename
    for phrase in phrases:
        assert phrase in text, f"{filename}: missing {phrase}"
    print(filename, f"paragraphs={len(paragraphs)}", f"tables={len(document.tables)}")

print("Private small-phone maintenance documents verified")
