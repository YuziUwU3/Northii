from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_release(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    document = Document(path)
    if any(paragraph.text.strip() == title for paragraph in document.paragraphs):
        print(f"Skipped existing section: {title}")
        return
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"Updated {filename}")


append_release(
    "AI开发项目_项目说明文档.docx",
    "v865｜主屏时间隐藏保持原布局（2026-08-10）",
    [
        "v865 基于已发布的 v864 main（f6d3a10db1277bd1fccf7dd39a0ac842a8add150）。完整保留 v864 的共同生活长期上下文、到家承接、伴生设备和此前全部功能；本次只修复主屏时间与日期关闭后的布局回流。",
        "偏好中的“主屏幕时间和日期”关闭后，时间文字和日期文字会隐藏，但原时间栏仍保留原高度。天气、组件、App 图标、分页点和底部 Dock 都保持开启时的坐标，不再整体上移，也不会因为主屏内容高度变化而在底部露出黑色区域。",
        "实现上不再从 renderHome 输出中删除 home-premium-head，而是始终渲染同一个时间栏；关闭时仅增加 home-clock-hidden 类，使用 visibility:hidden 和 pointer-events:none 隐藏内容并禁用点击。常规高度继续保留 67px，矮屏断点继续保留 61px。",
        "验证结果：app.js 语法检查通过；时间开关、主屏高级主题、版本与 Service Worker 专项通过；完整 Node 自动化回归 387/387 通过。本地隔离缓存浏览器确认 v865 矮屏时间栏高度为 61px、主屏滚动区从 61px 开始。",
    ],
)

append_release(
    "AI开发项目_Bug记录模板.docx",
    "v865 Bug 记录｜关闭主屏时间后整体上移并露出底部黑块（2026-08-10）",
    [
        "现象：关闭“主屏幕时间和日期”后，时间栏原来占据的空间被一起删除，天气、组件、App 图标和底部区域整体向上移动；在部分手机比例下，主屏底部会露出一段黑色区域。用户要求只去掉时间显示，其他布局必须和开启时完全一致。",
        "根因：v861 用条件模板决定是否输出 home-premium-head。关闭时整个 header DOM 不存在，而主屏采用纵向 flex 布局，home-scroll 又是 flex:1；时间栏被删除后滚动区重新分配高度，引发整页回流。此前测试只断言“关闭后没有 header”，没有验证开关前后的布局槽位和坐标一致。",
        "修复：home-premium-head 改为始终渲染；关闭时添加 home-clock-hidden，通过 visibility:hidden 保留布局空间，通过 pointer-events:none 避免透明区域仍可点击，并设置 aria-hidden。原 67px 高度及 max-height:760px 下的 61px 高度均不改变。",
        "防复发：home-clock-preference.test.mjs 现在明确禁止条件删除 header，要求隐藏类同时具备 visibility:hidden 与 pointer-events:none；premium-theme.test.mjs 同步保护动态隐藏类下的原时间字体与结构。完整 Node 自动化回归 387/387 通过。",
    ],
)

append_release(
    "AI开发项目_Bug修改规范.docx",
    "新增强制规范｜视觉开关不得意外删除布局槽位（v865 起）",
    [
        "当需求明确为“隐藏元素但页面其他部分不能移动”时，禁止通过条件模板、display:none 或删除 DOM 节点实现。必须保留原布局槽位，只隐藏约定的可见内容，并同时处理点击、焦点和无障碍状态。",
        "在 flex、grid、绝对定位与安全区混合的主屏中，元素是否存在会改变剩余空间分配。修改视觉开关时必须检查常规高度和所有紧凑高度断点，至少比较被隐藏元素高度、下一内容区起点、Dock 位置与页面底部背景。",
        "视觉回归测试不能只检查文字是否消失。必须同时断言隐藏方式不会触发布局回流，并保留原主题、组件、图标、分页点、底栏和安全区行为；全量回归通过后才可发布。",
    ],
)

append_release(
    "AI开发项目_新聊天启动说明.docx",
    "新聊天接手状态｜v865 主屏时间隐藏保持原布局（2026-08-10）",
    [
        "唯一仓库仍为 C:\\Users\\pc\\Documents\\小手机\\phone-work，唯一分支 main，远端 origin/main。禁止创建分支、worktree、强推或使用旧目录；一次只允许一个聊天修改和发布。",
        "v865 以 v864 提交 f6d3a10db1277bd1fccf7dd39a0ac842a8add150 为基线，完整保留共同生活长期上下文与到家承接。本次只修复关闭主屏时间后布局整体上移和底部黑块。",
        "“主屏幕时间和日期”关闭时仍保留 home-premium-head 的 67px 布局槽位，矮屏保留 61px；home-clock-hidden 只使用 visibility:hidden 和 pointer-events:none。后续不得改回条件删除 header 或 display:none。",
        "v865 本地验证为专项通过、完整 Node 回归 387/387。后续新聊天必须从最新 origin/main 开始，保留动态隐藏类及布局回归断言，并实际核对线上 APP_VER、Shell 和 Service Worker 版本一致。",
    ],
)

print("Updated v865 maintenance documents")
