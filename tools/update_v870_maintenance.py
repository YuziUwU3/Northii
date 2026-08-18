from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_release(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    document = Document(path)
    if any(paragraph.text.strip() == title for paragraph in document.paragraphs):
        print(f"Skipped existing section: {title}")
        return
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"Updated {filename}")


append_release(
    "AI开发项目_项目说明文档.docx",
    "v870｜主辅模型统一 API 路线与就近保存（2026-08-10）",
    [
        "设置中的路线一至路线四现在各自同时保存两组配置：主聊天模型保留接口地址、API Key、模型名、随机度和回复长度；同一路线内新增辅助模型配置，保存辅助模型的可选接口地址、API Key 和模型名。辅助地址或 Key 留空时，仍按原规则复用当前路线的主聊天接口；辅助模型名留空时，仍回退为使用主模型。",
        "在设置页点击路线会先捕获当前页面的主聊天与辅助模型输入，再一起回填目标路线。聊天、游戏和线下场景里的快速路线切换同样同时应用两组配置；它只影响下一次请求，不中断已经发出的请求，也不改变语音、识图、联网或共同生活页面自己的模型选择规则。",
        "聊天模型与辅助模型标题右上角各新增一个“保存”按钮，两个按钮执行同一个统一保存动作，不需要再滚动到页面底部。原有底部“保存设置”继续保留并兼容，会把当前两组模型写入同一条路线，同时保存页面中的其他设置。",
        "旧数据迁移采用无损策略：历史 chatRoutes 中没有 aux 子配置时，初始化会把升级前全局 S.settings.aux 的现有值复制到每一条旧路线；已经带 aux 的新路线（包括明确留空）按各自值保留。这样升级后立即切换旧路线不会突然清空用户原有的辅助模型。",
        "验证结果：app.js 语法检查通过；API 路线、旧数据迁移、两处就近保存按钮、快速路线切换、设置分页、版本壳与 Service Worker 专项通过；完整 Node 自动化回归 408/408 通过。正式 v870 页面可加载且应用控制台无错误；邀请码授权层存在时未绕过授权操作用户存档。",
    ],
)

append_release(
    "AI开发项目_Bug记录模板.docx",
    "v870 Bug 记录｜API 路线只保存主模型且保存入口过远（2026-08-10）",
    [
        "现象：路线一至路线四只能保存主聊天模型。辅助模型始终只有一份全局配置，因此切换主路线后仍沿用原辅助设置，无法让一条路线完整代表一套主辅模型组合；模型区域较长时还必须滚动到最底部才能使用原“保存设置”。",
        "根因：chatRouteCopy、chatRouteCaptureForm、chatRouteFillForm、chatRouteSwitch 和 chatRouteQuickSwitch 只处理主模型的扁平字段；S.settings.aux 仅在 saveSettings 中单独保存，没有进入 chatRoutes。界面标题也没有就近保存入口。",
        "修复：路线对象在保留主模型扁平字段兼容旧消费者的同时新增 aux 子对象；统一的捕获、应用、回填和快速切换链路同时处理主聊天与辅助模型。两个模型标题右上角新增同一个 chatRouteSaveCurrent 保存入口，底部 saveSettings 也改为在辅助配置更新后写入当前路线。",
        "迁移保护：不能把旧路线缺失 aux 当成用户主动留空。chatRoutesInit 对没有 aux 属性的历史路线使用升级前全局辅助配置作回填；只要路线已经显式存在 aux，就保留它自己的值。测试覆盖多条旧路线，确认升级和切换均不丢失辅助模型。",
        "防复发与结果：api-routes.test.mjs 同时验证主辅保存、主辅回填、点击当前路线不恢复旧值、快速切换、两处标题保存按钮和旧路线迁移；设置分页、缓存版本及完整 408/408 回归均通过。未修改 API Key 内容、模型请求协议、计费、服务端函数或数据库。",
    ],
)

append_release(
    "AI开发项目_Bug修改规范.docx",
    "新增强制规范｜一条 API 路线必须原子保存主辅模型（v870 起）",
    [
        "当产品把主聊天模型与辅助模型定义为同一条 API 路线时，保存、切换、快速切换、表单回填和当前路线摘要必须使用同一个路线对象，不能只切主模型而继续沿用另一条路线的辅助模型。",
        "路线结构升级必须区分“历史字段缺失”和“用户明确留空”。历史路线没有 aux 属性时应从升级前真实全局辅助配置迁移；已经显式保存 aux:{base:'',key:'',model:''} 时必须保留为空，不能用全局值反向覆盖。迁移必须幂等。",
        "主模型仍可保持旧扁平字段以兼容现有调用者，但应用路线时必须分别写入 S.settings.chat 与 S.settings.aux；禁止把 aux 混进主模型请求参数，禁止改变辅助地址/Key 留空时复用当前主路线的既有规则。",
        "新增就近保存按钮时，按钮必须保存用户理解中的完整配置单元，并与底部总保存使用同一数据语义。不得让两个位置保存出不同内容；切换前必须先捕获当前两组输入，避免未保存编辑被目标路线覆盖。",
        "API 路线回归至少覆盖：旧路线无损迁移、显式空辅助不被覆盖、当前路线主辅共同保存、目标路线主辅共同回填、快速切换、点击当前路线保持正在编辑的值、保存按钮数量和底部总保存兼容。任何测试或文档中不得写入真实 API Key。",
    ],
)

append_release(
    "AI开发项目_新聊天启动说明.docx",
    "新聊天接手状态｜v870 主辅模型统一 API 路线（2026-08-10）",
    [
        "固定仓库 C:\\Users\\pc\\Documents\\小手机\\phone-work，固定分支 main，远端 origin/main。禁止新分支、worktree、旧目录和强推。开始前 clean、fetch、ff-only；发布前运行相关测试与完整回归，版本只在最新线上基础上加 1。",
        "API 路线入口集中在 chatMainCopy、chatAuxCopy、chatRouteCopy、chatRoutesInit、chatRouteCaptureForm、chatRouteApply、chatRouteFillForm、chatRouteSwitch、chatRouteSaveCurrent 与 chatRouteQuickSwitch。路线主字段仍是 base/key/model/temp/maxTokens，辅助模型保存在同一路线的 aux:{base,key,model}。",
        "旧路线兼容边界：chatRoutesInit 只在路线完全缺少 aux 属性时使用原 S.settings.aux 回填；显式 aux 空对象经规范化后必须继续为空。切换路线必须同时更新 S.settings.chat 和 S.settings.aux，辅助地址/Key 留空仍复用当前主路线。",
        "界面边界：聊天模型和辅助模型标题右上角各有一个“保存”，两者都调用 chatRouteSaveCurrent；底部 saveSettings 继续保存所有设置并把主辅两组写入当前路线。不得移除底部总保存，也不得让两个标题按钮只保存各自一半。",
        "v870 验证基线为完整 Node 408/408。后续修改模型路线必须至少运行 api-routes、in-session-api-routes、settings-pagination、cache-version、emergency-recovery 以及全量 tests/*.test.mjs，并核对 APP_VER、HTML 资源参数、Shell 与 Service Worker BUILD 一致。",
    ],
)

print("Updated v870 maintenance documents")
