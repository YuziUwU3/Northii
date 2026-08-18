from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_release(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    document = Document(path)
    if any(paragraph.text.strip() == title for paragraph in document.paragraphs):
        print(f"Skipped existing section: {title}")
        return
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"Updated {filename}")


append_release(
    "AI开发项目_项目说明文档.docx",
    "v881｜真机同步抗卡死与锁定真值分层（2026-08-11）",
    [
        "本次从 v880 升级为 v881 · 真机同步抗卡死。仓库原生 CompanionSyncView 的自动 5 秒同步不再启动 DeviceActivity/FamilyActivityData 重型直读；只有用户明确点击真实数据同步时才读取使用量，并由 8 秒竞速超时保护。超时后会丢弃本轮使用量、退出“同步中”，继续上传锁定配置、电量、健康和位置等兼容快照，避免一个读取接口拖住全部数据。",
        "伴生 App 新增“重新授权屏幕使用时间”显式按钮；从后台回到前台时只重新检查 AuthorizationCenter 状态，再进行普通同步，不在场景切换时自动弹授权请求。锁定/解锁仍在执行时检查 Family Controls 权限，失效时明确要求用户重新授权。",
        "ManagedSettingsStore 的读回只代表配置集合已经写入，不代表 iOS 已经对目标 App 实际显示屏蔽页。原生回执与网页状态改为“设备配置含锁/未含锁”；网页不再显示“设备回报已锁”或“新快照确认锁定”，并要求以真机打开目标 App 时出现系统屏蔽页作为最终验收。",
        "崩溃文件 PhoneCompanionTest-2026-08-11-014151.ips 已定位为 2026-08-11 01:41:51 的 FRONTBOARD 0x8BADF00D scene-update watchdog。主线程在 App 转入后台时卡在 VectorKit/MKMapView 的 flushTileLoads、flushTileDecodeQueues 与 geo::TaskQueue::barrierSync，超过 10 秒后被系统杀死；这不是网络断开，也不是内存 Jetsam。当前 Windows 仓库没有 MapKit/MKMapView 源码，说明 Mac 实际工程包含仓库外代码，未获得实际源码前禁止猜改。",
        "两份 JetsamEvent-2026-08-10 日志分别是 coreduetd per-process-limit 和系统 vnode-limit，PhoneCompanionTest 当时只是 suspended 或未出现，不能作为 01:41 卡退根因。Mac 实际工程 /Users/zoushijie/Documents/AppleProjects/PhoneCompanionTest 仍需完整比较、编译和真机验证；本轮只完成仓库源代码与网页修复，不能写成已安装到 iPhone。",
        "缓存链已统一到 881。新增 companion-native-resilience 回归覆盖前台授权复检、显式重新授权、8 秒读取上限、超时降级上传、自动同步不触发重型读取，以及配置状态不得冒充真实屏蔽。Node 全量回归 tests/*.test.mjs 为 438/438，git diff --check 无错误，仅 Windows LF→CRLF 提示。",
        "把小手机网页嵌入一个原生 iOS 容器可以减少应用切换，但不能消除 Family Controls capability、用户授权、签名、ManagedSettings、DeviceActivity Monitor/Report Extension 等系统要求。稳定的私人 App 方向应是“一个原生容器 + Web 界面 + 原生桥接 + 屏幕使用时间扩展”，但必须与正在审核的公开 North 分开规划，不能把嵌入当作本次故障的快捷修复。",
    ],
)

append_release(
    "AI开发项目_Bug记录模板.docx",
    "v881 Bug 记录｜同步常驻、授权失效、假锁定确认与后台卡退（2026-08-11）",
    [
        "现象：伴生 App 会长期停在“同步中”，今日屏幕使用时间显示读取失败或 0 秒；授权失效后锁定命令无法真实拦截 App，重新授权后才恢复；网页却可能显示设备已经确认锁定。另有一次 App 在 01:41 转入后台时卡退。",
        "同步根因：CompanionSyncService.synchronize() 在处理命令后会等待 FamilyActivityData.shared.installedApplications 和 DeviceActivityData.activityData(using: .live) 的完整异步序列，原实现没有超时；完整快照只能在该 await 返回后上传，因此使用量读取卡住会连带阻塞电量、健康、位置和普通快照。5 秒循环还可能自动触发该重型读取，后续同步因 syncInFlight 等待而继续失败。",
        "授权根因：prepareDataAccess() 只在视图初始任务中请求授权，scenePhase 恢复 active 时没有刷新 Family Controls 状态。崩溃/系统回收后的实际授权变化不会及时呈现在 UI，用户只能靠重新授权恢复。修复后前台恢复只复检状态，重新请求必须由用户点击明确按钮。",
        "假确认根因：锁定完成逻辑读取 effectiveLockedTokens()，但该集合仍来源于本 App 写入的 ManagedSettings store 和本地账本。这只能证明配置被保存，Apple 没有提供读取“目标 App 当前是否真的被屏蔽”的接口；网页再用后续快照把 configured token 显示为“确认锁定”，形成循环自证。修复后所有文案只陈述配置，真机打开目标 App 才是执行验收。",
        "卡退根因：PhoneCompanionTest-2026-08-11-014151.ips 明确记录 EXC_CRASH/SIGKILL、FRONTBOARD 0x8BADF00D、scene-update watchdog transgression；主线程停在 MKBasicMapView _didEnterBackground → VKMapView flushTileLoads → VectorKit barrierSync。两份 8 月 10 日 JetsamEvent 与此无关。仓库没有对应 MapKit 文件，因此本轮没有修改地图代码。",
        "修复：取消自动使用量直读；手动读取增加 8 秒超时并在失败时继续兼容快照；前台复检授权并增加显式重新授权按钮；锁定失败提示改为本地配置写入失败，成功提示改为配置已写入且需真机验证；网页状态与说明同步去除假确认。",
        "失败方案与禁止项：不得把网络重试当成 0x8BADF00D 修复；不得用 ManagedSettings 配置读回证明真实屏蔽；不得让 DeviceActivity 直读无限等待或阻塞其他遥测；不得在 scenePhase 自动弹权限框；不得根据 Windows 仓库猜写 Mac 工程中的 MapKit 生命周期代码；不得宣称 iOS 26 主 App 直读在中国区客户安装中稳定可用。",
        "自动验收：新增专项 4/4，全量 438/438，版本与缓存一致，diff 检查无错误。未完成：Mac 实际工程差异比较、Swift/iOS 26 编译、MapKit 后台修复、重新签名安装、真机授权恢复、真实锁定/解锁和使用量超时验收。",
    ],
)

append_release(
    "AI开发项目_Bug修改规范.docx",
    "新增强制规范｜真机读取必须有上限，配置不得冒充执行（v881 起）",
    [
        "任何 FamilyActivityData、DeviceActivityData、HealthKit、定位或反向地理编码读取都不得无限等待并独占完整同步。重型读取必须有明确时间上限、取消路径和降级结果；锁定命令与基础快照必须先行或能在读取失败后继续上传。高频轮询不得自动启动重型使用量直读。",
        "Family Controls 权限必须在 App 回到前台和执行控制命令时复检。场景切换只能刷新状态，不得未经用户操作自动请求授权；重新授权必须由可见按钮触发，并在失败时给出明确恢复说明。",
        "ManagedSettingsStore、持久账本和快照中的 token 都属于“配置/设备回报”，不是实际执行证明。UI、日志、角色提示和服务器回执禁止使用“系统确认锁定”“新快照确认锁定”等表述。唯一有效的锁定验收是打开目标 App 后出现系统屏蔽页；解锁验收是目标 App 可正常打开。",
        "收到 iOS .ips 时必须先区分 watchdog、普通异常和 Jetsam。0x8BADF00D 要沿主线程与生命周期栈定位阻塞；不能因为现象发生在同步页面就归因于网络。无源码命中的框架栈必须先取得实际 Xcode 工程并比较，禁止在仓库中添加猜测性替代代码。",
        "Windows 静态测试不能替代 Swift 编译。涉及 @MainActor、TaskGroup、Sendable、Family Controls 或 iOS 新 API 的修改，在发布真机包前必须在目标 Mac 工程编译，并记录 Xcode/iOS 版本、entitlements、App Group、Extension target 与签名状态。",
        "把网页嵌入原生 App 只能改变界面承载方式，不能绕过 Screen Time 权限和扩展。若推进私人整合 App，必须单独规划原生桥接与 Extension，并与公开 North 的 Bundle ID、审核用途、数据和权限边界隔离。",
    ],
)

append_release(
    "AI开发项目_新聊天启动说明.docx",
    "新聊天接手状态｜v881 真机同步抗卡死（2026-08-11）",
    [
        "固定仓库：C:\\Users\\pc\\Documents\\小手机\\phone-work；固定分支 main。当前待发布版本为 v881 · 真机同步抗卡死，Node 全量回归 438/438。仓库改动包含原生同步超时/降级、授权恢复入口、锁定真值文案和缓存升级。",
        "已完成的源代码行为：自动 5 秒同步不再读取 DeviceActivity live 使用量；手动使用量读取最多等待 8 秒，超时仍上传其他真实数据；前台恢复复检 Family Controls；用户可点“重新授权屏幕使用时间”；锁定状态只叫设备配置含锁，必须真机打开目标 App 验证。",
        "已确认的卡退证据：PhoneCompanionTest-2026-08-11-014151.ips 是 FRONTBOARD 0x8BADF00D，主线程在后台切换时卡在 MKMapView/VectorKit flushTileLoads 超过 10 秒；两份 JetsamEvent 不是该卡退。Windows 仓库完全没有 MapKit 源码，不得继续猜改。",
        "下一步必须先取得或在 Mac 打开 /Users/zoushijie/Documents/AppleProjects/PhoneCompanionTest 的完整实际工程，比较仓库 Swift 与真实文件，找到 MKMapView 所在文件和 scenePhase/background 生命周期代码。先保存用户现有改动，再合入 v881，同步检查 Monitor/Report Extension、App Group、Family Controls entitlement 和签名。",
        "Mac 验收顺序：一，编译所有 targets；二，安装并显式重新授权；三，点击手动同步验证 8 秒内成功或超时后回到空闲且电量/健康/位置仍上传；四，锁定 3 个目标 App 并逐个打开确认系统屏蔽页；五，解锁并逐个打开；六，前后台 20 次确认不再出现 MapKit watchdog；七，导出新的崩溃日志与同步截图。",
        "私人整合 App 尚未开始。可以采用一个原生容器内嵌小手机 Web UI、通过原生桥接连接 Family Controls 和 Screen Time extensions，但这不会减少 entitlement、授权和签名工作；必须与公开 North 审核版本分开，等本轮真机稳定后再立项。",
        "本轮未在 Mac 编译、未安装 iPhone 包、未修改真实 Mac 工程、未修复仓库外 MapKit 文件，也未完成真机验证。后续接手者不得把 438/438 的 Node 结果写成 iPhone 已修复。",
    ],
)

print("Updated v881 maintenance documents")

for path in DOCS.glob("*.docx"):
    document = Document(path)
    changed = False
    for paragraph in document.paragraphs:
        cleaned = paragraph.text
        for invisible in ("\u200b", "\u200c", "\u200d", "\ufeff"):
            cleaned = cleaned.replace(invisible, "")
        if cleaned != paragraph.text:
            for run in paragraph.runs:
                run.text = ""
            paragraph.add_run(cleaned)
            changed = True
    if changed:
        document.save(path)
        print(f"Removed invisible characters from {path.name}")
