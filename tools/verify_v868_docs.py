from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
EXPECTED = {
    "AI开发项目_项目说明文档.docx": (
        "v868｜共同生活位置与一起外出同步（2026-08-10）",
        ["together-away", "共同生活位置|准确地点", "一起回家不会生成", "同一份 c.sched", "页面顶部显示实时年月日"],
    ),
    "AI开发项目_Bug记录模板.docx": (
        "v868 Bug 记录｜地点不跟随与一起外出被误判为单独外出（2026-08-10）",
        ["place/placeAt", "要不要回卧室", "禁止把静态预览", "明天”说成周日"],
    ),
    "AI开发项目_Bug修改规范.docx": (
        "新增强制规范｜共同生活阶段、同场关系与位置必须分离（v868 起）",
        ["同场关系决定能否继续面对面", "未来计划", "不得自动转微信", "不是界面直接切换线上状态"],
    ),
    "AI开发项目_新聊天启动说明.docx": (
        "新聊天接手状态｜v868 共同生活位置与一起外出（2026-08-10）",
        ["cohabTogetherScene", "一起外出|准确地点", "旧约会格式零变化", "不得新增独立共同生活作息副本"],
    ),
}


for filename, (marker, required) in EXPECTED.items():
    document = Document(DOCS / filename)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    text = "\n".join(paragraphs)
    assert paragraphs.count(marker) == 1, f"duplicate or missing v868 section in {filename}"
    assert paragraphs[-1], f"blank ending in {filename}"
    for phrase in required:
        assert phrase in text, f"missing {phrase!r} in {filename}"
    print(filename, f"paragraphs={len(paragraphs)}", f"tables={len(document.tables)}")

print("v868 maintenance documents verified")
