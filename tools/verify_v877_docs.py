from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
EXPECTED = {
    "AI开发项目_项目说明文档.docx": (
        "v877｜聊天页固定视口与输入框恢复（2026-08-10）",
        ["screen.height", "chatbg", "inputbar", "422/422"],
    ),
    "AI开发项目_Bug记录模板.docx": (
        "v877 Bug 记录｜物理屏幕高度覆盖导致聊天输入框落到屏幕外（2026-08-10）",
        ["north-standalone-shell", "screen.height", "失败方案记录", "422/422"],
    ),
    "AI开发项目_Bug修改规范.docx": (
        "新增强制规范｜全屏聊天只能内部滚动，根页面不得随消息增长（v877 起）",
        ["min-height: 0", "flex: 0 0 auto", "安卓独立主屏"],
    ),
    "AI开发项目_新聊天启动说明.docx": (
        "新聊天接手状态｜v877 固定视口发布路线（2026-08-10）",
        ["phone-work", "origin HEAD:main", "v870", "422/422"],
    ),
}


for filename, (title, required) in EXPECTED.items():
    document = Document(DOCS / filename)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    text = "\n".join(paragraphs)
    assert paragraphs.count(title) == 1, f"missing or duplicate heading in {filename}"
    assert paragraphs[-1], f"blank ending in {filename}"
    for phrase in required:
        assert phrase in text, f"missing {phrase!r} in {filename}"
    print(filename, f"paragraphs={len(paragraphs)}", f"tables={len(document.tables)}")

print("v877 maintenance documents verified")
