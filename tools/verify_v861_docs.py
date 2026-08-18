from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
EXPECTED = {
    "AI开发项目_项目说明文档.docx": [
        "v861｜角色手机密码同步与主屏时间开关",
        "v861 发布前完整 Node 自动化回归为 378/378 通过",
        "完整保留 v860 的伴生推送上线记录",
    ],
    "AI开发项目_Bug记录模板.docx": [
        "v861 Bug 记录｜角色口头改密码后“查他手机”无法解锁",
        "完整项目自动化 378/378 通过",
        "c.spy.pwd",
    ],
    "AI开发项目_Bug修改规范.docx": [
        "AI 口头状态必须与真实状态一致（v861 起）",
        "模型回复→命令处理→持久化→实际消费者读取",
        "一次只允许一个聊天修改和发布",
    ],
    "AI开发项目_新聊天启动说明.docx": [
        "v861 密码同步与主屏时间开关",
        "v860 的伴生推送上线与固定仓库交接记录必须完整保留",
        "禁止创建新分支、worktree 或使用旧目录",
    ],
}


for name, markers in EXPECTED.items():
    path = DOCS / name
    with ZipFile(path) as archive:
        assert archive.testzip() is None, f"corrupt zip member in {name}"
        assert "word/document.xml" in archive.namelist(), f"missing document.xml in {name}"
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for marker in markers:
        assert marker in text, f"missing marker {marker!r} in {name}"
    assert text.count(markers[0]) == 1, f"duplicate v861 section in {name}"
    print(f"verified {name}: {len(document.paragraphs)} paragraphs, {len(document.tables)} tables")
