from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
EXPECTED = {
    "AI开发项目_项目说明文档.docx": (
        "v881｜真机同步抗卡死与锁定真值分层（2026-08-11）",
        ["8 秒", "0x8BADF00D", "设备配置含锁", "438/438", "原生容器"],
    ),
    "AI开发项目_Bug记录模板.docx": (
        "v881 Bug 记录｜同步常驻、授权失效、假锁定确认与后台卡退（2026-08-11）",
        ["syncInFlight", "MKBasicMapView", "两份 8 月 10 日 JetsamEvent", "不得让 DeviceActivity", "未完成"],
    ),
    "AI开发项目_Bug修改规范.docx": (
        "新增强制规范｜真机读取必须有上限，配置不得冒充执行（v881 起）",
        ["高频轮询不得", "场景切换", "唯一有效的锁定验收", "Windows 静态测试不能替代", "私人整合 App"],
    ),
    "AI开发项目_新聊天启动说明.docx": (
        "新聊天接手状态｜v881 真机同步抗卡死（2026-08-11）",
        ["phone-work", "438/438", "MKMapView", "Mac 验收顺序", "不得把 438/438"],
    ),
}


for filename, (title, required) in EXPECTED.items():
    document = Document(DOCS / filename)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    text = "\n".join(paragraphs)
    assert paragraphs.count(title) == 1, f"missing or duplicate heading in {filename}"
    assert paragraphs[-1], f"blank ending in {filename}"
    for phrase in required:
        assert phrase in text, f"missing {phrase!r} in {filename}"
    print(filename, f"paragraphs={len(paragraphs)}", f"tables={len(document.tables)}")

print("v881 maintenance documents verified")
