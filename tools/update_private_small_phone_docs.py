from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_section(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    document = Document(path)
    if any(paragraph.text.strip() == title for paragraph in document.paragraphs):
        print(f"Skipped existing section: {title}")
        return
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"Updated {filename}")


append_section(
    "AI开发项目_项目说明文档.docx",
    "私人「小手机」真实 App 总纲与基础工程（2026-08-11）",
    [
        "用户最终确认：公开版继续叫 North，供其他用户使用且不修改当前审核中的版本；私人版固定叫「小手机」，只给本人安装，沿用猫咪图标，只修改名称；网页版继续供安卓和浏览器使用。私人版必须是完整、独立、可安装的真实 iOS App，不是网页链接、PWA 或只打开远程网址的壳。",
        "三入口继续共用后端、数据结构和主要业务代码。私人 iOS App 在构建时把仓库现有网页核心打包为本地资源，由 WKWebView 通过 loadFileURL 加载，再用版本化原生桥接入 Family Controls、Device Activity、Managed Settings、Keychain 和后台同步；生成的 PhoneWeb 目录不是第二套源码，不允许人工分叉维护。",
        "本人设备最终只允许私人「小手机」控制。公开 North 与私人「小手机」不得同时拥有同一设备的控制权；后端必须记录 controllerKind、控制器实例、租约版本和幂等键，旧控制器命令必须拒绝。迁移时还要废止公开 North 对本人设备的凭据并重新核验私人版授权，不能只隐藏界面入口。",
        "永久锁定、当日使用时间、快照版本和真实回执仍是最高验收条件：刷新、同步、重连、跨日和旧快照都无权解锁；采集时间与上传时间必须分开；界面明确显示命令已发送、设备已收到、设备已执行。收到真实执行回执前不得显示成功。",
        "本轮已建立 docs/maintenance/私人小手机App_唯一总纲.md 和 native/private-small-phone 基础目录，包括 SwiftUI 入口、本地资源 WKWebView、v1 原生桥契约、共享资源清单与暂存脚本。尚未伪造 Xcode 工程；Bundle ID、App Group、Apple Team、entitlements 和扩展 Target 必须从 Mac 上真实工程确认。",
        "这不是推倒重做：既有共同生活、查岗去重、模型回退、旁白、跨渠道时间轴和各 App 安全区全部保留。后续只把已经验证的原生伴生能力迁入私人主 App，并按真机证据修复问题。",
    ],
)

append_section(
    "AI开发项目_Bug记录模板.docx",
    "私人 App 转向风险记录｜双控制器、工程缺失与方向误解（2026-08-11）",
    [
        "现象与风险：公开 North 当前仍可能保留本人设备的控制链路，同时又准备建立私人「小手机」App；若两边都能下发或恢复 Managed Settings 规则，就会出现命令竞争、状态互相覆盖和无法判断真实执行者。此前交接文档还保留“私人版名称尚未确定”“先修稳伴生 App 再立项”等旧表述，容易让新聊天继续走错方向。",
        "根因：公开版、私人版和网页版的产品边界没有集中成唯一总纲；控制设备的资格也只有配对凭据，没有明确的单控制器租约。另一个高风险缺口是 Windows 仓库只有导出的 Swift 文件，没有 Mac 上完整 .xcodeproj、扩展 Target、签名配置和真实 MapKit 文件，因此无法在当前仓库编译或对 0x8BADF00D 崩溃作可靠源码修复。",
        "证据边界：PhoneCompanionTest-2026-08-11-014151.ips 指向进入后台时 MapKit/VectorKit 主线程看门狗退出；两份 JetsamEvent 不能证明同一根因。没有真实 Mac 工程前禁止把网络重试或猜测的地图生命周期修改写成已修复。",
        "当前处理：建立唯一总纲，明确私人版名称、真实 App 形态、共享核心和 North/小手机互斥；新增只加载安装包本地资源的 SwiftUI/WKWebView 骨架与版本化桥接。未修改公开 North、未猜 Bundle ID、未声称已编译或已安装。",
        "尚未完成：取得并核对 Mac 完整 Xcode 工程；迁移 Screen Time 扩展与伴生能力；实现服务端控制器租约；修复真实 MapKit 生命周期；完成真机签名、授权、断网、后台、跨日、重复命令和长时间稳定性测试。",
    ],
)

append_section(
    "AI开发项目_Bug修改规范.docx",
    "新增强制规范｜私人 App 单一总纲与单设备单控制器（2026-08-11 起）",
    [
        "私人版固定名为「小手机」，必须是本地打包业务核心并接入原生 Screen Time 能力的真实 iOS App。禁止把网页链接、PWA、远程首页或只改图标名称的 Web 壳标记为私人 App 完成。",
        "公开 North 与私人「小手机」禁止同时控制同一台设备。服务端控制命令必须校验控制器类型、实例、租约版本和幂等键；客户端界面隐藏、设备名称不同或使用不同 Bundle ID 都不能替代服务端互斥。控制权迁移必须有废止旧凭据和重新授权的真机步骤。",
        "网页核心只维护一份。iOS 构建可以把共享文件复制到生成目录，但生成目录必须可重建且不得人工编辑。原生桥必须版本化、白名单化，未知 action 默认拒绝；网页不得自行假装原生执行成功。",
        "Bundle ID、App Group、Apple Team、entitlements、扩展 ID 和签名配置必须从真实 Xcode 工程与开发者账号读取，禁止猜写。没有 Mac 完整工程和对应源码时，静态测试不得替代 Swift 编译，也不得猜修日志中命中的 MapKit/VectorKit 生命周期。",
        "新需求与旧交接冲突时，必须先更新唯一总纲并明确哪些旧表述失效。后续接手者先读 docs/maintenance/私人小手机App_唯一总纲.md，再读历史记录；不得把历史阶段顺序当成用户当前决定。",
    ],
)

append_section(
    "AI开发项目_新聊天启动说明.docx",
    "新聊天最高优先级｜私人「小手机」真实 App 已启动（2026-08-11）",
    [
        "先读 docs/maintenance/私人小手机App_唯一总纲.md。用户已经确认：私人版固定叫「小手机」，是完整真实 iOS App；公开 North 保持审核版本不动，并停止控制本人设备；网页版继续保留。旧文中“名称未定”“稳定后再立项”等冲突内容失效。",
        "当前已完成的是基础源码，不是真机成品：native/private-small-phone 已有 SwiftUI 入口、本地资源 WKWebView、v1 原生桥、共享资源清单和暂存脚本；专项测试验证不会加载远程首页，也不会维护第二套网页核心。",
        "当前阻塞证据在 Mac：必须取得 /Users/zoushijie/Documents/AppleProjects/PhoneCompanionTest 的完整工程，核对 .xcodeproj、所有 Screen Time 扩展、entitlements、App Group、签名以及 MapKit 文件。Windows 导出文件不足以完成编译和崩溃修复。",
        "下一步顺序：导入真实 Mac 工程；建立私人 Target/扩展且不改审核 North；迁移现有伴生控制能力；实现单设备单控制器租约；接入账号、Keychain 和恢复；最后做跨日、断网、后台、刷新、重复命令与长时间真机测试。",
        "不可破坏：永久锁定、当日使用时间、发送/收到/执行三阶段回执、共同生活去重、主副模型规则、旁白、真实时间轴和逐 App 安全区。不得把 Node 测试通过写成 iPhone 已完成。",
    ],
)

print("Private small-phone maintenance sections updated")
