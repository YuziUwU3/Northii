from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1] / "docs" / "maintenance"
MARKER = "2026-08-11 v894／私人 1.0.19 真实查看触发、网页横幅与滚动恢复快照"

ENTRIES = {
    "AI开发项目_项目说明文档.docx": [
        "网页核心升级为 v894，私人 iOS 候选升级为 1.0.19 (19)，原生桥契约仍为 11。角色只有在本次真实读取已经取得 readSessionId 后，才能谈论电量、屏幕使用、逐 App 用量、步数、睡眠、心率、心电/HRV 或位置。",
        "补齐自然语言触发：角色主动说要看步数、睡眠、屏幕、某个 App、电量或健康时，会先启动真实读取并显示当前项目；没有读取凭证时，‘没刷新、没同步、没戴手表’等猜测会被替换为先实际读取。低电量等自动化也只能谈本次授权的类别。",
        "普通网页查岗与真实 iPhone 读取重新分流。没有 intent 的网页定时/主动查岗继续执行微信聊天、朋友圈、抖音、X、钱包、购物、浏览器和线下记录的经典逐项顶部横幅；只有明确真实数据 intent 才进入外置读取。APP 共用入口，已同步规避同类误分流。",
        "恢复安全快照改为滚动到最新保存时间。较旧但数据更丰富的副本会另存为历史完整快照，不再阻止当前快照更新；恢复候选优先按 savedAt 最新排序，同时间再按完整度排序，因此不再长期停留在 8 月 9 日。",
        "网页侧仅修复经典查岗横幅与恢复快照日期，没有把 APP 的锁定、真实 HealthKit/Screen Time、APNs 等能力移植进网页。Mac 编译、签名和真机读取/通知验收仍须在新独立工程目录完成。",
    ],
    "AI开发项目_Bug记录模板.docx": [
        "现象一：网页版角色查手机时不再出现‘正在查看微信、朋友圈、抖音’的顶部逐项横幅；同一入口在 APP 也可能受到影响。现象二：设置里的本机恢复候选一直显示 2026/08/09，即使之后仍在使用和保存。",
        "横幅根因：doSpyView 的默认 focus 是‘查岗’，companionRoleExternalFocus 又把‘查岗/手机/概览’判断为外置数据，导致无 intent 的经典网页查岗先进入真实 iPhone 拉取；拉取失败即提前返回，doSpyViewCore 的经典步骤未执行。横幅 DOM、CSS 和步骤列表并未被删除。",
        "恢复根因：queueRecoverySnapshot 只要旧快照 score 更高就拒绝任何较小的新快照；recoveryCollectCandidates 又先按 score、后按 savedAt 排序。旧快照聊天更多时会永久压住新日期。",
        "修复：外置拉取增加 opts.intent 硬门槛；普通查岗直接进入经典逐项横幅。恢复写入拒绝的只剩乱序旧写入；旧且更丰富的快照先归档到 __recovery_history_state，再把 __recovery_state 滚动为当前。候选改为先按 savedAt 排序。",
        "验证：新增普通网页查岗不进入 native telemetry、经典微信/朋友圈/抖音/浏览器步骤存在、当前快照滚动且旧丰富快照归档、乱序旧写入被拒绝、最新候选排序等回归。Windows 自动化通过不等于 Xcode 真机通过。",
    ],
    "AI开发项目_Bug修改规范.docx": [
        "同一个关键词可同时属于虚拟网页能力和真实设备能力时，不能只按文本路由；必须要求明确的结构化 intent/来源标记。默认流程不得因‘查岗、手机、概览’等泛词误入外置设备分支。",
        "可见进度必须跟真实执行同生同灭。若读取分支提前失败，不能让另一个应执行的本地流程无声消失；测试应同时断言入口门槛、进度 DOM、逐项步骤和最终核心函数调用。",
        "安全恢复既要保留较完整旧副本，也要反映当前保存时间。不得用完整度分数永久冻结滚动快照；采用‘当前滚动副本 + 历史完整副本’，候选默认按时间最新优先，只有时间相同时再比较完整度。",
        "恢复写入必须拒绝 savedAt 倒退的乱序任务；页面启动完成后的完整 hydration 必须重新排队当前快照，确保版本更新后旧冻结状态能自行修复。",
    ],
    "AI开发项目_新聊天启动说明.docx": [
        "当前候选基线：网页 v894、私人 iOS 1.0.19 (19)、原生桥契约 11。接手时先确认网页普通查岗仍显示微信→朋友圈→抖音等逐项横幅，明确真实数据请求才显示电量→屏幕→逐 App→健康→位置。",
        "恢复数据验收：更新后完整打开页面并等待启动完成，再进入设置扫描恢复数据；本机安全快照应显示本次保存时间。若当前数据比旧档少，旧档仍可作为‘历史完整快照’保留，但不应排在更新日期之前。",
        "角色真实读取验收必须以顶部逐项进度和本次 readSessionId 为证；没有实际触发读取时，角色不能先说步数/屏幕没刷出来，也不能猜测未戴手表。",
        "Mac 只能打开新交付的 SmallPhone_v894_RealReadFix 独立目录中的 PhoneCompanionTest.xcodeproj。Windows 静态测试不能替代五 Target 编译、HealthKit/Screen Time 权限和前后台/强退通知真机验收。",
    ],
}


def main() -> None:
    for filename, items in ENTRIES.items():
        path = ROOT / filename
        doc = Document(path)
        changed = not any(MARKER in paragraph.text for paragraph in doc.paragraphs)
        if changed:
            doc.add_heading(MARKER, level=1)
            for item in items:
                doc.add_paragraph(item, style="List Bullet")
            doc.save(path)
        with ZipFile(path) as archive:
            assert archive.testzip() is None
            assert "[Content_Types].xml" in archive.namelist()
        check = Document(path)
        assert sum(MARKER in paragraph.text for paragraph in check.paragraphs) == 1
        print(f"{'updated' if changed else 'unchanged'}: {path.name}")


if __name__ == "__main__":
    main()
