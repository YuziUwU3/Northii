from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_release(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    document = Document(path)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)


append_release(
    "AI开发项目_项目说明文档.docx",
    "v857｜伴生极速回执与持久快照（2026-08-09）",
    [
        "当前前端版本统一升级为 v857。本次只调整 Apple 伴生设备同步链路，不改变 Screen Time 授权、HealthKit 授权、已选 App、App Groups、角色控制权限或内外 App 绑定规则。",
        "命令链路改为排队后立即请求 phone-companion-push，并把推送结果记录为 pushed 或 queued。网页不再静默吞掉 Edge Function 超时、未部署或无令牌结果；推送不可用时命令仍留在服务器队列，伴生 App 下次运行会补执行。",
        "iPhone 新替换包位于 PhoneCompanion/v857伴生极速回执与持久快照_微信传输包_2026-08-09。APNs 后台唤醒改走 synchronizeCommandsOnly：先拉取命令、应用 ManagedSettings、写回回执，再上传轻量控制快照；不等待 DeviceActivity 报告、HealthKit 刷新或反向地理编码。前台完整同步保持不变。",
        "轻量控制快照带 controlOnly 标记。网页只接受其中的锁定、解锁和限额状态，保留上一次真实使用时长、位置、足迹与健康数据，避免快速回执把真实动态数据覆盖成 0 或空值。",
        "快照采用分层语义：动态时长可以显示为较早，但上一次已知状态永久保留；锁定或解锁只有在服务器回执之后又收到设备快照且内容一致时才进入 confirmed，一旦确认不会仅因两分钟过去而退回未知。",
        "网页命令回执轮询提前到约 1.2 秒、4.5 秒和 12 秒；伴生页常规轮询下限从 6 秒缩短到 4 秒。全项目 376 项自动化测试通过，0 失败。",
    ],
)

append_release(
    "AI开发项目_Bug记录模板.docx",
    "v857 Bug 记录｜后台同步慢、快照过期与锁定回执等待（2026-08-09）",
    [
        "现象：伴生 App 留在后台时，网页命令经常要等到手动刷新、重新打开或退出再进入后才执行；锁定按钮等待回执时间长；两分钟后已确认状态会被标成快照过期；长同步偶发失败。",
        "根因一：线上 Edge Functions 列表只有 phone-ai、phone-license 和 phone-role-push，phone-companion-push 实际返回 HTTP 404。网页原逻辑对该请求 catch 后直接忽略，导致命令只能依赖 iOS 下次前台轮询。",
        "根因二：旧后台唤醒先刷新 HealthKit，再进入完整 synchronize；完整同步还可能等待正在执行的同步、读取 DeviceActivity 报告、反向解析最多八个足迹地址，最后才处理和回执命令。因此短暂的 iOS 后台执行窗口会被重任务抢占。",
        "根因三：网页把控制确认和动态数据新鲜度都绑定在同一个两分钟阈值上。即使锁定已收到服务器回执并被回执后的设备快照确认，两分钟后仍会降级为 stale。",
        "修复：新增原生命令快车道与命令串行锁；后台唤醒不刷新健康、不读取使用报告、不解析地址。完成命令后上传 controlOnly 控制快照。网页合并控制快照时保留旧动态字段，并让 verified confirmed 状态持久化；没有新命令的旧快照显示为上次已知状态而不是假装实时。",
        "服务器核验：phone_companion_pull_commands、phone_companion_register_push_token、phone_companion_push_snapshot 三个 RPC 均已存在并能对无效凭据安全返回 null/false，说明所需迁移已经生效；phone-companion-push 仍为 HTTP 404。",
        "失败尝试：Supabase 网页编辑器的名称输入和 Monaco 代码输入在多次 Playwright、剪贴板粘贴及可见控件输入中均被控制台网络超时中断。遵守两次失败后停止重复策略的规则，未继续盲点部署，也未改动现有三个线上函数。函数源码已保存在 supabase/functions/phone-companion-push/index.ts，等待一次人工 Dashboard 部署。",
        "验证：app.js 语法检查通过；48 项伴生/恢复专项测试通过；全项目 376 项测试通过，0 失败。Windows 环境没有 Xcode/Swift 编译器，原生包需在 Mac Xcode 替换两个主 App 文件后完成真机编译与锁定/解锁验收。",
    ],
)

append_release(
    "AI开发项目_Bug修改规范.docx",
    "新增强制规范｜伴生命令与动态快照必须分层（v857 起）",
    [
        "后台通知的第一优先级是完成可逆、幂等的设备命令并回执。不得在命令之前等待 DeviceActivity 报告、HealthKit 查询、反向地理编码、图片处理或其他可延后的动态数据任务。",
        "前台完整同步与后台命令同步必须使用独立的并发边界，并对拉取、执行、回执命令段做进程内串行化，防止同一 pending 命令被两个同步任务重复执行。",
        "控制快照必须显式标记 controlOnly。消费者只能用它更新锁定、解锁、限额等控制字段；不得用缺失或默认的 0 覆盖上一次真实使用时长、位置、足迹或健康数据。",
        "命令确认与动态数据新鲜度不得共用一个过期结论。只有回执后的设备快照与命令一致时才能确认；确认后可永久保留为上一次已验证状态，同时把时长、位置等动态数据单独标注为较早。",
        "APNs 只是唤醒提示，不是设备回执。网页必须读取并记录推送接口结果，但 pushed 不得把命令标为 completed；无令牌、节流、超时或函数不可用时必须保留服务器队列并给出可理解的降级说明。",
        "iOS 被用户从多任务界面强制划掉后，后台通知通常不能唤醒 App。产品提示必须如实说明这一系统限制，不得承诺永久后台运行或绝对即时。",
        "发布伴生改动必须至少验证：推送函数是否真实存在、RPC 是否存在、后台命令快车道、并发去重、controlOnly 合并、回执后快照确认、确认跨新鲜度阈值保持、完整测试套件和 Mac 真机编译。",
    ],
)

append_release(
    "AI开发项目_新聊天启动说明.docx",
    "新聊天接手状态｜v857 本地修复完成，服务器函数待部署（2026-08-09）",
    [
        r"网页工作目录：C:\Users\pc\Documents\小手机\phone-release-web-v835-avatar。原生替换包：C:\Users\pc\Documents\小手机\PhoneCompanion\v857伴生极速回执与持久快照_微信传输包_2026-08-09。",
        "v857 已完成：网页推送结果可见、快速回执轮询、持久已确认控制状态、旧动态数据保留；iPhone 后台唤醒改为命令优先，不再等待使用报告、健康刷新和地址解析。",
        "测试状态：app.js 语法检查通过；全项目 376 项自动化测试通过，0 失败。原生 Swift 仅做了静态契约测试，仍需在 Mac Xcode 编译。",
        "服务器状态：所需三个 RPC 已在线；phone-companion-push 仍返回 HTTP 404。Supabase Dashboard 代码输入连续被网络超时中断，禁止继续重复同一路径。下一步只需在 Edge Functions 里通过 Editor 新建 phone-companion-push，粘贴仓库 supabase/functions/phone-companion-push/index.ts 并部署。现有 APNs secrets 已在项目中，不要重建或显示其内容。",
        "真机步骤：部署函数后，在 Mac Xcode 用 v857 包的 CompanionSyncView.swift 和 PhoneCompanionTestApp.swift 替换主 App 同名文件并安装到原 iPhone；不要重置授权、App 选择、App Groups 或 Bundle ID。",
        "验收顺序：伴生 App 退到后台但不要强制划掉；网页锁定真实 App；数秒内应锁定并收到回执后的新快照；等待两分钟确认仍保留；再解锁；最后测试手动同步无需停留在伴生页。",
    ],
)

print("Updated v857 maintenance documents")
