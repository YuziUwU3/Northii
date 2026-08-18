from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_release(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    document = Document(path)
    if any(paragraph.text.strip() == title for paragraph in document.paragraphs):
        print(f"Skipped existing section: {title}")
        return
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"Updated {filename}")


append_release(
    "AI开发项目_项目说明文档.docx",
    "v866｜共同生活记忆与 iOS 底部适配（2026-08-10）",
    [
        "v866 基于已发布的 v865 main（cb87d74ae54a7091888bddcdf47f14988f0399fb）。共同生活页面顶部新增可原地展开的专属设置：最近上下文 10～80 条、自动总结 0～40 轮、旧总结引用 1～20 条、总结方式（分段重点/一整段）、对话模型路线和总结模型路线。设置按共同生活对象独立保存，修改时不重渲染聊天输入区。",
        "共同生活自动总结和“立即总结”只写入 S.cohabitation.homes[角色ID].summaries。它不复制到微信记忆列表，不写入 S.offline 的单次线下约会记录，也不写入角色 summaries。共同生活页面可直接查看全部总结；线上/线下同步开启时，微信模型和共同生活模型只按当前话题从该唯一存储临时读取相关条目，关闭同步后微信不读取。可见微信气泡、共同生活记录和单次约会记录始终互不复制。",
        "分段模式按旧约会的记忆原则生成 1～8 条独立第一人称重点，内容少时允许只生成 1 条，禁止凑数；整段模式生成一条完整第一人称记忆。两种方式都排除普通起身、走路、拿东西、环境、姿势和第三人称旁白，只保留新增事实、重要话语、约定、情绪或关系变化。自动总结和立即总结共用当前页面选择的方式与总结模型路线。",
        "部分 iOS 独立主屏/应用切换快照出现底部黑带，推断为旧 WebKit 在 standalone 竖屏下可用高度与动态视口高度不一致。v866 仅在 @supports(-webkit-touch-callout:none) + display-mode:standalone + portrait 范围内增加 -webkit-fill-available 底部补齐；不恢复全局 visualViewport 高度脚本，不改变普通 Safari、Android、桌面和本来正常的 iPhone 布局。",
        "验证结果：app.js 语法检查通过；共同生活设置、上下文、状态交接、独立记忆、分段/整段入口、iOS standalone 视口、移动端启动等专项 24/24 通过；完整 Node 自动化回归 392/392 通过。390×844 浏览器核验中，设置展开无横向溢出，聊天区仍保留 385px 高度；正式页面 html、body、phone、screen 的底部均为 844px。真实问题机仍需在线上安装态复核系统应用切换快照。",
    ],
)

append_release(
    "AI开发项目_Bug记录模板.docx",
    "v866 Bug 记录｜共同生活设置不可随手调整、记忆边界不清及部分 iOS 底部黑带（2026-08-10）",
    [
        "现象一：共同生活的上下文条数、自动总结间隔、总结引用数量和模型路线依赖后台或全局设置，用户在共同生活聊天中不能随手调整，也不能直接查看或立即总结共同生活记忆。记忆如果复用微信或旧约会存储，会造成格式污染、重复召回和删除边界不清。",
        "修复一：每个共同生活对象新增独立 settings；设置面板固定在聊天页顶部、可折叠且修改不触发整页 render。共同生活总结只进入 d.summaries，查看和立即总结均在共同生活页完成。微信仅在同步开关开启时临时检索该唯一存储，不把结果复制进微信记忆或单次约会。自动与手动总结共用分段/整段模式和总结模型路线。",
        "现象二：少数 iPhone 从系统应用切换界面查看独立主屏 PWA 时，页面底部出现一段黑色横带；其他 iPhone 未出现。截图特征是页面内容和 home indicator 已结束，但 WebApp 快照容器仍有未被页面根背景覆盖的可用高度。",
        "根因判断与修复二：这是设备相关的旧 WebKit standalone 竖屏可用高度差异，而不是统一的页面内容高度问题。修复限定在 WebKit standalone portrait，使用 -webkit-fill-available 补齐 html/body/phone/screen 的最小高度；禁止用全局 resize/visualViewport 脚本反复改根高度，以免破坏正常手机、键盘恢复和滚动。桌面 390×844 几何验证四层根容器均精确到底且无横向溢出，真实问题机仍须在线上安装态确认快照。",
        "防复发：新增 cohabitation-inline-settings.test.mjs 和 ios-standalone-viewport.test.mjs，并扩展共同生活与版本回归。测试明确禁止共同生活总结写入 c.summaries、offData 或 S.messages，要求微信和面对面回复只从 d.summaries 检索；同时要求 iOS 兼容必须是 CSS 限域方案，不得恢复全局动态视口脚本。完整回归 392/392。",
    ],
)

append_release(
    "AI开发项目_Bug修改规范.docx",
    "新增强制规范｜长期场景设置、独立记忆库与设备定向视口修复（v866 起）",
    [
        "长期场景的上下文条数、总结轮数、总结引用量、总结方式和模型路线不得写死，也不得只能到全局设置修改。应在当前长期场景聊天页提供可折叠、可持久化的原地设置；设置按角色/场景隔离，控件变更不得整页重渲染或破坏正在输入的内容。",
        "共同生活、单次线下约会和微信必须使用不同的可见记录与总结存储。共同生活总结只允许写入该共同生活 home 的 summaries；禁止复制进微信长期记忆、角色 summaries、S.messages 或 S.offline。跨线上线下需要互通时，只能在同步开关开启后从共同生活唯一存储按话题临时检索，禁止做第二份持久副本。",
        "长期记忆总结必须支持用户主动触发，并允许整段或分段。分段不得为了达到条数而重复或编造；内容不足时允许单条。无论哪种方式，都必须以角色第一人称保存，排除普通动作流水账、环境和第三人称旁白，只保留有长期价值的事实、话语、约定、情绪与关系变化。",
        "只在少数设备出现的视口问题必须优先采用能力检测、显示模式和方向共同限定的 CSS 修复。禁止在没有跨设备证据时修改所有设备的根高度，也禁止恢复全局 visualViewport/resize 高度写入。至少验证正常手机尺寸下 html、body、phone、screen 同底、无横向溢出、键盘与滚动回归不退化；问题设备必须保留上线后的安装态复核项。",
    ],
)

append_release(
    "AI开发项目_新聊天启动说明.docx",
    "新聊天接手状态｜v866 共同生活专属设置、独立记忆与 iOS 底部补齐（2026-08-10）",
    [
        "固定仓库仍为 C:\\Users\\pc\\Documents\\小手机\\phone-work，固定分支 main，远端 origin/main。禁止创建分支、worktree、强推或使用旧目录；修改前必须 clean、fetch、ff-only，发布前必须完整测试并再次核对远端。v866 的起始基线为 v865 提交 cb87d74ae54a7091888bddcdf47f14988f0399fb。",
        "共同生活页内设置保存在 S.cohabitation.homes[id].settings：contextLimit、summaryRounds、summaryMemoryLimit、summaryMode、replyRoute、summaryRoute。共同生活记忆唯一存储为同一个 home 的 summaries；页面有“查看共同生活记忆”和“立即总结”。不得把这些总结复制到微信记忆、角色 summaries 或原来的 S.offline 单次约会。",
        "微信与共同生活的互通规则：只有“约会中同步到线上”开启时，微信构建隐藏上下文才按当前话题临时读取共同生活 summaries；关闭时微信不读取。共同生活模型始终可以读取自己的记忆。可见聊天和旁白格式继续隔离，不得把共同生活动作复制为微信气泡，也不得改坏旧约会既有第三人称旁白与普通台词格式。",
        "iOS 底部黑带修复只允许保留当前的 WebKit + standalone + portrait CSS 限域，不得改成全局高度脚本。v866 本地验证为专项 24/24、完整 Node 回归 392/392；390×844 正式页面根层全部覆盖到底且设置面板无横向溢出。发布后仍需在用户提供的真实问题 iPhone 安装态应用切换快照中确认。",
    ],
)

print("Updated v866 maintenance documents")
