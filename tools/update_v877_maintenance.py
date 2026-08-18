from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_release(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    document = Document(path)
    if any(paragraph.text.strip() == title for paragraph in document.paragraphs):
        print(f"Skipped existing section: {title}")
        return
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"Updated {filename}")


append_release(
    "AI开发项目_项目说明文档.docx",
    "v877｜聊天页固定视口与输入框恢复（2026-08-10）",
    [
        "本次只修复全屏外壳高度回归，不回退 v870 至 v876 的双模型路线保存、角色密码、主屏组件开关、共同生活、North 审核支持页面等既有功能。发布基线升级为 v877 · 聊天视口与输入框恢复。",
        "全屏结构重新固定为三层约束：html、body、phone、screen 不得超过当前可用视口；app 与 page 负责裁切；chatbg 作为唯一消息滚动区，inputbar 保持 flex: 0 0 auto 固定占位。长聊天只增长消息区的 scrollHeight，不再增长整个页面。",
        "删除 v873 引入的物理屏幕高度补齐脚本及 north-standalone-shell 高度覆盖。该脚本会在安卓独立主屏把 screen.height 当成可用视口，状态栏和系统导航区差值会把整页拉长，导致聊天输入框落到屏幕外。",
        "底部黑缝不再通过拉长应用内容处理。当前保留 100dvh 与根容器背景覆盖，并增加 max-height: 100% 和 overflow: hidden 的硬边界；禁止重新引入 screen.height、全局 visualViewport 高度同步或键盘 resize 拉高逻辑。",
        "专项视口测试覆盖外壳上限、page 裁切、消息区内部滚动、输入框不可收缩以及禁止物理屏幕高度覆盖；完整 Node 回归 422/422 通过。",
    ],
)

append_release(
    "AI开发项目_Bug记录模板.docx",
    "v877 Bug 记录｜物理屏幕高度覆盖导致聊天输入框落到屏幕外（2026-08-10）",
    [
        "现象：安卓独立主屏打开长聊天后，页面总高度明显超过手机可用屏幕，右侧出现整页滚动条，聊天底部输入框和功能区落到屏幕外。用户只能看到消息列表，无法正常发送消息。",
        "根因：v873 为修复少数 iOS 主屏底部黑缝，加入了 north-standalone-shell 动态高度脚本。检测条件同时接受 display-mode: standalone 和触摸设备，安卓 PWA 也会命中；脚本在 screen.height 与 innerHeight 差值不超过 180px 时采用物理屏幕高度，并把该高度强制写给 html、body、phone、screen。安卓状态栏与系统导航区不属于网页可用视口，因此外壳被额外拉长。",
        "修复：完全删除动态物理屏幕高度脚本、CSS 变量和强制高度选择器；不恢复旧的 -webkit-fill-available 回退。根外壳仅跟随当前可用视口，并通过 max-height: 100% 与 overflow: hidden 封住整页增长；消息继续由 chatbg 内部滚动，inputbar 保持独立固定行。",
        "回退与失败方案记录：不能整体回退到 v869 或 v872，否则会丢失 v870 至 v876 的已发布功能；也不能仅把脚本限制为触摸独立主屏，因为安卓仍会命中；不能继续比较 screen.height 与 innerHeight 后猜测安全区，也不能恢复全局 visualViewport/resize 高度同步，这些方法会在状态栏、系统导航和软键盘变化时再次拉长或压短页面。",
        "验证：ios-standalone-viewport.test.mjs 断言根外壳上限、app/page 裁切、chatbg 内滚动、inputbar 固定占位，并禁止 screen.height、north-shell-height 和全局视口同步脚本；版本与 Service Worker 对齐测试及完整 Node 回归最终 422/422 通过。",
    ],
)

append_release(
    "AI开发项目_Bug修改规范.docx",
    "新增强制规范｜全屏聊天只能内部滚动，根页面不得随消息增长（v877 起）",
    [
        "全屏应用的 html、body、phone、screen 必须以当前网页可用视口为上限并裁切溢出；任何功能页不得通过内容高度、物理屏幕高度或安全区猜测扩大根页面。",
        "聊天页必须维持 app/page 为纵向 flex 容器，消息区具备 flex: 1、min-height: 0 和 overflow-y: auto，输入区具备 flex: 0 0 auto。新增顶部卡片、提示栏、礼物卡或模型状态条时，只能压缩消息区，不能把输入区推到视口外。",
        "禁止用 screen.height 作为网页 shell 高度，也禁止重新加入全局 visualViewport、resize 或键盘监听来持续写根高度。若需要修复 iOS 安全区，只能使用明确的 WebKit/standalone 范围和不会改变内容布局高度的视觉覆盖方案，并同时做安卓独立主屏回归。",
        "每次修改 PWA 高度、安全区、状态栏或主屏底部样式，回归至少覆盖：普通浏览器、安卓独立主屏、iOS 独立主屏、长聊天、输入框显示、软键盘打开/关闭、根页面无整体滚动，以及消息区仍可独立滚动。",
    ],
)

append_release(
    "AI开发项目_新聊天启动说明.docx",
    "新聊天接手状态｜v877 固定视口发布路线（2026-08-10）",
    [
        "固定仓库为 C:\\Users\\pc\\Documents\\小手机\\phone-work，固定分支 main，远端 origin/main。只走这一条发布路线：先 fetch 与 ff-only 对齐，确认工作区，完成最小修改和测试，提交后执行 git push origin HEAD:main；禁止创建新分支、worktree、旧版本目录、强推或覆盖远端。",
        "当前版本基线为 v877 · 聊天视口与输入框恢复。v870 的主模型与辅助模型统一路线保存、两个模型卡片右上角保存按钮，以及 v871 至 v876 的后续功能全部保留；后续修复不得通过整体回退抹掉这些版本。",
        "v873 的 north-standalone-shell、--north-shell-height 和 screen.height 补齐方案已经判定为回归根因并删除。以后遇到底部黑缝，不得再次用物理屏幕高度拉长 html/body/phone/screen，也不得恢复全局 visualViewport 高度同步。",
        "聊天页稳定契约：根外壳最多一屏且 overflow hidden；app/page 裁切；chatbg 内部滚动；inputbar 固定占位。任何新增聊天卡片、顶部状态栏或底部面板都必须在这一契约内工作。",
        "发布前必须运行完整 tests/*.test.mjs 并确认版本、缓存、Service Worker 对齐；维护文档四件套继续追加，不删除历史记录。v877 完整 Node 回归为 422/422。",
    ],
)

print("Updated v877 maintenance documents")
