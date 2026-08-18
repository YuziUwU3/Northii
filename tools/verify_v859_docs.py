from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
EXPECTED = {
    "AI开发项目_项目说明文档.docx": ["v859｜主动联系防复读与单一发布路线", "377 项自动化测试全部通过"],
    "AI开发项目_Bug记录模板.docx": ["v859 Bug 记录｜主动消息极短复读与跨路径重复", "没有用旧 v857 文件覆盖 v858"],
    "AI开发项目_Bug修改规范.docx": ["主动联系防重复与发布路线收敛", "项目长期只保留一个正式发布路线"],
    "AI开发项目_新聊天启动说明.docx": ["v859 单一发布路线", "唯一正式发布分支：main"],
}

for name, markers in EXPECTED.items():
    path = DOCS / name
    with ZipFile(path) as archive:
        assert archive.testzip() is None, f"corrupt zip member in {name}"
        assert "word/document.xml" in archive.namelist(), f"missing document.xml in {name}"
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for marker in markers:
        assert marker in text, f"missing marker {marker!r} in {name}"
    assert text.count(markers[0]) == 1, f"duplicate v859 section in {name}"
    print(f"verified {name}: {len(document.paragraphs)} paragraphs, {len(document.tables)} tables")
