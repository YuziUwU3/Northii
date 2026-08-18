from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1] / "docs" / "maintenance"
MARKER = "2026-08-11 v892／私人 1.0.17"


COMMON = [
    "统一产品边界：私人“小手机”App 内的网页核心是唯一主界面，原生桥是同一个 App 的真实设备能力，不再把本机能力做成“内置/外置”两个产品模式。控制当前 iPhone 时直接走原生桥；只有控制另一台伴生 iPhone 时才走云端。",
    "本机读取与控制：原生桥契约为 10，device.snapshot 返回当前 iPhone 的屏幕使用时间、健康、定位、电量和锁定账本等快照；device.command 直接执行锁定、解锁和限额。只有 stage = executed 才能显示成功，queued/sent 都不能冒充执行成功。",
    "主动联系：共同生活模式允许全天主动消息和主动来电。每次必须从最近一次会话互动结束后随机静默 30–60 分钟，绝不能固定为 30 分钟；正在聊天、通话或线下场景时不得插入。主动联系间隔为 0 时只按每日随机条数分布，正数时只按固定分钟间隔，两套调度不得同时触发。",
    "语境与防复读：用户说去忙时可自然询问进展或分享角色自己的日常，不得责问为什么不回；用户说睡觉、午睡或给出睡眠时长后，在预计醒来前禁止主动消息和来电。主动消息与主动电话是新事件，不得重新回答、改写或复述上一条用户消息。",
    "后台通知：服务端每条消息分别发送带默认声音的 APNs；通知服务扩展使用角色头像构造通信发送者，点击后打开对应聊天或来电页。iOS 顶层通知仍可能显示 App 图标；当前主动来电是有声音的通知并跳转 App 内电话，不得描述成持续响铃的 CallKit/VoIP 真来电。",
    "服务端要求：应用 supabase/migrations/202608110003_phone_role_push_all_day_random_idle.sql，并重新部署 supabase/functions/phone-role-push/index.ts。未部署时，强退后的全天主动联系、30–60 分钟静默与睡眠安静期不算生效。",
    "验证边界：Windows 静态与自动化测试为 458/458 通过，app.js 语法检查通过；Mac 上的五 Target 编译、签名、Family Controls/Managed Settings 真机执行、通知扩展、后台/强退 APNs 和长期稳定性仍待验证。",
]


DOC_CONTENT = {
    "AI开发项目_项目说明文档.docx": (
        "本轮项目状态与产品定义",
        COMMON,
    ),
    "AI开发项目_Bug修改规范.docx": (
        "新增强制修改规范",
        [
            "同一 App 内已有本机原生能力时，必须优先接本机桥，不得为了复用旧云链路再次制造“内置/外置”双模式。云端只作远程设备与后台任务链路。",
            "任何设备控制都必须区分 queued、sent、received、executed、failed；没有 executed 真回执不得写成功。锁定账本的安全优先级不得被刷新、旧快照、跨日或网络失败降低。",
            "主动联系只有一套最终调度。会话后随机 30–60 分钟静默、固定联系间隔、每日上限是三个独立约束：间隔为 0 时不启用固定间隔；正数时禁用随机条数计时器，只保留每日上限。",
            "主动内容生成前必须判断正在聊天/通话/线下场景和忙碌/睡眠语境。睡眠安静期是硬禁止；忙碌只能自然询问进展或分享日常；消息和电话均不得回答或复述上一条用户消息。",
            "后台消息必须逐条发送并带声音；头像使用通信通知扩展实现。不得承诺替换 iOS 顶层 App 图标，也不得用普通 APNs 假装 CallKit/VoIP 持续来电。",
            "每次改私人 iOS 候选版，必须同步网页版本、Service Worker 缓存、所有入口 cache 参数、桥契约、Xcode marketing/build version、唯一 PhoneWeb.bundle、测试和四份维护文档。",
            "完成声明必须以 Mac 编译和真机测试为界。Windows 的 458/458 只证明静态契约与网页回归，不得改写为真机已通过。",
        ],
    ),
    "AI开发项目_Bug记录模板.docx": (
        "Bug 记录：同一 App 能力割裂、主动联系冲突与后台通知不完整",
        [
            "现象：网页已经装进私人 iOS App，但角色仍把真实设备能力当作外置伴生模式；共同生活主动联系受到在线模式抑制，时间窗、每日条数与固定间隔可能冲突；后台通知不能完整表达角色发送者和来电入口。",
            "根因：网页层沿用“远程伴生手机”单一路径，没有先识别当前容器的原生桥；主动联系把共居、在线和间隔条件叠成多个调度入口；睡眠/忙碌信息只在对话文本中，没有形成服务端安静期；通知只按普通单条 App alert 处理。",
            "修复：新增 device.snapshot 和 device.command 本机直连，只有 executed 才算控制成功；私人 App 隐藏配对/演示/双数据源入口。共同生活启用本地与服务端主动联系，统一为全天，每次会话后随机静默 30–60 分钟；间隔 0 与正数采用互斥调度。加入忙碌/睡眠判断、上一句防复读和主动电话新事件提示。APNs 改为逐条、有声、通信发送者头像及点击路由。",
            "服务端变更：新增 202608110003_phone_role_push_all_day_random_idle.sql，扩展 quiet_until_at 和活动触碰函数；phone-role-push Edge Function 同步执行 30 分钟硬下限、随机 30–60 分钟计划、睡眠安静期和每条通知发送。",
            "自动化结果：node --check app.js 通过；node --test tests/*.test.mjs 为 458/458 通过。覆盖本机桥、共同生活、主动联系、服务端推送、通知路由、版本与资源契约。",
            "尚未验证：Windows 无法编译 Xcode。必须在 Mac 编译五个 Target，并在真机验证真实读取、锁定/解锁/限额、前台/后台/强退逐条通知、头像、声音、点击来电、睡眠静默和多次 30–60 分钟调度。",
        ],
    ),
    "AI开发项目_新聊天启动说明.docx": (
        "下一次聊天接手基线",
        [
            "先读本文件、项目说明、Bug 修改规范、Bug 记录和 docs/maintenance/私人小手机App_唯一总纲.md；当前基线是网页 v892、私人 iOS 1.0.17 (17)、原生桥契约 10。",
            "最高优先级仍是同一私人 App 内的真实读取与锁定/解锁/限额。先核验 device.snapshot、device.command、stage = executed 和唯一 PhoneWeb.bundle，不要先做边缘界面美化，也不要恢复内置/外置双模式。",
            "主动联系验收必须使用“会话停止后随机 30–60 分钟”，不是固定 30 分钟。间隔填 0 时按每日随机条数；填正数时按固定分钟间隔；正在聊天/通话/线下、预计睡眠期间禁止主动联系。",
            "部署检查：应用 202608110003 迁移并重新部署 phone-role-push。没有服务端部署时，App 强退后的消息/来电和睡眠安静期不能宣称完成。",
            "真机顺序：Mac 编译五 Target → 覆盖安装并保留数据 → 授予 Family Controls、通知、定位和健康权限 → 测本机读取与 executed 锁定/解锁 → 测前台、后台、强退的逐条有声通知与头像/点击路由 → 测忙碌、睡眠和 30–60 分钟随机主动联系。",
            "已知系统边界：通信通知可在支持位置显示角色头像，但 iOS 顶层仍可能保留 App 图标；普通 APNs 来电提醒不是 CallKit/VoIP 持续响铃。不要承诺系统不允许的表现。",
            "Windows 自动化当前为 458/458，通过不等于真机通过；完成后必须把真实失败日志和结果追加回四份 Word 文档，不得覆盖历史。",
        ],
    ),
}


def append_release(path: Path, heading: str, paragraphs: list[str]) -> bool:
    doc = Document(path)
    if any(MARKER in paragraph.text for paragraph in doc.paragraphs):
        return False

    spacer = doc.add_paragraph()
    spacer.add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading(MARKER, level=1)
    doc.add_heading(heading, level=2)
    for text in paragraphs:
        doc.add_paragraph(text, style="List Bullet")

    doc.save(path)
    return True


def main() -> None:
    for filename, (heading, paragraphs) in DOC_CONTENT.items():
        path = ROOT / filename
        changed = append_release(path, heading, paragraphs)
        print(f"{'updated' if changed else 'unchanged'}: {path.name}")


if __name__ == "__main__":
    main()
