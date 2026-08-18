from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
EXPECTED = {
    "AI开发项目_项目说明文档.docx": (
        "v871｜共同生活模型与 API 路线分离（2026-08-10）",
        ["四个选择器", "本次共同生活请求", "replyApiRoute", "412/412"],
    ),
    "AI开发项目_Bug记录模板.docx": (
        "v871 Bug 记录｜共同生活把模型与 API 路线混成一个选择（2026-08-10）",
        ["replyModel", "routeIndex", "S.settings.chatRouteActive", "412/412"],
    ),
    "AI开发项目_Bug修改规范.docx": (
        "新增强制规范｜模型选择与 API 路线必须正交（v871 起）",
        ["禁止把主模型写成主线路", "请求级 routeIndex", "路由稳定", "真实 API Key"],
    ),
    "AI开发项目_新聊天启动说明.docx": (
        "新聊天接手状态｜v871 共同生活模型与 API 路线分离（2026-08-10）",
        ["replyApiRoute", "chatRequestRoute", "四个清晰区域", "412/412"],
    ),
}

SUPPLEMENTS = {
    "AI开发项目_项目说明文档.docx": (
        "v871 补充｜共同生活查看手机只交付一次回复（2026-08-10）",
        ["只负责启动查看", "等待 cohabRunPhoneInspection", "22/22"],
    ),
    "AI开发项目_Bug记录模板.docx": (
        "v871 Bug 记录｜共同生活查看手机后同时回复两次（2026-08-10）",
        ["同一回合两次回答", "cohabApplyPhoneTags", "inspection 标记"],
    ),
    "AI开发项目_Bug修改规范.docx": (
        "新增强制规范｜异步查看动作必须单回合交付（v871 起）",
        ["只能有一个成为用户可见回复", "inspection／deferred", "不得回放动作前台词"],
    ),
    "AI开发项目_新聊天启动说明.docx": (
        "新聊天接手补充｜共同生活手机查看单回合规则（2026-08-10）",
        ["inspectionOwner:'offAI'", "opt.schedule=false", "completed inspection reaction"],
    ),
}


for filename, (marker, required) in EXPECTED.items():
    document = Document(DOCS / filename)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    text = "\n".join(paragraphs)
    assert paragraphs.count(marker) == 1, f"duplicate or missing v871 section in {filename}"
    assert paragraphs[-1], f"blank ending in {filename}"
    for phrase in required:
        assert phrase in text, f"missing {phrase!r} in {filename}"
    supplement, supplement_required = SUPPLEMENTS[filename]
    assert paragraphs.count(supplement) == 1, f"duplicate or missing v871 supplement in {filename}"
    for phrase in supplement_required:
        assert phrase in text, f"missing supplement phrase {phrase!r} in {filename}"
    print(filename, f"paragraphs={len(paragraphs)}", f"tables={len(document.tables)}")

print("v871 maintenance documents verified")
