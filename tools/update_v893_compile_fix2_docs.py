from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1] / "docs" / "maintenance"
MARKER = "2026-08-11 v893／私人 1.0.18 Xcode 编译修正 2"

ENTRIES = {
    "AI开发项目_项目说明文档.docx": [
        "第二张 Mac 截图仍显示 fileprivate 红错，证明 Mac 实际编译路径仍含旧 CompanionSyncView.swift；新修正版改用全新 SmallPhone_v893_CompileFix2 工程目录，禁止覆盖进旧目录混编。",
        "补齐 ContentView 对 iOS 26 approvedWithDataAccess 的显式 switch 分支；RoleNotificationService Target 显式链接 Intents 与 UserNotifications，消除通知扩展常见的链接缺项。",
        "定位反查迁移到 iOS 26 MapKit MKReverseGeocodingRequest；角标清理只使用 setBadgeCount，移除 iOS 17 已弃用回退。Mac 链接器最终结果仍必须由全新目录重新编译确认。",
    ],
    "AI开发项目_Bug修改规范.docx": [
        "同名 Xcode 工程连续覆盖可能保留旧源码路径或形成嵌套目录。若修正后的源码仍报完全相同的旧错误，应交付全新外层目录并要求直接打开其中的 xcodeproj，不得继续盲目覆盖。",
        "通知服务扩展使用 Intents 与 UserNotifications 时，工程校验必须确认两个系统 Framework 同时出现在该扩展的 Frameworks build phase；所有 PBX 对象引用、plist 和 entitlements 需在打包前静态校验。",
    ],
    "AI开发项目_Bug记录模板.docx": [
        "现象：Mac 仍出现 registerPushTokenIfAvailable 的旧 fileprivate 错误，同时新增 Command Ld failed；ContentView 还显示 Switch must be exhaustive。",
        "证据判断：当前仓库与第一修正版压缩包均已不存在 fileprivate 声明，因此相同旧错误来自 Mac 仍在编译旧目录或新旧文件混合，并非该修正行失效。Command Ld 未展开，不能虚构具体符号错误。",
        "修复：补齐 approvedWithDataAccess 分支；为 RoleNotificationService 显式链接 Intents/UserNotifications；迁移两项已弃用 API；以全新独立工程目录交付，避免覆盖混编。",
        "验证：PBX 对象引用、五个 Target、14 个 Swift 源文件、10 个 plist/entitlements、访问级别和 Framework 链接均已静态检查；Mac 仍须重新编译验证实际链接器输出。",
    ],
    "AI开发项目_新聊天启动说明.docx": [
        "当前应使用 SmallPhone_v893_CompileFix2 全新目录中的 PhoneCompanionTest.xcodeproj，禁止再从旧 XcodeProject 打开或把新目录拖进旧目录。",
        "若全新目录编译后仍有 Command Ld failed，必须展开该项并记录 Undefined symbols、duplicate symbols 或 framework not found 的完整下一层内容；只有顶层一句不能确定剩余根因。",
    ],
}


def append_entry(path: Path, items: list[str]) -> bool:
    doc = Document(path)
    if any(MARKER in paragraph.text for paragraph in doc.paragraphs):
        return False
    doc.add_heading(MARKER, level=1)
    for item in items:
        doc.add_paragraph(item, style="List Bullet")
    doc.save(path)
    return True


def main() -> None:
    for filename, items in ENTRIES.items():
        path = ROOT / filename
        changed = append_entry(path, items)
        with ZipFile(path) as archive:
            assert archive.testzip() is None
            assert "[Content_Types].xml" in archive.namelist()
        doc = Document(path)
        marker_count = sum(MARKER in paragraph.text for paragraph in doc.paragraphs)
        assert marker_count == 1
        print(
            f"{'updated' if changed else 'unchanged'}: {path.name}; "
            f"marker={marker_count}; paragraphs={len(doc.paragraphs)}"
        )


if __name__ == "__main__":
    main()
