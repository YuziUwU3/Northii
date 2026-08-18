from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
EXPECTED = {
    "AI开发项目_项目说明文档.docx": ("v878｜苹果主屏幕安全区手动适配（2026-08-10）", ["苹果主屏幕适配", "Android standalone", "424/424"]),
    "AI开发项目_Bug记录模板.docx": ("v878 Bug 记录｜iOS 主屏幕灵动岛遮挡与底部黑区（2026-08-10）", ["灵动岛", "screen.height", "424/424"]),
    "AI开发项目_Bug修改规范.docx": ("新增强制规范｜苹果安全区适配必须手动、独立且不影响安卓（v878 起）", ["north-ios-home-safe", "100dvh", "Android standalone"]),
    "AI开发项目_新聊天启动说明.docx": ("新聊天接手状态｜v878 苹果主屏幕适配发布路线（2026-08-10）", ["phone-work", "origin HEAD:main", "424/424"]),
}


for filename, (title, required) in EXPECTED.items():
    document = Document(DOCS / filename)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    text = "\n".join(paragraphs)
    assert paragraphs.count(title) == 1, f"missing or duplicate heading in {filename}"
    assert paragraphs[-1], f"blank ending in {filename}"
    for phrase in required:
        assert phrase in text, f"missing {phrase!r} in {filename}"
    print(filename, f"paragraphs={len(paragraphs)}", f"tables={len(document.tables)}")

print("v878 maintenance documents verified")
