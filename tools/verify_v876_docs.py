from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
EXPECTED = {
    "AI开发项目_项目说明文档.docx": (
        "v876｜North App Store 审核支持与隐私页面（2026-08-10）",
        ["north-support.html", "north-privacy.html", "Unlisted App", "v876 · North 审核支持页面", "422/422"],
    ),
    "AI开发项目_Bug记录模板.docx": (
        "v876 Bug 记录｜公开审核页面会被 Service Worker 错换成小手机主页（2026-08-10）",
        ["request.mode === 'navigate'", "event.respondWith", "失败方案记录", "422/422"],
    ),
    "AI开发项目_Bug修改规范.docx": (
        "新增强制规范｜公开审核文档必须绕开应用外壳导航回退（v876 起）",
        ["Service Worker", "固定文件名或明确白名单", "service_role"],
    ),
    "AI开发项目_新聊天启动说明.docx": (
        "新聊天接手状态｜v876 North 非公开 App 审核资料（2026-08-10）",
        ["phone-work", "com.qianyi.PhoneCompanionTest", "用户推送", "north-(support|privacy).html"],
    ),
}


for filename, (title, required) in EXPECTED.items():
    document = Document(DOCS / filename)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    text = "\n".join(paragraphs)
    assert paragraphs.count(title) == 1, f"missing or duplicate heading in {filename}"
    assert paragraphs[-1], f"blank ending in {filename}"
    for phrase in required:
        assert phrase in text, f"missing {phrase!r} in {filename}"
    print(filename, f"paragraphs={len(paragraphs)}", f"tables={len(document.tables)}")

print("v876 maintenance documents verified")
