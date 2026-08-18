from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_release(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    document = Document(path)
    if any(paragraph.text.strip() == title for paragraph in document.paragraphs):
        print(f"Skipped existing section: {title}")
        return
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"Updated {filename}")


def append_note(filename: str, text: str) -> None:
    path = DOCS / filename
    document = Document(path)
    if any(paragraph.text.strip() == text for paragraph in document.paragraphs):
        print(f"Skipped existing note in {filename}")
        return
    document.add_paragraph(text)
    document.save(path)
    print(f"Added schedule note to {filename}")


append_release(
    "AI开发项目_项目说明文档.docx",
    "v868｜共同生活位置与一起外出同步（2026-08-10）",
    [
        "v868 将共同生活的“是否在同一现场”“当前生活阶段”和“当前地点”拆成明确状态。角色在家时，顶部可显示“在家 · 在卧室/书房/客厅”等；双方一起外出时显示“一起外出 · 具体地点”，仍保留共同生活的面对面输入与回复，不误切到微信。只有角色单独上班、外出或回家路上时，页面才提示去微信联系。",
        "地点只在动作已经真实发生后更新。模型优先输出隐藏的 [共同生活位置|准确地点] 标签；兼容逻辑只识别“回了卧室、走进书房”等完成式动作，忽略“要不要回卧室、想去书房、没有去客厅”等询问、建议、想象和否定，避免只凭地点关键词误移动。地点作为共同生活 home 的单一现场位置持久化，不改变旧约会的旁白、台词和逐字显示格式。",
        "双方一起外出使用独立 together-away 阶段，并用 [共同生活状态|一起外出|准确地点] 与 [共同生活状态|一起回家|玄关] 更新。一起回家不会生成“角色单独到家”的横幅或额外进门承接；角色此前确实独自离开后从微信到家，才进入原有线上到线下承接队列。该区分防止共同出行被误判为异地线上聊天。",
        "回归要求覆盖：实际移动与仅讨论移动的边界、地点标签持久化、顶部状态文字、一起外出保持面对面、一起回家不生成单人到家横幅、角色单独外出仍切换微信，以及旧共同生活上下文、独立记忆、模型路线、一次性约会格式和版本缓存链路。",
    ],
)

append_release(
    "AI开发项目_Bug记录模板.docx",
    "v868 Bug 记录｜地点不跟随与一起外出被误判为单独外出（2026-08-10）",
    [
        "现象：共同生活里双方已经回卧室或进入书房，顶部仍停留在“在家”；双方一起外出也只能落入“外出中”，页面随后要求去微信联系。这样会让同一空间的面对面场景错误断开，并使模型、界面和用户看到的生活状态互相矛盾。",
        "根因：旧实现把 activity 同时当作生活阶段、状态文案和地点，没有独立地点字段，也没有“一起外出”阶段。若简单扫描“卧室、书房”等词，又会把询问、建议和否定误当成已经移动；若复用单人到家判断，一起回家还会错误生成单独到家横幅。",
        "修复：新增 place/placeAt 和 together-away，统一由共同生活状态机驱动顶部标签、输入权限与微信切换。优先解析结构化位置/一起外出标签，标签缺失时只对完成式可见动作做保守兜底。到家队列增加前置阶段判断，together-away 回家不进入单人承接。",
        "防复发：自动化测试必须同时验证正例与反例，不能只看状态文字出现；至少包含“回了卧室”会移动、“要不要回卧室”不会移动、“一起外出”仍可面对面回复、“一起回家”不出现单人到家横幅、“角色单独外出”仍进入微信。禁止把静态预览或提示词存在当成完整链路成功。",
    ],
)

append_release(
    "AI开发项目_Bug修改规范.docx",
    "新增强制规范｜共同生活阶段、同场关系与位置必须分离（v868 起）",
    [
        "共同生活状态不得再用一个自由文本字段同时表示阶段、是否同场和位置。阶段负责在家/单独上班/单独外出/回家路上/一起外出；同场关系决定能否继续面对面；位置只记录已真实抵达的卧室、书房、客厅或外出地点。界面、模型提示和输入权限必须读取同一份持久化状态。",
        "地点更新必须以已经发生的动作或结构化状态标签为准。禁止仅因回答中出现地点名就移动；询问、建议、未来计划、想象、条件句和否定句都不能改变位置。兼容自然语言推断只能作为窄兜底，并必须配套未移动反例测试。",
        "双方一起外出仍属于同一面对面现场，不得自动转微信；双方一起回家不得生成角色单独到家的横幅或重复进门回复。只有角色与用户不在同一空间时，才显示微信入口；只有从单独离家阶段真正到家时，才生成线上到线下承接。",
        "任何共同生活改动都不得改变一次性旧约会的第三人称旁白、普通台词、逐字速度和可见记录格式。验收必须执行正式代码链路测试、共同生活专项测试和完整回归，不能以静态页面、测试按钮或单元片段代替端到端边界核对。",
    ],
)

append_release(
    "AI开发项目_新聊天启动说明.docx",
    "新聊天接手状态｜v868 共同生活位置与一起外出（2026-08-10）",
    [
        "固定仓库为 C:\\Users\\pc\\Documents\\小手机\\phone-work，固定分支 main，远端 origin/main。禁止创建新分支、worktree、强推或使用旧目录。开始前必须 clean、fetch、ff-only；发布前执行相关测试和完整回归，版本号只在当前线上版本基础上加 1。",
        "共同生活位置保存在 S.cohabitation.homes[id].place/placeAt，phase 允许 home、work、returning、away、together-away。cohabTogetherScene 是面对面权限边界：home 与 together-away 为 true；work、returning、away 为 false。顶部状态、微信状态和共同生活输入必须共同读取这套状态。",
        "模型协议：真实移动后使用 [共同生活位置|准确地点]；双方一起外出使用 [共同生活状态|一起外出|准确地点]；双方一起回家使用 [共同生活状态|一起回家|玄关]；角色单独外出才使用 [共同生活状态|外出|状态]。不能根据“要不要、想、准备、没有”等未完成语句更新地点。",
        "发布前重点复核：一起外出时共同生活仍可输入和生成回复；一起回家不出现单人回家横幅；角色单独外出仍显示微信入口；旧约会格式零变化；共同生活的上下文、独立总结与模型路线继续按页面设置工作。",
    ],
)

append_note(
    "AI开发项目_项目说明文档.docx",
    "v868 同时把线上角色原有作息表接入共同生活：共同生活设置区可直接打开并修改同一份 c.sched，周一至周五按用户设置的上午/下午工作与午休时间执行，周六、周日默认休息。每轮共同生活提示均提供准确年月日、星期、时间以及昨天/今天/明天的星期关系。到上班时间后由角色先在面对面现场完成准备、告别和真实出门，再输出上班状态；程序不允许只因钟点到达就跳过动作强制切线上。页面顶部显示实时年月日、星期和时间。",
)
append_note(
    "AI开发项目_Bug记录模板.docx",
    "补充现象与修复：角色曾在 2026-08-10 周一把“明天”说成周日，且共同生活没有复用线上作息。根因是共同生活只收到不含星期关系的时间片段，线上 activitySpec/whereNow 也未区分工作日与周末。v868 明确注入今天/昨天/明天星期，作息计算增加周一至周五与周末分支，并在共同生活内提供同一份作息表入口；状态切换仍以角色实际出门为准。",
)
append_note(
    "AI开发项目_Bug修改规范.docx",
    "日期、星期和相对日期必须从同一个真实时间戳计算，禁止让模型自行猜测“明天星期几”。线上与共同生活必须复用角色同一份可编辑作息数据，不得复制第二套固定配置。作息到点是角色必须执行生活动作的约束，不是界面直接切换线上状态的触发器；只有实际离开同一空间并持久化上班/外出状态后，才显示微信入口。",
)
append_note(
    "AI开发项目_新聊天启动说明.docx",
    "v868 作息与时间边界：共同生活设置中的“作息时间表”直接调用角色已有 schedSet/saveSched；roleSchedulePrompt 与 timeAwarenessPrompt 每轮注入真实工作日、星期关系和当前活动。不得新增独立共同生活作息副本，也不得写成到点自动切微信；模型必须先生成真实出门并输出上班状态标签。顶部 cohabLiveTime 每秒读取真实年月日、星期和时间。",
)

print("Updated v868 maintenance documents")
