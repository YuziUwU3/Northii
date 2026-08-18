from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_section(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    document = Document(path)
    if any(paragraph.text.strip() == title for paragraph in document.paragraphs):
        print(f"Skipped existing section: {title}")
        return
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"Updated {filename}")


append_section(
    "AI开发项目_项目说明文档.docx",
    "私人「小手机」真实 Xcode 工程已导入（2026-08-11）",
    [
        "用户提供的 AppleProjects.zip 已核验完整：AppleProjects/PhoneCompanionTest 包含真实 .xcodeproj、主 App、PhoneCompanionReport、PhoneCompanionMonitor、PhoneCompanionShield、RoleNotificationService、4 份 entitlements 和当前未提交的 Mac 源码。旧 PhoneCompanionTest.zip 只有初始工程，不能再作为基线。",
        "真实工程已过滤 .git、xcuserdata、旧 ZIP 和 macOS 元数据后导入 native/private-small-phone/XcodeProject。项目确认 Team 32H9FL4NK7、主 Bundle ID com.qianyi.PhoneCompanionTest、App Group group.com.qianyi.PhoneCompanionTest 及五个 Target。私人副本沿用这些已经签名可用的标识，只把显示名称改为「小手机」；因此在本人设备上会替换同 Bundle ID 的旧 North/伴生开发包，天然不能同时安装。审核中的公开提交没有修改。",
        "私人主 App 已改为加载安装包内 PhoneWeb.bundle 的完整网页核心，并保留原生设备管理入口。网页资源由共享源码生成，不维护第二套前端。原生桥使用白名单 action 和版本号，未知 action 默认拒绝。",
        "已把 v881 的永久锁定账本、当日快照字段、授权复检和真实状态措辞合入私人真实工程，同时修正两个只在真实工程中才能确认的问题：使用量超时改为不等待输家的异步流竞速；scene 进入 inactive 时移除 SwiftUI Map，定位采集与地图渲染分离。",
        "当前仍不是已安装成品：Windows 只能完成结构、资源和静态回归，下一步必须在 Mac 编译全部五个 Target、签名安装，并进行前后台、锁定、解锁、使用时间、断网、跨日和长时间真机验收。",
    ],
)

append_section(
    "AI开发项目_Bug记录模板.docx",
    "真实工程根因补充｜假超时与 MapKit 后台看门狗（2026-08-11）",
    [
        "新证据：真实 ContentView.swift 在定位授权且已有位置后始终挂载 SwiftUI Map；scenePhase 变化只在回到 active 时刷新授权，没有在 inactive/background 阶段释放地图。该实现与崩溃日志中 MKBasicMapView _didEnterBackground、VectorKit flushTileLoads/barrierSync 主线程超时完全对应。",
        "地图修复：私人副本只在 scenePhase == .active 时创建 Map。scene 先进入 inactive 时 SwiftUI 会移除地图资源，避免带着实时地图进入后台；LocationManager 仍独立记录定位，不因地图卸载而停止。此修复必须用前后台至少 20 次真机回归确认。",
        "第二根因：v881 的 fetchTodayDirectUsageWithTimeout 使用 withTaskGroup 竞争读取和 8 秒 sleep，但结构化任务组退出作用域前仍等待所有子任务结束。若 DeviceActivity live 序列不响应取消，界面仍可能一直停在同步中，所以原专项测试只证明存在超时分支，未证明运行时真的能返回。",
        "超时修复：改用 AsyncStream 接收读取任务和定时任务的第一个结果，拿到结果后立即结束流并取消两者，不再等待卡住的输家；真实读取在写入 latestDirectUsageSnapshot 和 UI 状态前再次检查 Task.isCancelled，晚到数据不得覆盖新快照。",
        "验证边界：已完成资源暂存、plist/entitlements 解析和 Node 静态回归；尚未在 Mac 编译，不能把静态通过写成 iPhone 已修复。",
    ],
)

append_section(
    "AI开发项目_Bug修改规范.docx",
    "新增强制规范｜超时必须真实返回，重型 UI 必须随 scene 释放（2026-08-11 起）",
    [
        "异步超时测试不能只搜索 Task.sleep、timeout 常量或 group.cancelAll。若使用 withTaskGroup，必须证明作用域退出不会继续等待不响应取消的子任务；对于可能卡住的系统异步序列，应使用能返回首个结果且不等待输家的边界，并阻止被取消任务晚到写状态。",
        "MapKit、相机、视频、实时渲染等重型 UI 不得在 scene inactive/background 时无条件常驻。收到真实 watchdog 日志后，应沿主线程栈命中的资源建立 scenePhase 生命周期；释放 UI 资源不能顺带停止仍被产品要求保留的独立数据采集。",
        "从用户 Mac 导入工程时必须保留未提交工作树内容，过滤 .git、xcuserdata、旧压缩包和系统元数据；不得用初始 Git commit 覆盖 Xcode 中标记 A/M 的当前文件。导入后记录 Target、Team、Bundle ID、App Group 和 entitlements 的真实值。",
        "Windows 静态测试不替代 Xcode 编译。涉及 AsyncStream、MainActor、DeviceActivity、MapKit 或同步目录资源包的修改，交付前必须在目标 Mac 编译所有 Target，并在真机验证后台切换和超时退出。",
    ],
)

append_section(
    "AI开发项目_新聊天启动说明.docx",
    "新聊天接手状态｜真实私人 Xcode 工程已入库（2026-08-11）",
    [
        "真实 Mac 基线已经取得，不再要求用户重复打包。正确来源是 AppleProjects.zip 内的 AppleProjects/PhoneCompanionTest；旧 PhoneCompanionTest.zip 和 PhoneCompanionTest 2/3 都不是当前完整工程。",
        "私人真实工程位于 native/private-small-phone/XcodeProject，包含五个 Target。显示名称已改为「小手机」，主界面加载本地 PhoneWeb.bundle，原生设备管理通过右上角入口打开。网页核心仍从仓库根目录生成。",
        "已合入 v881 锁定与快照修复，并修正真实工程确认的 MapKit 后台看门狗和结构化任务组假超时。禁止把原来的 withTaskGroup 竞速恢复回来，也禁止让 Map 在 scene 非 active 时常驻。",
        "下一步必须在 Mac 打开私人副本，先编译全部 Target，再签名安装。真机顺序：确认小手机替换旧同 ID 包；重新授权；测试使用量 8 秒退出；锁定/解锁各 3 个 App；前后台 20 次；断网重试；跨日和旧快照；最后导出崩溃日志。",
        "当前 Windows 全量测试和 DOCX 验证完成后才能提交，但即使全部通过，也只能写成“私人 Xcode 源码已准备”，不能写成“真机已完成”。",
    ],
)

append_section(
    "AI开发项目_项目说明文档.docx",
    "私人 App 首次真机安装与本地网页白屏修复（2026-08-11）",
    [
        "私人「小手机」已在真实 iPhone 成功安装，并因沿用 com.qianyi.PhoneCompanionTest 覆盖同 Bundle ID 的旧伴生开发包。原生屏幕时间管理页、既有授权和伴生数据仍可见，说明五个 Target 与原生能力没有因主界面改造丢失。",
        "私人网页主界面第一次打开为白屏。Xcode 真机日志明确记录 WebPageProxy::maybeInitializeSandboxExtensionHandle: url is not inside resource directory url、Could not create a sandbox extension、Ignoring request to load this main resource because it is outside the sandbox。该问题不是网络失败，也不是用户数据丢失。",
        "修复方案保持单一共享核心：构建脚本继续复制同一份小手机.html，同时生成只作原生入口的 PhoneWeb.bundle/index.html；WKWebView 从 index.html 的标准化 URL 直接推导父目录，并把同一父目录传给 allowingReadAccessTo，保证主资源在授权目录之内。导航加载失败时改为显示可见诊断页，不再保留无提示白屏。",
        "用户现有 North备份_2026-08-10.json 已只读校验有效：文件约 4.5 MB，JSON 完整，包含 55 个顶层数据块。备份含聊天、角色、设置及密钥类私人数据，只能用于本次本地迁移；必须先复测主页面成功加载，再预览并导入，原备份不得删除或被空容器覆盖。",
        "当前完成边界：真机安装和根因取证已完成，白屏修复已进入源码和静态回归；仍须重新生成安装包、在 Mac 编译并二次安装，确认主页面与相对资源全部显示后才能开始数据迁移。",
    ],
)

append_section(
    "AI开发项目_Bug记录模板.docx",
    "真机根因补充｜WKWebView 本地资源越过沙盒导致白屏（2026-08-11）",
    [
        "现象：私人 App 可以安装和进入，右上角原生管理入口可打开完整屏幕时间页面，但内置小手机网页区域全白。原生功能可用而网页不可见，不能误判为网络、Family Controls 授权或备份损坏。",
        "决定性日志：url is not inside resource directory url；Could not create a sandbox extension for PhoneCompanionTest.app；WebProcessProxy::hasAssumedReadAccessToURL: no access；Ignoring request to load this main resource because it is outside the sandbox。WebKit 在创建网页进程沙盒扩展时拒绝了主文件。",
        "根因：旧实现分别通过嵌套 Bundle 查询中文主文件和 PhoneWeb.bundle 目录，再交给 loadFileURL。真机 WebKit 计算出的主资源 URL 与读取授权目录不满足同一标准化父子路径约束，因此加载请求在任何 HTML 或 JavaScript 执行前就被拒绝。原错误页使用 about:blank，而导航策略又只允许 file URL，错误页也可能被取消，最终只剩白屏。",
        "修复：生成 ASCII 入口 index.html；从该入口的 standardizedFileURL 直接调用 deletingLastPathComponent 得到唯一读取目录；loadFileURL 只使用这对同源路径。导航策略允许 about 错误页，并实现 didFail/didFailProvisionalNavigation 可见诊断。",
        "防回归：专项测试必须同时断言 index.html 被生成、主文件与 readAccessURL 由同一 fileURL 推导、allowingReadAccessTo 使用该目录、about 错误页可加载以及 provisional failure 有诊断处理。最终结论仍以新包真机日志不再出现 outside the sandbox 且首页可见为准。",
    ],
)

append_section(
    "AI开发项目_Bug修改规范.docx",
    "新增强制规范｜WKWebView 本地资源与沙盒读取范围（2026-08-11 起）",
    [
        "WKWebView 加载安装包本地文件时，main file URL 与 allowingReadAccessTo 目录必须从同一个标准化文件 URL 推导；禁止分别通过不同 Bundle、字符串拼接或未经标准化的路径独立计算两者。读取范围只授予资源文件的必要父目录，不得扩大到用户目录或任意外部路径。",
        "打包入口统一使用 ASCII 文件名 index.html。中文业务源文件仍是小手机.html，构建时复制生成入口，禁止人工维护两份网页源码。打包测试必须确认 index.html、脚本、样式、图片和 vendor 文件同时存在。",
        "所有本地网页容器必须实现 didFail 和 didFailProvisionalNavigation，并显示不会删除数据的可见诊断页；导航白名单必须允许自身 file URL 与 about 错误页。任何白屏都必须先取得 Xcode/WebKit 原始日志，不能按网络问题盲改。",
        "Windows 静态测试只能验证路径契约和打包内容。每次修改本地资源加载方式后必须在真实 iPhone 重新安装，检查控制台没有 outside the sandbox、首页立即可见、相对资源完整、原生管理入口可用，再开始备份导入。",
    ],
)

append_section(
    "AI开发项目_新聊天启动说明.docx",
    "新聊天接手状态｜私人 App 已首装，白屏沙盒修复待二次真机确认（2026-08-11）",
    [
        "私人「小手机」已成功安装到真实 iPhone并覆盖同 Bundle ID 的旧伴生包。原生屏幕时间管理页、授权和既有伴生数据仍在；当前不是从头重做，也不允许把方向改回网页链接或公开 North 控制本人设备。",
        "首装唯一阻塞是私人网页主界面白屏。真实日志已锁定 WKWebView 本地资源沙盒路径不匹配：url is not inside resource directory url / outside the sandbox。禁止把它归因于网络或要求用户重复授权。",
        "源码已改为 PhoneWeb.bundle/index.html、同源标准化父目录读取权限及可见加载失败页。下一步先运行专项和全量测试、重新打包提交，再让用户在 Mac 用新包覆盖安装；确认网页首页显示后再导入备份。",
        "用户有效备份 North备份_2026-08-10.json 约 4.5 MB、JSON 完整、55 个顶层数据块。不要在聊天中输出备份内容、密钥或私人消息；不要删除原文件；不要在白屏容器中提前导入。",
        "二次真机验收顺序：新包覆盖安装；确认首页非白屏和 Xcode 无 outside the sandbox；打开原生管理页确认屏幕时间；导入备份并核对角色/聊天/设置数量；最后再测锁定、回执、前后台和同步稳定性。",
    ],
)

print("Private Xcode import maintenance sections updated")
