from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
EXPECTED = {
    "AI开发项目_项目说明文档.docx": (
        "v880｜持久锁定与当日同步（2026-08-11）",
        ["持久锁账本", "desiredLocked", "usageRevision", "434/434", "跨渠道时间轴", "真机验收"],
    ),
    "AI开发项目_Bug记录模板.docx": (
        "v880 Bug 记录｜刷新误解锁、假手动解锁与跨日旧使用量（2026-08-11）",
        ["rebuildDailyLimitMonitoring", "manualUnlock", "61/61", "上一句话", "不得被文档写成已经通过"],
    ),
    "AI开发项目_Bug修改规范.docx": (
        "新增强制规范｜锁定意图与使用量快照必须分层（v880 起）",
        ["reportedLocked", "显式解锁", "Monitor Extension", "红灯测试", "跨渠道消息强制规范", "独立 App 安全区强制规范"],
    ),
    "AI开发项目_新聊天启动说明.docx": (
        "新聊天接手状态｜v880 持久锁定与当日同步（2026-08-11）",
        ["phone-work", "434/434", "最新用户/角色锚点", "连续刷新 10 次", "私人 App 阶段尚未开始"],
    ),
}


for filename, (title, required) in EXPECTED.items():
    document = Document(DOCS / filename)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    text = "\n".join(paragraphs)
    assert paragraphs.count(title) == 1, f"missing or duplicate heading in {filename}"
    assert paragraphs[-1], f"blank ending in {filename}"
    for phrase in required:
        assert phrase in text, f"missing {phrase!r} in {filename}"
    print(filename, f"paragraphs={len(paragraphs)}", f"tables={len(document.tables)}")

print("v880 maintenance documents verified")
